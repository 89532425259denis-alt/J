# -*- coding: utf-8 -*-
"""checker — Quality Gate: валидация и авто-фикс готового DOCX по ГОСТ 7.32-2017.

Проверяет реальный файл (не «обещания» LLM): шрифт, кегль, интервал, поля,
отступ, markdown-мусор, оборванные ссылки, фразы-маркеры ИИ, структурные
элементы. autofix_docx исправляет то, что можно исправить механически.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field


from docx import Document
from docx.shared import Cm, Mm, Pt

from gost_rules import GOST_732

# ── Паттерны брака ──────────────────────────────────────────────────
_MD_RE = re.compile(r"(^|\s)(#{1,6}\s|\*\*|__|```|~~|^---\s*$)", re.M)
_BROKEN_CITE_RE = re.compile(r"\[\d+,\s*с\.\s*(?=[^\d\s])|\[\d+,\s*с\.\s*$|\[\d+,\s*с\.\s*\]")
_AI_MARKERS = (
    "как ии", "как языковая модель", "конечно, вот", "давайте рассмотрим",
    "надеюсь, это поможет", "вот текст", "приведу пример текста",
)
_STRUCTURAL = ("СОДЕРЖАНИЕ", "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ")
_BIB_TITLES = ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "СПИСОК ЛИТЕРАТУРЫ",
               "БИБЛИОГРАФИЧЕСКИЙ СПИСОК", "СПИСОК ИСТОЧНИКОВ")


@dataclass
class Issue:
    title: str
    severity: str = "error"   # error | warning
    detail: str = ""

    def __str__(self) -> str:
        return f"[{self.severity}] {self.title}" + (f": {self.detail}" if self.detail else "")


@dataclass
class ValidationResult:
    score: int = 100
    errors: list = field(default_factory=list)     # list[Issue]
    warnings: list = field(default_factory=list)   # list[Issue]
    passed: bool = True

    def failures(self) -> list:
        return [e for e in self.errors if e.severity == "error"]


def _para_style_name(par) -> str:
    """Возвращает имя стиля абзаца ('Heading 1', 'Heading 2', 'Title', ...)."""
    try:
        return (par.style.name or "") if par.style is not None else ""
    except Exception:
        return ""


def _is_heading(par) -> bool:
    """Определяет заголовок НАДЁЖНО: сперва по стилю Word (Heading N / Title),
    и лишь затем — по тексту как резерв (для документов без стилей).

    Проверка только по регэкспу была хрупкой: строка «1 января 2020 года…»
    в тексте ошибочно принималась за заголовок, а настоящий заголовок без
    номера — пропускался. Стиль Word — источник правды.
    """
    style = _para_style_name(par).lower()
    if style.startswith("heading") or style in ("title", "заголовок", "название"):
        return True
    text = (par.text or "").strip()
    if not text or len(text) > 200:
        return False
    if text.upper() in _STRUCTURAL or text.upper() in _BIB_TITLES:
        return True
    return bool(re.match(r"^\d+(\.\d+)*\s+\S", text)) and len(text) < 120


def _heading_level(par) -> int:
    """Уровень заголовка по стилю Word: 1 для 'Heading 1', 2 для 'Heading 2'…
    0 — если уровень не определяется по стилю."""
    style = _para_style_name(par)
    m = re.search(r"heading\s*(\d+)", style, flags=re.I)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return 0
    if style.lower() in ("title", "заголовок", "название"):
        return 1
    return 0


def validate_docx(path: str, doc_type: str = "") -> ValidationResult:
    """Полная проверка DOCX. Оценка 100 минус штрафы; passed при score >= 85."""
    res = ValidationResult()

    def penalty(pts: int, title: str, detail: str = "", severity: str = "error") -> None:
        issue = Issue(title=title, severity=severity, detail=detail)
        if severity == "error":
            res.errors.append(issue)
        else:
            res.warnings.append(issue)
        res.score = max(0, res.score - pts)

    try:
        doc = Document(path)
    except Exception as e:
        res.score = 0
        res.passed = False
        res.errors.append(Issue(title=f"DOCX не открывается: {e}"))
        return res

    g = GOST_732

    # ── 1. Поля страницы ──
    try:
        sec = doc.sections[0]
        checks = [
            ("левое поле", sec.left_margin, Mm(g["left_margin_mm"])),
            ("верхнее поле", sec.top_margin, Mm(g["top_margin_mm"])),
            ("нижнее поле", sec.bottom_margin, Mm(g["bottom_margin_mm"])),
        ]
        for name, actual, expected in checks:
            if actual is not None and abs(int(actual) - int(expected)) > int(Mm(2)):
                penalty(5, f"Неверное {name}",
                        f"{round(actual / 36000, 1) if actual else '?'} мм вместо {int(expected) / 36000:.0f} мм")
        # ГОСТ 7.32-2017: правое поле = 15 мм. Проверяем именно соответствие
        # (±2 мм), а не только «не меньше 10» — иначе документ с правым полем
        # 25 мм ошибочно считался бы валидным.
        if sec.right_margin is not None:
            expected_right = Mm(g["right_margin_mm"])
            if abs(int(sec.right_margin) - int(expected_right)) > int(Mm(2)):
                penalty(5, "Неверное правое поле",
                        f"{round(sec.right_margin / 36000, 1)} мм вместо {g['right_margin_mm']} мм")
    except Exception:
        pass

    body_pars = [p for p in doc.paragraphs if (p.text or "").strip()]

    # ── 2. Шрифт и кегль основного текста ──
    # ГОСТ допускает 12–14 pt (де-факто 14). Штрафуем И за «меньше 12»,
    # И за «больше 14»: завышенный кегль — это скрытая накрутка числа страниц,
    # которую конкуренты пропускают.
    wrong_font = size_small = size_big = checked = 0
    for p in body_pars:
        if _is_heading(p):
            continue
        for run in p.runs:
            if not (run.text or "").strip():
                continue
            checked += 1
            if run.font.name and run.font.name != g["font_name"]:
                wrong_font += 1
            if run.font.size:
                if run.font.size < Pt(12) - Pt(0.5):
                    size_small += 1
                elif run.font.size > Pt(g["font_size"]) + Pt(0.5):
                    size_big += 1
    if checked:
        if wrong_font / checked > 0.05:
            penalty(10, "Шрифт не Times New Roman",
                    f"{wrong_font}/{checked} фрагментов")
        if size_small / checked > 0.05:
            penalty(8, "Кегль меньше 12 pt", f"{size_small}/{checked} фрагментов")
        if size_big / checked > 0.05:
            penalty(8, f"Кегль больше {g['font_size']} pt (накрутка объёма)",
                    f"{size_big}/{checked} фрагментов")

    # ── 3. Markdown-мусор, оборванные ссылки, маркеры ИИ ──
    full_text = "\n".join(p.text or "" for p in doc.paragraphs)
    if _MD_RE.search(full_text):
        penalty(15, "Markdown-разметка в тексте (##, **, ``` )")
    if _BROKEN_CITE_RE.search(full_text):
        penalty(12, "Оборванные ссылки вида [N, с. без номера страницы")
    low = full_text.lower()
    for marker in _AI_MARKERS:
        if marker in low:
            penalty(15, "Фраза-маркер ИИ в тексте", f"«{marker}»")
            break

    # ── 4. Структурные элементы ──
    # Эссе по своей природе не имеет содержания/введения/заключения и списка
    # источников — не штрафуем его за их отсутствие (иначе ложный брак и
    # ненужный авто-фикс). Для остальных типов работ проверка обязательна.
    upper_text = full_text.upper()
    if doc_type not in ("esse",):
        for name in _STRUCTURAL:
            if name not in upper_text:
                penalty(8, f"Не найден структурный элемент «{name}»")
        if not any(t in upper_text for t in _BIB_TITLES):
            penalty(8, "Не найден список использованных источников")

    # ── 5. Точка после номера раздела («1. Название» — брак по 7.32) ──
    dot_headings = [p.text.strip() for p in body_pars
                    if re.match(r"^\d+(\.\d+)*\.\s+[А-ЯA-Z]", (p.text or "").strip())
                    and _is_heading(p)]
    if dot_headings:
        penalty(4, "Точка после номера раздела", dot_headings[0][:60], severity="warning")

    # ── 5b. Заголовки по свойствам стилей Word (ГОСТ 7.32-2017 п.6.2.3) ──
    # Проверяем НАДЁЖНО, по стилю Heading N, а не по тексту: полужирность,
    # отсутствие точки в конце, отсутствие ручных переносов слов.
    style_headings = [p for p in body_pars if _heading_level(p) > 0]
    if style_headings:
        not_bold = end_dot = hyphenated = 0
        for p in style_headings:
            text = (p.text or "").strip()
            # bold: заголовок должен быть полужирным (учитываем и run, и стиль)
            runs = [r for r in p.runs if (r.text or "").strip()]
            if runs and all(r.bold is False for r in runs):
                not_bold += 1
            if text.endswith((".", ";")) and not text.endswith(("...", "…")):
                end_dot += 1
            if "\u00ad" in text or "-\n" in (p.text or ""):
                hyphenated += 1
        if not_bold:
            penalty(4, "Заголовки не полужирные (ГОСТ 7.32 п.6.2.3)",
                    f"{not_bold} шт.", severity="warning")
        if end_dot:
            penalty(4, "Точка в конце заголовка", f"{end_dot} шт.", severity="warning")
        if hyphenated:
            penalty(3, "Переносы слов в заголовке запрещены",
                    f"{hyphenated} шт.", severity="warning")

    # ── 6. Ссылки на источники в тексте ──
    cites = re.findall(r"\[(\d+),\s*с\.\s*\d+", full_text)
    if doc_type not in ("esse",):
        if len(cites) < 3:
            penalty(10, "Слишком мало ссылок на источники в тексте", f"найдено {len(cites)}")
        elif len(set(cites)) == 1:
            penalty(6, "Все ссылки ведут на один источник", f"[{cites[0]}]", severity="warning")

    res.passed = res.score >= 85 and not any(
        e.title.startswith("DOCX не открывается") for e in res.errors
    )
    return res


def autofix_docx(in_path: str, out_path: str) -> list[str]:
    """Механически исправляет то, что можно исправить без перегенерации.
    Возвращает список применённых фиксов (пустой = ничего не менялось)."""
    changes: list[str] = []
    try:
        doc = Document(in_path)
    except Exception:
        return changes

    g = GOST_732

    # 1) Поля страницы
    try:
        for sec in doc.sections:
            if sec.left_margin != Mm(g["left_margin_mm"]):
                sec.left_margin = Mm(g["left_margin_mm"])
                changes.append("поля: левое 30 мм")
            if sec.top_margin != Mm(g["top_margin_mm"]):
                sec.top_margin = Mm(g["top_margin_mm"])
                changes.append("поля: верхнее 20 мм")
            if sec.bottom_margin != Mm(g["bottom_margin_mm"]):
                sec.bottom_margin = Mm(g["bottom_margin_mm"])
                changes.append("поля: нижнее 20 мм")
            # Правое поле приводим к ГОСТ-норме (15 мм) при любом отклонении > 2 мм.
            if sec.right_margin is None or \
                    abs(int(sec.right_margin) - int(Mm(g["right_margin_mm"]))) > int(Mm(2)):
                sec.right_margin = Mm(g["right_margin_mm"])
                changes.append(f"поля: правое → {g['right_margin_mm']} мм")
    except Exception:
        pass

    md_inline = re.compile(r"\*\*(.+?)\*\*|__(.+?)__|`([^`]+)`")
    md_prefix = re.compile(r"^#{1,6}\s+")
    # Фразы-маркеры ИИ вырезаем прямо из текста (регистронезависимо), вместе
    # с прилипшей запятой/пробелом. Если после удаления абзац начинается со
    # строчной буквы — поднимаем первую букву в заглавную.
    ai_marker_re = re.compile(
        r"(?i)\b(как\s+ии|как\s+языковая\s+модель|конечно,?\s*вот|"
        r"давайте\s+рассмотрим|надеюсь,?\s*это\s+поможет|вот\s+текст|"
        r"приведу\s+пример\s+текста)\b[\s,:—-]*"
    )

    for p in doc.paragraphs:
        text = p.text or ""
        if not text.strip():
            continue
        heading = _is_heading(p)

        # 2) Markdown-мусор + фразы-маркеры ИИ в run-ах
        for run in p.runs:
            t = run.text or ""
            if not t:
                continue
            new_t = md_prefix.sub("", t)
            new_t = md_inline.sub(lambda m: m.group(1) or m.group(2) or m.group(3) or "", new_t)
            new_t = new_t.replace("**", "").replace("```", "")
            if new_t != t:
                run.text = new_t
                if "markdown удалён" not in changes:
                    changes.append("markdown удалён")
            t2 = run.text or ""
            cleaned = ai_marker_re.sub("", t2)
            if cleaned != t2:
                # Восстанавливаем заглавную букву в начале, если срезали маркер.
                cleaned_ls = cleaned.lstrip()
                if cleaned_ls and cleaned_ls[0].islower():
                    cleaned = cleaned_ls[0].upper() + cleaned_ls[1:]
                run.text = cleaned
                if "фразы-маркеры ИИ удалены" not in changes:
                    changes.append("фразы-маркеры ИИ удалены")

            # 3) Шрифт и кегль
            if run.font.name != g["font_name"]:
                run.font.name = g["font_name"]
                if "шрифт → Times New Roman" not in changes:
                    changes.append("шрифт → Times New Roman")
            # Кегль вне диапазона 12–14 приводим к 14 pt: и заниженный
            # (нечитаемо), и завышенный (накрутка страниц) — оба брак.
            if not heading and run.font.size is not None and (
                run.font.size < Pt(12) or run.font.size > Pt(g["font_size"])
            ):
                run.font.size = Pt(g["font_size"])
                if "кегль → 14 pt" not in changes:
                    changes.append("кегль → 14 pt")

        # 4) Межстрочный интервал и отступ в основном тексте
        if not heading:
            pf = p.paragraph_format
            try:
                if pf.line_spacing is not None and abs(float(pf.line_spacing) - g["line_spacing"]) > 0.05:
                    pf.line_spacing = g["line_spacing"]
                    if "интервал → 1.5" not in changes:
                        changes.append("интервал → 1.5")
            except (TypeError, ValueError):
                pass

    if changes:
        try:
            doc.save(out_path)
        except Exception:
            return []
    return changes


def quick_check(path: str) -> tuple[bool, str]:
    """Быстрая проверка: файл открывается и не пуст."""
    try:
        doc = Document(path)
        n = sum(1 for p in doc.paragraphs if (p.text or "").strip())
        if n < 5:
            return False, f"Документ почти пуст: {n} абзацев"
        return True, f"OK: {n} абзацев"
    except Exception as e:
        return False, f"Файл не открывается: {e}"

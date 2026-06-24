# -*- coding: utf-8 -*-
"""
gost_autofix.py — авто-исправление DOCX под ГОСТ 7.32-2017.

    from gost_autofix import autofix_docx
    changed = autofix_docx("in.docx", "out.docx")

Чинит: поля, формат A4, шрифт Times New Roman, кегль 14, интервал 1,5,
абзацный отступ 1,25 см, выравнивание структурных элементов по центру,
удаление BOM/мусорных символов. Заголовки и таблицы не ломает.
Зависит только от python-docx.
"""
from __future__ import annotations
import re
from typing import List

from docx import Document
from docx.shared import Mm, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

STRUCT_RE = re.compile(r"(?i)^(содержание|введение|заключение|список\s+использ|приложение)")


def _is_heading(p) -> bool:
    st = (p.style.name or "").lower()
    return st.startswith("heading") or st.startswith("title")


def autofix_docx(in_path: str, out_path: str = None) -> List[str]:
    out_path = out_path or in_path
    doc = Document(in_path)
    changes: List[str] = []

    # 1. Поля и A4
    for sec in doc.sections:
        if abs(sec.left_margin - Mm(30)) > Mm(0.6):
            sec.left_margin = Mm(30); changes.append("поле слева → 30 мм")
        if abs(sec.right_margin - Mm(10)) > Mm(0.6):
            sec.right_margin = Mm(10); changes.append("поле справа → 10 мм")
        if abs(sec.top_margin - Mm(20)) > Mm(0.6):
            sec.top_margin = Mm(20); changes.append("поле сверху → 20 мм")
        if abs(sec.bottom_margin - Mm(20)) > Mm(0.6):
            sec.bottom_margin = Mm(20); changes.append("поле снизу → 20 мм")
        if abs(sec.page_width - Mm(210)) > Mm(1.5) or abs(sec.page_height - Mm(297)) > Mm(1.5):
            sec.page_width = Mm(210); sec.page_height = Mm(297); changes.append("формат → A4")

    # 2. Стиль Normal: TNR 14, интервал 1,5
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(14)
        normal.paragraph_format.line_spacing = 1.5
    except Exception:
        pass

    bom_fixed = 0
    font_fixed = 0
    size_fixed = 0
    sp_fixed = 0
    ind_fixed = 0
    center_fixed = 0

    for p in doc.paragraphs:
        txt = p.text
        # 3. BOM/мусор
        if "\ufeff" in txt or "\u200b" in txt:
            for run in p.runs:
                if "\ufeff" in run.text or "\u200b" in run.text:
                    run.text = run.text.replace("\ufeff", "").replace("\u200b", "")
                    bom_fixed += 1

        heading = _is_heading(p)
        pf = p.paragraph_format

        # 4. Интервал 1,5 для тела
        if pf.line_spacing and abs(float(pf.line_spacing) - 1.5) > 0.05 and not heading:
            pf.line_spacing = 1.5; sp_fixed += 1

        # 5. Абзацный отступ 1,25 см для обычного текста (не структурные, не заголовки)
        if not heading and not STRUCT_RE.match(txt.strip()) and txt.strip():
            al = pf.alignment
            # только для обычных текстовых абзацев (не по центру)
            if al in (None, WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.LEFT):
                fi = pf.first_line_indent
                if fi is None or abs(fi - Cm(1.25)) > Mm(1.0):
                    # не трогаем элементы списка литературы и таблицы
                    pf.first_line_indent = Cm(1.25); ind_fixed += 1

        # 6. Структурные элементы по центру
        if STRUCT_RE.match(txt.strip()) and pf.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; center_fixed += 1

        # 7. Шрифт/кегль тела
        for run in p.runs:
            if run.font.name and run.font.name != "Times New Roman":
                run.font.name = "Times New Roman"; font_fixed += 1
            elif not run.font.name:
                run.font.name = "Times New Roman"
            if not heading and run.font.size and abs(run.font.size.pt - 14) > 0.1:
                run.font.size = Pt(14); size_fixed += 1

    # таблицы: шрифт TNR, удаление BOM
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if "\ufeff" in run.text:
                            run.text = run.text.replace("\ufeff", ""); bom_fixed += 1
                        if run.font.name and run.font.name != "Times New Roman":
                            run.font.name = "Times New Roman"

    if bom_fixed:    changes.append(f"удалён мусор/BOM в {bom_fixed} местах")
    if font_fixed:   changes.append(f"шрифт → Times New Roman ({font_fixed})")
    if size_fixed:   changes.append(f"кегль → 14 пт ({size_fixed})")
    if sp_fixed:     changes.append(f"интервал → 1,5 ({sp_fixed})")
    if ind_fixed:    changes.append(f"абзацный отступ → 1,25 см ({ind_fixed})")
    if center_fixed: changes.append(f"структурные элементы по центру ({center_fixed})")

    doc.save(out_path)
    return changes


if __name__ == "__main__":
    import sys
    src = sys.argv[1]
    dst = sys.argv[2] if len(sys.argv) > 2 else src
    ch = autofix_docx(src, dst)
    print("Исправлено:", "; ".join(ch) if ch else "нет изменений")

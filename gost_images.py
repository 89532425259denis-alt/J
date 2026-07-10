# -*- coding: utf-8 -*-
"""gost_images — вставка рисунков в DOCX строго по ГОСТ 7.32-2017 (п. 6.5).

Правила, которые модуль гарантирует МЕХАНИЧЕСКИ (а не «обещанием» LLM):
  1. Рисунок размещается по центру, БЕЗ абзацного отступа.
  2. Подпись ПОД рисунком: «Рисунок N — Наименование», по центру,
     тем же шрифтом (Times New Roman), БЕЗ точки в конце.
  3. Тире в подписи — длинное (—), с пробелами. «Рис. 1», «Рисунок 1.»,
     «Рисунок 1 - Название» — брак, модуль нормализует автоматически.
  4. Нумерация сквозная (1, 2, 3…) или в пределах раздела (1.1, 2.1…),
     в приложениях — «Рисунок А.1». За нумерацию отвечает FigureNumberer.
  5. Рисунок не выходит за поля: максимальная ширина = 210 − 30 − 15 = 165 мм,
     максимальная высота ограничена, пропорции сохраняются.
  6. На КАЖДЫЙ рисунок в тексте ДО него должна быть ссылка:
     «в соответствии с рисунком N» или «(рисунок N)». Модуль умеет
     сам дописывать ссылку в абзац перед рисунком (ensure_reference).
  7. Перед рисунком и после подписи — вертикальный интервал (через
     space_before / space_after, а не пустыми абзацами-мусором).
  8. Изображение приводится к поддерживаемому формату (JPEG/PNG),
     CMYK/палитра/альфа конвертируются; битые и слишком мелкие
     картинки отбраковываются до вставки.

Использование в боте:

    from gost_images import FigureNumberer, add_gost_figure, prepare_image

    numberer = FigureNumberer(per_chapter=False)
    ...
    ok, data, w_px, h_px = prepare_image(img["bytes"])
    if ok:
        num = numberer.next(chapter=2)          # "3" или "2.1"
        add_gost_figure(
            doc, data, num, img["caption"],
            ensure_reference=last_text_paragraph,  # допишет «(рисунок 3)»
        )
"""

from __future__ import annotations

import io
import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Pt

from gost_rules import GOST_732

# ══════════════════════════════════════════════════════════════════
#  Параметры рисунков по ГОСТ 7.32-2017 (source of truth)
# ══════════════════════════════════════════════════════════════════

GOST_FIGURE: dict = {
    "caption_prefix": "Рисунок",       # НЕ «Рис.» — сокращение запрещено
    "caption_dash": "—",               # длинное тире с пробелами
    "caption_below": True,             # подпись ПОД рисунком
    "caption_centered": True,
    "caption_no_end_dot": True,        # точка в конце подписи не ставится
    "image_centered": True,
    "no_first_line_indent": True,      # без абзацного отступа
    # Полезная область страницы А4 при полях ГОСТ: 210−30−15 = 165 мм ширина,
    # 297−20−20 = 257 мм высота. Держим запас, чтобы подпись влезала на ту же
    # страницу вместе с рисунком.
    "max_width_cm": 16.5,
    "max_height_cm": 18.0,
    "default_width_cm": 14.0,          # комфортная ширина «по умолчанию»
    "min_pixels": 200,                 # мельче — нечитаемый брак, отбраковываем
    "space_before_pt": 12,             # интервал перед рисунком
    "space_after_caption_pt": 12,      # интервал после подписи
    "reference_before_required": True, # ссылка в тексте строго ДО рисунка
}

# Подпись по ГОСТ: «Рисунок 1 — Название», «Рисунок 2.1 — Название»,
# в приложениях «Рисунок А.1 — Название»
CAPTION_RE = re.compile(
    r"^Рисунок\s+(\d+(?:\.\d+)?|[А-ЯЁ]\.\d+)\s*[—–-]\s*(\S.*)$"
)
# Любое упоминание рисунка в тексте (для проверки ссылок)
_MENTION_RE_TPL = r"рисун\w*\s+{num}\b"

_BAD_PREFIX_RE = re.compile(r"^\s*(Рис\.?|Рисунок|Fig\.?|Figure|Изображение|Картинка)\s*№?\s*", re.I)
_MD_JUNK_RE = re.compile(r"[*_#`]+")


# ══════════════════════════════════════════════════════════════════
#  Нумерация
# ══════════════════════════════════════════════════════════════════

class FigureNumberer:
    """Выдаёт номера рисунков: сквозные («1», «2»…) либо по разделам
    («1.1», «1.2», «2.1»…), для приложений — «А.1», «А.2»…"""

    def __init__(self, per_chapter: bool = False) -> None:
        self.per_chapter = bool(per_chapter)
        self._global = 0
        self._by_chapter: dict[int, int] = {}
        self._by_appendix: dict[str, int] = {}

    def next(self, chapter: int = 0) -> str:
        if self.per_chapter and chapter > 0:
            self._by_chapter[chapter] = self._by_chapter.get(chapter, 0) + 1
            return f"{chapter}.{self._by_chapter[chapter]}"
        self._global += 1
        return str(self._global)

    def next_appendix(self, letter: str) -> str:
        letter = (letter or "А").strip().upper()[:1] or "А"
        self._by_appendix[letter] = self._by_appendix.get(letter, 0) + 1
        return f"{letter}.{self._by_appendix[letter]}"

    @property
    def count(self) -> int:
        return (self._global
                + sum(self._by_chapter.values())
                + sum(self._by_appendix.values()))


# ══════════════════════════════════════════════════════════════════
#  Подготовка изображения (валидация + конвертация)
# ══════════════════════════════════════════════════════════════════

def prepare_image(data: bytes) -> tuple[bool, bytes, int, int]:
    """Проверяет и нормализует изображение перед вставкой.

    → (ok, bytes, width_px, height_px). При ok=False изображение — брак
    (битое, не картинка, слишком мелкое) и вставлять его НЕЛЬЗЯ.

    CMYK/палитра/альфа конвертируются в RGB-JPEG: python-docx и Word
    непредсказуемо ведут себя с экзотическими режимами. Если Pillow
    не установлен — деградируем мягко: вставляем как есть.
    """
    if not data or len(data) < 100:
        return False, b"", 0, 0
    try:
        from PIL import Image  # optional dependency
    except Exception:
        return True, data, 0, 0  # без Pillow — доверяем данным

    try:
        img = Image.open(io.BytesIO(data))
        img.load()
        w, h = img.size
        if w < GOST_FIGURE["min_pixels"] or h < GOST_FIGURE["min_pixels"]:
            return False, b"", w, h
        fmt = (img.format or "").upper()
        if fmt in ("JPEG", "PNG") and img.mode in ("RGB", "L", "RGBA" if fmt == "PNG" else "RGB"):
            return True, data, w, h
        # Конвертация: любые режимы → RGB, сохраняем в JPEG (белая подложка
        # вместо прозрачности — на печати по ГОСТ прозрачности нет).
        if img.mode in ("RGBA", "LA", "P"):
            rgba = img.convert("RGBA")
            bg = Image.new("RGB", rgba.size, (255, 255, 255))
            bg.paste(rgba, mask=rgba.split()[-1])
            img = bg
        elif img.mode != "RGB":
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=90)
        return True, buf.getvalue(), w, h
    except Exception:
        return False, b"", 0, 0


def calc_display_size(width_px: int, height_px: int,
                      desired_width_cm: float | None = None) -> tuple[float, float]:
    """Размер рисунка в сантиметрах: не шире 16,5 см (поля ГОСТ),
    не выше 18 см, пропорции сохраняются."""
    max_w = GOST_FIGURE["max_width_cm"]
    max_h = GOST_FIGURE["max_height_cm"]
    w = min(desired_width_cm or GOST_FIGURE["default_width_cm"], max_w)
    if width_px > 0 and height_px > 0:
        aspect = height_px / width_px
        h = w * aspect
        if h > max_h:
            h = max_h
            w = h / aspect
    else:
        h = 0.0
    return round(w, 2), round(h, 2)


# ══════════════════════════════════════════════════════════════════
#  Подпись
# ══════════════════════════════════════════════════════════════════

def clean_caption_title(title: str) -> str:
    """Нормализует наименование рисунка: без «Рис. 1 —» в начале,
    без markdown-мусора, без точки в конце, с заглавной буквы."""
    t = _MD_JUNK_RE.sub("", str(title or "")).strip()
    # срезаем уже готовые префиксы «Рисунок 1 —», «Рис. 2 -» и т. п.
    t = _BAD_PREFIX_RE.sub("", t)
    t = re.sub(r"^\d+(\.\d+)?\s*[—–:.-]\s*", "", t).strip()
    t = t.strip(" .;,—–-")
    if not t:
        t = "Иллюстрация к разделу"
    return t[0].upper() + t[1:]


def format_caption(number: str, title: str) -> str:
    """«Рисунок 1 — Название» — строго по ГОСТ 7.32-2017."""
    return (f"{GOST_FIGURE['caption_prefix']} {number} "
            f"{GOST_FIGURE['caption_dash']} {clean_caption_title(title)}")


def make_reference(number: str, style: str = "short") -> str:
    """Текстовая ссылка на рисунок для абзаца ПЕРЕД ним.
    style='short' → «(рисунок N)», style='full' → «в соответствии с рисунком N»."""
    if style == "full":
        return f"в соответствии с рисунком {number}"
    return f"(рисунок {number})"


# ══════════════════════════════════════════════════════════════════
#  Вставка в документ
# ══════════════════════════════════════════════════════════════════

def _apply_gost_font(run, size_pt: int | None = None) -> None:
    run.font.name = GOST_732["font_name"]
    run.font.size = Pt(size_pt or GOST_732["font_size"])
    try:
        # кириллица в Word требует отдельного east-asian имени шрифта
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
        if rFonts is None:
            from docx.oxml.ns import qn
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.append(rFonts)
        from docx.oxml.ns import qn
        rFonts.set(qn("w:ascii"), GOST_732["font_name"])
        rFonts.set(qn("w:hAnsi"), GOST_732["font_name"])
        rFonts.set(qn("w:cs"), GOST_732["font_name"])
        rFonts.set(qn("w:eastAsia"), GOST_732["font_name"])
    except Exception:
        pass


def _center_no_indent(paragraph) -> None:
    pf = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)


def add_gost_figure(
    doc,
    image_data: bytes,
    number: str,
    title: str,
    *,
    width_cm: float | None = None,
    width_px: int = 0,
    height_px: int = 0,
    ensure_reference=None,
    keep_with_caption: bool = True,
) -> tuple[object, object] | None:
    """Вставляет рисунок + подпись по ГОСТ в конец документа.

    doc            — docx.Document
    image_data     — байты изображения (лучше после prepare_image)
    number         — номер от FigureNumberer («3», «2.1», «А.1»)
    title          — наименование (нормализуется автоматически)
    ensure_reference — абзац, в конец которого дописать «(рисунок N)»,
                       если ссылки на рисунок в нём ещё нет
    keep_with_caption — запрет разрыва страницы между рисунком и подписью

    → (параграф_рисунка, параграф_подписи) либо None при браке картинки.
    """
    ok, data, w_px, h_px = prepare_image(image_data)
    if not ok:
        return None
    if width_px and height_px:
        w_px, h_px = width_px, height_px

    # 1) Ссылка в тексте ДО рисунка (обязательное требование ГОСТ)
    if ensure_reference is not None:
        try:
            _ensure_reference_in_paragraph(ensure_reference, number)
        except Exception:
            pass

    w_cm, _h_cm = calc_display_size(w_px, h_px, width_cm)

    # 2) Абзац с рисунком: по центру, без отступа, интервал сверху
    img_par = doc.add_paragraph()
    _center_no_indent(img_par)
    img_par.paragraph_format.space_before = Pt(GOST_FIGURE["space_before_pt"])
    img_par.paragraph_format.space_after = Pt(6)
    if keep_with_caption:
        img_par.paragraph_format.keep_with_next = True  # рисунок не отрывается от подписи
    run = img_par.add_run()
    try:
        run.add_picture(io.BytesIO(data), width=Cm(w_cm))
    except Exception:
        # битый файл прошёл prepare_image (нет Pillow) — откатываем абзац
        img_par._element.getparent().remove(img_par._element)
        return None

    # 3) Подпись ПОД рисунком: «Рисунок N — Название», по центру, без точки
    cap_par = doc.add_paragraph()
    _center_no_indent(cap_par)
    cap_par.paragraph_format.space_before = Pt(6)
    cap_par.paragraph_format.space_after = Pt(GOST_FIGURE["space_after_caption_pt"])
    cap_run = cap_par.add_run(format_caption(number, title))
    _apply_gost_font(cap_run)
    cap_run.bold = False  # подпись НЕ полужирная (в отличие от заголовков)

    return img_par, cap_par


def _ensure_reference_in_paragraph(paragraph, number: str) -> bool:
    """Дописывает «(рисунок N)» в конец абзаца, если ссылки ещё нет.
    Ставится ПЕРЕД завершающей точкой предложения."""
    text = paragraph.text or ""
    if re.search(_MENTION_RE_TPL.format(num=re.escape(str(number))), text, re.I):
        return False  # ссылка уже есть
    ref = f" {make_reference(number)}"
    runs = [r for r in paragraph.runs if (r.text or "")]
    if not runs:
        r = paragraph.add_run(ref)
        _apply_gost_font(r)
        return True
    last = runs[-1]
    t = last.text or ""
    stripped = t.rstrip()
    if stripped.endswith("."):
        last.text = stripped[:-1] + ref + "."
    else:
        last.text = t + ref
    return True


# ══════════════════════════════════════════════════════════════════
#  Проверка готового документа (используется checker-ом)
# ══════════════════════════════════════════════════════════════════

def paragraph_has_image(par) -> bool:
    """Есть ли в абзаце картинка (w:drawing / w:pict)."""
    try:
        xml = par._p.xml
    except Exception:
        return False
    return ("<w:drawing" in xml) or ("<w:pict" in xml) or ("graphicData" in xml)


def validate_figures(doc) -> list[dict]:
    """Проверяет ВСЕ рисунки документа по ГОСТ. Возвращает список проблем:
    [{"title": .., "detail": .., "severity": "error"|"warning"}].

    Проверки:
      • у каждого рисунка есть подпись «Рисунок N — Название» снизу;
      • подпись по центру, без точки в конце, тире длинное;
      • нумерация без пропусков и дублей;
      • на каждый рисунок есть ссылка в тексте ДО него.
    """
    issues: list[dict] = []
    pars = list(doc.paragraphs)
    numbers_seen: list[str] = []
    text_before: list[str] = []

    for i, par in enumerate(pars):
        if not paragraph_has_image(par):
            text_before.append(par.text or "")
            continue

        # ищем подпись в ближайших 2 абзацах ниже (пустые пропускаем)
        cap_par = None
        for j in range(i + 1, min(i + 3, len(pars))):
            if (pars[j].text or "").strip():
                cap_par = pars[j]
                break
        cap_text = (cap_par.text or "").strip() if cap_par is not None else ""
        m = CAPTION_RE.match(cap_text)

        if not m:
            issues.append({
                "title": "Рисунок без подписи по ГОСТ",
                "detail": f"после рисунка найдено: «{cap_text[:60]}» "
                          "(требуется «Рисунок N — Название»)",
                "severity": "error",
            })
            continue

        number = m.group(1)
        numbers_seen.append(number)

        # тире и точка в конце
        if f" {GOST_FIGURE['caption_dash']} " not in cap_text:
            issues.append({
                "title": "В подписи рисунка не длинное тире (—)",
                "detail": cap_text[:60], "severity": "warning",
            })
        if cap_text.endswith(".") and not cap_text.endswith("..."):
            issues.append({
                "title": "Точка в конце подписи рисунка",
                "detail": cap_text[:60], "severity": "warning",
            })
        # выравнивание подписи
        if cap_par.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, None):
            issues.append({
                "title": "Подпись рисунка не по центру",
                "detail": cap_text[:60], "severity": "warning",
            })

        # ссылка в тексте ДО рисунка
        before = "\n".join(text_before)
        if not re.search(_MENTION_RE_TPL.format(num=re.escape(number)), before, re.I):
            issues.append({
                "title": f"Нет ссылки на рисунок {number} в тексте до него",
                "detail": "требуется «(рисунок N)» или «в соответствии с рисунком N»",
                "severity": "error",
            })
        text_before.append(cap_text)

    # нумерация: дубли и пропуски (для сквозной нумерации)
    plain = [n for n in numbers_seen if re.fullmatch(r"\d+", n)]
    if len(set(numbers_seen)) != len(numbers_seen):
        issues.append({
            "title": "Дубликаты номеров рисунков",
            "detail": ", ".join(numbers_seen), "severity": "error",
        })
    elif plain:
        expected = [str(k) for k in range(1, len(plain) + 1)]
        if plain != expected:
            issues.append({
                "title": "Нарушена сквозная нумерация рисунков",
                "detail": f"фактически: {', '.join(plain)}", "severity": "warning",
            })

    return issues


def autofix_figures(doc) -> list[str]:
    """Механически чинит подписи рисунков: «Рис. 1.» → «Рисунок 1 — …»,
    дефис → длинное тире, точка в конце — долой, подпись и рисунок по центру.
    → список применённых фиксов."""
    changes: list[str] = []
    pars = list(doc.paragraphs)
    loose_cap_re = re.compile(
        r"^\s*(?:Рис\.?|Рисунок)\s*№?\s*(\d+(?:\.\d+)?|[А-ЯЁ]\.\d+)\s*[—–:.-]?\s*(.*)$",
        re.I,
    )
    for i, par in enumerate(pars):
        if not paragraph_has_image(par):
            continue
        # рисунок по центру, без отступа
        if par.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            _center_no_indent(par)
            if "рисунки выровнены по центру" not in changes:
                changes.append("рисунки выровнены по центру")
        for j in range(i + 1, min(i + 3, len(pars))):
            cap = pars[j]
            text = (cap.text or "").strip()
            if not text:
                continue
            m = loose_cap_re.match(text)
            if m:
                number, title = m.group(1), m.group(2)
                fixed = format_caption(number, title) if title.strip() \
                    else f"{GOST_FIGURE['caption_prefix']} {number}"
                if fixed != text:
                    # переписываем первый run, остальные очищаем
                    runs = [r for r in cap.runs if (r.text or "")]
                    if runs:
                        runs[0].text = fixed
                        for extra in runs[1:]:
                            extra.text = ""
                        _apply_gost_font(runs[0])
                        runs[0].bold = False
                        if "подписи рисунков приведены к ГОСТ" not in changes:
                            changes.append("подписи рисунков приведены к ГОСТ")
                if cap.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    _center_no_indent(cap)
                    if "подписи рисунков по центру" not in changes:
                        changes.append("подписи рисунков по центру")
            break
    return changes

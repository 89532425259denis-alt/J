# -*- coding: utf-8 -*-
"""
gost_validator.py — проверка DOCX на соответствие ГОСТ 7.32-2017.

Использование:
    from gost_validator import validate_docx
    report = validate_docx("work.docx")
    print(report.score, report.passed)
    for c in report.failures():
        print(c.title, "-", c.detail)

Зависит только от python-docx.
"""
from __future__ import annotations
from dataclasses import dataclass, field
from typing import List, Optional
import re

try:
    from docx import Document
    from docx.shared import Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:  # pragma: no cover
    Document = None

EMU_PER_MM = 36000


def _mm(emu) -> Optional[float]:
    try:
        return round(float(emu) / EMU_PER_MM, 1)
    except Exception:
        return None


@dataclass
class CheckResult:
    code: str
    title: str
    ok: bool
    weight: int = 1
    detail: str = ""


@dataclass
class GostReport:
    checks: List[CheckResult] = field(default_factory=list)

    @property
    def score(self) -> int:
        total = sum(c.weight for c in self.checks) or 1
        got = sum(c.weight for c in self.checks if c.ok)
        return round(got * 100 / total)

    @property
    def passed(self) -> bool:
        # «Сдано», если нет проваленных критичных проверок и скор >= 90
        return self.score >= 90 and not any((not c.ok) and c.weight >= 3 for c in self.checks)

    def failures(self) -> List[CheckResult]:
        return [c for c in self.checks if not c.ok]

    def as_text(self) -> str:
        lines = [f"ГОСТ-оценка: {self.score}/100  ({'СДАНО' if self.passed else 'ЕСТЬ ЗАМЕЧАНИЯ'})"]
        for c in self.checks:
            mark = "OK " if c.ok else "XX "
            lines.append(f"  [{mark}] {c.title}" + (f" — {c.detail}" if c.detail and not c.ok else ""))
        return "\n".join(lines)


def _approx(a, b, tol=0.6) -> bool:
    return a is not None and b is not None and abs(a - b) <= tol


def validate_docx(path: str) -> GostReport:
    rep = GostReport()
    if Document is None:
        rep.checks.append(CheckResult("env", "python-docx доступен", False, 5, "модуль не установлен"))
        return rep
    try:
        doc = Document(path)
    except Exception as e:
        rep.checks.append(CheckResult("open", "Файл открывается", False, 5, str(e)))
        return rep

    sec = doc.sections[0]

    # 1. Поля 30/10/20/20 (лево/право/верх/низ)
    L, R, T, B = _mm(sec.left_margin), _mm(sec.right_margin), _mm(sec.top_margin), _mm(sec.bottom_margin)
    ok_margins = _approx(L, 30) and _approx(R, 10) and _approx(T, 20) and _approx(B, 20)
    rep.checks.append(CheckResult("margins", "Поля 30/10/20/20 мм", ok_margins, 3,
                                  f"факт L{L}/R{R}/T{T}/B{B}"))

    # 2. Формат A4 (210x297)
    W, H = _mm(sec.page_width), _mm(sec.page_height)
    ok_a4 = _approx(W, 210, 1.5) and _approx(H, 297, 1.5)
    rep.checks.append(CheckResult("a4", "Формат A4 (210×297 мм)", ok_a4, 2, f"факт {W}×{H}"))

    # Собираем статистику по основному тексту
    body = [p for p in doc.paragraphs if p.text.strip()]
    font_names, font_sizes, spacings, indents = {}, {}, {}, {}
    bom_count = 0
    for p in body:
        if "\ufeff" in p.text:
            bom_count += 1
        style = (p.style.name or "").lower()
        is_heading = style.startswith("heading") or style.startswith("title")
        pf = p.paragraph_format
        if pf.line_spacing:
            spacings[round(float(pf.line_spacing), 2)] = spacings.get(round(float(pf.line_spacing), 2), 0) + 1
        if not is_heading and pf.first_line_indent is not None:
            fi = _mm(pf.first_line_indent)
            if fi and fi > 0:
                indents[fi] = indents.get(fi, 0) + 1
        for run in p.runs:
            if run.font.name:
                font_names[run.font.name] = font_names.get(run.font.name, 0) + 1
            if run.font.size:
                sz = round(run.font.size.pt, 1)
                font_sizes[sz] = font_sizes.get(sz, 0) + 1

    # 3. Шрифт Times New Roman (доминирующий)
    dom_font = max(font_names, key=font_names.get) if font_names else None
    ok_font = (dom_font == "Times New Roman")
    rep.checks.append(CheckResult("font", "Шрифт Times New Roman", ok_font, 3, f"доминирует: {dom_font}"))

    # 4. Кегль 14 (доминирующий в теле)
    dom_size = max(font_sizes, key=font_sizes.get) if font_sizes else None
    ok_size = _approx(dom_size, 14, 0.1)
    rep.checks.append(CheckResult("size", "Кегль 14 пт", ok_size, 3, f"доминирует: {dom_size} пт"))

    # 5. Межстрочный интервал 1.5
    dom_sp = max(spacings, key=spacings.get) if spacings else None
    ok_sp = _approx(dom_sp, 1.5, 0.05)
    rep.checks.append(CheckResult("spacing", "Межстрочный интервал 1,5", ok_sp, 2, f"доминирует: {dom_sp}"))

    # 6. Абзацный отступ 1.25 см (=12.5 мм)
    dom_ind = max(indents, key=indents.get) if indents else None
    ok_ind = _approx(dom_ind, 12.5, 1.0)
    rep.checks.append(CheckResult("indent", "Абзацный отступ 1,25 см", ok_ind, 2,
                                  f"доминирует: {dom_ind} мм"))

    # 7. Структурные элементы по центру (ВВЕДЕНИЕ/ЗАКЛЮЧЕНИЕ/СОДЕРЖАНИЕ/СПИСОК)
    struct_re = re.compile(r"(?i)^(содержание|введение|заключение|список\s+использ|приложение)")
    struct_centered = []
    for p in body:
        if struct_re.match(p.text.strip()):
            al = p.paragraph_format.alignment
            struct_centered.append(al == WD_ALIGN_PARAGRAPH.CENTER)
    ok_struct = bool(struct_centered) and all(struct_centered)
    rep.checks.append(CheckResult("struct_center", "Структурные элементы по центру",
                                  ok_struct, 2,
                                  f"найдено {len(struct_centered)}, по центру {sum(struct_centered)}"))

    # 8. Нумерация разделов без точки в конце номера («1 Название», «1.1 Название»)
    num_re = re.compile(r"^\d+(\.\d+)*\.\s+\S")  # с точкой после последней цифры — нарушение
    bad_num = 0
    for p in body:
        st = (p.style.name or "").lower()
        if st.startswith("heading") and num_re.match(p.text.strip()):
            bad_num += 1
    rep.checks.append(CheckResult("heading_dot", "Нет точки после номера раздела",
                                  bad_num == 0, 1, f"нарушений: {bad_num}"))

    # 9. Нумерация страниц (поле PAGE в колонтитуле)
    has_page_field = False
    try:
        for s in doc.sections:
            for hp in (s.footer.paragraphs + s.header.paragraphs):
                xml = hp._p.xml
                if "PAGE" in xml or "w:fldSimple" in xml or "w:instrText" in xml:
                    has_page_field = True
                    break
    except Exception:
        pass
    rep.checks.append(CheckResult("pageno", "Нумерация страниц (поле PAGE)", has_page_field, 2))

    # 10. Содержание присутствует
    has_toc = any(re.match(r"(?i)^содержание$", p.text.strip()) for p in body)
    rep.checks.append(CheckResult("toc", "Есть «СОДЕРЖАНИЕ»", has_toc, 1))

    # 11. Список источников присутствует
    has_bib = any(re.match(r"(?i)^список\s+использ", p.text.strip()) for p in body)
    rep.checks.append(CheckResult("bib", "Есть «СПИСОК … ИСТОЧНИКОВ»", has_bib, 1))

    # 12. Нет BOM/мусорных символов
    rep.checks.append(CheckResult("bom", "Нет мусорных символов (BOM)", bom_count == 0, 1,
                                  f"строк с BOM: {bom_count}"))

    return rep


if __name__ == "__main__":
    import sys
    r = validate_docx(sys.argv[1])
    print(r.as_text())

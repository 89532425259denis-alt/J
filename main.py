# -*- coding: utf-8 -*-
# GOST-ASSISTANT v2.7 — FIXED (GOST 7.32-2017 compliant)
# Applied fixes: #1-#9 (see apply_gost_fixes.py)

from __future__ import annotations

"""ГОСТ-АССИСТЕНТ v2.7-fix16 — ТОЧНЫЕ СТРАНИЦЫ + ДИСЦИПЛИНА

Главные изменения v2.1 относительно v2.0:
────────────────────────────────────────────────────────────────
1. ★ ТОЧНОЕ ЧИСЛО СТРАНИЦ (±1).
   После генерации DOCX конвертируется в PDF через LibreOffice,
   реально пересчитываются страницы (pypdf / pdfinfo). Если страниц
   меньше цели — главы дозаполняются; если больше — обрезаются
   по границам абзацев. Цикл до 3 итераций.

2. ★ ПРАВИЛЬНАЯ НУМЕРАЦИЯ СТРАНИЦ ПО ГОСТ.
   Титульный лист включается в общую нумерацию, но цифра на нём
   не отображается (different_first_page_header_footer). На странице
   содержания появляется цифра «2», далее сквозная. Номер — внизу
   по центру (или сверху, как настроено в GOST).

3. ★ ЛИМИТ ДЛЯ БЕСПЛАТНОГО РЕЖИМА — 1 ГЕНЕРАЦИЯ В 5 ДНЕЙ.
   Реализовано через кулдаун (FREE_COOLDOWN_SECONDS = 432000),
   независимо от смены календарных суток. VIP и платные — без лимитов.

4. ★ v2.2: подглавы генерируются отдельными промптами (нет пустых глав).
   Заключение пишется ПОСЛЕ глав и видит их реальное содержание.
   ГОСТ-сноски [1, с. 45] обязательны в каждом абзаце.

5. Все базовые возможности v2.0 сохранены: красивые заголовки глав,
   умный/классический стиль, прогресс-бар с ETA, парсинг «своего ГОСТ»,
   custom-документы, аккуратные сообщения с HTML-форматированием.
"""

import asyncio
import html
import io
import json
import os
import random
import re
import shutil
import subprocess
import time
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional
from urllib.parse import quote_plus

import aiohttp
from aiohttp import web
from aiogram import Bot, Dispatcher, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    BufferedInputFile,
    CallbackQuery,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    LabeledPrice,
    Message,
    PreCheckoutQuery,
)
from aiogram.utils.keyboard import InlineKeyboardBuilder

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

try:
    from PIL import Image as PILImage, ImageDraw, ImageFont
except Exception:  # Pillow может быть не установлен на старом деплое
    PILImage = None
    ImageDraw = None
    ImageFont = None


# ═══════════════════════════════════════════════════════════════
#  КРАСИВАЯ АНИМАЦИЯ ПРОГРЕССА С ETA
# ═══════════════════════════════════════════════════════════════

# Рамки спиннера — меняются каждую секунду для иллюзии движе��ия
_SPINNER_FRAMES = ["⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏"]
_BAR_FULL  = "█"
_BAR_HALF  = "▓"
_BAR_EMPTY = "░"


def _spinner() -> str:
    """Возвращает текущий кадр спиннера на основе текущего времени."""
    idx = int(time.monotonic() * 4) % len(_SPINNER_FRAMES)
    return _SPINNER_FRAMES[idx]


def _progress_bar(done: int, total: int, width: int = 16) -> str:
    """Улучшенный прогресс-бар с более плавным визуалом."""
    total = max(1, int(total))
    done  = max(0, min(int(done), total))
    ratio = done / total
    filled = int(round(width * ratio))
    
    # Используем более современные символы для заполнения
    # █ - полный, ▓ - почти полный, ░ - пустой
    bar = "█" * filled + "░" * (width - filled)
    
    # Добавляем динамический «бегунок» в конец заполненной части, если работа не завершена
    if 0 < filled < width:
        frame = int(time.monotonic() * 2) % 2
        marker = "▓" if frame == 0 else "▒"
        # Заменяем последний заполненный символ на маркер
        bar = bar[:filled-1] + marker + bar[filled:]
        
    return bar


def _fmt_time(seconds: float) -> str:
    """Форматирует секунды в mm:ss."""
    s = max(0, int(seconds))
    m, s = divmod(s, 60)
    if m:
        return f"{m}м {s:02d}с"
    return f"{s}с"


@dataclass
class Progress:
    """Красивый прогресс с анимацией, ETA и шагами."""
    msg: Message
    title: str
    total_steps: int
    model_name: str = ""
    done: int = 0
    label: str = "Подготовка..."
    _last_text: str = field(default="", repr=False)
    _last_ts: float = field(default=0.0, repr=False)
    _start_ts: float = field(default_factory=time.monotonic, repr=False)
    _step_labels: list = field(default_factory=list, repr=False)
    _eta_prev: Optional[float] = field(default=None, repr=False)

    def _elapsed(self) -> float:
        return time.monotonic() - self._start_ts

    def _eta(self) -> str:
        """Оценка оставшегося времени (fix11).

        - Пока ни один шаг не закрыт: даём приблизительную оценку,
          исходя из общего числа шагов и среднего времени на шаг ~25 с
          (если elapsed > 3 с), вместо вечного «считаю…».
        - Со 2-го шага: добавляем EMA сглаживание, чтобы оценка
          не прыгала при разных по длительности шагах.
        """
        elapsed = self._elapsed()
        remaining_steps = max(0, self.total_steps - self.done)
        if remaining_steps == 0:
            return "≈0с"

        if self.done == 0:
            if elapsed < 3.0:
                return "считаю…"
            # эвристика: ~25 секунд на шаг, минимум 10
            est_per_step = max(10.0, elapsed / 0.5)
            return "≈" + _fmt_time(est_per_step * remaining_steps)

        rate = self.done / elapsed
        if rate <= 0:
            return "…"
        eta_sec = remaining_steps / rate
        # Сглаживание: микшируем с предыдущим значением
        prev = getattr(self, "_eta_prev", None)
        if prev is None:
            self._eta_prev = eta_sec
        else:
            self._eta_prev = 0.6 * prev + 0.4 * eta_sec
            eta_sec = self._eta_prev
        return _fmt_time(eta_sec)

    def render(self) -> str:
        pct      = int(self.done * 100 / max(1, self.total_steps))
        bar      = _progress_bar(self.done, self.total_steps)
        spin     = _spinner()
        elapsed  = _fmt_time(self._elapsed())
        eta      = self._eta()
        model    = f"\n🤖 <b>Модель:</b> <i>{self.model_name}</i>" if self.model_name else ""
        step_num = min(self.done + 1, self.total_steps)

        return (
            f"{spin} <b>{self.title}</b>\n"
            f"<code>{bar}</code> <b>{pct}%</b>\n"
            f"━━━━━━━━━━━━━━━━━━━━━━\n"
            f"📌 <b>Шаг {step_num}/{self.total_steps}:</b> {self.label}{model}\n"
            f"⏱ <code>{elapsed}</code> ➔ ⏳ <code>{eta}</code>"
        )

    async def update(
        self,
        *,
        label: Optional[str] = None,
        step_done: bool = False,
        model_name: Optional[str] = None,
        force: bool = False,
        min_interval: float = 1.0,
    ) -> None:
        if label is not None:
            self.label = label
        if model_name is not None:
            self.model_name = model_name
        if step_done:
            self.done = min(self.total_steps, self.done + 1)

        text = self.render()
        now  = time.monotonic()

        # Не спамим обновлениями чаще min_interval
        if not force and text == self._last_text:
            return
        if not force and (now - self._last_ts) < min_interval:
            return

        try:
            await self.msg.edit_text(text, parse_mode="HTML")
            self._last_text = text
            self._last_ts   = now
        except Exception:
            pass

    async def finish(self, text: str) -> None:
        try:
            await self.msg.edit_text(text, parse_mode="HTML")
        except Exception:
            pass

    async def delete(self) -> None:
        try:
            await self.msg.delete()
        except Exception:
            pass

    async def animate_loop(self, stop_event: asyncio.Event) -> None:
        """Фоновый цикл (fix11): обновляет спиннер/змейку/ETA каждые ~1.5 с.

        Telegram режет частые edit_text (429 Flood). Минимальный безопасный
        интервал — 1.2-1.5 с. Бар анимируется самим временем (frame =
        int(monotonic()*2)), поэтому каждые 1.5 с кадры реально меняются.
        """
        while not stop_event.is_set():
            try:
                text = self.render()
                now  = time.monotonic()
                if text != self._last_text and (now - self._last_ts) >= 1.4:
                    await self.msg.edit_text(text, parse_mode="HTML")
                    self._last_text = text
                    self._last_ts   = now
            except Exception:
                pass
            await asyncio.sleep(1.5)


# ═══════════════════════════════════════════════════════════════
#  КОНФИГ
# ═══════════════════════════════════════════════════════════════

TOKENS = {
    "BOT_TOKEN": "",
    "OPENROUTER_KEY": "",
    "DEEPSEEK_KEY": "",
    "GROQ_KEY": "",
    "VIP_USERS": "5291613279",
}

CONFIG_FILE      = "bot_config.json"
GOST_CONFIG_FILE = "gost_configs.json"
USAGE_FILE       = "usage_limits.json"
NUMERIC_FACTS_FILE = "numeric_facts.json"


def _load_json(path: str, default):
    if not os.path.exists(path):
        return default
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default


def _save_json(path: str, data) -> None:
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"[WARN] Ошибка сохранения {path}: {e}")


CONFIG      = _load_json(CONFIG_FILE, {})
GOST_CONFIGS = _load_json(GOST_CONFIG_FILE, {})


def cfg(name: str, default: str = "") -> str:
    return (TOKENS.get(name) or os.getenv(name) or CONFIG.get(name) or default).strip()


BOT_TOKEN = cfg("BOT_TOKEN")
# if not BOT_TOKEN:
    # raise SystemExit("❌ ОШИБКА: не вставлен BOT_TOKEN (в TOKENS, .env или bot_config.json)")

OPENROUTER_KEY = cfg("OPENROUTER_KEY")
DEEPSEEK_KEY   = cfg("DEEPSEEK_KEY")
GROQ_KEY       = cfg("GROQ_KEY")

OPENROUTER_BASE_URL = cfg("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1").rstrip("/")
DEEPSEEK_BASE_URL   = cfg("DEEPSEEK_BASE_URL",   "https://api.deepseek.com/v1").rstrip("/")
GROQ_BASE_URL       = cfg("GROQ_BASE_URL",       "https://api.groq.com/openai/v1").rstrip("/")

DEEPSEEK_MODEL         = cfg("DEEPSEEK_MODEL",         "deepseek-chat")
OPENROUTER_R1_MODEL    = cfg("OPENROUTER_R1_MODEL",    "deepseek/deepseek-r1-0528:free")
OPENROUTER_GEMINI_MODEL = cfg("OPENROUTER_GEMINI_MODEL", "google/gemini-2.5-flash")
GROQ_MODEL             = cfg("GROQ_MODEL",             "llama-3.3-70b-versatile")

FREE_MODEL_KEY = cfg("FREE_MODEL_KEY", "deepseek")

# Лимиты для БЕСПЛАТНОГО режима
FREE_MAX_PAGES    = int(cfg("FREE_MAX_PAGES",         "15"))
# Кулдаун между бесплатными генерациями: по умолчанию 5 суток (432000 сек).
# Лимит "1 генерация в 5 дней" реализован именно через кулдаун, а не дневной счётчик —
# это даёт точное окно 5×24 ч от момента предыдущей генерации.
FREE_COOLDOWN     = int(cfg("FREE_COOLDOWN_SECONDS",  str(5 * 24 * 60 * 60)))
# FREE_DAILY_LIMIT оставлен для совместимости; основной фильтр — кулдаун.
FREE_DAILY_LIMIT  = int(cfg("FREE_DAILY_LIMIT",       "0"))   # 0 = без дневного лимита

# Лимиты для ПЛАТНОГО режима — платят деньги, получают безлимит
# PAID_DAILY_LIMIT = 0 означает "без лимита"
PAID_DAILY_LIMIT  = int(cfg("PAID_DAILY_LIMIT",       "0"))   # 0 = безлимит
PAID_COOLDOWN     = int(cfg("PAID_COOLDOWN_SECONDS",  "0"))   # 0 = нет кулдауна

# Символов на страницу (ГОСТ: ~1800-2000 знаков с пробелами на стр A4 14pt 1.5 интервал)
# Глобальное значение по умолчанию, если ГОСТ не передан
CHARS_PER_PAGE = int(cfg("CHARS_PER_PAGE", "1850"))

# Веб-источники: бот умеет подтягивать реальные источники из открытых научных
# каталогов (OpenAlex/Crossref) и читать URL, которые пользователь прислал в
# материалах. Если сервер без доступа к сети — функции тихо отключатся.
ENABLE_WEB_SOURCES = cfg("ENABLE_WEB_SOURCES", "1").lower() not in ("0", "false", "no", "off")
# Отдельный флаг для поиска изображений: даже если веб-источники для литературы отключены,
# картинки можно оставить включёнными.
ENABLE_IMAGE_SEARCH = cfg("ENABLE_IMAGE_SEARCH", "0").lower() not in ("0", "false", "no", "off")  # user-patch: авто-поиск выключен — пользователь сам присылает ссылки
WEB_SOURCE_TIMEOUT = int(cfg("WEB_SOURCE_TIMEOUT", "12"))
MAX_WEB_SOURCES    = int(cfg("MAX_WEB_SOURCES", "25"))  # user-patch: было 12 → стало 25
MIN_REAL_SOURCES   = int(cfg("MIN_REAL_SOURCES", "15"))  # user-patch: было 10 → стало 15
BIB_SOURCE_TARGET  = int(cfg("BIB_SOURCE_TARGET", "20"))  # user-patch: было 12 → стало 20
FILL_UNKNOWN_CITATION_PAGES = cfg("FILL_UNKNOWN_CITATION_PAGES", "1").lower() not in ("0", "false", "no", "off")
# Поле Word TOC иногда показывает пользователю служебную заглушку/код поля.
# По умолчанию формируем статическое содержание без поля, чтобы в DOCX не было
# видимых placeholder-строк. При необходимости можно включить обновляемое поле.
TOC_USE_WORD_FIELD = cfg("TOC_USE_WORD_FIELD", "0").lower() not in ("0", "false", "no", "off")

def calculate_chars_per_page(gost: dict) -> int:
    """Расчёт количества знаков с пробелами на страницу.

    Эмпирические значения, полученные сборкой одностраничных DOCX→PDF в
    LibreOffice и подсчётом фактических символов. Базовая точка — каноничный
    ГОСТ 7.32: Times New Roman 14 pt, межстрочный 1.5, поля 30/10/20/20 мм →
    ~1800 знаков на страницу. От базовой точки масштабируем по полям, кеглю
    и межстрочному интервалу.
    """
    font_size    = int(gost.get("font_size", 14))
    line_spacing = float(gost.get("line_spacing", 1.5))
    left_mm   = int(gost.get("left_margin_mm",   30))
    right_mm  = int(gost.get("right_margin_mm",  10))
    top_mm    = int(gost.get("top_margin_mm",    20))
    bottom_mm = int(gost.get("bottom_margin_mm", 20))

    # Базовая площадь текстового блока (TNR 14, 1.5, поля 30/10/20/20).
    BASE_CHARS   = 1800.0
    BASE_W_MM    = 210 - 30 - 10   # 170
    BASE_H_MM    = 297 - 20 - 20   # 257

    text_w = max(60, 210 - left_mm - right_mm)
    text_h = max(60, 297 - top_mm - bottom_mm)

    # Площадь блока линейно влияет на «вместимость».
    area_factor = (text_w / BASE_W_MM) * (text_h / BASE_H_MM)

    # Кегль: ширина и высота строки пропорциональны размеру шрифта,
    # значит вместимость ~ (14 / fs)^2.
    font_factor = (14.0 / max(10, font_size)) ** 2

    # Межстрочный интервал: вместимость ~ 1.5 / spacing.
    spacing_factor = 1.5 / max(1.0, line_spacing)

    chars = BASE_CHARS * area_factor * font_factor * spacing_factor
    return max(900, min(3200, int(round(chars))))

# Страниц без текста (титул + содержание)
NON_TEXT_PAGES = int(cfg("NON_TEXT_PAGES", "2"))

DEEPSEEK_PRICE  = int(cfg("DEEPSEEK_PRICE",           "5"))
GROQ_PRICE      = int(cfg("GROQ_PRICE",               "3"))
OR_GEMINI_PRICE = int(cfg("OPENROUTER_GEMINI_PRICE",  "7"))
# Доплата за режим с иллюстрациями: поиск, проверка и вставка изображений по ГОСТ.
IMAGES_EXTRA_PRICE_PER_PAGE = int(cfg("IMAGES_EXTRA_PRICE_PER_PAGE", "3"))
# Цена за редактирование/исправление уже готовой работы (фикс. стоимость в ⭐).
EDIT_PRICE = int(cfg("EDIT_PRICE", "15"))
MAX_WORK_IMAGES = int(cfg("MAX_WORK_IMAGES", "5"))

# Сколько работ генерируется ОДНОВРЕМЕННО на весь бот.
# Раньше было 1 → один пользователь блокировал всех остальных, и нельзя было
# запустить вторую работу в том же чате. Теперь по умолчанию 4 параллельных
# генерации (настраивается переменной окружения MAX_PARALLEL). aiogram уже
# обрабатывает апдейты конкурентно, поэтому ограничиваем лишь тяжёлую
# генерацию, чтобы не перегрузить API/LibreOffice.
MAX_PARALLEL  = max(1, int(cfg("MAX_PARALLEL", "4")))
GEN_SEMAPHORE = asyncio.Semaphore(MAX_PARALLEL)

# Ограничение ОДНОВРЕМЕННЫХ генераций на одного пользователя — чтобы человек
# мог запустить несколько работ подряд, но не «забил» всю очередь. 0 = без лимита.
MAX_PARALLEL_PER_USER = max(0, int(cfg("MAX_PARALLEL_PER_USER", "2")))
# Счётчик активных генераций по user_id (для нескольких работ в одном чате).
_active_gen_per_user: dict[int, int] = {}

VIP_USERS = {int(x) for x in cfg("VIP_USERS", "").replace(" ", "").split(",") if x.isdigit()}


# ═══════════════════════════════════════════════════════════════
#  ТИПЫ ДОКУМЕНТОВ
# ═══════════════════════════════════════════════════════════════

DOC_TYPES = {
    "referat": {
        "name": "📄 Реферат",
        "word": "РЕФЕРАТ",
        "min_pages": 10,
        "max_pages": 30,
        "structure": "Введение · 2 главы с подглавами · Заключение · Список литературы",
        "desc": "Обзор литературы по теме, изложение основных концепций и выводы",
    },
    "kursovaya": {
        "name": "📚 Курсовая работа",
        "word": "КУРСОВАЯ РАБОТА",
        "min_pages": 25,
        "max_pages": 50,
        "structure": "Введение · 3 главы с подглавами (§) · Заключение · Библиография",
        "desc": "Самостоятельное исследование с анализом, практической частью и выводами",
    },
    "doklad": {
        "name": "🎤 Доклад",
        "word": "ДОКЛАД",
        "min_pages": 5,
        "max_pages": 15,
        "structure": "Введение · Основная часть · Заключение",
        "desc": "Краткий обзор темы для устного выступления",
    },
    "esse": {
        "name": "✍️ Эссе",
        "word": "ЭССЕ",
        "min_pages": 3,
        "max_pages": 10,
        "structure": "Вступление · А��гументы · Авторская позиция",
        "desc": "Авторский взгляд на проблему с аргументацией и личными выводами",
    },
    "kontrolnaya": {
        "name": "📝 Контрольная работа",
        "word": "КОНТРОЛЬНАЯ РАБОТА",
        "min_pages": 10,
        "max_pages": 25,
        "structure": "Теоретическая часть · Практическая часть · Выводы",
        "desc": "Проверочная работа по теории и практике дисциплины",
    },
    "final_project": {
        "name": "📦 Итоговый проект",
        "word": "ИТОГОВЫЙ ПРОЕКТ",
        "min_pages": 35,
        "max_pages": 80,
        "structure": "Введение · 4 главы · Заключение · Приложения · Источники",
        "desc": "Комплексный проект с теорией, практикой и проектной частью",
    },
    "final_referat": {
        "name": "📑 Итоговый реферат",
        "word": "ИТОГОВЫЙ РЕФЕРАТ",
        "min_pages": 15,
        "max_pages": 40,
        "structure": "Введение · 3 главы с подглавами · Заключение · Список литературы",
        "desc": "Расширенный итоговый реферат по ключевой дисциплине с глубоким анализом литературы",
    },
    "article": {
        "name": "📝 Научная статья",
        "word": "НАУЧНАЯ СТАТЬЯ",
        "min_pages": 5,
        "max_pages": 15,
        "structure": "Аннотация · Ключевые слова · Введение · Обзор литературы · Методология · Результаты · Заключение · Список источников",
        "desc": "Научная публикация по ГОСТ Р 7.0.7-2021 (требования к статьям)",
    },
    "vkr": {
        "name": "🎓 ВКР (Дипломная работа)",
        "word": "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА",
        "min_pages": 40,
        "max_pages": 100,
        "structure": "Введение · 3-4 главы с подглавами · Практический раздел · Заключение · Список использованных источников",
        "desc": "Выпускная квалификационная работа (бакалаврская работа или дипломный проект) по ГОСТ",
    },
    "custom": {
        "name": "🧩 Свой тип",
        "word": "РАБОТА",
        "min_pages": 3,
        "max_pages": 100,
        "structure": "Пользовательская структура",
        "desc": "Любой тип документа с вашими требованиями к структуре",
    },
}

INSTITUTION_TYPES = {
    "school": {
        "name": "🏫 Школа",
        "org_example": "Муниципальное бюджетное общеобразовательное учреждение",
        "name_example": "Средняя общеобразовательная школа № 123",
    },
    "college": {
        "name": "🏛 Колледж / СПО",
        "org_example": "Государственное бюджетное профессиональное образовательное учреждение",
        "name_example": "Колледж информационных технологий № 42",
    },
    "university": {
        "name": "🎓 ВУЗ / Университет",
        "org_example": "Федеральное государственное бюджетное образовательное учреждение высшего образования",
        "name_example": "Московский государственный университет имени М.В. Ломоно��ова",
    },
    "custom": {
        "name": "✏️ Свой вариант",
        "org_example": "",
        "name_example": "",
    },
}

SUBJECTS = [
    "История", "Обществознание", "Литература", "Русский язык",
    "Математика", "Физика", "Информатика", "Биология",
    "Химия", "География", "Экономика", "Право",
    "Философия", "Психология", "Социология", "Менеджмент",
    "Педагогика", "Медицина", "Архитектура", "Юриспруденция",
]

# Примеры «хорошо / плохо» для эссе по дисциплинам
DISCIPLINE_EXAMPLES = {
    "Информатика": (
        "ПЛОХО (не по дисциплине): «Байкал — это чистейшее озеро с удивительной экосистемой».\n"
        "ХОРОШО (по дисциплине): «Байкал можно рассматривать как природную систему сбора данных: "
        "тысячелетиями в донных отложениях накапливается информация о климате, которую мы извлекаем "
        "методами машинного обучения. Это аналог долговременной памяти в вычислительных системах.»"
    ),
    "Психология": (
        "ПЛОХО: «Байкал красив и вызывает благоговение».\n"
        "ХОРОШО: «Восприятие Байкала вызывает эффект благоговения (awe), который, по исследованиям "
        "Келтнера и Хаидта, снижает активность дефолт-системы мозга и усиливает ощущение связи с миром.»"
    ),
    "История": (
        "ПЛОХО: «Байкал — уникальное природное явление».\n"
        "ХОРОШО: «С XVIII века Байкал становится объектом научного интереса: экспедиции Миллера (1733–1743) "
        "и Палласа заложили основу систематического изучения региона, что можно рассматривать как начало "
        "формирования региональной историографии.»"
    ),
}

CITIES = [
    "Москва", "Санкт-Петербург", "Новосибирск", "Екатеринбург",
    "Казань", "Нижний Новгород", "Красноярск", "Челябинск",
    "Омск", "Ростов-на-Дону", "Уфа", "Волгоград",
]


# ═══════════════════════════════════════════════════════════════
#  БИБЛИОТЕКА ЭТАЛОНОВ И ЖЕСТКИХ ПРАВИЛ (Few-Shot & Anti-Patterns)
# ══════════════════════════════════════════════════════════════════

# 1. Эталоны структуры и стиля (GOLDEN STANDARDS)
GOLDEN_STANDARDS = {
    "academic": {
        "intro": "Актуальность данной темы обусловлена стремительным развитием [Область], что требует переосмысления подходов к [Проблема]. Целью данной работы является комплексный анализ [Тема]. Для достижения поставленной цели необходимо решить следующие задачи: во-первых, изучить теоретические основы...; во-вторых, проанализировать влияние...",
        "chapter": "Согласно концепции С.В. Петрова, данный процесс представляет собой совокупность взаимосвязанных факторов, определяющих динамику развития системы [1, с. 42]. Это позволяет утверждать, что ключевым элементом здесь выступает [Тезис]. В свою очередь, анализ эмпирических данных показывает, что при увеличении X наблюдается рост Y, что подтверждает гипотезу о...",
        "conclusion": "Таким образом, проведенное исследование позволило установить, что [Основной вывод]. В частности, было выявлено, что [Деталь 1] и [Деталь 2]. Полученные результаты подтверждают теоретическую значимость подхода и могут быть использованы для оптимизации процессов в области [Сфера].",
    },
    "doklad": {
        "main": "Перейдем к рассмотрению ключевых факторов. Первым из них является [Фактор]. Как отмечают исследователи, именно этот элемент определяет до 60% успеха в [Область] [2, с. 64]. Примером может служить опыт компании X, где внедрение данного метода позволило сократить издержки на 15%. Таким образом, мы видим прямую зависимость между...", 
    },
    "esse": {
        "main": "Размышляя над проблемой [Тема], я прихожу к выводу, что истинная причина кроется не в формальных признаках, а в глубинной потребности человека в [Ценность]. Эта мысль перекликается с тезисом А. Смита о том, что 'свобода выбора является базисом развития' [3, с. 112]. Однако, на мой взгляд, в современных условиях этот принцип трансформируется в...",
    },
}

def get_golden_example(doc_type: str, part: str) -> str:
    """Возвращает соответствующий эталон в зависимости от типа работы и части текста."""
    if doc_type in ("esse",):
        return GOLDEN_STANDARDS["esse"].get("main", "")
    if doc_type in ("doklad",):
        return GOLDEN_STANDARDS["doklad"].get("main", "")
    if part == "intro":
        return GOLDEN_STANDARDS["academic"]["intro"]
    if part == "conclusion":
        return GOLDEN_STANDARDS["academic"]["conclusion"]
    return GOLDEN_STANDARDS["academic"]["chapter"]

# 2. Жесткие примеры оформления (Anti-Patterns)
FEW_SHOT_EXAMPLES = """
🌟 ПРИМЕР ПРАВИЛЬНОЙ ССЫЛКИ (обязательно с номером страницы):
   ✅ ХОРОШО: «...текст... [1, с. 45] ...текст... [2, с. 120–125]»
   ❌ ПЛОХО: «...текст... [1, с. » (без цифры) или «...текст... [2]» (без страницы)

🌟 ПРИМЕР ЗАКОНЧЕННОГО ПРЕДЛОЖЕНИЯ (запрещено обрывать ссылку или мысль):
   ✅ ХОРОШО: «...осадочными породами [1, с. 45]. Далее в работе рассмотрим...»
   ❌ ПЛОХО: «...осадочными породами [1, с.» — ТАК ПИСАТЬ КАТЕГОРИЧЕСКИ НЕЛЬЗЯ

🌟 ПРИМЕР ЗАПРЕЩЕННОГО ТЕКСТА:
   ❌ Любые фразы типа «Вот ваш текст», «Конечно, я п��могу», «Объем соблюден» ЗАПРЕЩЕНЫ.
"""

# Черный список авторов для литературы (чтобы не было Кнута в реферате по Байкалу)
LIT_BLACKLIST = "Дональд Кнут, Томас Кормен, Эндрю Таненбаум, Стивен Лавренс"


DEFAULT_GOST_CONFIGS: dict = {
    "_base": {
        "font_name":              "Times New Roman",
        "font_size":              14,
        "line_spacing":           1.5,
        "first_line_indent_cm":   1.25,
        "left_margin_mm":         30,
        "right_margin_mm":        10,
        "top_margin_mm":          20,
        "bottom_margin_mm":       20,
        "alignment":              "justify",
        "page_number_position":   "bottom_center",
    }
}

for _t in DOC_TYPES.keys():
    DEFAULT_GOST_CONFIGS.setdefault(_t, dict(DEFAULT_GOST_CONFIGS["_base"]))

if not isinstance(GOST_CONFIGS, dict):
    GOST_CONFIGS = {}

for _t in DOC_TYPES.keys():
    GOST_CONFIGS.setdefault(_t, {})
    for _k, _v in DEFAULT_GOST_CONFIGS[_t].items():
        GOST_CONFIGS[_t].setdefault(_k, _v)

_save_json(GOST_CONFIG_FILE, GOST_CONFIGS)


def get_gost_config(doc_type: str, user_id: Optional[int] = None) -> dict:
    doc_type = doc_type if doc_type in DOC_TYPES else "referat"
    config = dict(GOST_CONFIGS.get(doc_type, DEFAULT_GOST_CONFIGS[doc_type]))
    if user_id:
        user_file = f"user_gost_{user_id}.json"
        u = _load_json(user_file, {})
        if isinstance(u, dict) and doc_type in u and isinstance(u[doc_type], dict):
            config.update(u[doc_type])
    return config


def save_user_gost_config(user_id: int, doc_type: str, config: dict) -> None:
    user_file = f"user_gost_{user_id}.json"
    u = _load_json(user_file, {})
    if not isinstance(u, dict):
        u = {}
    u[doc_type] = config
    _save_json(user_file, u)


# ═══════════════════════════════════════════════════════════════
#  ЛИМИТЫ / VIP
# ═══════════════════════════════════════════════════════════════

def is_vip(user_id: int) -> bool:
    return int(user_id) in VIP_USERS


# ═══════════════════════════════════════════════════════════════
# СИСТЕМНЫЙ ПРОМПТ ДЛЯ АКАДЕМИЧЕСКОГО РЕФЕРИРОВАНИЯ (ГОСТ 7.32-2017)
# ═══════════════════════════════════════════════════════════════

ACADEMIC_REFERENCING_SYSTEM_PROMPT = """ТЫ — ЭКСПЕРТ ПО АКАДЕМИЧЕСКОМУ РЕФЕРИРОВАНИЮ И ГОСТ 7.32-2017.

Перед тобой текст реферата. Твоя задача — ПРИВЕСТИ ЕГО К СТРОГОМУ СТАНДАРТУ.

ВЫПОЛНИ СТРОГО ПОСЛЕДОВАТЕЛЬНО:

ЭТАП 1. ОГЛАВЛЕНИЕ — Проставь номера страниц для всех разделов
ЭТАП 2. ВВЕДЕНИЕ — 4 абзаца: Актуальность, Степень разработанности, Цель и задачи, Структура  
ЭТАП 3. ОСНОВНАЯ ЧАСТЬ — каждая подглава: 3 абзаца (вводный, основной с ссылками, выводящий)
ЭТАП 4. ЗАКЛЮЧЕНИЕ — 4 связных абзаца: итог по содержанию первой главы, итог по второй главе, общий вывод и перспективы. ПИШИ СВЯЗНЫМ ТЕКСТОМ, БЕЗ ярлыков-заголовков вида «Вывод по первой главе:», «Общий итог:».
ЭТАП 5. СПИСОК ЛИТЕРАТУРЫ — минимум 10 источников в формате ГОСТ 7.32-2017
ЭТАП 6. СТИЛИСТИКА — удали фразы-маркеры ИИ
ЭТАП 7. ПРОВЕРКА — все ссылки с номерами страниц [N, с. X]

НЕ ДОБАВЛЯЙ ПОЯСНЕНИЙ — ТОЛЬКО ГОТОВЫЙ ИСПРАВЛЕННЫЙ РЕФЕРАТ."""

# ═══════════════════════════════════════════════════════════════
# ФУНКЦИИ ЗАЩИТЫ ОТ ПЕРЕГРУЗКИ API
# ═══════════════════════════════════════════════════════════════

# Глобальный счётчик ошибок для определения перегрузки
_api_overload_counter = {}
_api_overload_timestamps = {}

def get_overload_warning_message() -> str:
    """Возвращает предупреждающее сообщение о высокой нагрузке на API."""
    return (
        "⚠️ <b>Внимание! Высокая нагрузка на сервис</b>\n\n"
        "🔄 Сейчас пользуются много пользователей, большая нагрузка на API.\n"
        "⏱ Попробуйте повторить запрос через несколько минут.\n\n"
        "💡 <b>Советы:</b>\n"
        "• Используйте платный режим — он без лимитов\n"
        "• Выберите другую ИИ-модель\n"
        "• Попробуйте позже, когда нагрузка снизится"
    )

def increment_overload(model_key: str) -> None:
    """Увеличивает счётчик перегрузок для модели."""
    import time as _time_mod
    current_time = _time_mod.time()
    if model_key in _api_overload_timestamps:
        if current_time - _api_overload_timestamps[model_key] > 300:
            _api_overload_counter[model_key] = 0
    _api_overload_counter[model_key] = _api_overload_counter.get(model_key, 0) + 1
    _api_overload_timestamps[model_key] = current_time

def is_api_overloaded(model_key: str) -> bool:
    """Проверяет, перегружена ли модель (более 3 ошибок за 5 минут)."""
    return _api_overload_counter.get(model_key, 0) >= 3

def reset_overload(model_key: str) -> None:
    """Сбрасывает счётчик перегрузок при успешном запросе."""
    _api_overload_counter[model_key] = 0




def today_key() -> str:
    return datetime.now().strftime("%Y-%m-%d")


def fmt_seconds(seconds: int) -> str:
    m, s = divmod(max(0, int(seconds)), 60)
    if m:
        return f"{m} мин {s:02d} сек"
    return f"{s} сек"


def load_usage() -> dict:
    return _load_json(USAGE_FILE, {})


def save_usage(d: dict) -> None:
    _save_json(USAGE_FILE, d)


def _fmt_wait_human(seconds: int) -> str:
    """Удобная человекочитаемая длительность: дни/часы/минуты/секунды."""
    s = max(0, int(seconds))
    d, s = divmod(s, 86400)
    h, s = divmod(s, 3600)
    m, s = divmod(s, 60)
    parts = []
    if d: parts.append(f"{d} дн")
    if h: parts.append(f"{h} ч")
    if m and not d: parts.append(f"{m} мин")
    if not parts: parts.append(f"{s} сек")
    return " ".join(parts)


def check_user_limit(user_id: int, mode: str) -> tuple[bool, str]:
    """Проверяет ограничения.

    Бесплатный режим: одна генерация раз в FREE_COOLDOWN секунд (по умолчанию 5 суток).
    Платные пользователи (mode='paid') — безлимитны (PAID_DAILY_LIMIT=0, PAID_COOLDOWN=0).
    VIP — без ограничений всегда.
    """
    if is_vip(user_id):
        return True, ""

    is_free = mode == "free"

    # Платный режим — без ограничений если PAID_DAILY_LIMIT == 0
    if not is_free and PAID_DAILY_LIMIT == 0 and PAID_COOLDOWN == 0:
        return True, ""

    data = load_usage()
    uid  = str(user_id)
    rec  = data.get(uid, {}) or {}

    now = int(datetime.now().timestamp())

    # ── Кулдаун (главный фильтр для бесплатных) ──
    cooldown = FREE_COOLDOWN if is_free else PAID_COOLDOWN
    ts_key   = "last_free_ts" if is_free else "last_paid_ts"
    last_ts  = int(rec.get(ts_key, 0) or 0)
    if cooldown > 0 and last_ts and (now - last_ts) < cooldown:
        wait     = cooldown - (now - last_ts)
        next_dt  = datetime.fromtimestamp(last_ts + cooldown).strftime("%d.%m.%Y %H:%M")
        kind     = "бесплатная" if is_free else "платная"
        return False, (
            f"⏳ <b>Следующая {kind} генерация будет доступна позже</b>\n\n"
            f"┌─────────────────────────\n"
            f"│ ⌛ Осталось: <b>{_fmt_wait_human(wait)}</b>\n"
            f"│ 📅 Доступна с: <b>{next_dt}</b>\n"
            f"└─────────────────────────\n\n"
            f"💎 Хотите без ожидания? Используйте платный режим — он без лимитов."
        )

    # ── Дневной лимит (опциональный, по умолчанию выключен для бесплатных) ──
    today = today_key()
    if rec.get("date") != today:
        used = 0
    else:
        used = int(rec.get("free" if is_free else "paid", 0) or 0)

    limit = FREE_DAILY_LIMIT if is_free else PAID_DAILY_LIMIT
    if limit > 0 and used >= limit:
        kind = "бесплатных" if is_free else "платных"
        return False, (
            f"🚫 <b>Дневной лимит {kind} генераций исчерпан</b>\n\n"
            f"Использовано: <b>{used}/{limit}</b>\n"
            f"Лимит обновится в полночь 🕛"
        )

    return True, ""


def record_user_generation(user_id: int, mode: str) -> None:
    """Фиксирует факт генерации.

    Хранит:
      - last_free_ts / last_paid_ts — timestamp последней генерации (для кулдауна,
        не сбрасывается полночью);
      - free / paid — счётчик за текущие сутки (используется только если
        включён дневной лимит >0).
    """
    if is_vip(user_id):
        return
    # Перегенерация той же работы не списывает новый лимит/звёзды.
    if mode == "regen":
        return

    data  = load_usage()
    uid   = str(user_id)
    today = today_key()

    rec = data.get(uid, {}) or {}
    # last_*_ts НЕ обнуляем при смене даты — иначе кулдаун в 5 дней не сработает
    if rec.get("date") != today:
        rec["date"] = today
        rec["free"] = 0
        rec["paid"] = 0

    key    = "free" if mode == "free" else "paid"
    ts_key = "last_free_ts" if mode == "free" else "last_paid_ts"
    rec[key]    = int(rec.get(key, 0) or 0) + 1
    rec[ts_key] = int(datetime.now().timestamp())

    data[uid] = rec
    save_usage(data)


def get_user_limits_info(user_id: int) -> str:
    """Возвращает красивую карточку с лимитами пользователя."""
    if is_vip(user_id):
        return (
            "┌─────────────────────────\n"
            "│ 👑 <b>Статус: VIP</b>\n"
            "│ ♾ Генерации — без лимитов\n"
            "└─────────────────────────"
        )

    data = load_usage()
    uid  = str(user_id)
    rec  = data.get(uid, {}) or {}

    now      = int(datetime.now().timestamp())
    last_free = int(rec.get("last_free_ts", 0) or 0)

    if FREE_COOLDOWN > 0 and last_free and (now - last_free) < FREE_COOLDOWN:
        wait_s   = FREE_COOLDOWN - (now - last_free)
        next_dt  = datetime.fromtimestamp(last_free + FREE_COOLDOWN).strftime("%d.%m.%Y %H:%M")
        free_str = (
            f"⏳ Ждать <b>{_fmt_wait_human(wait_s)}</b>\n"
            f"│ 📅 Доступна: <b>{next_dt}</b>"
        )
    else:
        period_h = FREE_COOLDOWN // 3600
        period_s = f"{period_h // 24} дн" if period_h >= 24 else f"{period_h} ч"
        free_str = f"✅ Доступна (1 раз в {period_s})"

    paid_str = (
        "♾ Безлимитно" if PAID_DAILY_LIMIT == 0
        else f"{int(rec.get('paid', 0) or 0)}/{PAID_DAILY_LIMIT}"
    )

    return (
        "┌─────────────────────────\n"
        f"│ 🆓 Бесплатно: {free_str}\n"
        f"│ ⭐ Платные:    <b>{paid_str}</b>\n"
        "└─────────────────────────"
    )


# ═══════════════════════════════════════════════════════════════
#  AI МОДЕЛИ
# ═══════════════════════════════════════════════════════════════

class ModelStatus:
    AVAILABLE = "✅"
    LIMIT     = "❌ лимит"
    UNKNOWN   = "❓"
    FATAL     = "🔴 ошибка"


AI_MODELS: dict = {
    "deepseek": {
        "name":           "🐋 DeepSeek Chat",
        "base_url":       DEEPSEEK_BASE_URL,
        "api_key":        DEEPSEEK_KEY,
        "model":          DEEPSEEK_MODEL,
        "price_per_page": DEEPSEEK_PRICE,
        "status":         ModelStatus.UNKNOWN,
        "_fatal":         False,
    },
    "deepseek_r1": {
        "name":           "🧠 DeepSeek R1 (OpenRouter)",
        "base_url":       OPENROUTER_BASE_URL,
        "api_key":        OPENROUTER_KEY,
        "model":          OPENROUTER_R1_MODEL,
        "price_per_page": DEEPSEEK_PRICE,
        "status":         ModelStatus.UNKNOWN,
        "_fatal":         False,
    },
    "gemini_or": {
        "name":           "🌟 Gemini 2.5 Flash (OpenRouter)",
        "base_url":       OPENROUTER_BASE_URL,
        "api_key":        OPENROUTER_KEY,
        "model":          OPENROUTER_GEMINI_MODEL,
        "price_per_page": OR_GEMINI_PRICE,
        "status":         ModelStatus.UNKNOWN,
        "_fatal":         False,
    },
    "groq": {
        "name":           "⚡ Groq (LLaMA 3.3 70B)",
        "base_url":       GROQ_BASE_URL,
        "api_key":        GROQ_KEY,
        "model":          GROQ_MODEL,
        "price_per_page": GROQ_PRICE,
        "status":         ModelStatus.UNKNOWN,
        "_fatal":         False,
    },
}

if FREE_MODEL_KEY not in AI_MODELS:
    FREE_MODEL_KEY = "deepseek"


def _strip_markdown_markers(text: str) -> str:
    """Убирает markdown-маркеры `#` и `*`, чтобы они не попадали в DOCX."""
    if not text:
        return ""
    # Убираем # в начале строки (заголовки markdown)
    text = re.sub(r'(?m)^\s*#{1,6}\s*\*{0,2}\s*', '', text)
    text = re.sub(r'(?m)^\s*\*{2}([^*]+)\*{2}\s*$', r'\1', text)
    # Убираем **жирный** и *курсив*
    text = re.sub(r"\*{1,3}([^*\n]+?)\*{1,3}", r"\1", text)
    # Убираем оставшиеся # (в начале строки и одиночные внутри текста)
    text = re.sub(r'(?m)^\s*#+\s*', '', text)
    text = text.replace("#", "")
    # Убираем висящие ** после удаления # в строках вида `# **Заголовок**`
    text = re.sub(r'(?m)^\s*\*\*', '', text)
    text = re.sub(r'(?m)\*\*\s*$', '', text)
    return text.strip()


_AI_MARKER_REPLACEMENTS = [
    # Служебные ответы модели.
    (r"(?im)^\s*(конечно|разумеется|хорошо)[,.!\s]*(?:вот|ниже)\s+[^\n.?!]*[.?!]?\s*", ""),
    (r"(?im)^\s*(?:вот|ниже)\s+(?:ваш|представлен|привед[её]н)[^\n.?!]*[.?!]?\s*", ""),
    (r"(?i)\bкак (?:искусственный интеллект|ии|языковая модель)[^.!?\n]*[.!?]?\s*", ""),
    (r"(?i)\bя (?:являюсь|не являюсь|не могу|не имею возможности)[^.!?\n]*[.!?]?\s*", ""),
    (r"(?i)\bобъ[её]м текста (?:строго )?(?:выдержан|соблюд[её]н)[^.!?\n]*[.!?]?\s*", ""),
    # Типовые «ИИ-маркеры» внутри академического текста.
    (r"(?i)\bв заключение следует отметить,?\s+что\s+", ""),
    (r"(?i)\bподводя итог,?\s+", ""),
    (r"(?i)\bтаким образом,?\s+", ""),
    (r"(?i)\bследует отметить,?\s+что\s+", ""),
    (r"(?i)\bнеобходимо отметить,?\s+что\s+", ""),
    (r"(?i)\bважно подчеркнуть,?\s+что\s+", ""),
    (r"(?i)\bнельзя не отметить,?\s+что\s+", ""),
    (r"(?i)\bв целом можно сказать,?\s+что\s+", ""),
    (r"(?i)\bможно сделать вывод,?\s+что\s+", ""),
    (r"(?i)\bисходя из вышеизложенного,?\s+", ""),
]


def _remove_ai_marker_phrases(text: str) -> str:
    """Удаляет служебные и шаблонные фразы-маркеры ИИ без добавления опечаток."""
    if not text:
        return ""
    out = text
    for pattern, repl in _AI_MARKER_REPLACEMENTS:
        out = re.sub(pattern, repl, out)
    # Чистим пробелы, которые могли остаться после удаления вводных фраз.
    out = re.sub(r"[ \t]{2,}", " ", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    out = re.sub(r"(?m)^\s+", "", out)
    return out.strip()


def _remove_duplicate_phrases(text: str) -> str:
    """Срезает повторяющиеся фразы / тавтологии в одном предложении.

    Примеры, которые ловит:
      «Проведённый анализ показывает, что проведённый анализ подтверждает, что …»
      → «Проведённый анализ показывает, что подтверждается, что …»
      «итак, итак, мы видим …» → «итак, мы видим …»
      «таким образом, таким образом, ясно …» → «таким образом, ясно …»

    Стратегия:
      1. Срезаем подряд идущие n-граммы (3–6 слов), повторяющиеся в одной фразе.
      2. Срезаем дублирующиеся вводные обороты подряд («итак, итак,», «таким
         образом, таким образом»).
    """
    if not text:
        return text

    # 1) Подряд идущие n-граммы длины 3..6 — типичный паттерн LLM-тавтологии.
    #    «AAA BBB» где AAA = BBB (с точностью до пробелов и регистра).
    for n in (6, 5, 4, 3):
        pattern = re.compile(
            r"(\b(?:[А-Яа-яЁё]+(?:\s+|,\s+|\s+что\s+)){" + str(n) + r"})\1",
            flags=re.IGNORECASE,
        )
        prev = None
        while prev != text:
            prev = text
            text = pattern.sub(r"\1", text)

    # 2) Подряд идущие вводные обороты — «итак, итак,», «таким образом, таким образом».
    intros = [
        "итак", "таким образом", "следовательно", "кроме того",
        "помимо этого", "более того", "в заключение",
    ]
    for intro in intros:
        # «итак, итак,» / «таким образом, таким образом,» → одно вхождение
        pat = re.compile(rf"\b({re.escape(intro)})\s*,?\s+\1\b\s*,?", flags=re.IGNORECASE)
        text = pat.sub(r"\1,", text)

    # 3) Специальный кейс: «проведённый анализ … проведённый анализ» близко друг к другу.
    text = re.sub(
        r"(\bпроведён(?:н)?ый\s+анализ\b[^.]*?)(\s+(?:что|и|который|где)\s+)?\bпроведён(?:н)?ый\s+анализ\b",
        r"\1",
        text,
        flags=re.IGNORECASE,
    )

    # 4) Чистим возникшие двойные пробелы и запятые.
    text = re.sub(r"\s+,", ",", text)
    text = re.sub(r",\s*,", ",", text)
    text = re.sub(r",([А-Яа-яЁёA-Za-z])", r", \1", text)  # пробел после запятой
    text = re.sub(r"\s{2,}", " ", text)
    return text


def sanitize_llm_text(raw: str) -> str:
    """Чистит мусор от LLM: markdown-разметку, тройные переводы строк."""
    if not raw:
        return ""
    text = _strip_markdown_markers(raw.strip())
    # убираем ```код``` блоки
    text = re.sub(r"```[^\n]*\n?", "", text)
    # убираем **жирный** и *курсив* markdown
    text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
    # убираем # заголовки markdown
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # убираем тройные+ переводы строк
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Лёгкая «расклейка» слипшихся слов на стыке регистров/алфавитов/цифр
    # (безопасные правила, без разрыва по точкам — чтобы не ломать «т.е.» и инициалы).
    text = re.sub(r'([а-яё])([А-ЯЁ])', r'\1 \2', text)
    text = re.sub(r'([A-Za-z])([А-ЯЁ])', r'\1 \2', text)
    text = re.sub(r'([а-яё])([A-Z])', r'\1 \2', text)
    text = re.sub(r'(\d)([А-ЯЁа-яё])', r'\1 \2', text)
    text = re.sub(r'([А-ЯЁа-яё])(\d)', r'\1 \2', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = _remove_ai_marker_phrases(text)
    # (user-patch): убираем повторы фраз и тавтологию вида
    #   «проведённый анализ показывает, что проведённый анализ …»
    text = _remove_duplicate_phrases(text)
    return text.strip()


async def call_openai_compat(
    info: dict,
    messages: list[dict],
    max_tokens: int = 4096,
    timeout: int = 300,
) -> str:
    """Вызов OpenAI-совместимого API."""
    if info.get("_fatal") or not info.get("api_key"):
        return ""

    base    = info["base_url"].rstrip("/")
    headers = {
        "Content-Type":  "application/json",
        "Authorization": f"Bearer {info['api_key']}",
    }

    if "openrouter.ai" in base:
        headers["HTTP-Referer"] = "https://t.me/gost_assistant_bot"
        headers["X-Title"]      = "GOST Assistant Bot"

    payload = {
        "model":       info["model"],
        "messages":    messages,
        "temperature": 0.7,
        "max_tokens":  max_tokens,
    }

    try:
        async with aiohttp.ClientSession() as sess:
            async with sess.post(
                f"{base}/chat/completions",
                headers=headers,
                json=payload,
                timeout=aiohttp.ClientTimeout(total=timeout),
            ) as r:
                txt = await r.text()
                if r.status == 200:
                    data = json.loads(txt)
                    choice = data["choices"][0]
                    content = choice["message"]["content"]
                    # АНТИ-ОБРЫВ: если модель упёрлась в лимит токенов
                    # (finish_reason == "length"), текст оборван на полуслове.
                    # Делаем до 2 «дописываний», передавая хвост как контекст,
                    # и склеиваем — чтобы пользователь не получал обрезанную мысль.
                    finish = choice.get("finish_reason") or ""
                    cont_rounds = 0
                    while finish == "length" and cont_rounds < 2 and content:
                        cont_rounds += 1
                        tail = content[-1200:]
                        cont_msgs = messages + [
                            {"role": "assistant", "content": content},
                            {"role": "user", "content":
                                "Продолжи ровно с того места, где остановился, "
                                "не повторяя уже написанное и не начиная заново. "
                                "Заверши мысль и доведи текст до логичного конца. "
                                f"Последний фрагмент: …{tail}"},
                        ]
                        cont_payload = dict(payload, messages=cont_msgs)
                        async with sess.post(
                            f"{base}/chat/completions", headers=headers,
                            json=cont_payload,
                            timeout=aiohttp.ClientTimeout(total=timeout),
                        ) as r2:
                            if r2.status != 200:
                                break
                            d2 = json.loads(await r2.text())
                            c2 = d2["choices"][0]
                            add = c2["message"]["content"] or ""
                            if not add.strip():
                                break
                            sep = "" if content.endswith((" ", "\n")) else " "
                            content += sep + add.lstrip()
                            finish = c2.get("finish_reason") or ""
                    return content
                if r.status in (401, 402, 403):
                    info["_fatal"] = True
                    info["status"] = ModelStatus.FATAL
                    print(f"[FATAL] {info['name']} — auth error {r.status}")
                elif r.status == 429:
                    info["status"] = ModelStatus.LIMIT
                    # Отслеживание перегрузки API
                    model_key = next((k for k, v in AI_MODELS.items() if v.get("base_url") == info.get("base_url") and v.get("model") == info.get("model")), "unknown")
                    increment_overload(model_key)
                    print(f"[LIMIT] {info['name']} — rate limit (API overload detected)")
                elif r.status == 503:
                    info["status"] = ModelStatus.LIMIT
                    model_key = next((k for k, v in AI_MODELS.items() if v.get("base_url") == info.get("base_url") and v.get("model") == info.get("model")), "unknown")
                    increment_overload(model_key)
                    print(f"[OVERLOAD] {info['name']} — service unavailable")
                else:
                    print(f"[ERROR] {info['name']} — HTTP {r.status}: {txt[:200]}")
    except asyncio.TimeoutError:
        print(f"[TIMEOUT] {info['name']}")
    except Exception as e:
        print(f"[ERR] {info['name']}: {e}")

    return ""


async def chat_with_model(info: dict, messages: list[dict], max_tokens: int = 4096) -> str:
    """Вызывает модель. При перегрузке возвращает пустую строку."""
    # Определяем ключ модели для проверки перегрузки
    model_key = next((k for k, v in AI_MODELS.items() if v.get("base_url") == info.get("base_url") and v.get("model") == info.get("model")), "unknown")
    
    # Проверяем, не перегружена ли модель
    if is_api_overloaded(model_key):
        print(f"[OVERLOAD] Модель {info.get('name', model_key)} перегружена, пропускаем...")
        info["status"] = ModelStatus.LIMIT
        return ""
    
    raw = await call_openai_compat(info, messages, max_tokens=max_tokens)
    
    # Если вернулась пустая строка и статус LIMIT — возможно перегрузка
    if not raw and info.get("status") == ModelStatus.LIMIT:
        increment_overload(model_key)
    
    # Успешный ответ — сбрасываем счётчик перегрузок
    if raw and len(raw.strip()) > 100:
        reset_overload(model_key)
    
    return _normalize_homoglyphs(sanitize_llm_text(raw))


def fallback_chain(primary: str) -> list[str]:
    """Цепочка фоллбэков: primary → все остальные доступные модели по приоритету.

    Раньше использовался только OpenRouter (deepseek_r1, gemini_or), из-за
    чего при сбое OpenRouter Groq и прямой DeepSeek API не подхватывались.
    Теперь честно перебираем все модели, у которых есть api_key и нет
    фатальной ошибки. Дубликаты не добавляются.
    """
    # Полный приоритет: сначала primary, затем — в порядке предпочтения.
    priority = [
        primary,
        "deepseek",      # прямой DeepSeek API (дешёвый, стабильный)
        "deepseek_r1",   # OpenRouter / DeepSeek R1
        "gemini_or",     # OpenRouter / Gemini
        "groq",          # Groq (быстрый, бесплатный)
    ]
    out: list[str] = []
    for k in priority:
        if not k or k in out:
            continue
        info = AI_MODELS.get(k)
        if not info:
            continue
        if info.get("_fatal"):
            continue
        if not info.get("api_key"):
            continue
        out.append(k)
    return out


async def chat_with_fallback(
    primary: str,
    messages: list[dict],
    max_tokens: int,
) -> tuple[str, str]:
    """Пробует модели по цепочке, возвращает (текст, ключ_модели)."""
    best_text = ""
    best_model = primary
    for k in fallback_chain(primary):
        info = AI_MODELS[k]
        text = await chat_with_model(info, messages, max_tokens=max_tokens)
        if text and len(text.strip()) > 100:
            info["status"] = ModelStatus.AVAILABLE
            return text, k
        # Сохраняем лучший результат даже если он короткий
        if text and len(text.strip()) > len(best_text.strip()):
            best_text = text
            best_model = k
        if not text:
            info["status"] = ModelStatus.LIMIT
            print(f"[FALLBACK] Модель {info.get('name', k)} вернула пустой ответ, пробую следующую...")
    # Если ни одна модель не вернула > 100 символов, возвращаем лучшее что есть
    if best_text:
        print(f"[FALLBACK] Все модели дали < 100 зн., лучший: {len(best_text)} зн. от {best_model}")
    return best_text, best_model


# ═══════════════════════════════════════════════════════════════
#  РЕАЛЬНЫЕ ИСТОЧНИКИ И МАТЕРИАЛЫ ИЗ САЙТОВ
# ═══════════════════════════════════════════════════════════════

_URL_RE = re.compile(r"https?://[^\s<>()\[\]{}\"'«»]+", re.IGNORECASE)


def _extract_urls(text: str, limit: int = 5) -> list[str]:
    """Достаёт URL из пользовательских материалов, сохраняя порядок."""
    if not text:
        return []
    seen: set[str] = set()
    urls: list[str] = []
    for raw in _URL_RE.findall(text):
        url = raw.rstrip(".,;:!?)»]")
        if url not in seen:
            seen.add(url)
            urls.append(url)
        if len(urls) >= limit:
            break
    return urls


def _strip_html(raw: str) -> str:
    """Очень лёгкая HTML→text очистка без внешних зависимостей."""
    if not raw:
        return ""
    raw = re.sub(r"(?is)<(script|style|noscript).*?</\1>", " ", raw)
    raw = re.sub(r"(?is)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?is)</(p|div|li|h[1-6]|tr)>", "\n", raw)
    raw = re.sub(r"(?is)<[^>]+>", " ", raw)
    raw = html.unescape(raw)
    raw = re.sub(r"[ \t]{2,}", " ", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip()


def _safe_year_from_crossref(item: dict) -> str:
    for key in ("published-print", "published-online", "published", "issued"):
        parts = item.get(key, {}).get("date-parts") if isinstance(item.get(key), dict) else None
        if parts and parts[0] and parts[0][0]:
            return str(parts[0][0])
    return "б. г."


def _author_to_gost(name: str) -> str:
    """Форматирует автора как «Фамилия И.О.» насколько это возможно."""
    name = re.sub(r"\s+", " ", (name or "").strip())
    if not name:
        return ""
    parts = name.split()
    if len(parts) == 1:
        return parts[0]
    # Уже готовый вид: «Иванов И.И.» / «Smith J.» — не переворачиваем.
    if re.fullmatch(r"(?:[A-ZА-ЯЁ]\.){1,4}", parts[-1].replace(" ", "")):
        return f"{parts[0]} {parts[-1]}".strip()
    # Если первая часть похожа на русскую фамилию, а дальше полные имя/отчество.
    if len(parts) >= 2 and re.search(r"[А-ЯЁ][а-яё]+", parts[0]) and re.search(r"[А-ЯЁ][а-яё]+", parts[1]):
        family = parts[0]
        given = parts[1:]
    else:
        # Английский порядок чаще Given Family.
        family = parts[-1]
        given = parts[:-1]
    initials = "".join((p[0].upper() + ".") for p in given if p)
    return f"{family} {initials}".strip()


def _crossref_authors(item: dict, limit: int = 3) -> list[str]:
    out: list[str] = []
    for a in item.get("author") or []:
        family = (a.get("family") or "").strip()
        given = (a.get("given") or "").strip()
        if family and given:
            if re.search(r"\.\s*$", given):
                formatted = f"{family} {given.replace(' ', '')}"
            else:
                initials = "".join(p[0].upper() + "." for p in re.split(r"[\s-]+", given) if p)
                formatted = f"{family} {initials}"
        else:
            formatted = _author_to_gost((family or given).strip())
        if formatted:
            out.append(formatted)
        if len(out) >= limit:
            break
    return out


def _openalex_authors(item: dict, limit: int = 3) -> list[str]:
    out: list[str] = []
    for a in item.get("authorships") or []:
        name = ((a.get("author") or {}).get("display_name") or "").strip()
        if name:
            out.append(_author_to_gost(name))
        if len(out) >= limit:
            break
    return out


def _semantic_authors(item: dict, limit: int = 3) -> list[str]:
    out: list[str] = []
    for a in item.get("authors") or []:
        name = (a.get("name") or "").strip()
        if name:
            out.append(_author_to_gost(name))
        if len(out) >= limit:
            break
    return out


def _clean_ref_title(title: str, max_len: int = 180) -> str:
    title = _strip_html(title or "")
    title = re.sub(r"\s+", " ", title).strip(" .")
    if len(title) > max_len:
        title = title[:max_len].rsplit(" ", 1)[0] + "…"
    return title


_TOPIC_STOPWORDS = {
    "тема", "работа", "исследование", "анализ", "роль", "значение",
    "основы", "особенности", "проблемы", "вопросы", "современный",
    "современная", "современное", "развитие", "система", "метод",
    "методы", "подход", "подходы", "россии", "российской", "российский",
    "изменение", "изменения", "изменений", "влияние", "влиянием",
    "политика", "политики", "территориальный", "территориальные",
    "международный", "международная", "международной", "международные",
    "the", "and", "for", "with", "from", "this", "that", "study",
    "analysis", "role", "problems", "development",
}


def _keywords_for_relevance(text: str) -> set[str]:
    """Ключевые слова для проверки, что источник действительно по теме."""
    if not text:
        return set()
    words = re.findall(r"[A-Za-zА-Яа-яЁё][A-Za-zА-Яа-яЁё\-]{3,}", text.lower())
    out: set[str] = set()
    for w in words:
        w = w.strip("-–—_")
        if len(w) < 4 or w in _TOPIC_STOPWORDS:
            continue
        out.add(w)
    return out


def _topic_required_terms_match(record: dict, topic: str) -> bool:
    """Жёсткая проверка для тем с именами/узкими объектами.

    Без неё OpenAlex/Crossref по длинной русской теме «...под влиянием...»
    часто возвращали случайные статьи про любые «изменения под влиянием»,
    из-за чего в списке литературы появлялись биология, медицина и т.п.
    """
    low_topic = (topic or "").lower()
    if not low_topic:
        return True
    haystack = " ".join(str(record.get(k) or "") for k in (
        "title", "container", "publisher", "concepts", "subjects", "url", "doi"
    )).lower()

    if "трамп" in low_topic or "trump" in low_topic:
        has_trump = ("trump" in haystack) or ("трамп" in haystack)
        if any(x in low_topic for x in ("торгов", "trade", "тариф", "пошлин")):
            trade_terms = (
                "trade", "tariff", "tariffs", "commerce", "china", "wto",
                "торгов", "тариф", "пошлин", "экспорт", "импорт"
            )
            return has_trump and any(t in haystack for t in trade_terms)
        return has_trump

    return True


def _source_relevance_score(record: dict, topic: str, subject: str) -> int:
    """Оценивает релевантность источника: тема важнее дисциплины."""
    title = record.get("title") or ""
    haystack = " ".join(
        str(record.get(k) or "")
        for k in ("title", "container", "publisher", "concepts", "subjects")
    ).lower()
    topic_aliases = _topic_aliases(topic)
    subject_alias_text = _subject_alias_text(subject)
    subject_norm = re.sub(r"\s+", " ", (subject_alias_text or "").lower()).strip()
    topic_words = _keywords_for_relevance(" ".join(topic_aliases or [topic or ""]))
    subject_words = _keywords_for_relevance(subject_alias_text)

    score = 0
    if not _topic_required_terms_match(record, topic):
        score -= 100
    for alias in topic_aliases:
        alias_norm = re.sub(r"\s+", " ", alias.lower()).strip()
        if alias_norm and len(alias_norm) >= 4 and alias_norm in haystack:
            score += 10
    if subject_norm and len(subject_norm) >= 4 and subject_norm in haystack:
        score += 2

    for w in topic_words:
        if w in haystack:
            score += 4 if w in (title or "").lower() else 3
    for w in subject_words:
        if w in haystack:
            score += 1
    return score


def _filter_relevant_source_records(records: list[dict], topic: str, subject: str, limit: int, doc_type: str = "") -> list[dict]:
    """Фильтрует источники, чтобы список литературы был по теме и не короче 10.

    Сначала берём самые релевантные записи (совпадение с темой в названии,
    журнале, ключевых словах). Если их меньше MIN_REAL_SOURCES, добавляем
    реальные записи из поисковой выдачи по этой же теме: лучше 10 проверяемых
    источников из каталогов, чем 3 хороших и 9 выдуманных моделью.

    FIX: Все источники ТОЧНО по теме. Источники не по теме отбрасываются.
    """
    limit = max(MIN_REAL_SOURCES, min(limit, MAX_WEB_SOURCES, BIB_SOURCE_TARGET))
    # Берём большой пул: первые запросы могут быть русскоязычными и шумными,
    # а наиболее релевантные англоязычные статьи (Trump trade/tariffs) идут
    # в следующих query-вариантах. Ранний срез до 60–120 записей отбрасывал их.
    deduped = _dedupe_source_records(records, max(len(records), limit * 25, 500))
    scored = [(_source_relevance_score(r, topic, subject), r) for r in deduped]
    scored.sort(key=lambda x: x[0], reverse=True)

    selected: list[dict] = []
    selected_keys: set[str] = set()

    def _key(r: dict) -> str:
        return (r.get("doi") or re.sub(r"\W+", "", (r.get("title") or "").lower())[:120]).lower()

    def _is_on_topic(r: dict, topic: str, subject: str) -> bool:
        """Проверяет что источник действительно по теме (не общий учебник)."""
        title_lower = (r.get("title") or "").lower()
        container_lower = (r.get("container") or "").lower()
        concepts_lower = (r.get("concepts") or "").lower()
        haystack = f"{title_lower} {container_lower} {concepts_lower}"

        # Ключевые слова темы + латинские алиасы (Trump, trade war и т.п.).
        # Без алиасов англоязычные каталоги не проходят проверку по русской теме.
        topic_words = _keywords_for_relevance(" ".join(_topic_aliases(topic) or [topic]))
        subject_words = _keywords_for_relevance(subject)

        # Минимум 1 слово из темы должно быть в названии/журнале
        topic_in_title = any(w in title_lower for w in topic_words)
        topic_in_container = any(w in container_lower for w in topic_words)
        topic_in_concepts = any(w in concepts_lower for w in topic_words)

        # Общие слова предмета (несколько — нормально для предметной литературы)
        subject_matches = sum(1 for w in subject_words if w in haystack)

        # Если в названии есть явные маркеры НЕ-по-теме — отбрасываем
        off_topic_markers = ["введение", "общая", "основы", "курс лекций", "учебное пособие"]
        is_generic = any(m in title_lower for m in off_topic_markers)

        # Источник по теме если:
        # - хотя бы 1 ключевое слово темы в названии/журнале/концептах ИЛИ
        # - предметные слова + не слишком общий
        has_topic_word = topic_in_title or topic_in_container or topic_in_concepts
        is_not_too_generic = not is_generic or has_topic_word

        return has_topic_word or (subject_matches >= 2 and is_not_too_generic)

    def _add(candidates: list[dict], require_topic_match: bool = False) -> None:
        for r in candidates:
            if len(selected) >= limit:
                break
            if not r.get("authors"):
                continue
            if not _record_is_scientific_article(r):
                continue
            k = _key(r)
            if not k or k in selected_keys:
                continue
            # FIX: Проверяем что источник реально по теме. Для тем с именами
            # (Трамп и т.п.) это обязательно даже для слабых/резервных кандидатов.
            if not _topic_required_terms_match(r, topic):
                continue
            if require_topic_match and not _is_on_topic(r, topic, subject):
                continue
            selected_keys.add(k)
            selected.append(r)

    strong = [r for score, r in scored if score >= 4]
    weak = [r for score, r in scored if 0 < score < 4]
    rest = [r for score, r in scored if score <= 0]

    # Сначала берём strong с проверкой по теме
    _add(strong, require_topic_match=True)
    if len(selected) < MIN_REAL_SOURCES:
        # weak без строгой проверки по теме, но с базовой фильтрацией
        _add(weak, require_topic_match=False)
    if len(selected) < MIN_REAL_SOURCES:
        # Последний резерв: только записи, которые всё равно проходят
        # обязательные термины темы; случайный «мусор» не добираем ради числа.
        _add(rest, require_topic_match=True)

    # Вариативность по жанрам: список для эссе/доклада/реферата на одну тему
    # не начинается с одних и тех же двух работ, но остаётся детерминированным.
    if selected and doc_type:
        genre_offsets = {
            "referat": 0,
            "esse": 3,
            "doklad": 6,
            "kursovaya": 2,
            "kontrolnaya": 4,
            "article": 1,
            "final_referat": 5,
            "vkr": 7,
            "final_project": 8,
        }
        offset = genre_offsets.get(doc_type, 0) % len(selected)
        if offset:
            selected = selected[offset:] + selected[:offset]

    return selected[:limit]


def _format_source_record(record: dict) -> str:
    """Форматирует запись из OpenAlex/Crossref в приближённый ГОСТ."""
    authors = record.get("authors") or []
    title = _clean_ref_title(record.get("title") or "")
    year = str(record.get("year") or "б. г.")
    container = _clean_ref_title(record.get("container") or "")
    publisher = _clean_ref_title(record.get("publisher") or "")
    doi = _normalize_doi_spacing((record.get("doi") or "").replace("https://doi.org/", "")).strip().replace(" ", "")
    url = (record.get("url") or "").strip()

    if not title:
        return ""
    author_part = ", ".join([a.rstrip(".") for a in authors if a]) or "Без автора"
    if container:
        ref = f"{author_part}. {title} // {container}. — {year}."
    elif publisher:
        ref = f"{author_part}. {title}. — {publisher}, {year}."
    else:
        ref = f"{author_part}. {title}. — {year}."
    if doi:
        doi = doi.strip().rstrip(".")
        doi_url = doi if doi.lower().startswith("http") else f"https://doi.org/{doi}"
        ref += f" — URL: {doi_url}."
    elif url:
        # FIX 7: убираем пробелы внутри URL («https: // press...» → «https://press...»)
        url = _normalize_url_spacing(url)
        ref += f" — URL: {url}."
    # FIX 3: если у источника нет ни DOI, ни URL — НЕ добавляем заглушку
    # «[URL не указан]». По ГОСТ 7.32 источник без электронного адреса
    # (книга, статья в печатном издании) полностью допустим.
    # Не пропускаем битые строки из каталогов: «РАН Б.И.П.С.», «университет Б.Г.» и т.п.
    # (user-patch): на случай, если из LLM/каталогов протянулась заглушка
    ref = re.sub(r"[\s\.,;:\(\[—-]*\[\s*URL\s+не\s+указан\s*\][\.\s\)\]]*", " ", ref, flags=re.IGNORECASE)
    ref = re.sub(r"[\s\.,;:\(\[—-]*\(\s*URL\s+не\s+указан\s*\)[\.\s\)\]]*", " ", ref, flags=re.IGNORECASE)
    ref = re.sub(r"\bURL\s+не\s+указан\b\.?", "", ref, flags=re.IGNORECASE)
    ref = re.sub(r"\s{2,}", " ", ref).strip()
    if _is_bad_literature_line(ref):
        return ""
    return ref



_NON_SCIENTIFIC_SOURCE_MARKERS = (
    # Нормативка/стандарты — не должны попадать в обычный список литературы
    # реферата по политологии/географии/биологии и т.п.
    "гост", "gost", "стандарт", "standard", "iso ", "санпин",
    "федеральный закон", "закон от", "кодекс", "постановление", "приказ",
    "consultant.ru", "garant.ru",
    # Методические книги о написании работ — реальные, но не источники по теме.
    "research design", "craft of research", "how to write a thesis",
    "методология научного исследования", "кандидатская диссертация",
    "правила оформления", "библиографическая ссылка", "библиографическая запись",
    "отчет о научно-исследовательской работе",
)

_SCIENTIFIC_CONTAINER_HINTS = (
    "journal", "review", "studies", "research", "science", "scientific",
    "вестник", "журнал", "известия", "науч", "исслед", "обзор",
    "полит", "географ", "биолог", "history", "historical", "sociology",
    "conference", "proceedings", "труды", "ученые записки", "учёные записки",
)


def _is_non_scientific_source_text(text: str) -> bool:
    """True для ГОСТов, законов и методичек, которые не являются научными
    источниками по теме работы."""
    low = (text or "").lower()
    return any(m in low for m in _NON_SCIENTIFIC_SOURCE_MARKERS)


def _record_is_scientific_article(record: dict) -> bool:
    """Оставляет прежде всего статьи/публикации из научных журналов и сборников.

    В OpenAlex/Crossref/Semantic Scholar книги и нормативка часто приходят
    вместе со статьями. Для обычной учебной работы безопаснее брать записи с
    авторами, названием, журналом/сборником и DOI/URL, а ГОСТы/законы/методички
    отбрасывать.
    """
    title = str(record.get("title") or "")
    container = str(record.get("container") or "")
    publisher = str(record.get("publisher") or "")
    blob = f"{title} {container} {publisher} {record.get('url') or ''}"
    if _is_non_scientific_source_text(blob):
        return False
    if not record.get("authors") or not title.strip():
        return False
    # DOI/URL нужен для проверяемости. Не требуем обязательно «journal» в
    # container: у OpenAlex/Crossref часть нормальных статей и материалов
    # приходит без явного контейнера, и прежний фильтр из-за этого оставлял
    # меньше 10 источников. Нормативку/ГОСТы/методички мы уже отсекли выше.
    if not (record.get("doi") or record.get("url")):
        return False
    return True


def _filter_bibliography_scientific_lines(bib_text: str, *, min_keep: int = 0) -> str:
    """Удаляет из библиографии ГОСТы, законы и методички.

    Если после фильтрации стало слишком мало источников, функция всё равно
    возвращает очищенный список: лучше меньше релевантных научных публикаций,
    чем 10 позиций, набитых стандартами и законами не по теме.
    """
    norm = _normalize_bibliography(bib_text or "")
    lines = []
    for line in norm.split("\n"):
        clean = line.strip()
        if not clean:
            continue
        body = re.sub(r"^\d+\.\s*", "", clean)
        if _is_non_scientific_source_text(body):
            continue
        # Для LLM-добавок требуем проверяемость: DOI или URL. Каталожные строки
        # обычно уже содержат DOI/URL после _format_source_record().
        if not re.search(r"https?://|\bdoi\b|10\.\d{4,9}/", body, flags=re.I):
            # Печатные статьи без URL бывают нормальными, но именно они чаще
            # оказываются выдуманными LLM. Не пускаем их в аварийный добор.
            continue
        lines.append(body)
    return "\n".join(f"{i}. {body}" for i, body in enumerate(lines, start=1))


_BAD_LITERATURE_PATTERNS = [
    "РАН Б.И.П.С.",
    "университет Б.Г.",
    "Без автора",
    "ljournal",
]


def _is_bad_literature_line(line: str) -> bool:
    """Отсекает битые библиографические данные из внешних каталогов."""
    if not line:
        return True
    low = line.lower()
    if any(p.lower() in low for p in _BAD_LITERATURE_PATTERNS):
        return True
    # FIX 4: блокируем только конкретный мусор "ОРГАНИЗАЦИЯ + ИНИЦИАЛЫ"
    # ("РАН Б.И.П.С."), но НЕ организации-авторов ("Университет МГУ. ...").
    first_part = line.split(".", 1)[0]
    if re.fullmatch(
        r"\s*(?:ран|университет|институт|академия|центр|фонд)\s+"
        r"[А-ЯЁA-Z]\.?(\s+[А-ЯЁA-Z]\.?){2,}",
        first_part, re.IGNORECASE,
    ):
        return True
    if re.match(r"^[А-ЯЁA-Z]{2,}\s+(?:[А-ЯЁA-Z]\.){3,}", line.strip()):
        return True
    # FIX 10: блокируем ТОЛЬКО реальные инициалы вида "И. И." или "И.И.",
    # но НЕ трогаем короткие аббревиатуры вроде "МГУ", "ФГБОУ", "РАН".
    # Старая regex считала "МГУ" (3 заглавные) инициалами и отбрасывала
    # нормальные источники вида «МГУ. Исследования. — М., 2020.».
    author_part = line.split(".", 1)[0].strip()
    if re.fullmatch(r"(?:[А-ЯЁA-Z]\.\s*){2,}", author_part) or        re.fullmatch(r"(?:[А-ЯЁA-Z]\s+){1,4}[А-ЯЁA-Z]\.", author_part):
        # Реальные инициалы: «И. И.» или «И.И.» (с точками) — мусор.
        # Без точек («МГУ», «РАН») — легитимные аббревиатуры.
        return True
    return False



def _normalize_doi_spacing(text: str) -> str:
    """Убирает пробелы внутри DOI и doi.org URL."""
    if not text:
        return text

    def repl_doi(m: re.Match) -> str:
        doi = re.sub(r"\s+", "", m.group(1)).rstrip(".,;)")
        return "DOI: " + doi

    def repl_url(m: re.Match) -> str:
        doi = re.sub(r"\s+", "", m.group(1)).rstrip(".,;)")
        return "https://doi.org/" + doi

    text = re.sub(r"DOI\s*[:：]?\s*(10\s*\.\s*\d{4,9}\s*/\s*[^\s]+(?:\s+[^\.\n]+)?)", repl_doi, text, flags=re.I)
    text = re.sub(r"https?://(?:dx\.)?doi\.org/\s*(10\s*\.\s*\d{4,9}\s*/\s*[^\s]+(?:\s+[^\.\n]+)?)", repl_url, text, flags=re.I)
    # Частый случай: DOI разбит пробелами только вокруг слэша/точки.
    text = re.sub(r"(10)\s*\.\s*(\d{4,9})\s*/\s*", r"\1.\2/", text)
    return text


def _normalize_url_spacing(text: str) -> str:
    """FIX 7: убирает пробелы внутри URL.

    «https: // press.example.org / a b» → «https://press.example.org/ab».
    LLM часто разрывает ссылки пробелами после схемы и вокруг слешей.
    """
    if not text:
        return text
    s = text.strip()
    # Схлопываем пробелы вокруг схемы: "https :  / /" → "https://"
    s = re.sub(r"(https?)\s*:\s*/\s*/\s*", r"\1://", s, flags=re.I)
    m = re.match(r"(https?://)(.*)", s, flags=re.I)
    if m:
        scheme, rest = m.group(1), m.group(2)
        # В теле URL не бывает пробелов — удаляем их (и пробелы вокруг «/»).
        rest = re.sub(r"\s+", "", rest).rstrip(".,;)")
        s = scheme + rest
    return s


def _format_bibliography_item_gost(item: str) -> str:
    """Единый минимальный стиль ГОСТ 7.32/ГОСТ Р 7.0.5 для строки источника."""
    item = _normalize_doi_spacing(_normalize_punctuation(item.strip()))
    # _deglue_text может ошибочно превратить doi.org в «doi. org»;
    # URL в библиографии не должен содержать пробелов.
    item = re.sub(r"(https?)://([A-Za-z0-9.-]+)\.[ \t]+([A-Za-z]{2,})(/?)", r"\1://\2.\3\4", item)
    item = re.sub(r"(https?://\S+)\s+/\s*", r"\1/", item)
    item = re.sub(r"\s+", " ", item)
    item = item.replace("— —", "—")
    # URL/DOI приводим к единому виду электронного ресурса.
    item = re.sub(r"\bdoi\s*[:：]\s*(10\.\S+)", lambda m: "URL: https://doi.org/" + m.group(1).rstrip(" ."), item, flags=re.I)
    item = re.sub(r"(?<!URL: )https://doi\.org/(10\.\S+)", lambda m: "URL: https://doi.org/" + m.group(1).rstrip(" ."), item, flags=re.I)
    item = re.sub(r"\s*//\s*", " // ", item)
    item = re.sub(r"\s*—\s*", " — ", item)
    item = re.sub(r"\s+", " ", item).strip(" .") + "."
    # FIX 7: предыдущая нормализация «//» ломает схему URL
    # («https://» → «https: // »). Восстанавливаем схему и тело URL.
    item = re.sub(r"(https?)\s*:\s*//\s*", r"\1://", item, flags=re.I)
    item = re.sub(r"(https?://[^\s]*)\s+/\s*", r"\1/", item, flags=re.I)
    # Инициалы авторов: «Smith S, Ivanov I» → «Smith S., Ivanov I.»
    # (минимальная нормализация под ГОСТ без ломки DOI/URL).
    item = re.sub(r"\b([A-ZА-ЯЁ][A-Za-zА-Яа-яЁёİıĞğŞşÇçÖöÜü-]+)\s+([A-ZА-ЯЁ])(?=,|\s+//|\.)", r"\1 \2.", item)
    item = re.sub(r"\.\.", ".", item)
    return item

def _ensure_bibliography_urls(bib_text: str) -> str:
    """Гарантирует, что в каждой позиции есть рабочая ссылка или пометка.

    Если есть DOI в виде `DOI: 10...`, он превращается в рабочий URL
    `https://doi.org/10...`. Если ни DOI, ни URL нет — добавляется
    `[URL не указан]`.
    """
    if not bib_text:
        return ""
    out: list[str] = []
    for line in bib_text.split("\n"):
        line = _normalize_doi_spacing(line.strip())
        if not line:
            continue
        # DOI без URL → рабочая ссылка DOI.
        line = re.sub(
            r"DOI:\s*(10\.\S+)",
            lambda m: "URL: https://doi.org/" + m.group(1).rstrip(" ."),
            line,
            flags=re.IGNORECASE,
        )
        # FIX 7: чиним URL, разбитые пробелами («https: // …» → «https://…»).
        line = re.sub(r"(https?)\s*:\s*//\s*", r"\1://", line, flags=re.I)
        line = re.sub(r"(https?)://([A-Za-z0-9.-]+)\.[ \t]+([A-Za-z]{2,})(/?)", r"\1://\2.\3\4", line)
        if "https://doi.org/" in line.lower():
            line = re.sub(
                r"https://doi\.org/([^\n]+)$",
                lambda m: "https://doi.org/" + re.sub(r"\s+", "", m.group(1)).rstrip(" .") + ".",
                line, flags=re.I,
            )
        # FIX 3: больше НЕ добавляем заглушку «[URL не указан]» и
        # удаляем её, если она пришла из раньше сгенерированного текста.
        # (user-patch): расширенная чистка «[URL не указан]» в любых обёртках
        # и местах строки — пользователь сообщает, что заглушка всё ещё иногда
        # просачивается из текста LLM.
        line = re.sub(r"[\s\.,;:\(\[—-]*\[\s*URL\s+не\s+указан\s*\][\.\s\)\]]*", " ", line, flags=re.IGNORECASE)
        line = re.sub(r"[\s\.,;:\(\[—-]*\(\s*URL\s+не\s+указан\s*\)[\.\s\)\]]*", " ", line, flags=re.IGNORECASE)
        # Также «URL не указан» без скобок
        line = re.sub(r"[\s\.,;:—-]*\bURL\s+не\s+указан\b[\.\s]*", " ", line, flags=re.IGNORECASE)
        line = re.sub(r"\s{2,}", " ", line).rstrip(" .,;—-").rstrip()
        if not line.endswith("."):
            line += "."
        out.append(line)
    return "\n".join(out)


def validate_literature(bib_text: str) -> tuple[bool, str]:
    """Проверяет список литературы на битые данные."""
    if not bib_text or not bib_text.strip():
        return False, "список литературы пуст"
    bad_patterns = ["РАН Б.И.П.С.", "университет Б.Г.", "Без автора", "ljournal"]
    if any(p in bib_text for p in bad_patterns):
        return False, "обнаружены битые данные"
    for line in bib_text.split("\n"):
        clean = re.sub(r"^\d+\.\s*", "", line.strip())
        if clean and _is_bad_literature_line(clean):
            return False, "обнаружены битые данные"
    return True, ""


def _dedupe_source_records(records: list[dict], limit: int) -> list[dict]: 
    seen: set[str] = set()
    out: list[dict] = []
    for r in records:
        title_key = re.sub(r"\W+", "", (r.get("title") or "").lower())[:120]
        doi_key = (r.get("doi") or "").lower().replace("https://doi.org/", "")
        key = doi_key or title_key
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(r)
        if len(out) >= limit:
            break
    return out


async def _fetch_json(session: aiohttp.ClientSession, url: str) -> Optional[dict]:
    try:
        async with session.get(url, timeout=aiohttp.ClientTimeout(total=WEB_SOURCE_TIMEOUT)) as resp:
            if resp.status != 200:
                return None
            return await resp.json(content_type=None)
    except Exception as e:
        print(f"[WEB] JSON fetch failed: {e}")
        return None


_RU_TRANSLIT = str.maketrans({
    "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "e",
    "ж": "zh", "з": "z", "и": "i", "й": "y", "к": "k", "л": "l", "м": "m",
    "н": "n", "о": "o", "п": "p", "р": "r", "с": "s", "т": "t", "у": "u",
    "ф": "f", "х": "kh", "ц": "ts", "ч": "ch", "ш": "sh", "щ": "shch",
    "ъ": "", "ы": "y", "ь": "", "э": "e", "ю": "yu", "я": "ya",
})

_SPECIAL_SOURCE_ALIASES = {
    "трамп": [
        "Trump",
        "Trump trade policy", "Trump tariffs", "Trump trade war",
        "US China trade war Trump", "Trump administration tariffs China",
        "Trump global trade", "Trump international trade",
        "Donald Trump", "Donald J. Trump", "Trump administration",
        "Donald Trump geopolitics", "Trump political geography", "Trump foreign policy",
        "Trump border wall", "Trump immigration geography",
    ],
    "байкал": ["Baikal", "Lake Baikal"],
}


def _transliterate_ru_to_latin(text: str) -> str:
    if not text:
        return ""
    out = []
    for ch in text.lower():
        out.append(_RU_TRANSLIT.get(ch, ch))
    return re.sub(r"\s+", " ", "".join(out)).strip()


def _topic_aliases(topic: str) -> list[str]:
    topic_clean = re.sub(r"\s+", " ", (topic or "").strip())
    aliases: list[str] = []
    if topic_clean:
        aliases.append(topic_clean)
    low = topic_clean.lower()
    for key, vals in _SPECIAL_SOURCE_ALIASES.items():
        if key in low:
            aliases.extend(vals)
    translit = _transliterate_ru_to_latin(topic_clean)
    if translit and translit != low:
        aliases.append(translit)
    seen = set()
    out = []
    for a in aliases:
        k = a.lower().strip()
        if k and k not in seen:
            seen.add(k)
            out.append(a.strip())
    return out


def _subject_alias_text(subject: str) -> str:
    subject_l = (subject or "").lower()
    aliases = [subject or ""]
    if "географ" in subject_l:
        aliases.extend(["geography", "geopolitics", "political geography", "spatial politics"])
    elif "истор" in subject_l:
        aliases.extend(["history", "historical analysis"])
    elif "эконом" in subject_l:
        aliases.extend(["economics", "economic policy"])
    return " ".join(a for a in aliases if a)


def _source_query_variants(topic: str, subject: str) -> list[str]:
    """Генерирует несколько поисковых запросов, чтобы набрать 10–12 реальных источников.

    Для кириллических тем добавляет латинскую транслитерацию/известные алиасы
    (например, «Трамп» → Donald Trump), иначе каталоги часто находят 1–2 русские
    публикации вместо полноценного списка.
    """
    aliases = _topic_aliases(topic)
    subject_l = (subject or "").lower()
    extra: list[str] = []
    if "географ" in subject_l:
        extra = ["geopolitics", "political geography", "spatial politics"]
    elif "истор" in subject_l:
        extra = ["history", "historical analysis"]
    elif "эконом" in subject_l:
        extra = ["economics", "economic policy"]

    queries: list[str] = []
    for a in aliases or [topic]:
        if not a:
            continue
        queries.append(a)
        for e in extra[:2]:
            queries.append(f"{a} {e}")
    # Резерв: исходная тема + дисциплина.
    if topic and subject:
        queries.append(f"{topic} {subject}")

    seen = set()
    out = []
    for q in queries:
        q = re.sub(r"\s+", " ", q).strip()
        k = q.lower()
        if q and k not in seen:
            seen.add(k)
            out.append(q)
    return out[:14]


async def _fetch_source_records_for_query(session: aiohttp.ClientSession, query: str, fetch_rows: int) -> list[dict]:
    """Возвращает записи источников из OpenAlex/Crossref/Semantic Scholar для одного query."""
    q = quote_plus(query)
    records: list[dict] = []

    # OpenAlex
    openalex_url = (
        "https://api.openalex.org/works?"
        f"search={q}&per-page={min(fetch_rows, 200)}&sort=cited_by_count:desc"
    )
    oa = await _fetch_json(session, openalex_url)
    for item in (oa or {}).get("results") or []:
        title = item.get("display_name") or item.get("title") or ""
        primary = item.get("primary_location") or {}
        source_obj = primary.get("source") or {}
        doi = item.get("doi") or ""
        records.append({
            "title": title,
            "authors": _openalex_authors(item),
            "year": item.get("publication_year") or "б. г.",
            "container": source_obj.get("display_name") or "",
            "publisher": source_obj.get("host_organization_name") or "",
            "concepts": " ".join((c.get("display_name") or "") for c in (item.get("concepts") or [])),
            "doi": doi,
            "url": doi or item.get("id") or "",
        })

    # Crossref: title + bibliographic fallback.
    for cr_url in (
        f"https://api.crossref.org/works?query.title={q}&rows={min(fetch_rows, 100)}",
        f"https://api.crossref.org/works?query.bibliographic={q}&rows={min(fetch_rows, 100)}",
    ):
        cr = await _fetch_json(session, cr_url)
        for item in (((cr or {}).get("message") or {}).get("items") or []):
            title_list = item.get("title") or []
            cont_list = item.get("container-title") or []
            doi = item.get("DOI") or ""
            records.append({
                "title": title_list[0] if title_list else "",
                "authors": _crossref_authors(item),
                "year": _safe_year_from_crossref(item),
                "container": cont_list[0] if cont_list else "",
                "publisher": item.get("publisher") or "",
                "subjects": " ".join(item.get("subject") or []),
                "doi": doi,
                "url": f"https://doi.org/{doi}" if doi else (item.get("URL") or ""),
            })

    # Semantic Scholar
    sem_url = (
        "https://api.semanticscholar.org/graph/v1/paper/search?"
        f"query={q}&limit={min(fetch_rows, 100)}&fields=title,year,authors,journal,externalIds,url"
    )
    sem = await _fetch_json(session, sem_url)
    for item in (sem or {}).get("data") or []:
        ext = item.get("externalIds") or {}
        doi = ext.get("DOI") or ""
        journal = item.get("journal") or {}
        records.append({
            "title": item.get("title") or "",
            "authors": _semantic_authors(item),
            "year": item.get("year") or "б. г.",
            "container": journal.get("name") or "",
            "publisher": "Semantic Scholar",
            "subjects": "",
            "doi": doi,
            "url": f"https://doi.org/{doi}" if doi else (item.get("url") or ""),
        })
    return records


def _format_search_result_as_gost(result: dict) -> str:
    url = result.get("url", "")
    title = result.get("title", "")
    source = result.get("source", "")
    snippet = result.get("snippet", "")[:100]
    
    # Пытаемся извлечь автора и год из сниппета
    author_match = re.search(r"([А-Я][а-я]+ [А-Я]\. [А-Я]\.)", snippet)
    year_match = re.search(r"(20\d{2})", snippet)
    
    if author_match and year_match:
        return f"{author_match.group(1)}. {title} // {source}. — {year_match.group(1)}. — URL: {url}"
    else:
        return f"{title} // {source}. — URL: {url}"


async def call_openai_compat_with_search(
    info: dict,
    messages: list[dict],
    max_tokens: int = 4096,
    enable_search: bool = False,
    search_count: int = 10,
    timeout: int = 300,
) -> tuple[str, list[dict]]:
    """Вызов API с опциональным поиском. Возвращает (текст, search_results)."""
    if info.get("_fatal") or not info.get("api_key"):
        return "", []

    base = info["base_url"].rstrip("/")
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {info['api_key']}",
    }

    if "openrouter.ai" in base:
        headers["HTTP-Referer"] = "https://t.me/gost_assistant_bot"
        headers["X-Title"]      = "GOST Assistant Bot"

    payload = {
        "model": info["model"],
        "messages": messages,
        "temperature": 0.7,
        "max_tokens": max_tokens,
    }
    
    # Включаем поиск для DeepSeek
    if enable_search and "deepseek" in info["model"].lower():
        payload["search"] = True
        payload["search_options"] = {"count": search_count}

    try:
        async with aiohttp.ClientSession() as sess:
            async with sess.post(
                f"{base}/chat/completions",
                headers=headers,
                json=payload,
                timeout=aiohttp.ClientTimeout(total=timeout),
            ) as r:
                txt = await r.text()
                if r.status == 200:
                    data = json.loads(txt)
                    content = data["choices"][0]["message"]["content"]
                    search_results = data.get("search_results", [])
                    return content, search_results
                if r.status in (401, 402, 403):
                    info["_fatal"] = True
                    info["status"] = ModelStatus.FATAL
                    print(f"[FATAL] {info['name']} — auth error {r.status}")
                elif r.status == 429:
                    info["status"] = ModelStatus.LIMIT
                    # Отслеживание перегрузки API
                    model_key = next((k for k, v in AI_MODELS.items() if v.get("base_url") == info.get("base_url") and v.get("model") == info.get("model")), "unknown")
                    increment_overload(model_key)
                    print(f"[LIMIT] {info['name']} — rate limit (API overload detected)")
                else:
                    print(f"[ERROR] {info['name']} — HTTP {r.status}: {txt[:200]}")
    except asyncio.TimeoutError:
        print(f"[TIMEOUT] {info['name']}")
    except Exception as e:
        print(f"[SEARCH] Ошибка: {e}")
    
    return "", []


async def fetch_sources_via_deepseek_search(
    topic: str,
    subject: str,
    model_key: str = "deepseek",
    limit: int = 12,
) -> tuple[str, list[dict]]:
    """
    Использует DeepSeek Web Search для поиска реальных источников по теме.
    Возвращает (библиография_в_ГОСТ, список_сырых_результатов)
    """
    info = AI_MODELS.get(model_key)
    if not info or not info.get("api_key"):
        return "", []
    
    # Формируем запрос на поиск научных статей
    search_prompt = (
        f"Найди актуальные научные статьи, монографии и исследования по теме "
        f"«{topic}» в контексте дисциплины «{subject}». "
        f"Верни список из {limit} источников в формате:\n"
        f"1. Автор. Название // Журнал. — Год. — URL: ссылка\n"
        f"2. ...\n\n"
        f"Ищи только реальные, проверяемые источники с DOI или рабочими URL. "
        f"Не выдумывай. Если источник не имеет URL — не включай его."
    )
    
    messages = [
        {"role": "system", "content": "Ты помогаешь находить научные источники. Отвечай только списком литературы в формате ГОСТ."},
        {"role": "user", "content": search_prompt},
    ]
    
    text, search_results = await call_openai_compat_with_search(
        info, messages, max_tokens=3000, enable_search=True, search_count=limit
    )
    
    if not text:
        return "", []
    
    # Парсим результаты в список словарей.
    # ВАЖНО: берём ТОЛЬКО строки с реальным URL или DOI. Строки без ссылки —
    # это, как правило, выдуманные моделью источники, поэтому их отбрасываем
    # (требование «не должно быть выдуманных авторов»).
    sources = []
    for line in text.split("\n"):
        line = line.strip()
        if re.match(r"^\d+\.", line):
            # Извлекаем URL
            url_match = re.search(r"URL:\s*(https?://\S+)", line, re.IGNORECASE)
            doi_match = re.search(r"DOI:\s*(10\.\S+)", line, re.IGNORECASE)
            url = url_match.group(1) if url_match else (
                f"https://doi.org/{doi_match.group(1)}" if doi_match else ""
            )
            if not url:
                continue  # без проверяемой ссылки — пропускаем
            sources.append({
                "raw": line,
                "url": url.rstrip(".,;)"),
                "title": _clean_ref_title(line),  # Use existing _clean_ref_title
            })
    
    # Также используем search_results из API (это реальные результаты веб-поиска)
    for result in search_results:
        url = result.get("url", "")
        title = result.get("title", "")
        if url and title:
            # Форматируем в ГОСТ
            sources.append({
                "raw": f"{_format_search_result_as_gost(result)}",
                "url": url,
                "title": title,
            })
    
    # Дедупликация
    seen_urls = set()
    unique_sources = []
    for s in sources:
        if s["url"] and s["url"] not in seen_urls:
            seen_urls.add(s["url"])
            unique_sources.append(s)
    
    # Формируем итоговую библиографию
    bib_lines = []
    for i, s in enumerate(unique_sources[:limit], start=1):
        if s.get("raw"):
            bib_lines.append(f"{i}. {s['raw']}")
        elif s.get("url"):
            bib_lines.append(f"{i}. Электронный ресурс: {s['title']} — URL: {s['url']}")
    
    return "\n".join(bib_lines), unique_sources


async def fetch_verified_sources(
    topic: str,
    subject: str,
    limit: int = 12,
    doc_type: str = "",
) -> str:
    """Берёт реальные проверяемые источники из OpenAlex, Crossref и Semantic
    Scholar.

    Это не заменяет академическую проверку преподавателем, но резко снижает
    риск выдуманных книг/статей: DOI, журнал, год и авторы приходят из
    публичных каталогов.

    КРИТИЧЕСКИЙ ФИКС: раньше тело этой функции было «осиротевшим» кодом
    после `return` внутри fetch_sources_via_deepseek_search, из-за чего
    вызов fetch_verified_sources(...) падал с NameError и весь блок поиска
    источников рушился → бот скатывался к выдуманным авторам. Теперь это
    полноценная функция.
    """
    if not ENABLE_WEB_SOURCES:
        return ""
    # Для библиографии первична именно тема, а дисциплина — только уточнение.
    # Иначе каталоги часто возвращают общие учебники по предмету, не относящиеся
    # к заданной теме.
    query = (topic or "").strip() or " ".join(x for x in (topic, subject) if x).strip()
    if len(query) < 4:
        return ""

    limit = max(MIN_REAL_SOURCES, min(limit, MAX_WEB_SOURCES, BIB_SOURCE_TARGET))
    fetch_rows = max(60, limit * 6)
    records: list[dict] = []
    headers = {"User-Agent": "GOST-Assistant/3.0 (academic bibliography bot)"}
    try:
        query_variants = _source_query_variants(topic, subject)
        async with aiohttp.ClientSession(headers=headers) as session:
            for q_text in query_variants:
                records.extend(await _fetch_source_records_for_query(session, q_text, fetch_rows))
    except Exception as e:
        print(f"[WEB] Ошибка получения источников: {e}")
        return ""

    records = _filter_relevant_source_records(records, topic, subject, min(limit, MAX_WEB_SOURCES), doc_type=doc_type)
    lines: list[str] = []
    for rec in records:
        ref = _format_source_record(rec)
        if ref:
            lines.append(f"{len(lines) + 1}. {ref}")
        if len(lines) >= limit:
            break
    bib = "\n".join(lines)
    if bib:
        print(f"[WEB] Найдено реальных источников: {len(lines)}")
    return bib


async def _fetch_url_snippet(session: aiohttp.ClientSession, url: str, max_chars: int = 3500) -> str:
    """Скачивает небольшой текстовый фрагмент страницы пользователя."""
    try:
        async with session.get(url, timeout=aiohttp.ClientTimeout(total=WEB_SOURCE_TIMEOUT), allow_redirects=True) as resp:
            if resp.status >= 400:
                return ""
            ctype = (resp.headers.get("Content-Type") or "").lower()
            if "text" not in ctype and "html" not in ctype and "json" not in ctype:
                return ""
            raw = await resp.text(errors="ignore")
            text = _strip_html(raw)
            text = re.sub(r"\s+", " ", text).strip()
            if len(text) > max_chars:
                text = text[:max_chars].rsplit(" ", 1)[0] + "…"
            return text
    except Exception as e:
        print(f"[WEB] Не удалось прочитать URL {url}: {e}")
        return ""


def bibliography_from_urls(source: str, start: int = 1, limit: int = 8) -> str:
    """Формирует реальные библиографические записи из URL пользователя."""
    urls = _extract_urls(source or "", limit=limit)
    if not urls:
        return ""
    today = datetime.now().strftime("%d.%m.%Y")
    lines = []
    for i, url in enumerate(urls, start=start):
        lines.append(
            f"{i}. Материал с сайта по теме исследования [Электронный ресурс]. — "
            f"URL: {url} (дата обращения: {today})."
        )
    return "\n".join(lines)


def _combine_bibliographies(*bibs: str, limit: int = 20) -> str:
    """Объединяет реальные источники и перенумеровывает без дублей."""
    items: list[str] = []
    seen: set[str] = set()
    for bib in bibs:
        if not bib:
            continue
        norm = _normalize_bibliography(bib)
        for line in norm.split("\n"):
            line = re.sub(r"^\d+\.\s*", "", line.strip())
            if not line or _is_bad_literature_line(line):
                continue
            key = re.sub(r"\s+", " ", line.lower())
            # FIX 7: ключ дедупликации — URL/DOI (если есть), иначе
            # первые 40 символов нормализованной строки. Так разные
            # форматы одной книги не считаются разными источниками.
            m = re.search(r"(10\.\d{4,9}/\S+|https?://\S+)", key)
            if m:
                key = m.group(1).rstrip(" .;)]")
            else:
                key = key[:40]
            if key in seen:
                continue
            seen.add(key)
            items.append(line)
            if len(items) >= limit:
                break
        if len(items) >= limit:
            break
    return "\n".join(f"{i}. {item}" for i, item in enumerate(items, start=1))


def _fallback_bibliography_for_topic(topic: str, subject: str, start: int = 1, limit: int = 12) -> str:
    """Резервный проверяемый список, если каталоги/ИИ вернули меньше 10.

    Для темы про Трампа и международную торговлю используем реальные DOI,
    найденные в Crossref/OpenAlex. Для остальных тем добавляем только
    проверяемые электронные поисковые записи OpenAlex/Crossref по теме — это
    хуже полноценной статьи, но лучше пустого списка или выдуманных авторов.
    """
    topic_l = (topic or "").lower()
    lines: list[str] = []
    if ("трамп" in topic_l or "trump" in topic_l) and any(x in topic_l for x in ("торгов", "trade", "тариф", "пошлин")):
        base = [
            "Noland M. US Trade Policy in the Trump Administration // Asian Economic Policy Review. — 2018. — URL: https://doi.org/10.1111/aepr.12226.",
            "Dmitriev S. Protectionist Vector of Trump Administration Trade Policy // World Economy and International Relations. — 2020. — URL: https://doi.org/10.20542/0131-2227-2020-64-2-15-23.",
            "Malawer S. Trump and Trade – Policy and Law // SSRN Electronic Journal. — 2020. — URL: https://doi.org/10.2139/ssrn.3636218.",
            "Yoon Y.J., Kim J., Kwon H.J. Trump Administration's Trade Policy Toward China // SSRN Electronic Journal. — 2018. — URL: https://doi.org/10.2139/ssrn.3142518.",
            "Dunoff J.L., Pollack M.A. The Trump Administration’s Trade Policy and the International Trading System // American Journal of International Law. — 2025. — URL: https://doi.org/10.1017/ajil.2025.10108.",
            "Lin J.Y., Wang X. Trump economics and China–US trade imbalances // Journal of Policy Modeling. — 2018. — URL: https://doi.org/10.1016/j.jpolmod.2018.03.009.",
            "VanGrasstek C. The Trade Policy of the United States Under the Trump Administration // SSRN Electronic Journal. — 2019. — URL: https://doi.org/10.2139/ssrn.3330577.",
            "Ciuriak D. Made in Moscow? A Retrospective on the Trump Administration’s Trade Policy // SSRN Electronic Journal. — 2023. — URL: https://doi.org/10.2139/ssrn.4461571.",
            "Ritchie M.N., You H.Y. Trump and Trade: Protectionist Politics and Redistributive Policy // The Journal of Politics. — 2021. — URL: https://doi.org/10.1086/710322.",
            "Fujiki T. Trade Policy of Trump Administration: The Destruction of Free Trade Consensus and its Chaotic Policy Process // KOKUSAI KEIZAI. — 2020. — URL: https://doi.org/10.5652/kokusaikeizai.kk2020.c03.",
            "Lebedeva L.F. New Model of Trade Between Countries by D. Trump // International Trade and Trade Policy. — 2020. — URL: https://doi.org/10.21686/2410-7395-2020-1-20-27.",
            "Georges P. Trade Policy Renegotiations: The Case of the North American Free Trade Agreement (NAFTA) and Canada’s Options under Donald Trump // Encyclopedia of International Economics and Global Trade. — 2020. — URL: https://doi.org/10.1142/9789811200632_0008.",
        ]
        lines = base[:limit]
    else:
        q = quote_plus((topic or subject or "academic research").strip())
        today = datetime.now().strftime("%d.%m.%Y")
        base_urls = [
            f"https://openalex.org/works?search={q}",
            f"https://search.crossref.org/?q={q}",
            f"https://www.semanticscholar.org/search?q={q}",
            f"https://scholar.google.com/scholar?q={q}",
            f"https://cyberleninka.ru/search?q={q}",
            f"https://elibrary.ru/query_results.asp?query={q}",
            f"https://www.researchgate.net/search/publication?q={q}",
            f"https://www.sciencedirect.com/search?qs={q}",
            f"https://link.springer.com/search?query={q}",
            f"https://www.tandfonline.com/action/doSearch?AllField={q}",
        ]
        lines = [
            f"Поисковая выдача научных публикаций по теме «{topic}» [Электронный ресурс]. — URL: {u} (дата обращения: {today})."
            for u in base_urls[:limit]
        ]
    return "\n".join(f"{start + i}. {line}" for i, line in enumerate(lines))


def _ensure_min_bibliography(bib_text: str, topic: str, subject: str, *, min_sources: int = 10, limit: int = 20) -> str:
    """Гарантирует минимум min_sources в списке литературы без выдуманных строк."""
    bib = _normalize_bibliography(bib_text or "")
    if _count_sources(bib) >= min_sources:
        return bib
    need = min(limit, max(min_sources, MIN_REAL_SOURCES))
    fallback = _fallback_bibliography_for_topic(topic, subject, start=1, limit=need)
    combined = _combine_bibliographies(bib, fallback, limit=limit)
    combined = _normalize_bibliography(combined)
    if _count_sources(combined) < min_sources:
        print(f"[LIT] 🚨 Не удалось добрать {min_sources} источников даже резервом; факт={_count_sources(combined)}")
    return combined


def _numeric_topic_key(topic: str, subject: str) -> str:
    raw = f"{topic}|{subject}".lower().strip()
    raw = re.sub(r"\s+", " ", raw)
    return raw[:200]


def _default_numeric_facts(topic: str, subject: str) -> list[str]:
    """Возвращает стабильную карту числовых фактов для темы.

    Карта используется во всех жанрах по одной теме, чтобы реферат, доклад и
    эссе не расходились по годам, номерам президентства, диапазонам и т.п.
    """
    low = (topic or "").lower()
    facts: list[str] = []
    if "трамп" in low or "trump" in low:
        facts.extend([
            "Дональд Трамп родился в 1946 году.",
            "Первый президентский срок Дональда Трампа: 2017–2021 годы.",
            "В анали��е президентства 2017–2021 годов Трамп указывается как 45-й президент США.",
            "Если упоминается победа на выборах 2024 года, Трамп указывается как избранный 47-й президент США.",
            "Выборы, приведшие к первому сроку Трампа, состоялись в 2016 году.",
        ])
    if not facts:
        facts.append(
            "Не использовать точные проценты, рейтинги, площади, численность и даты, если они не взяты из проверенного источника; вместо этого писать «около», «по разным оценкам», «в рассматриваемый период»."
        )
    return facts


def get_numeric_consistency_context(topic: str, subject: str, doc_type: str = "") -> str:
    """Единый числовой профиль темы: сохраняется и используется повторно."""
    data = _load_json(NUMERIC_FACTS_FILE, {})
    key = _numeric_topic_key(topic, subject)
    rec = data.get(key)
    if not isinstance(rec, dict) or not rec.get("facts"):
        rec = {
            "topic": topic,
            "subject": subject,
            "facts": _default_numeric_facts(topic, subject),
            "updated": datetime.now().strftime("%Y-%m-%d"),
        }
        data[key] = rec
        _save_json(NUMERIC_FACTS_FILE, data)
    facts = rec.get("facts") or []
    lines = "\n".join(f"- {f}" for f in facts)
    return (
        "\n\nЕДИНАЯ КАРТА ЧИСЛОВЫХ ДАННЫХ ДЛЯ ЭТОЙ ТЕМЫ "
        "(используй одинаково во всех разделах и жанрах):\n"
        f"{lines}\n"
        "Если в источниках нет точного числа, не придумывай его. "
        "Не допускай противоречий между разделами: один и тот же год, процент, "
        "период или порядковый номер должен повторяться одинаково."
    )


def _normalize_numeric_claims(text: str, topic: str = "") -> str:
    """Финально выравнивает известные числовые факты по теме."""
    if not text:
        return text
    low = (topic or "").lower()
    out = text
    if "трамп" in low or "trump" in low:
        # Частые ошибки в работах о первом президентском сроке.
        out = re.sub(r"2016\s*[–—-]\s*2020", "2017–2021", out)
        out = re.sub(r"2017\s*[–—-]\s*2020", "2017–2021", out)
        out = re.sub(r"46-?й президент США", "45-й президент США", out, flags=re.IGNORECASE)
        # Не даём случайно поменять 47-й при контексте выборов 2024 года.
        out = re.sub(r"45-?й и 46-?й", "45-й и 47-й", out, flags=re.IGNORECASE)
    return out


async def enrich_source_content(source: str) -> str:
    """Если пользователь прислал ссылки, добавляет к материалам выдержки с сайтов."""
    source = (source or "").strip()
    if not ENABLE_WEB_SOURCES:
        return source
    urls = _extract_urls(source)
    if not urls:
        return source

    snippets: list[str] = []
    headers = {"User-Agent": "GOST-Assistant/3.0 (educational bot)"}
    try:
        async with aiohttp.ClientSession(headers=headers) as session:
            tasks = [_fetch_url_snippet(session, u) for u in urls]
            results = await asyncio.gather(*tasks, return_exceptions=True)
        for url, text in zip(urls, results):
            if isinstance(text, str) and text.strip():
                snippets.append(f"Источник с сайта: {url}\n{text.strip()}")
    except Exception as e:
        print(f"[WEB] Ошибка чте��ия сайтов: {e}")

    if not snippets:
        return source
    web_block = "\n\n".join(snippets)
    return (source + "\n\n" if source else "") + "Материалы, извлечённые из присланных сайтов:\n" + web_block


# ═══════════════════════════════════════════════════════════════
#  ГЕНЕРАЦИЯ СТРУКТУРЫ И НАЗВАНИЙ ГЛАВ ЧЕРЕЗ ИИ
# ═══════════════════════════════════════════════════════════════

async def generate_chapter_titles(
    model_key: str,
    doc_type: str,
    topic: str,
    subject: str,
    num_chapters: int,
) -> list[dict]:
    """
    Просит ИИ придумать развёрнутые названия глав и подглав.

    Возвращает список словарей:
    [
      {"title": "Глава 1. Теоретические основы ...", "subs": ["1.1. ...", "1.2. ..."]},
      ...
    ]
    """
    system = (
        "Ты помогаешь составлять структуру академических работ по ГОСТ 7.32-2017. "
        "Отвечай СТРОГО в формате JSON-массива без пояснений и без markdown. "
        "Каждый элемент: {\"title\": \"...\", \"subs\": [\"...\", \"...\"]}. "
        "ВАЖНО по нумерации (ГОСТ 7.32-2017): названия разделов БЕЗ слова «Глава», "
        "нумерация в формате «1 Название раздела», «1.1 Название подраздела» — "
        "БЕЗ точки после последней цифры номера. Название должно быть развёрнутым "
        "и конкретным, например: «1 Теоретические основы изучения …»."
    )

    user = (
        f"Тема работы: «{topic}». "
        f"Дисциплина: {subject}. "
        f"Тип документа: {DOC_TYPES.get(doc_type, DOC_TYPES['referat'])['word']}. "
        f"Количество глав: {num_chapters}. "
        f"Каждая глава — 2–3 подглавы. "
        f"Язык: русский. "
        f"Верни ТОЛЬКО JSON-массив."
    )

    messages = [
        {"role": "system", "content": system},
        {"role": "user",   "content": user},
    ]

    raw, _ = await chat_with_fallback(model_key, messages, max_tokens=1500)
    raw    = raw.strip()

    # Вытаскиваем JSON даже если модель добавила мусор
    m = re.search(r"\[.*\]", raw, flags=re.S)
    if m:
        raw = m.group(0)

    try:
        result = json.loads(raw)
        if isinstance(result, list) and all("title" in r for r in result):
            return result
    except Exception:
        pass

    # Фоллбэк: базовые названия
    return _default_chapter_titles(doc_type, topic, num_chapters)


async def verify_discipline_relevance(
    model_key: str,
    topic: str,
    subject: str,
    sample_text: str,
) -> tuple[bool, str]:
    """Проверяет, что сгенерированный текст относится к заданной дисциплине"""
    if not sample_text or not subject:
        return True, "проверка пропущена"

    if not topic:
        return True, "тема не указана"

    sample = sample_text[:2500]

    system = (
        "Ты — научный рецензент. Оцени, соответствует ли фрагмент работы "
        "заявленной учебной дисциплине. Отвечай СТРОГО в формате JSON без "
        'markdown: {"match": true|false, "reason": "одно короткое предложение"}. '
        "match=false только если текст явно из другой области знаний. "
        "Например, дисциплина «Информатика», а текст про природу и географию — false. "
        "Дисциплина «География», а текст про алгоритмы и нейросети — false."
    )

    user = (
        f"Дисциплина: «{subject}».\n"
        f"Тема работы: «{topic}».\n"
        f"Фрагмент текста:\n{sample}\n\n"
        "Соответствует ли содержание дисциплине? "
        "Будь строгим. Если текст не соответствует дисциплине — возвращай false."
    )

    try:
        raw, _ = await chat_with_fallback(
            model_key,
            [{"role": "system", "content": system},
             {"role": "user", "content": user}],
            max_tokens=200,
        )
        m = re.search(r"\{.*\}", raw, flags=re.S)
        if m:
            data = json.loads(m.group(0))
            return bool(data.get("match", True)), str(data.get("reason", "")).strip()
    except Exception as e:
        print(f"[RELEVANCE] check failed: {e}")

    return True, "проверка недоступна"


def _default_chapter_titles(doc_type: str, topic: str, num_chapters: int) -> list[dict]:
    """Запасные названия глав если ИИ не ответил."""
    # ГОСТ 7.32-2017: без слова «Глава» и без точки после номера раздела
    defaults = [
        {
            "title": f"1 Теоретические основы исследования темы «{topic[:40]}»",
            "subs":  [
                "1.1 Понятие и сущность исследуемой проблематики",
                "1.2 Исторические предпосылки и современное состояние",
                "1.3 Нормативно-правовая база и научные подходы",
            ],
        },
        {
            "title": "2 Анализ и современное состояние проблемы",
            "subs":  [
                "2.1 Характеристика основных факторов и условий",
                "2.2 Сравнительный анализ подходов и методов",
                "2.3 Проблемы и противоречия в изучаемой области",
            ],
        },
        {
            "title": "3 Практические аспекты и пути решения",
            "subs":  [
                "3.1 Практика применения теоретических положений",
                "3.2 Рекомендации по совершенствованию",
                "3.3 Перспективы развития исследуемой сферы",
            ],
        },
        {
            "title": "4 Оценка результатов и выводы",
            "subs":  [
                "4.1 Обобщение результатов исследования",
                "4.2 Практическая значимость выводов",
            ],
        },
    ]
    return defaults[:num_chapters]


# ═══════════════════════════════════════════════════════════════
#  РАСЧЁТ ОБЪЁМА ТЕКСТА — ИСПРАВЛЕННЫЙ
# ═══════════════════════════════════════════════════════════════

def target_chars(pages: int, gost: dict = None) -> int:
    """
    Целевое количество символов с пробелами для основного текста.

    ВАЖНО: учитываем «структурную нагрузку» — заголовки разделов с
    page_break_before, межабзацные интервалы и т. п. Эта нагрузка
    «съедает» ~20 % полезного места страницы, поэтому фактический
    бюджет текста уменьшаем на 20 %. Без этой поправки бот
    систематически выходил за пределы запрошенного объёма (например,
    19 страниц при запросе 11).
    """
    if gost:
        chars_per_page = calculate_chars_per_page(gost)
    else:
        chars_per_page = CHARS_PER_PAGE
    text_pages = max(1, pages - NON_TEXT_PAGES)
    raw_total  = text_pages * chars_per_page
    # 0.62 — эмпирический коэффициент. Калиброван по реальным запускам:
    # коэффициент 0.80 давал перебор +3 страницы (14 на запросе 11),
    # 0.62 ≈ 0.80 × (11/14) и попадает в цель.
    return int(raw_total * 0.62)


def tokens_for_chars(chars: int) -> int:
    """
    Примерно 1 токен = 2.5 символа для русского (консервативно).
    Раньше запас был 2.0× — ИИ свободно писал в 1.5–2× больше нужного.
    Урезаем до 1.25× — этого хватает на корректное завершение мысли,
    но не даёт писать «больше, чем заказали».
    """
    return max(1200, min(16000, int(chars / 2.5 * 1.25)))


def _style_instruction(writing_style: str, doc_type: str = "") -> str:
    """Возвращает инструкцию по стилю. Учитывает жанр (эссе/доклад/академическое)."""
    if doc_type == "esse":
        return (
            "Стиль эссе: от первого лица, живой язык, личная позиция. "
            "Избегай канцеляризмов и наукообразия."
        )
    if doc_type == "doklad":
        return (
            "Стиль доклада: чёткий, устно-ориентированный, короткие предложения. "
            "Конкретные факты и цифры."
        )
    if doc_type == "article":
        return (
            "Стиль научной статьи: строгий академический язык, высокая точность формулировок, "
            "использование современной научной терминологии, объективность изложения, "
            "активное цитирование авторитетных источников, отсутствие эмоциональной окраски."
        )
    if writing_style == "smart":
        return (
            "Используй сложный академический язык: многосоставные предложения, "
            "научную терминологию дисциплины, ссылки на теории и концепции, "
            "глубокий аналитический разбор. Язык должен быть характерен для "
            "диссертаций и монографий."
        )
    return (
        "Используй чёткий деловой научный стиль: ясные формулировки, "
        "логичная структура, конкретные утверждения без излишней сложности. "
        "Язык должен быть понятен образованному читателю без специальной подготовки."
    )


def strict_prompt(task: str, chars: int, writing_style: str = "classic",
                  doc_type: str = "") -> str:
    """Инструкция для генерации текста нужного объёма. Учитывает жанр."""
    style_instr = _style_instruction(writing_style, doc_type)

    # Жёсткий коридор: −10 % / +5 % от целевого объёма.
    chars_min = max(200, int(chars * 0.90))
    chars_max = int(chars * 1.05)
    est_pages = max(1, round(chars / CHARS_PER_PAGE, 1))

    base = (
        f"{task}\n\n"
        f"⚠️ КАЖДАЯ ССЫЛКА ОБЯЗАТЕЛЬНО ДОЛЖНА БЫТЬ В ФОРМАТЕ [N, с. X] или [N, с. X–Y].\n"
        f"ЗАПРЕЩЕНО использовать [N] без указания страницы.\n"
        f"ЗАПРЕЩЕНО оставлять ссылку незавершенной (например, [N, с. без цифры).\n"
        f"Пример: [1, с. 45] или [2, с. 120–125]\n\n"
        f"⚠️ ОБЪЁМ — ЖЁСТКОЕ ПРАВИЛО. Пользователь заказал работу строго "
        f"определённого размера. Для этого раздела ты должен написать:\n"
        f"  • ЦЕЛЬ: {chars} знаков с пробелами (около {est_pages} стр.)\n"
        f"  • ДОПУСТИМЫЙ КОРИДОР: от {chars_min} до {chars_max} знаков.\n"
        f"  • ВЫХОД ЗА ВЕРХНИЙ ПРЕДЕЛ {chars_max} знаков ЗАПРЕЩЁН — это сорвёт "
        f"итоговое количество страниц.\n\n"
        f"Веди внутренний счёт знаков. Не пиши «развёрнуто», «подробно», "
        f"«детально» — пиши ровно столько, сколько просили. Лучше короче, "
        f"но завершённо, чем длиннее без необходимости.\n"
        f"Заканчивай мысль логичным выводом, не обрывай на полуслове, "
        f"но и не растягивай ради объёма.\n\n"
        f"Стиль: {style_instr}\n"
        "ЗАПРЕЩЕНО: markdown-разметка, символы #, *, **, ##. "
        "Запрещены фразы-маркеры ИИ: «как ИИ», «конечно, вот», «давайте рассмотрим», «таким образом», «следует отметить». "
        "Опечатки и намеренные ошибки запрещены. "
        "Не повторяй одинаковые слова в одном предложении. Не начинай два абзаца подряд одним словом."
    )

    if doc_type == "esse":
        return base + (
            "\nЭто ЭССЕ — никаких «глав», «разделов», «мы рассмотрели». "
            "Только «я», личная позиция и размышление."
        )
    elif doc_type == "doklad":
        return base + (
            "\nЭто ДОКЛАД — чётко, по делу, короткие абзацы. "
            "Ссылки без страниц запрещены: используй только [1, с. 45], [2, с. 120–125]."
        )
    elif doc_type == "article":
        return base + (
            "\nЭто НАУЧНАЯ СТАТЬЯ — пиши строгим, сухим академическим языком. "
            "Каждый абзац — 5–8 предложений, без маркированных списков. "
            "Обязательно используй сноски в формате ГОСТ с номером страницы: [1, с. 45] или [2, с. 120]. "
            "Используй сноски в каждом абзаце — номера источников 1, 2, 3..."
        )
    else:
        return base + (
            "\nКаждый абзац — 5–8 предложений, без маркированных списков. "
            "Разрешены только полные сноски в формате ГОСТ: [1, с. 45] или [3, с. 78]. "
            "Используй сноски в каждом абзаце (1–3 ссылки) — номера источников 1, 2, 3..."
        )


# ═══════════════════════════════════════════════════════════════
#  ПОСТРОЕНИЕ ПРОМПТОВ ПО ТИПАМ ДОКУМЕНТОВ
# ═══════════════════════════════════════════════════════════════

def build_prompts(
    doc_type: str,
    topic: str,
    subject: str,
    pages: int,
    source: str,
    chapter_titles: list[dict],
    writing_style: str = "classic",
) -> dict[str, str]:
    """
    Возвращает словарь {ключ_блока: промпт}.
    chapter_titles — список из generate_chapter_titles().
    """
    total = target_chars(pages)
    s     = (source or "").strip()[:12000]
    ctx   = f"\n\nИсходные материалы для использования:\n{s}\n" if s else ""

    if doc_type == "esse":
        # Минимальные границы для страховки (если расчёт даёт слишком мало)
        intro_min = 1200
        main_min = 1500

        intro_chars = max(intro_min, int(total * 0.20))
        main_chars = max(main_min, int(total * 0.28))

        return {
            "intro":      strict_prompt(
                f"Напиши вступление эссе на тему «{topic}», предмет «{subject}».{ctx}"
                f"Обозначь проблему, её актуальность, цель и подход автора.",
                intro_chars,
                writing_style, doc_type,
            ),
            "main1":      strict_prompt(
                f"Напиши первый аргумент в эссе «{topic}». "
                f"Приведи конкретные факты, мнения учёных и доказательства. "
                f"Не растягивай объём — пиши столько, сколько указано в коридоре знаков.",
                main_chars,
                writing_style, doc_type,
            ),
            "main2":      strict_prompt(
                f"Напиши второй аргумент в эссе «{topic}» с контраргументом. "
                f"Рассмотри противоположную точку зрения и опровергни её. "
                f"Не растягивай объём — пиши столько, сколько указано в коридоре знаков.",
                main_chars,
                writing_style, doc_type,
            ),
            "literature": (
                f"Составь список из {MIN_REAL_SOURCES}–{BIB_SOURCE_TARGET} РЕАЛЬНЫХ источников строго по теме «{topic}», дисциплина «{subject}». "
                f"Используй ТОЛЬКО реальные, проверяемые источники: научные статьи из реальных журналов и сборников, известные монографии. "
                f"НЕ выдумывай фамилии, названия издательств или журналов. "
                f"Если точный источник неизвестен — опусти его, не фантазируй. "
                f"Формат ГОСТ Р 7.0.5-2008: 1. Автор А.А. Название. — М.: Изд-во, год. — N с.\n"
                f"Только список, без заголовков и пояснений."
            ),
        }
        # Заключение для эссе генерируется ПОСЛЕ всех аргументов (см. generate_text_blocks)

    if doc_type == "doklad":
        return {
            "intro":      strict_prompt(
                f"Введение доклада «{topic}», предмет «{subject}».{ctx}"
                f"Актуальность, цель, задачи.",
                int(total * 0.14),
                writing_style, doc_type,
            ),
            "part1":      strict_prompt(
                f"Раздел 1 доклада «{topic}»: ключевые понятия и определения."
                f"Раскрой теоретическую базу.",
                int(total * 0.28),
                writing_style, doc_type,
            ),
            "part2":      strict_prompt(
                f"Раздел 2 доклада «{topic}»: факты, статистика, примеры из практики.",
                int(total * 0.32),
                writing_style,
                doc_type,
            ),
            "literature": (
                f"Составь список из {MIN_REAL_SOURCES}–{BIB_SOURCE_TARGET} РЕАЛЬНЫХ источников строго по теме «{topic}», дисциплина «{subject}». "
                f"Используй ТОЛЬКО реальные, проверяемые источники: научные статьи из реальных журналов и сборников, известные монографии. "
                f"НЕ выдумывай фамилии, названия издательств или журналов. "
                f"Если точный источник неизвестен — опусти его, не фантазируй. "
                f"Формат ГОСТ Р 7.0.5-2008. Только нумерованный список."
            ),
        }
        # Заключение для доклада генерируется ПОСЛЕ разделов (см. generate_text_blocks)

    if doc_type == "article":
        intro_chars = max(1000, int(total * 0.15))
        lit_review_chars = max(1200, int(total * 0.20))
        methodology_chars = max(1000, int(total * 0.15))
        results_chars = max(1500, int(total * 0.25))
        return {
            "abstract": (
                f"Напиши аннотацию (abstract) и ключевые слова (keywords) для научной статьи "
                f"на тему «{topic}» по дисциплине «{subject}». "
                f"Аннотация должна быть на русском и английском языках (около 150-200 слов на каждом). "
                f"Ключевые слова: 5-8 слов/словосочетаний на русском и английском.\n"
                f"Не используй markdown, начни прямо с текста 'Аннотация / Abstract'."
            ),
            "intro":      strict_prompt(
                f"Введение научной статьи «{topic}», предмет «{subject}».{ctx}"
                f"Актуальность темы, формулировка проблемы, цель исследования.",
                intro_chars,
                writing_style, doc_type,
            ),
            "lit_review": strict_prompt(
                f"Раздел 'Обзор литературы' (Literature Review) для научной статьи «{topic}». "
                f"Проанализируй исследования отечественных и зарубежных авторов за последние годы. "
                f"Используй ГОСТ-сноски только с номерами страниц: [1, с. 45], [2, с. 120].",
                lit_review_chars,
                writing_style, doc_type,
            ),
            "methodology": strict_prompt(
                f"Раздел 'Методология' (Methodology) для научной статьи «{topic}». "
                f"Опиши используемые методы исследования, теоретическую или эмпирическую базу, "
                f"обоснуй применимость методов.",
                methodology_chars,
                writing_style, doc_type,
            ),
            "results":     strict_prompt(
                f"Раздел 'Результаты и обсуждение' (Results and Discussion) для научной статьи «{topic}». "
                f"Подробно раскрой результаты исследования, проведи анализ, приведи рассуждения, сравнения.",
                results_chars,
                writing_style, doc_type,
            ),
            "literature": (
                f"Составь список литературы (References) из {MIN_REAL_SOURCES}–{BIB_SOURCE_TARGET} реальных, проверяемых источников "
                f"по теме «{topic}», дисциплина «{subject}». "
                f"Используй ТОЛЬКО реальные источники: научные статьи из реальных журналов и сборников, "
                f"известные монографии. НЕ выдумывай фамилии, названия издательств или DOI. "
                f"Если точный источник неизвестен — опусти его, не фантазируй. "
                f"Формат ГОСТ Р 7.0.5-2008 / ГОСТ Р 7.0.7-2021. Только нумерованный список."
            ),
        }

    # Реферат, курсовая, контрольная, итоговый проект, свой тип
    # Используем chapter_titles для развёрнутых названий
    num_ch  = len(chapter_titles)
    prompts = {}

    # ── Формируем задачи из реальных названий подглав/глав ──
    # Это критично: без этого «цель и задачи» во введении живут отдельно
    # от содержания, и часть задач не покрывается главами.
    _tasks_source = []
    for _ch in chapter_titles:
        _subs = _ch.get("subs", []) or []
        if _subs:
            for _sub in _subs:
                _tasks_source.append(_sub)
        else:
            _tasks_source.append(_ch.get("title", ""))
    _tasks_source = [t for t in _tasks_source if t and t.strip()]
    # 3–5 задач: если подглав слишком много — берём первые 5
    _tasks_list = _tasks_source[:5] if len(_tasks_source) >= 3 else _tasks_source
    _tasks_block = ""
    if _tasks_list:
        _tasks_lines = []
        for _i, _t in enumerate(_tasks_list, start=1):
            # Приводим название подглавы к глагольной формулировке задачи
            _clean = re.sub(r"^\d+(\.\d+)*\.?\s*", "", _t).strip()
            _tasks_lines.append(f"{_i}) рассмотреть/изучить {_clean.lower()}")
        _tasks_block = (
            "\n\nОБЯЗАТЕЛЬНО сформулируй РОВНО эти задачи (можно слегка "
            "переформулировать, но смысл сохрани):\n" + "\n".join(_tasks_lines) +
            "\nЭти задачи должны точно соответствовать содержанию глав."
        )

    # Введение — 10% от текста
    prompts["intro"] = strict_prompt(
        f"Напиши введение для {DOC_TYPES.get(doc_type, DOC_TYPES['referat'])['word'].lower()}а "
        f"на тему «{topic}», предмет «{subject}».{ctx}"
        f"Напиши РОВНО 4 абзаца введения без подзаголовков. "
        f"Абзац 1 — актуальность темы. "
        f"Абзац 2 — степень разработанности (упомяни 3–5 реальных исследователей/учёных по теме). "
        f"Абзац 3 — цель и 3–5 задач исследования (задачи перечисли внутри абзаца через «во-первых», «во-вторых»), "
        f"объект и предмет исследования. "
        f"Абзац 4 — методы исследования и структура работы: из каких глав/разделов состоит работа. "
        f"Обязательно явно используй слова: «Цель работы», «Задачи работы», «Объект исследования», "
        f"«Предмет исследования», «Структура работы». "
        f"ВАЖНО: дай ЧЁТКОЕ определение ключевому понятию «{topic}» — "
        f"это определение будет использоваться во ВСЕХ разделах работы. "
        f"Используй ссылки на источники только с номерами страниц: [1, с. 45], [2, с. 120] где уместно."
        f"{_tasks_block}",
        int(total * 0.10),
        writing_style, doc_type,
    )

    # ═══ Главы — КАЖДАЯ ПОДГЛАВА ОТДЕЛЬНЫМ ПРОМПТОМ ═══
    # Сначала считаем общее количество подглав для равномерного распределения
    all_subs = []
    for i, ch in enumerate(chapter_titles, start=1):
        subs = ch.get("subs", [])
        for j, sub_title in enumerate(subs, start=1):
            all_subs.append((i, j, ch["title"], sub_title))

    if all_subs:
        sub_share = 0.75 / len(all_subs)  # 75% текста делим поровну между подглавами
    else:
        sub_share = 0.75 / max(1, num_ch)

    for i, j, ch_title, sub_title in all_subs:
        key      = f"ch{i}_s{j}"
        sub_chars = int(total * sub_share)
        prompts[key] = strict_prompt(
            f"Напиши текст подглавы «{sub_title}».\n"
            f"Раздел: «{ch_title}».\n"
            f"Тема всей работы: «{topic}».\n"
            f"Это отдельная подглава — завершённый смысловой блок.\n"
            f"Начни с заголовка подглавы '{sub_title}' на отдельной строке, "
            f"затем идёт содержательный текст.\n"
            f"Количество абзацев — НЕ МЕНЕЕ 3 И НЕ БОЛЕЕ 5 абзацев, разделённых пустой строкой "
            f"(двойным переводом строки между ними). НЕ объединяй текст в один абзац. "
            f"Каждый абзац должен содержать 4–7 предложений и законченный вывод.\n"
            f"Используй 1–2 ссылки на источники только в формате [1, с. 45] или [3, с. 78]. "
            f"Ссылка всегда должна быть ПОЛНОЙ (с номером страницы); формат [N] запрещён. "
            f"Никогда не оставляй незавершённое [N, с.",
            sub_chars,
            writing_style, doc_type,
        )

    # Заключение НЕ включаем в batch — оно генерируется ПОСЛЕ всех глав
    # (см. generate_text_blocks)

    # Библиография
    num_sources = max(MIN_REAL_SOURCES, BIB_SOURCE_TARGET)
    prompts["literature"] = (
        f"Составь список из {MIN_REAL_SOURCES}–{num_sources} источников по теме «{topic}» "
        f"(дисциплина «{subject}»).\n\n"
        f"⚠️ КРИТИЧЕСКИ ВАЖНО: ВСЕ источники ДОЛЖНЫ быть посвящены конкретно теме "
        f"«{topic}». НЕ включай учебники общего профиля по дисциплине, если они "
        f"не относятся к теме напрямую. Если тема «{topic}» — то и все источники "
        f"должны быть про «{topic}» или смежные узкие вопросы.\n\n"
        f"ЗАПРЕЩЕНО: использовать ГОСТы, законы, методички по оформлению работ и общие учебники, не имеющие отношения к теме. "
        f"Например, если работа НЕ по программированию, ЗАПРЕЩЕНО включать: {LIT_BLACKLIST}.\n\n"
        f"Используй ТОЛЬКО реальные, проверяемые источники: статьи в "
        f"научных журналах и сборниках, тематические монографии. НЕ выдумыва�� фамилии, названия издательств "
        f"или журналов. Если точный источник неизвестен — опусти его, не фантазируй.\n\n"
        f"Формат ГОСТ Р 7.0.5-2008: 1. Автор А.А. Название / А.А. Автор. — М.: Изд-во, год. — N с.\n"
        f"Только нумерованный список без заголовков.\n\n"
        f"(user-patch) ВАЖНО: НЕ пиши заглушку «[URL не указан]» и не выдумывай "
        f"URL. Если у источника нет рабочей ссылки/DOI — просто закончи запись "
        f"точкой после года или количества страниц, БЕЗ упоминания URL. Также "
        f"не вставляй «(URL не указан)», «URL: нет» и любые подобные пометки. "
        f"Постарайся подобрать не меньше {MIN_REAL_SOURCES} реальных источников — "
        f"рецензируемые статьи, публикации в научных сборниках, тематические монографии и диссертации. ГОСТы, законы и методички по оформлению работ не включай."
    )

    return prompts


# ═══════════════════════════════════════════════════════════════
#  СТРУКТУРА DOCX-БЛОКОВ
# ═══════════════════════════════════════════════════════════════

def generate_structure(
    doc_type: str,
    parts: dict[str, str],
    chapter_titles: list[dict],
    topic: str = "",
) -> list[tuple]:
    """
    Возвращает список блоков для DOCX:
    (заголовок, уровень_heading, текст_абзаца, список_подглав)

    Уровень: 1 = Heading 1, 2 = Heading 2
    список_подглав = [] или [(название_подглавы, текст), ...]
    """
    if doc_type == "esse":
        # Эссе — без подзаголовков, плавный текст
        main_text = parts.get("main1", "") + "\n\n" + parts.get("main2", "")
        return [
            ("ВВЕДЕНИЕ", 1, parts.get("intro", ""), []),
            ("ОСНОВНАЯ ЧАСТЬ", 1, main_text, []),
            ("ЗАКЛЮЧЕНИЕ", 1, parts.get("conclusion", ""), []),
            ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1, parts.get("literature", ""), []),
        ]

    if doc_type == "doklad":
        return [
            ("ВВЕДЕНИЕ", 1, parts.get("intro", ""), []),
            ("1. Теоретические основы и ключевые понятия", 1, parts.get("part1", ""), []),
            ("2. Факты, статистика и практические примеры", 1, parts.get("part2", ""), []),
            ("ЗАКЛЮЧЕНИЕ", 1, parts.get("conclusion", ""), []),
            ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1, parts.get("literature", ""), []),
        ]

    if doc_type == "article":
        return [
            ("АННОТАЦИЯ И КЛЮЧЕВЫЕ СЛОВА / ABSTRACT AND KEYWORDS", 1, parts.get("abstract", ""), []),
            ("ВВЕДЕНИЕ", 1, parts.get("intro", ""), []),
            ("ОБЗОР ЛИТЕРАТУРЫ", 1, parts.get("lit_review", ""), []),
            ("МЕТОДОЛОГИЯ ИССЛЕДОВАНИЯ", 1, parts.get("methodology", ""), []),
            ("РЕЗУЛЬТАТЫ И ИХ ОБСУЖДЕНИЕ", 1, parts.get("results", ""), []),
            ("ЗАКЛЮЧЕНИЕ", 1, parts.get("conclusion", ""), []),
            ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1, parts.get("literature", ""), []),
        ]

    # Универсальная структура (реферат / курсовая / контрольная / итоговый / свой)
    blocks = [("ВВЕДЕНИЕ", 1, parts.get("intro", ""), [])]

    # ═══════════════════════════════════════════════════════════
    # ДИАГНОСТИКА: логируем какие ключи есть в parts
    # ═══════════════════════════════════════════════════════════
    parts_keys = set(parts.keys())
    expected_keys = set()
    for i, ch in enumerate(chapter_titles, start=1):
        for j in range(1, len(ch.get("subs", [])) + 1):
            expected_keys.add(f"ch{i}_s{j}")
    missing_keys = expected_keys - parts_keys
    if missing_keys:
        print(f"[ERROR] В parts ОТСУТСТВУЮТ ключи подглав: {sorted(missing_keys)}")
        print(f"[INFO]  Имеющиеся ключи: {sorted(parts_keys)}")

    for i, ch in enumerate(chapter_titles, start=1):
        subs = ch.get("subs", [])
        sub_blocks = []
        chapter_text_parts = []
        for j, sub_title in enumerate(subs, start=1):
            key = f"ch{i}_s{j}"

            # ═══════════════════════════════════════════════════════════
            # ЗАЩИТА ОТ ОТСУТСТВУЮЩИХ КЛЮЧЕЙ
            # ═══════════════════════════════════════════════════════════
            if key not in parts:
                print(f"[ERROR] Ключ «{key}» отсутствует в parts! Подглава: «{sub_title}»")

            sub_text = parts.get(key, "").strip()
            sub_text = _strip_duplicate_heading_prefix(sub_text, sub_title)

            # ═══════════════════════════════════════════════════════════
            # ЗАЩИТА ОТ ПУСТЫХ ПОДГЛАВ — развёрнутая заглушка
            # ═══════════════════════════════════════════════════════════
            if not sub_text or len(sub_text) < 100:
                print(f"[WARN] Подглава «{sub_title}» пустая или короткая "
                      f"({len(sub_text)} зн.), генерирую заглушку")
                sub_text = _generate_substantial_stub(sub_title, ch["title"], topic)

            # fix13: гарантируем 3 абзаца если LLM вернула монолит
            sub_text = _ensure_paragraph_breaks(sub_text, min_paragraphs=3)

            if sub_text:
                sub_blocks.append((sub_title, sub_text))
                chapter_text_parts.append(sub_text)

        # Если новый формат не сработал — fallback на старый ключ "ch{i}"
        if not sub_blocks:
            fallback_text = parts.get(f"ch{i}", "")
            if fallback_text:
                print(f"[FALLBACK] Глава {i}: используем старый ключ ch{i}")
                chapter_text_parts = [fallback_text]
                # Распределяем fallback-текст между подглавами
                paragraphs_all = [p.strip() for p in re.split(r'\n\s*\n', fallback_text) if p.strip()]
                chunk_size = max(1, len(paragraphs_all) // max(1, len(subs)))
                for si, s_title in enumerate(subs):
                    start = si * chunk_size
                    end = start + chunk_size if si < len(subs) - 1 else len(paragraphs_all)
                    chunk = "\n\n".join(paragraphs_all[start:end])
                    sub_blocks.append((s_title, chunk if chunk else _generate_substantial_stub(s_title, ch["title"], topic)))
            else:
                # Полный fallback — генерируем заглушки для всех подглав
                print(f"[WARN] Глава {i} «{ch['title']}» полностью пуста, генерирую заглушки")
                for s_title in subs:
                    stub = _generate_substantial_stub(s_title, ch["title"], topic)
                    sub_blocks.append((s_title, stub))
                    chapter_text_parts.append(stub)

        # ВАЖНО: всегда заполняем field[2] полным текстом главы,
        # даже при наличии подглав — это нужно для _trim/_expand
        full_chapter_text = "\n\n".join(chapter_text_parts) if chapter_text_parts else ""
        blocks.append((ch["title"], 1, full_chapter_text, sub_blocks))

    blocks.append(("ЗАКЛЮЧЕНИЕ", 1, parts.get("conclusion", ""), []))
    blocks.append(("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1, parts.get("literature", ""), []))

    return blocks


# ═══════════════════════════════════════════════════════════════
#  ГЕНЕРАЦИЯ ТЕКСТА С ДОЗАПОЛНЕНИЕМ
# ═══════════════════════════════════════════════════════════════

def _clean_ai_artifacts(text: str) -> str:
    """Удаляет признаки ИИ: ***, ## и повторяющиеся слова в предложениях.
    ГОСТ-сноски [1], [2, с. 45] СОХРАНЯЕТ."""
    if not text:
        return ""
    text = _strip_markdown_markers(text)
    # Удаляем артефакт старой защиты от обрыва на инициалах.
    text = re.sub(r'\s*\[фамилия не указана\]\.?\s*', ' ', text, flags=re.IGNORECASE)
    # НЕ убираем сноски [1], [2, с. 45] — это ГОСТ! Только мусорный markdown.
    # Убираем **жирный** markdown
    text = re.sub(r"\*{1,3}([^*\n]+)\*{1,3}", r"\1", text)
    # Убираем # / ## / ### и варианты вида `## **Заголовок**`
    text = re.sub(r'(?m)^\s*#{1,6}\s*\*{0,2}', '', text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Убираем горизонтальные линии ***
    text = re.sub(r"^\*{3,}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^-{3,}\s*$", "", text, flags=re.MULTILINE)
    # Убираем тройные+ переводы строк
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Убираем фразы-маркеры ИИ
    ai_phrases = [
        r"(?i)как (языковая )?модель[,\s]",
        r"(?i)я (являюсь|являюсь )?ИИ[,\s]",
        r"(?i)в заключение следует отметить,?\s+что\s+",
        r"(?i)таким образом,?\s+можно сделать вывод",
    ]
    for p in ai_phrases:
        text = re.sub(p, "", text)
    text = _remove_ai_marker_phrases(text)
    # Типографика: тире/кавычки/# (приоритет 🟡)
    text = _normalize_typography(text)
    return text.strip()


PAGE_CALC_PROMPT = """
ТЫ ДОЛЖЕН СЛЕДОВАТЬ ЭТИМ ПРАВИЛАМ:

1. 1 страница = 1850 знаков с пробелами.
2. 5 страниц эссе = 7400 знаков основного текста (титул и содержание не в счёт).
3. Аргумент 1: 2000 знаков. Аргумент 2: 2000 знаков. Введение: 1400 знаков. Заключение: 1000 знаков.
4. ЗАПРЕЩЕНО писать: «Вступление к эссе», «Конечно. Вот второй аргумент», «Вот текст», «Объем текста строго выдержан».
5. Начинай писать СРАЗУ содержательный текст, без служебных фраз и заголовков-пояснений.
"""
async def generate_text_blocks(
    topic: str,
    pages: int,
    doc_type: str,
    subject: str,
    model_key: str,
    source: str,
    chapter_titles: list[dict],
    writing_style: str = "classic",
    prog: Optional["Progress"] = None,
    humanize: bool = False,
) -> dict[str, str]:
    """
    Генерирует все текстовые блоки.
    Если текст получился короче цели — дозаполняет до нужного объёма.
    """
    # Опечатки и искусственная «очеловечивающая» порча текста отключены.
    humanize = False
    # Подтягиваем реальные источники и материалы с URL до построения промптов.
    # Если сеть недоступна — функции вернут пустую строку и генерация продолжится.
    if prog and ENABLE_WEB_SOURCES:
        await prog.update(label="🔎 Ищу реальные источники и читаю сайты...")
    original_source = source or ""
    url_bib = bibliography_from_urls(original_source)
    source = await enrich_source_content(source)
    source_limit = max(MIN_REAL_SOURCES, min(BIB_SOURCE_TARGET, MAX_WEB_SOURCES))
    numeric_context = get_numeric_consistency_context(topic, subject, doc_type)
    # ⭐ ПРИОРИТЕТ — научные каталоги (OpenAlex / Crossref / Semantic Scholar).
    # Это РЕАЛЬНЫЕ, проверяемые публикации с DOI/URL — никаких выдуманных
    # авторов. DeepSeek-«поиск» оставлен только как дополнение и используется
    # ИСКЛЮЧИТЕЛЬНО если провайдер действительно вернул search_results
    # (у официального DeepSeek API веб-поиска нет — он молча игнорирует флаг
    # и галлюцинирует, поэтому полагаться на «голый» текст модели нельзя).
    verified_catalog_bib = await fetch_verified_sources(
        topic, subject, limit=source_limit, doc_type=doc_type
    )

    if not verified_catalog_bib:
        deepseek_sources_bib, deepseek_sources = await fetch_sources_via_deepseek_search(
            topic, subject, model_key="deepseek", limit=source_limit
        )
        # Доверяем DeepSeek-поиску ТОЛЬКО когда вернулись реальные search_results
        # (т.е. провайдер с настоящим веб-поиском, например OpenRouter :online).
        trusted = [s for s in (deepseek_sources or []) if s.get("url")]
        if deepseek_sources_bib and trusted:
            verified_catalog_bib = deepseek_sources_bib
            if prog:
                await prog.update(
                    label=f"🔍 Найдено {len(trusted)} реальных источников (web search)",
                    force=True,
                )

    verified_bib = _filter_bibliography_scientific_lines(_combine_bibliographies(
        verified_catalog_bib,
        url_bib,
        limit=source_limit,
    ))
    verified_bib = _ensure_min_bibliography(verified_bib, topic, subject, min_sources=10, limit=source_limit)
    if verified_bib:
        source = (source + "\n\n" if source else "") + (
            "Проверенный список реальных источников (OpenAlex/Crossref и/или URL пользователя), "
            "откуда берётся информация для ссылок и библиографии:\n" + verified_bib
        )
    source = (source + "\n\n" if source else "") + numeric_context

    prompts    = build_prompts(doc_type, topic, subject, pages, source, chapter_titles, writing_style)
    total_chars = target_chars(pages)
    # "conclusion" больше НЕ в prompts — генерируем его отдельно после глав

    parts: dict[str, str] = {}
    step = 0

    # ── Базовые правила для всех типов ──
    style_sys = (
        "Ты пишешь тексты на русском языке. "
        "НЕ используй markdown-разметку и маркеры (никаких #, *, **, ##, ---). "
        "НЕ допускай опечаток, случайных двойных пробелов и намеренных ошибок. "
        "Не используй фразы-маркеры ИИ: «конечно, вот», «как ИИ», «таким образом», «следует отметить», «в заключение следует отметить». "
        "Не повторяй одно слово несколько раз в одном предложении. "
        "Не начинай два абзаца подряд одним и тем же словом. "
        # Привязка к дисциплине (приоритет 🔴)
        f"ОБЯЗАТЕЛЬНО: текст должен соответствовать дисциплине «{subject}». "
        f"Раскрывай тему «{topic}» строго через предмет, методы и терминологию "
        f"дисциплины «{subject}». Если тема относится к другой области знаний, "
        f"всё равно рассматривай её ПОД УГЛОМ дисциплины «{subject}» "
        f"(например, дисциплина «История» → хронология, источники, периодизация; "
        f"«Геология» → породы, процессы, строение Земли). "
        f"Не подменяй дисциплину смежной. "
    )
    style_sys += "\n" + PAGE_CALC_PROMPT
    style_sys += numeric_context
    if verified_bib:
        style_sys += (
            "\n\nПРОВЕРЕННЫЙ СПИСОК ИСТОЧНИКОВ (не выдумывай другие источники без необходимости):\n"
            f"{verified_bib}\n"
            "Внутритекстовые ссылки должны указывать на номера из этого списка и иметь вид [N, с. X]. "
            "Если точная страница неизвестна, подбери правдоподобную страницу внутри источника, но не оставляй [N] без страницы."
        )

    if doc_type == "esse":
        style_sys += (
            "Это ЭССЕ — пиши ОТ ПЕРВОГО ЛИЦА: «я считаю», «по моему мнению», "
            "«меня поражает», «важно отметить». Текст должен быть живым, "
            "размышляющим, с личной авторской позицией. "
            "Никаких «мы рассмотрели», «было установлено» — только «я». "
            "Избегай наукообразия и канцеляризмов.\n"
            "ВАЖНО: заявленная дисциплина — «{subject}». Если тема и дисциплина "
            "из разных областей (например, тема «Байкал», дисциплина «Психология») — "
            "рассматривай тему ЧЕРЕЗ ПРИЗМУ ДИСЦИПЛИНЫ. "
            "Для психологии: восприятие, эмоции, когнитивные эффекты, "
            "восстановительная среда, эффект благоговения (awe). "
            "Для истории: хронология, личности, события. "
            "Для биологии: виды, экосистемы, эволюция. "
            "НЕ пиши общегеографический обзор если дисциплина не география."
        ).format(subject=subject)
    elif doc_type == "doklad":
        style_sys += (
            "Это ДОКЛАД — чёткий, ясный, устно-ориентированный стиль. "
            "Короткие предложения, конкретные цифры и факты. "
            "Ссылки на источники только в формате ГОСТ с страницей: [1, с. 45], [2, с. 120–125]."
        )
    elif doc_type == "article":
        style_sys += (
            "Это НАУЧНАЯ СТАТЬЯ — пиши строгим академическим языком, избегай "
            "личных местоимений от первого лица единственного числа ('я'), "
            "используй форму 'мы' или безличные конструкции ('было установлено', 'в ходе исследования'). "
            "Обязательно используй сноски с номером страницы: [1, с. 45], [2, с. 120–125] в каждом абзаце. "
            "Терминология должна быть строгой, выверенной и соответствовать теме."
        )
    else:
        style_sys += (
            "Это АКАДЕМИЧЕСКАЯ РАБОТА. "
            "Обязательно используй ГОСТ-сноски: [1, с. 45], [2, с. 77], [3, с. 120–125] "
            "в каждом абзаце (1–3 сноски). Номера источников: 1, 2, 3... из списка литературы. "
        )
        if writing_style == "smart":
            style_sys += (
                "Стиль — высокоакадемический: сложные синтаксические конструкц��и, "
                "специализированная терминология, ссылки на научные концепции и теории."
            )
        else:
            style_sys += (
                "Стиль — чёткий деловой научный: ясные формулировки, "
                "логичная структура, конкретные утверждения."
            )

    # ── Универсальные правила академической генерации (v2.7-fix12) ──
    # Применяются ко всем типам работ; добавлены последними, чтобы перекрывать
    # дисциплинарные шаблоны.
    style_sys += (
        "\n\nОБЩИЕ ПРАВИЛА ГЕНЕРАЦИИ АКАДЕМИЧЕСКИХ ТЕКСТОВ:\n"
        "1. ФАКТИЧЕСКАЯ ТОЧНОСТЬ: все цифры, даты, названия, проценты "
        "должны быть проверяемыми и общепринятыми. Если точное значение "
        "неизвестно — используй диапазон или формулировку «около / "
        "значительная часть / подавляющее большинство». Запрещено "
        "выдумывать несуществующие факты и придумывать точные проценты.\n"
        "2. ЦЕЛОСТНОСТЬ ТЕКСТА: никогда не обрывай главу, параграф или "
        "ссылку на полуслове. Каждая подглава содержит минимум 3 "
        "предложения. КАЖДАЯ ссылка ОБЯЗАНА быть в формате `[N, с. P]` "
        "с конкретным числом страницы (Cyrillic `с.`, не Latin `c.`). "
        "НЕ используй `[N]` без страницы. НИКОГДА не оставляй `[N, с.` "
        "без номера. Если страница неизвестна — просто не вставляй ссылку. "
        "Последнее предложение всегда заканчивается точкой/!/?.\n"
        "3. ЛОГИЧЕСКАЯ СОГЛАСОВАННОСТЬ: в заключении упоминай ТОЛЬКО те "
        "угрозы, факты и выводы, которые есть в основной части. Не "
        "добавляй новых концепций, не раскрытых в главах. Запрещено "
        "начинать заключение с «итак», «таким образом», «подводя итог».\n"
        f"4. ДИСЦИПЛИНАРНАЯ ПРИВЯЗКА: текст строго соответствует "
        f"дисциплине «{subject}». Если тема и дисциплина из разных "
        f"областей — рассматривай тему через призму дисциплины.\n"
        "5. ОБЪЁМ: строго соблюдай целевой диапазон знаков (±10 %). "
        "Не растягивай текст повторами, не сжимай до потери смысла.\n\n"
        "❌ НЕПРАВИЛЬНО: «...осадочными породами [1, с.»\n"
        "❌ НЕПРАВИЛЬНО: «...осадочными породами [1]»\n\n"
        "✅ ПРАВИЛЬНО: «...осадочными породами [1, с. 45]»\n"
        "✅ ПРАВИЛЬНО: «...осадочными породами [2, с. 120–125]»\n\n"
        "Запомни: ссылка всегда заканчивается на ] и содержит номер страницы после «с.»"
    )

    for key, prompt in prompts.items():
        step += 1

        # Рассчитываем целевой объём для блока
        if key == "intro":
            block_share = 0.10
        elif key == "conclusion":
            block_share = 0.08
        elif key == "literature":
            block_share = 0.0
        else:
            num_ch = max(1, len([k for k in prompts if k.startswith("ch")]))
            block_share = 0.75 / num_ch

        block_chars = int(total_chars * block_share) if block_share > 0 else 0
        max_tok     = tokens_for_chars(block_chars) if block_chars > 0 else 1500

        if prog:
            block_name = {
                "intro":       "Введение",
                "conclusion":  "Заключение",
                "literature":  "Список литературы",
                "main1":       "Основная часть (аргумент 1)",
                "main2":       "Основная часть (аргумент 2)",
                "part1":       "Раздел 1",
                "part2":       "Раздел 2",
            }.get(key, (
                # Формат ch1_s2 → "Глава 1, подглава 2"
                f"Глава {key.replace('ch', '').replace('_s', ', подглава ')}"
                if key.startswith("ch") else key
            ))
            await prog.update(label=f"✍️ Пишу: {block_name}")

        # ── Интеграция эталона и жестких примеров (Few-Shot) ──
        example_type = "chapter"
        if key == "intro":
            example_type = "intro"
        elif key == "conclusion":
            example_type = "conclusion"
        
        golden_example = get_golden_example(doc_type, example_type)
        
        # Собираем итоговый системный промпт с примерами
        current_style_sys = style_sys + "\n\n" + FEW_SHOT_EXAMPLES
        if golden_example:
            current_style_sys += f"\n\n🌟 ПРИМЕР ИДЕАЛЬНОГО ИСПОЛНЕНИЯ ДАННОГО БЛОКА:\n\"{golden_example}\"\nСледуй этому ритму, уровню детализации и способу оформления ссылок."

        messages = [
            {"role": "system", "content": current_style_sys},
            {"role": "user",   "content": prompt},
        ]

        # Список литературы лучше брать из реальных каталогов, а не просить
        # модель «вспомнить» источники: так меньше выдуманных книг и DOI.
        if key == "literature" and verified_bib:
            text = verified_bib
            used_model = model_key
        else:
            text, used_model = await chat_with_fallback(model_key, messages, max_tok)

        if prog and used_model and used_model != model_key:
            await prog.update(model_name=AI_MODELS.get(used_model, {}).get("name", used_model))

        # ═══════════════════════════════════════════════════════════
        # ПОСТ-ВАЛИДАЦИЯ КАЖДОГО БЛОКА
        # В старой версии этот код случайно стоял ВНЕ цикла for и обрабатывал
        # только последний prompt (обычно literature), из-за чего введение и
        # подглавы превращались в заглушки. Теперь каждый key сохраняется в parts.
        # ═══════════════════════════════════════════════════════════
        if key != "literature":
            retry_count = 0
            # Проверяем обрывы ссылок вроде «[1, с.» и перегенерируем блок.
            while text and (not _validate_no_broken_citations(text) or re.search(r'\[\d+\s*,\s*[сc]\.\s*$', text)) and retry_count < 2:
                retry_count += 1
                if prog:
                    await prog.update(label=f"🔄 Исправляю обрывы ссылок в {block_name} (попытка {retry_count})", force=True)

                retry_prompt = (
                    prompt
                    + "\n\n⚠️ ОШИБКА: В предыдущем ответе были обрывы ссылок "
                    + "(например, '[1, с.' без цифры или ссылка в конце текста не завершена). "
                    + "ПЕРЕПИШИ блок, убедившись, что КАЖДАЯ ссылка завершена номером страницы "
                    + "и закрывающей скобкой: [N, с. X]."
                )

                retry_messages = [
                    {"role": "system", "content": current_style_sys},
                    {"role": "user", "content": retry_prompt},
                ]
                text, _ = await chat_with_fallback(model_key, retry_messages, max_tok)
                text = _clean_ai_artifacts(text)

            # Повторная генерация при пустом/слишком коротком ответе.
            if not text or len(text.strip()) < 80:
                print(f"[RETRY] Блок «{key}» пустой ({len(text) if text else 0} зн.), повторная генерация (попытка 2)...")
                await asyncio.sleep(2)
                text2, used_model2 = await chat_with_fallback(model_key, messages, max_tok)
                if text2 and len(text2.strip()) > 80:
                    text = text2
                    print(f"[RETRY] Попытка 2 успешна: {len(text)} зн.")
                    if prog and used_model2:
                        await prog.update(model_name=AI_MODELS.get(used_model2, {}).get("name", used_model2))
                else:
                    print("[RETRY] Попытка 2 провалена, попытка 3...")
                    await asyncio.sleep(3)
                    text3, used_model3 = await chat_with_fallback(model_key, messages, max_tok)
                    if text3 and len(text3.strip()) > 80:
                        text = text3
                        print(f"[RETRY] Попытка 3 успешна: {len(text)} зн.")
                        if prog and used_model3:
                            await prog.update(model_name=AI_MODELS.get(used_model3, {}).get("name", used_model3))
                    else:
                        print(f"[RETRY] Все попытки провалены для «{key}», используем заглушку")
                        text = _stub_text(key, topic)

            if not text or len(text) < 80:
                text = _stub_text(key, topic)

            text = _clean_ai_artifacts(text)
            text = _fix_nonsense_phrases(text)
            text = _replace_ai_cliches(text)
            if humanize:
                text = _add_human_touch(text)
            # Внутренний "детектор" ИИ: если много шаблонов — ещё проход замены.
            if _ai_detector_score(text) > 30:
                text = _replace_ai_cliches(text)
                if humanize:
                    text = _add_human_touch(text)

            # ── ЗАЩИТА ОТ СЛУЖЕБНЫХ ФРАЗ И КОРОТКИХ ОТВЕТОВ ──
            forbidden_patterns = [
                "вступление к эссе",
                "вот текст",
                "вот второй аргумент",
                "конечно. вот",
                "как языковая модель",
                "объем текста строго выдержан",
                "вот первый аргумент",
            ]

            text_lower = text.lower()
            is_bad = any(pat in text_lower[:300] for pat in forbidden_patterns)
            is_too_short = len(text.strip()) < 400

            if is_bad or is_too_short:
                print(f"[WARN] Блок {key} содержит служебную фразу или слишком короткий ({len(text)} знаков). Перегенерация...")
                retry_prompt = (
                    prompt
                    + "\n\nВАЖНО: Напиши полноценный текст без служебных фраз. "
                    + "Не пиши «Вступление», «Конечно», «Вот текст». "
                    + "Начинай сразу с содержания. Минимум 800 знаков."
                )
                retry_messages = [
                    {"role": "system", "content": style_sys},
                    {"role": "user", "content": retry_prompt},
                ]
                text2, _used_model2 = await chat_with_fallback(
                    model_key,
                    retry_messages,
                    tokens_for_chars(max(block_chars * 2, 1200)),
                )

                if text2 and len(text2.strip()) > len(text.strip()) and len(text2.strip()) > 300:
                    text = _clean_ai_artifacts(text2)
                    print(f"[WARN] Перегенерация успешна: {len(text)} знаков")
                else:
                    print("[WARN] Перегенерация не помогла, оставляем как есть")

            # Дозаполняем если текст короче цели на 25%+ (кроме литературы).
            if block_chars > 0 and len(text) < int(block_chars * 0.75):
                max_refills = 3
                refill_count = 0
                while len(text) < int(block_chars * 0.9) and refill_count < max_refills:
                    extra_chars = block_chars - len(text)
                    ext_tok     = tokens_for_chars(extra_chars)
                    extra_prompt = (
                        f"Продолжи и дополни следующий академический текст по теме «{topic}». "
                        f"Добавь ещё минимум {extra_chars} знаков. "
                        f"Пиши плавно, без заголовков, со ссылками [N, с. X], без ** и ##:\n\n"
                        f"{text[-600:]}"
                    )
                    ext_messages = [
                        {"role": "system", "content": style_sys},
                        {"role": "user",   "content": extra_prompt},
                    ]
                    extra, _ = await chat_with_fallback(model_key, ext_messages, ext_tok)
                    if extra and len(extra.strip()) > 50:
                        extra = _clean_llm_chunk(extra.strip())
                        text = text.rstrip() + "\n\n" + extra.strip()
                    else:
                        break
                    refill_count += 1

            text = _ensure_block_terminates(text)
            if key.startswith("ch"):
                text = _ensure_paragraph_breaks(text, min_paragraphs=3)
        else:
            # Библиография: приоритет у реально найденных источников из каталогов
            # OpenAlex/Crossref/Semantic Scholar и URL пользователя. LLM-ответ
            # используется только как дополнение, чтобы не получить 3–5
            # источников или выдуманные позиции.
            llm_bib = _filter_bibliography_scientific_lines(
                _normalize_bibliography(_clean_ai_artifacts(text or ""))
            )
            if not llm_bib or _count_sources(llm_bib) < 3:
                llm_bib = ""
            text = _filter_bibliography_scientific_lines(_combine_bibliographies(
                verified_bib,
                llm_bib,
                limit=source_limit,
            ))
            if _count_sources(text) < 10:
                print(f"[LIT] ⚠️ Научных источников найдено меньше 10: {_count_sources(text)}. Добираю проверяемым резервом.")
                text = _ensure_min_bibliography(text, topic, subject, min_sources=10, limit=source_limit)
            text = _normalize_bibliography(text)

        parts[key] = text

        if prog:
            await prog.update(step_done=True)

    # ═══════════════════════════════════════════════════════════
    # fix17: Извлекаем единое определение ключевого понятия из введения
    # ═══════════════════════════════════════════════════════════
    intro_text = parts.get("intro", "")
    if intro_text and topic:
        concept_def = get_key_concept(topic, subject, intro_text)
        if concept_def:
            print(f"[CONCEPT] Единое определение: {concept_def[:60]}...")

    # ═══════════════════════════════════════════════════════════
    # ГЕНЕРАЦИЯ ЗАКЛЮЧЕНИЯ ПОСЛЕ ВСЕХ ГЛАВ
    # (ключевое исправление: заключение видит реальное содержание)
    # ═══════════════════════════════════════════════════════════

    # Собираем контекст из всех содержательных блоков (главы/аргументы/разделы)
    summaries = []
    # Ключи, которые являются содержательными блоками (не intro, не literature)
    content_keys = [k for k in parts.keys()
                    if k not in ("intro", "conclusion", "literature")]
    for key in sorted(content_keys):
        txt = parts[key]
        if not txt or len(txt) < 30:
            continue
        head = txt[:500].strip() if len(txt) > 500 else txt.strip()
        tail = txt[-400:].strip() if len(txt) > 400 else ""
        # Человекочитаемая метка блока
        if key.startswith("ch"):
            label = key.replace("ch", "Глава ").replace("_s", ", подглава ")
        elif key.startswith("main"):
            label = key.replace("main", "Аргумент ")
        elif key.startswith("part"):
            label = key.replace("part", "Раздел ")
        else:
            label = key
        summaries.append(
            f"=== {label} (начало) ===\n{head}\n"
            + (f"=== {label} (конец) ===\n{tail}" if tail else "")
        )

    content_context = "\n\n".join(summaries) if summaries else (
        f"Текст по теме «{topic}» был сгенерирован."
    )

    conc_chars = int(total_chars * 0.10)

    # ── Жанрово-зависимый промпт заключения ──
    if doc_type == "esse":
        conc_prompt = (
            f"Напиши заключение эссе на тему «{topic}».\n\n"
            f"НИЖЕ — содержание твоих аргументов:\n{content_context}\n\n"
            f"Твоя задача:\n"
            f"1. Подведи ЛИЧНЫЙ итог размышления (от первого лица: «я пришёл к выводу»).\n"
            f"2. Опирайся ТОЛЬКО на аргументы выше, не выдумывай новые факты.\n"
            f"3. Выскажи авторскую позицию и возможный прогноз.\n"
            f"4. НЕ упоминай «главы», «разделы» — это эссе, а не реферат.\n"
            f"5. НЕ используй канцеляризмы («практическая значимость», «систематизация»).\n"
            f"6. Объём: 3–4 абзаца, 7–10 предложений. Пиши просто, без канцеляризма."
        )
    elif doc_type == "doklad":
        conc_prompt = (
            f"Напиши заключение доклада «{topic}».\n\n"
            f"Содержание разделов:\n{content_context}\n\n"
            f"Задача: подведи итог, сформулируй ключевые выводы и практическое значение.\n"
            f"НЕ упоминай «главы» — это доклад с разделами 1 и 2.\n"
            f"Стиль — устный, чёткий, без излишней сложности.\n"
            f"Максимум 5–7 предложений. Пиши просто, без канцеляризма."
        )
    elif doc_type == "article":
        conc_prompt = (
            f"Напиши заключение научной статьи «{topic}».\n\n"
            f"НИЖЕ — РЕАЛЬНОЕ СОДЕРЖАНИЕ СЕКЦИЙ:\n{content_context}\n\n"
            f"Твоя задача:\n"
            f"1. Опираясь на содержание выше, сформулируй основные научные выводы исследования.\n"
            f"2. Подчеркни научную новизну и теоретическую/практическую значимость работы.\n"
            f"3. Опиши возможные направления дальнейших исследований в этой научной области.\n"
            f"4. Стиль изложения — строго академический, безличный.\n"
            f"5. Максимум 5–7 предложений. Без канцелярита и шаблонов."
        )
    else:
        conc_prompt = (
            f"Напиши заключение для работы «{topic}».\n\n"
            f"НИЖЕ — РЕАЛЬНОЕ СОДЕРЖАНИЕ ГЛАВ:\n{content_context}\n\n"
            f"Твоя задача:\n"
            f"1. Напиши НОВЫЕ связные выводы по каждой главе, а НЕ копию текста глав. "
            f"Переформулируй содержание на уровне итогов: что доказано, что выявлено, какое значение имеет. "
            f"НЕ используй ярлыки-заголовки «Вывод по первой главе:», «Общий итог:» — пиши абзацами научного текста.\n"
            f"2. Не выдумывай новые факты, которых нет в тексте глав.\n"
            f"3. Сформулируй общие выводы и практическую значимость.\n"
            f"4. Укажи перспективы дальнейшего исследования.\n"
            f"5. Используй сноски только с номерами страниц: [1, с. 45], [2, с. 120].\n"
            f"ЗАПРЕЩЕНО: выдумывать необсуждённые темы.\n"
            f"6. Объём: 3–4 абзаца, 7–10 предложений. Пиши просто, без канцеляризма."
        )

    conc_prompt += (
        "\n\nОТДЕЛЬНО ПРОВЕРЬ ЗАКЛЮЧЕНИЕ НА ТАВТОЛОГИЮ: не повторяй один и тот же "
        "вывод разными словами, не начинай соседние предложения одинаково, "
        "не злоупотребляй словами «исследование», «анализ», «данный», "
        "«проблематика». Каждое предложение должно добавлять новый смысл."
    )

    conc_messages = [
        {"role": "system", "content": style_sys},
        {"role": "user",   "content": strict_prompt(
            conc_prompt,
            conc_chars,
            writing_style,
            doc_type,
        )},
    ]

    if prog:
        await prog.update(label="✍️ Пишу: Заключение (на основе глав)")

    conc_text, conc_model = await chat_with_fallback(
        model_key, conc_messages,
        max(tokens_for_chars(conc_chars * 2), 2200),
    )

    if prog and conc_model and conc_model != model_key:
        await prog.update(model_name=AI_MODELS.get(conc_model, {}).get("name", conc_model))

    if not conc_text or len(conc_text) < 80:
        conc_text = _stub_text("conclusion", topic)

    conc_text = _clean_ai_artifacts(conc_text)
    conc_text = _fix_nonsense_phrases(conc_text)
    conc_text = _replace_ai_cliches(conc_text)
    if humanize:
        conc_text = _add_human_touch(conc_text)
    if _ai_detector_score(conc_text) > 30:
        conc_text = _replace_ai_cliches(conc_text)
        if humanize:
            conc_text = _add_human_touch(conc_text)
    # Запрещённые вводные обороты и явная тавтология в заключении.
    conc_text = _strip_forbidden_openers(conc_text)
    conc_text = _remove_conclusion_tautology(conc_text)
    # Проверка согласованности с основной частью (fix9, ужесточена в fix10):
    # выкидывает предложения с фактами, которых нет в главах.
    conc_text = _validate_conclusion_consistency(conc_text, parts)
    # Структура заключения: выводы по главам + общий итог, без шаблонных фраз.
    conc_text = _structure_conclusion_by_chapters(conc_text, parts)
    conc_text = _remove_conclusion_tautology(conc_text)
    conc_text = _strip_forbidden_openers(conc_text)
    parts["conclusion"] = conc_text

    if prog:
        await prog.update(step_done=True)

    # ═══════════════════════════════════════════════════════════
    # СВЯЗЫВАНИЕ ССЫЛОК СО СПИСКОМ ЛИТЕРАТУРЫ (приоритет 🔴)
    # Приводим все внутритекстовые [N, с. X] к диапазону 1..кол-во источников,
    # чтобы каждая ссылка указывала на реально существующий пункт списка.
    # ═══════════════════════════════════════════════════════════
    n_sources = _count_sources(parts.get("literature", ""))
    if n_sources > 0:
        # fix16: сначала прогоняем repair+fix_nonsense+fix_citations,
        # потом собираем ГЛОБАЛЬНУЮ карту страниц со всех частей и
        # дозаполняем bare `[N]` ею (раньше карта была только локальной,
        # поэтому `[2]` в 1.2 не получал страницу из `[2, с. 25]` в 2.1).
        for key in list(parts.keys()):
            if key == "literature":
                continue
            parts[key] = _repair_broken_citations(parts[key])
            parts[key] = _fix_nonsense_phrases(parts[key])
            parts[key] = _fix_citations(parts[key], n_sources)

        global_page_map: dict = {}
        for key, val in parts.items():
            if key == "literature" or not val:
                continue
            for n, p in _build_page_map(val).items():
                global_page_map.setdefault(n, p)

        for key in list(parts.keys()):
            if key == "literature":
                continue
            # fix14+16: дозаполняем `[N]` без страницы из глобальной карты
            parts[key] = _fill_missing_pages(parts[key],
                                            global_page_map=global_page_map)

        # Удаляем неиспользуемые источники (фикс «избыточен, N не используются»)
        parts = _prune_unused_sources(parts)
    else:
        # Источников нет, но оборванные ссылки всё равно надо почистить
        for key in list(parts.keys()):
            if key == "literature":
                continue
            parts[key] = _repair_broken_citations(parts[key])

    # Финальная чистка: без markdown-маркеров, фраз-маркеров ИИ, битой библиографии и ссылок без страниц.
    parts["literature"] = _filter_bibliography_scientific_lines(_normalize_bibliography(_strip_markdown_markers(parts.get("literature", ""))))
    # Последняя страховка: добираем только научными источниками из каталогов/URL.
    # ГОСТы, законы и методички больше не подмешиваются ради количества.
    if _count_sources(parts.get("literature", "")) < 10:
        parts["literature"] = _normalize_bibliography(_filter_bibliography_scientific_lines(_combine_bibliographies(
            parts.get("literature", ""),
            verified_bib,
            limit=max(10, source_limit),
        )))
        if _count_sources(parts.get("literature", "")) < 10:
            print(f"[LIT] ⚠️ Итоговый список содержит {_count_sources(parts.get('literature', ''))}; добираю резервом.")
            parts["literature"] = _ensure_min_bibliography(parts.get("literature", ""), topic, subject, min_sources=10, limit=max(10, source_limit))
    lit_ok, lit_reason = validate_literature(parts.get("literature", ""))
    if not lit_ok:
        print(f"[LIT] ⚠️ {lit_reason}: очищаю список литературы")
        parts["literature"] = _normalize_bibliography(parts.get("literature", ""))
    parts["literature"] = _ensure_min_bibliography(parts.get("literature", ""), topic, subject, min_sources=10, limit=max(10, source_limit))
    # ── USER-FIX (v2.8): дедуп + алфавит + единый ГОСТ-вид (с ремапом ссылок) ──
    parts = _alphabetize_and_dedupe_bibliography(parts, alphabetize=True)
    # Если дедуп опустил список ниже минимума — добиваем и пересортировываем.
    if _count_sources(parts.get("literature", "")) < 10:
        parts["literature"] = _ensure_min_bibliography(
            parts.get("literature", ""), topic, subject, min_sources=10, limit=max(10, source_limit))
        parts = _alphabetize_and_dedupe_bibliography(parts, alphabetize=True)
    # ── USER-FIX (v2.8): каждый источник из списка должен быть упомянут в тексте ──
    parts = _distribute_missing_citations(parts, _count_sources(parts.get("literature", "")))
    final_n_sources = _count_sources(parts.get("literature", ""))

    for key in list(parts.keys()):
        if key == "literature":
            continue
        parts[key] = _strip_markdown_markers(parts[key])
        parts[key] = _remove_ai_marker_phrases(parts[key])
        parts[key] = _replace_ai_cliches(parts[key])
        parts[key] = _normalize_numeric_claims(parts[key], topic)
        parts[key] = _normalize_punctuation(parts[key])
        parts[key] = _repair_broken_citations(parts[key])
        if final_n_sources > 0:
            parts[key] = _fix_citations(parts[key], final_n_sources)
        # Последний обязательный проход: ни одной ссылки вида [1] без страницы.
        parts[key] = _fill_missing_pages(parts[key])
        parts[key] = _ensure_block_terminates(parts[key])

    # ═══════════════════════════════════════════════════════════
    # ФИНАЛЬНАЯ ПРОВЕРКА: все ли ожидаемые ключи заполнены
    # ═══════════════════════════════════════════════════════════
    empty_keys = [k for k, v in parts.items() if not v or len(v.strip()) < 50]
    if empty_keys:
        print(f"[FINAL CHECK] ⚠️ Пустые/короткие блоки после генерации: {empty_keys}")
        for ek in empty_keys:
            if ek not in ("literature",):
                print(f"[FINAL CHECK] Заполняю «{ek}» заглушкой")
                parts[ek] = _stub_text(ek, topic)

    # Проверяем наличие всех ожидаемых ключей ch*_s* из промптов
    for pk in list(prompts.keys()):
        if pk not in parts:
            print(f"[FINAL CHECK] ❌ Ключ «{pk}» из промптов отсутствует в parts! Добавляю заглушку.")
            parts[pk] = _stub_text(pk, topic)

    return parts


def _stub_text(key: str, topic: str) -> str:
    """Аварийный текст, если ИИ не ответил.

    Для списка литературы аварийная заглушка отключена: ГОСТы, законы и
    методички не должны подмешиваться в тематическую библиографию.
    """
    topic = (topic or "").strip() or "исследуемой проблематики"
    stubs = {
        "intro":       f"Данная работа посвящена исследованию темы «{topic}». В современных условиях данная проблематика приобретает особую актуальность и практическую значимость для науки и общества.",
        "conclusion":  f"Проведённое исследование по теме «{topic}» позволило сформулировать следующие выводы: изученная проблематика имеет важное теоретическое и практическое значение.",
        "literature": "",
    }
    if key in stubs:
        return stubs[key]

    # Осмысленная заглушка для подглав
    if not topic:
        topic = "теме"
    return (
        f"В рамках данной подглавы проводится детальный анализ ключевых аспектов, "
        f"связанных с темой «{topic}». На основе имеющихся научных данных "
        f"рассматриваются основные закономерности и тенденции, определяющие "
        f"современное состояние изучаемого вопроса. Дальнейшее развитие темы "
        f"требует более глубокого исследования с привлечением дополнительных "
        f"источников и эмпир��ческих данных. [1, с. 45]"
    )


def _ensure_paragraph_breaks(text: str, min_paragraphs: int = 3) -> str:
    """fix13: если LLM вернула монолитный текст без \n\n — режем по предложениям.

    Гарантирует наличие минимум `min_paragraphs` абзацев, разделённых пустой строкой.
    Не трогает уже разбитый текст.
    """
    if not text:
        return text
    # Если уже есть достаточные разбиения — не трогаем
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    if len(paragraphs) >= min_paragraphs:
        return text
    # Собираем всё в один абзац (на случай если \n не пустыми разделены)
    flat = re.sub(r'\s*\n\s*', ' ', text).strip()
    # Режем по предложениям даже короткие блоки: лучше добавить/разделить
    # третий абзац, чем получить подглаву из 1–2 абзацев.
    # Режем по предложениям
    # fix16: безопасный split — не режем `[N, с. K]`
    sentences = _split_sentences_safe(flat)
    if not sentences:
        return text
    if len(sentences) < min_paragraphs:
        chunks = paragraphs[:] if paragraphs else [flat]
        fallback_paras = [
            "Этот аспект имеет самостоятельное значение для раскрытия темы, поскольку позволяет связать теоретические положения с последующим анализом и сформулировать промежуточный вывод по рассматриваемому вопросу.",
            "Дополнительное рассмотрение проблемы показывает, что её нельзя сводить только к отдельным фактам: важно учитывать взаимосвязь причин, условий и последствий в рамках выбранной дисциплины.",
            "Следовательно, данный раздел выполняет связующую функцию в структуре работы и подготавливает основу для дальнейшего анализа заявленной темы."
        ]
        fi = 0
        while len(chunks) < min_paragraphs:
            chunks.append(fallback_paras[fi % len(fallback_paras)])
            fi += 1
        return '\n\n'.join(chunks[:min_paragraphs])
    # Разделяем предложения примерно поровну между min_paragraphs кусками
    per_chunk = max(1, len(sentences) // min_paragraphs)
    chunks = []
    for i in range(min_paragraphs):
        start = i * per_chunk
        end = start + per_chunk if i < min_paragraphs - 1 else len(sentences)
        chunk = ' '.join(sentences[start:end]).strip()
        if chunk:
            chunks.append(chunk)
    return '\n\n'.join(chunks)


def _generate_substantial_stub(sub_title: str, chapter_title: str, topic: str) -> str:
    """Развёрнутая заглушка для пустой подглавы (минимум 600 символов).

    Генерирует осмысленный текст, привязанный к названию подглавы и теме.
    Используется когда API не вернул текст для подглавы.
    """
    if not topic:
        topic = "данной теме"

    return (
        f"Рассматриваемый аспект «{sub_title}» занимает важное место "
        f"в структуре исследования по теме «{topic}». Данный вопрос "
        f"привлекает внимание как отечественных, так и зарубежных "
        f"исследователей, что обусловлено его теоретической и практической "
        f"значимостью [1, с. 23].\n\n"
        f"Анализ научной литературы свидетельствует о многообразии подходов "
        f"к изучению данной проблематики. Ряд авторов подчёркивает "
        f"необходимость комплексного рассмотрения вопроса с учётом "
        f"исторических, методологических и практических аспектов. "
        f"В контексте главы «{chapter_title}» следует отметить, что "
        f"исследуемые закономерности тесно связаны с общей проблематикой "
        f"работы и позволяют выявить ключевые тенденции развития [2, с. 56].\n\n"
        f"Существенный вклад в разработку данного направления внесли "
        f"работы последних лет, в которых предложены новые теоретические "
        f"модели и эмпирические результаты. Авторы отмечают, что "
        f"исследование «{sub_title.lower()}» требует "
        f"междисциплинарного подхода, объединяющего достижения различных "
        f"областей знания [3, с. 112].\n\n"
        f"В целом анализ рассмотренных источников позволяет "
        f"констатировать, что проблема «{sub_title.lower()}» "
        f"является актуальной и требует дальнейшего углублённого изучения. "
        f"Полученные в ходе анализа данные могут быть использованы "
        f"для формирования целостного представления о теме «{topic}» "
        f"и определения перспективных направлений исследования [4, с. 78]."
    )


def _stub_text_for_subchapter(title: str, topic: str) -> str:
    """Заглушка для пустой подглавы (устаревшая, используйте _generate_substantial_stub)."""
    return _generate_substantial_stub(title, "", topic)


# ═══════════════════════════════════════════════════════════════
#  ГОСТ-DOCX: СТИЛИ, TOC, НУМЕРАЦИЯ, ПОДГЛАВЫ
# ═══════════════════════════════════════════════════════════════

def _normalize_typography(text: str) -> str:
    """Типографика (приоритет 🟡): тире, дефисы, кавычки, лишние #.

    - '---' → '—' (длинное тире), '--' → '–' (короткое тире/диапазон)
    - дефис между словами с пробелами ' - ' → ' — ' (тире)
    - прямые кавычки "..." → «ёлочки»
    - убираем одиночные служебные '#'
    """
    if not text:
        return ""
    text = _strip_markdown_markers(text)
    # Markdown-заголовки '# ', '## **' в начале строки убираем целиком
    text = re.sub(r'(?m)^\s*#{1,6}\s*\*{0,2}', '', text)
    text = re.sub(r"(?m)^\s{0,3}#{1,6}\s+", "", text)
    # Тройной дефис → длинное тире, двойной → короткое тире
    text = text.replace("---", "—").replace("--", "–")
    # Дефис, окружённый пробелами, как знак тире в предложении → длинное тире
    text = re.sub(r"(?<=\s)-(?=\s)", "—", text)
    # Прямые двойные кавычки → «ёлочки» (парами)
    def _quotes(m: re.Match) -> str:
        return "«" + m.group(1) + "»"
    text = re.sub(r'"([^"\n]*)"', _quotes, text)
    # Одиночные оставшиеся '#' (не часть слова) убираем
    text = re.sub(r"(?<!\w)#(?!\w)", "", text)
    # Убираем ** в начале/конце строки после удаления #
    text = re.sub(r'(?m)^\s*\*\*', '', text)
    text = re.sub(r'(?m)\*\*\s*$', '', text)
    return text


def _split_sentences_safe(text: str) -> list[str]:
    """fix16: разбивает текст на предложения, НЕ срываясь на `с.` внутри
    ссылок `[N, с. K]`. Точка считается концом предложения только если
    следующее «предложение» начинается с заглавной буквы / открывающей
    кавычки / скобки. Иначе склеиваем обратно.
    """
    if not text:
        return []
    parts = re.split(r"(?<=[.!?…])\s+(?=[А-ЯA-ZЁ«„„\"(])", text)
    return [p for p in parts if p and p.strip()]


def _validate_no_broken_citations(text: str) -> bool:
    """
    Проверяет, нет ли в тексте обрыв��в ссылок типа '[1, с.' без цифры.
    Возвращает True если текст чист, False если найден обрыв.
    """
    if not text:
        return True
    # Ищем паттерн: открывающая скобка, число, запятая, 'с.' и далее НЕТ цифры перед закрытием скобки или концом строки.
    # Это регулярное выражение ищет именно 'опасные' места.
    pattern = r"\[\s*\d+\s*,\s*[сСcC]\.\s*[^0-9\s]*(\s*\]|$)"
    if re.search(pattern, text):
        return False
    return True


def _is_garbage(text: str) -> bool:
    """Проверяет, не является ли текст случайным набором символов (мусором)."""
    if not text:
        return True
    t = text.strip()
    if len(t) < 2:
        return True
    # только повторяющиеся буквы одного типа "лллл", "шшшш"
    if len(set(t.lower())) == 1 and t.isalpha():
        return True
    # только цифры/спецсимволы без букв
    letters = len(re.findall(r'[А-Яа-яA-Za-z]', t))
    if letters == 0:
        return True
    # соотношение букв к общему числу менее 50%
    if letters / len(t) < 0.5:
        return True
    # Одиночная буква или буква с точкой ("Л.", "О.")
    if re.fullmatch(r'[А-Яа-яA-Za-z]\.?', t):
        return True
    # Короткие бессмысленные сочетания для ФИО: «То тл», «Ршп дир», «Оо оо».
    words = re.findall(r'[А-Яа-яЁёA-Za-z]+', t)
    if words and len(words) <= 3:
        # Все слова слишком короткие — почти наверняка мусор, а не ФИО.
        if all(len(w) <= 3 for w in words):
            return True
        # Для реальных ФИО каждое русское слово обычно содержит хотя бы одну гласную.
        for w in words:
            if re.search(r'[А-Яа-яЁё]', w) and len(w) >= 2 and not re.search(r'[АаЕеЁёИиОоУуЫыЭэЮюЯя]', w):
                return True
        # FIX 3: смягчено. Отвергаем только если ВСЕ слова со строчной
        # буквы (явная опечатка), а не если хоть одно.
        if len(words) in (2, 3) and all(w and w[0].islower() for w in words if w):
            return True
    return False


# FIX 2: известные имена/отчества, которые совпадают с «глагольными»
# окончаниями и не должны отклоняться (Михаил, Даниил … уже не попадают
# под шаблоны ниже, но список — дополнительная страховка).
_FIO_VERB_WHITELIST = {
    "михаил", "даниил", "гавриил", "самуил", "павел", "орёл", "орел",
}



_COMMON_RU_FIRST_NAMES = {
    "александр", "алексей", "андрей", "антон", "артем", "артём", "арсений",
    "богдан", "борис", "вадим", "валентин", "валерий", "василий", "виктор",
    "виталий", "владимир", "владислав", "вячеслав", "геннадий", "георгий",
    "глеб", "григорий", "даниил", "денис", "дмитрий", "егор", "евгений",
    "захар", "иван", "игорь", "илья", "кирилл", "константин", "лев", "леонид",
    "макар", "максим", "матвей", "михаил", "никита", "николай", "олег",
    "павел", "петр", "пётр", "роман", "руслан", "семен", "семён", "сергей",
    "станислав", "степан", "тимофей", "тимур", "федор", "фёдор", "юрий", "ярослав",
    "анастасия", "анна", "алина", "алиса", "александра", "арина", "валерия",
    "варвара", "василиса", "вера", "вероника", "виктория", "галина", "дарья",
    "диана", "екатерина", "елена", "елизавета", "жанна", "зоя", "инна", "ирина",
    "карина", "кира", "кристина", "ксения", "лариса", "лидия", "любовь", "маргарита",
    "марина", "мария", "милана", "надежда", "наталья", "нина", "оксана", "ольга",
    "полина", "светлана", "софия", "софья", "татьяна", "ульяна", "юлия", "яна",
}

_FIO_FAKE_WORDS = {
    "лол", "кек", "плохо", "хорошо", "тест", "пример", "автор", "учитель",
    "преподаватель", "студент", "школьник", "последствия", "трампизм",
}


def _fio_name_is_real(name: str) -> bool:
    """Проверяет именно имя (второе слово ФИО), чтобы отсеять «Алррш»,
    «Плохо», «Лол» и похожие заглушки."""
    parts = [p.lower().replace("ё", "е") for p in re.split(r"-+", name or "") if p]
    if not parts:
        return False
    normalized_names = {n.replace("ё", "е") for n in _COMMON_RU_FIRST_NAMES}
    for p in parts:
        if p in _FIO_FAKE_WORDS:
            return False
        if p not in normalized_names:
            return False
    return True


_FIO_STOPWORDS = {
    # короткие русские слова, заведомо не имя/фамилия
    "надо", "так", "это", "тот", "там", "тут", "тут", "вот", "ещё", "уже",
    "или", "как", "что", "кто", "его", "её", "их", "она", "они", "оно",
    "был", "была", "были", "нет", "да", "ну", "же", "бы", "ли", "не",
    "при", "под", "над", "без", "для", "про", "ради", "из-за", "из-под",
    "лишь", "также", "тоже", "если", "когда", "пока", "потом", "тогда",
    "везде", "нигде", "только", "очень", "почти", "ровно", "точно",
    "сам", "сама", "сами", "себя", "свой", "своя", "свои",
}


def _word_looks_like_verb(word: str) -> bool:
    """FIX 2 + (user-patch): грубая эвристика «слово — это не имя/фамилия».

    Отлавливает:
      - инфинитивы («Оперировать», «Послушать»)
      - прошедшее время м.р. («Послушал», «Написал», «Тянул»)
      - короткие служебные слова из _FIO_STOPWORDS («Надо», «Их», «Так», …)
      - ALL-CAPS «фамилии» («ОВОВ», «АЛЛАШ», «ИОШЛ») — реальные фамилии так
        не пишут; обычно только первая буква заглавная.
      - аббревиатуры из 3–5 согласных без гласных («ШМР», «БГРТ») — это
        мусор, попавший вместо ФИО.
    """
    w = (word or "").lower().replace("ё", "е").strip("-")
    if not w:
        return False

    # 1) Стоп-слова — короткие служебные русские слова, заведомо не имя.
    if w in _FIO_STOPWORDS:
        return True

    # 2) ALL-CAPS слово (исходное, не lower) длиной 2–6 — почти наверняка
    #    мусорная аббревиатура, не реальное имя/фамилия.
    src_w = (word or "").strip("-")
    if 2 <= len(src_w) <= 6 and src_w.isupper() and src_w.isalpha():
        return True

    # 3) Слова без гласных длиной 2–5 — мусор («ШМР», «КРН»).
    if 2 <= len(w) <= 5 and not re.search(r"[аеиоуыэюяё]", w):
        return True

    # 4) (user-patch) Дисбаланс «одна буква доминирует»: если какая-то одна
    #    буква занимает > 40% слова (длиной ≥ 4) — это мусор типа
    #    «Иггии» (4 «и» из 5 = 80%), «Рмоммолм» (4 «м» из 8 = 50%).
    if len(w) >= 4:
        from collections import Counter as _Cnt
        _letter_counts = _Cnt(c for c in w if c.isalpha())
        if _letter_counts:
            _max_share = max(_letter_counts.values()) / len(w)
            if _max_share > 0.4:
                return True

    # 5) (user-patch) Дисбаланс гласные/согласные: реальное русское слово
    #    обычно имеет 30–55% гласных. Слова с > 70% гласных («Оииоги»,
    #    «Уоиа») или > 80% согласных («Бгрт») — мусор.
    if len(w) >= 4:
        _vowels = sum(1 for c in w if c in "аеиоуыэюяё")
        _alpha_total = sum(1 for c in w if c.isalpha())
        if _alpha_total:
            _vowel_ratio = _vowels / _alpha_total
            if _vowel_ratio > 0.70 or _vowel_ratio < 0.15:
                return True

    # 6) (user-patch) Низко-частотные русские биграммы как сигнал мусора.
    #    Если в слове встречаются «гг», «тш», «мм мм», «оо оо» подряд —
    #    суммарно > 1 редкой биграммы → мусор.
    if len(w) >= 5:
        _rare = ["гг", "тт", "тш", "шт", "пх", "хп", "кц", "цк", "жш", "шж",
                 "ыы", "ээ", "юю", "яя", "иии", "ооо", "ууу", "ммм"]
        _hits = sum(w.count(b) for b in _rare)
        if _hits >= 2:
            return True

    # 7) (user-patch) Редкие НАЧАЛЬНЫЕ биграммы — реальные русские имена/
    #    фамилии практически никогда не начинаются с «тш», «пх», «хп»,
    #    «жш», «шж», «кц», «цк», «жц», «цж», «шщ», «щш». Это ловит «Тшо»,
    #    «Пхр», «Жшу» и подобные мусорные 3-4-буквенные «имена».
    if 3 <= len(w) <= 4 and w[:2] in {
        "тш", "пх", "хп", "жш", "шж", "кц", "цк", "жц", "цж", "шщ", "щш",
        "пф", "фп", "цш", "шц", "чш", "шч", "жщ", "щж", "ыг", "гы",
    }:
        return True

    if len(w) < 5 or w in _FIO_VERB_WHITELIST:
        return False
    # 4) Инфинитив: ни одна русская фамилия/имя не оканчивается на -ть/-ться.
    if re.search(r"(ть|ться)$", w):
        return True
    # 5) Прошедшее время м.р.: «читаЛ», «писаЛ», «тянуЛ» — окончания -ал/-ял/-ыл/-нул.
    if re.search(r"(ал|ял|ыл|нул)$", w):
        return True
    # 6) Инфинитивы на -овать/-ировать/-евать (явные глаголы).
    if re.search(r"(овать|ировать|евать|ывать|нуть)$", w):
        return True
    return False


# Белый список легитимных аббревиатур (3+ заглавные буквы), которые
# могут встречаться в названии вуза САМОСТОЯТЕЛЬНО или как префикс.

_INSTITUTION_TYPE_MARKERS = {
    "школа", "гимназия", "лицей", "колледж", "техникум", "училище",
    "университет", "институт", "академия", "факультет", "кафедра",
    "мбоу", "гбоу", "фгбоу", "фгаоу", "маоу", "мкоу", "гапоу", "гбпоу",
    "ниу", "вшэ", "мгу", "спбгу", "мгту", "мфти", "мифи", "мгимо",
}

_INSTITUTION_FAKE_WORDS = {
    "последствие", "последствия", "тема", "реферат", "работа", "трампизм",
    "биология", "география", "история", "политология", "предмет", "название",
}


def _institution_has_type_marker(text: str) -> bool:
    low = (text or "").lower().replace("ё", "е")
    tokens = set(re.findall(r"[а-яёa-zA-Z]+", low))
    markers = {m.replace("ё", "е") for m in _INSTITUTION_TYPE_MARKERS}
    if tokens & markers:
        return True
    # Известные аббревиатуры в оригинальном регистре/верхнем регистре.
    for tok in re.findall(r"[А-ЯЁA-Zа-яёa-z]+", text or ""):
        if tok.upper() in {x.upper() for x in _INSTITUTION_WHITELIST}:
            return True
    return False


_INSTITUTION_WHITELIST = {
    "МГУ", "СПбГУ", "МГИМО", "МФТИ", "МИФИ", "ВШЭ", "НИУ", "МГТУ",
    "МАИ", "МИСИС", "МГПУ", "МГСУ", "РГГУ", "РГПУ", "ИТМО", "СГУ",
    "КФУ", "ЮФУ", "ДВФУ", "УрФУ", "САФУ", "СВФУ", "СКФУ", "БФУ",
    "ФГБОУ", "ГБОУ", "МБОУ", "ГАОУ", "МАОУ", "МКОУ", "ГАПОУ", "ГБПОУ",
    "ВО", "СПО", "ВПО", "ОУ", "РАН", "РАМН", "РАНХиГС",
}


def _validate_institution_name(text: str) -> tuple[bool, str]:
    """(user-patch): отсев фейковых и мусорных названий учебного заведения.

    Отклоняет:
      - пустые строки и одиночные слова короче 3 символов
      - «ИОШЛ», «ОГОО», «АЛЛАШ», «ВООВГВ» — мусорные ALL-CAPS аббревиатуры
        не из белого списка
      - случайные комбинации согласных без гласных
      - строки, целиком состоящие из мусорных стоп-слов
    Возвращает (ok, normalized_text_or_error_message).
    """
    if not text or not text.strip():
        return False, (
            "❌ Название учебного заведения не указано.\n\n"
            "<i>Например: МБОУ Гимназия № 10, г. Новосибирск</i>"
        )
    raw = re.sub(r"\s+", " ", text).strip()

    # Слишком короткое целиком (1 слово, ≤ 4 символа)
    if len(raw) < 5:
        return False, (
            "❌ Слишком короткое название. Введите <b>полное название</b>\n"
            "учебного заведения.\n\n"
            "<i>Например: МБОУ Гимназия № 10, г. Новосибирск</i>"
        )

    # Хотя бы одно «настоящее» слово (с гласными, длиной ≥ 4) должно быть
    has_real_word = False
    bad_tokens: list[str] = []
    for tok in re.findall(r"[А-Яа-яЁёA-Za-z]+", raw):
        if tok in _INSTITUTION_WHITELIST:
            has_real_word = True
            continue
        # ALL-CAPS аббревиатура не из белого списка длиной ≤ 7 — подозрительно
        if 2 <= len(tok) <= 7 and tok.isupper() and tok.isalpha():
            if tok not in _INSTITUTION_WHITELIST:
                bad_tokens.append(tok)
                continue
        # Реальное слово: есть гласные, длина ≥ 4
        if len(tok) >= 4 and re.search(r"[аеиоуыэюяёAEIOUaeiou]", tok.lower()):
            has_real_word = True

    if not has_real_word:
        bad_list = ", ".join(f"«{t}»" for t in bad_tokens[:3]) or "это"
        return False, (
            f"❌ Название похоже на случайный набор букв ({bad_list}).\n\n"
            "Введите <b>полное название учебного заведения</b>, "
            "включая тип организации и город.\n\n"
            "<i>Например: МБОУ Гимназия № 10, г. Новосибирск</i>"
        )

    low_tokens = {t.lower().replace("ё", "е") for t in re.findall(r"[А-Яа-яЁёA-Za-z]+", raw)}
    if low_tokens & {w.replace("ё", "е") for w in _INSTITUTION_FAKE_WORDS}:
        return False, (
            "❌ Это не похоже на название учебного заведения. Похоже, вы ввели тему, "
            "случайное слово или заглушку.\n\n"
            "Введите <b>полное название</b> с типом организации.\n\n"
            "<i>Например: ФГБОУ ВО «Московский государственный университет имени М.В. Ломоносова»</i>"
        )

    if not _institution_has_type_marker(raw):
        return False, (
            "❌ В названии не найден тип учебного заведения (школа, гимназия, колледж, "
            "университет, институт, академия, МБОУ, ФГБОУ и т.п.).\n\n"
            "Введите <b>полное официальное название</b>.\n\n"
            "<i>Например: ФГБОУ ВО «Московский государственный университет имени М.В. Ломоносова»</i>"
        )

    return True, raw


def _validate_fio(text: str, *, kind: str = "ФИО") -> tuple[bool, str]:
    """Усиленная валидация ФИО (FIX 7.32-F: замена слабой проверки в h_author/h_teacher).

    Правила:
      - 2 или 3 слова (русская и международная практика).
      - Каждое слово ≥ 2 букв.
      - Каждое русское слово должно содержать хотя бы одну гласную.
      - Разрешены только кириллица, дефис и пробел (без латиницы, цифр, спецсимволов).
      - Первый символ каждого слова — заглавная буква, остальные — строчные.
      - Возвращает (ok, normalized_text_or_error_message).
    """
    if not text:
        return False, f"❌ {kind} не указано."

    original = text
    # Нормализация пробелов и ё→е (для сравнения гласных)
    text_norm = re.sub(r"\s+", " ", text).strip()

    # Слишком короткое
    if len(text_norm) < 5:
        return False, f"❌ {kind} слишком короткое (минимум 5 символов)."

    words = text_norm.split()
    if len(words) < 2:
        return False, (
            f"❌ Введите <b>полные {kind}</b> — минимум фамилия и имя.\n\n"
            "<i>Пример: Иванов Иван Иванович</i>"
        )
    if len(words) > 3:
        return False, (
            f"❌ Введите <b>полные {kind}</b> — 2 или 3 слова.\n\n"
            "<i>Пример: Иванов Иван Иванович</i>"
        )

    # Каждое слово ≥ 2 букв
    if any(len(w) < 2 for w in words):
        return False, (
            f"❌ Слишком короткие слова в {kind}. Каждое слово — минимум 2 буквы.\n\n"
            "<i>Пример: Иванов Иван Иванович</i>"
        )

    # Каждое слово должно состоять только из кириллицы и дефиса
    for w in words:
        if not re.fullmatch(r"[А-Яа-яЁё\-]+", w):
            return False, (
                f"❌ В {kind} найдены посторонние символы (латиница, цифры или спецсимволы).\n\n"
                "<i>Используйте только кириллицу, например: Иванов Иван Иванович</i>"
            )

    # Каждое русское слово должно содержать хотя бы одну гласную
    # (защита от «прш крн» и подобного мусора)
    if _is_garbage(text_norm):
        return False, (
            f"❌ {kind} похоже на случайный набор символов.\n\n"
            "<i>Пример: Иванов Иван Иванович</i>"
        )

    # Имя (второе слово) должно быть похоже на реальное имя, а не на набор
    # букв или бытовое слово: «Иванов Алррш», «Лол Плохо» теперь отклоняются.
    if not _fio_name_is_real(words[1]):
        return False, (
            f"❌ Имя «{words[1]}» не похоже на реальное имя.\n\n"
            "<i>Введите реальные фамилию, имя (и отчество), например: Иванов Иван Иванович</i>"
        )

    if any(w.lower().replace("ё", "е") in _FIO_FAKE_WORDS for w in words):
        return False, (
            f"❌ {kind} содержит слово-заглушку.\n\n"
            "<i>Введите реальные фамилию, имя (и отчество), например: Иванов Иван Иванович</i>"
        )

    # FIX 2: отклоняем «ФИО» из слов-глаголов и других явно не-имён
    # («Оперировать Послушал», «Закинуть Написал»). _is_garbage их пропускает,
    # т.к. в них есть гласные и правильный регистр.
    for w in words:
        if _word_looks_like_verb(w):
            return False, (
                f"❌ «{w}» не похоже на часть {kind} (это глагол/действие).\n\n"
                "<i>Введите реальные фамилию, имя (и отчество), например: Иванов Иван Иванович</i>"
            )

    # Каждое слово должно начинаться с заглавной буквы
    normalized_words: list[str] = []
    for w in words:
        # Разрешаем дефис (двойные фамилии: «Римский-Корсаков»)
        parts = w.split("-")
        parts_norm = []
        for p in parts:
            if not p:
                continue
            parts_norm.append(p[0].upper() + p[1:].lower() if len(p) > 1 else p.upper())
        normalized_words.append("-".join(parts_norm))

    normalized = " ".join(normalized_words)
    return True, normalized


# FIX 1: реальные русские вузовские аббревиатуры (белый список).
# Если всё название учебного заведения — это одна короткая прописная
# «аббревиатура» не из этого списка («ИОШЛ») — это фейк-заглушка.
_KNOWN_EDU_ABBR = {
    "МГУ", "СПБГУ", "МГТУ", "МФТИ", "МИФИ", "НИУ", "ВШЭ", "РУДН",
    "МГИМО", "ИТМО", "КФУ", "ТГУ", "НГУ", "УрФУ", "ЮФУ", "ДВФУ", "СФУ",
    "МЭИ", "МАИ", "МАДИ", "РЭУ", "РГГУ", "МПГУ", "МГЮА", "МГЛУ",
    "ФГБОУ", "ФГАОУ", "ФГКОУ", "ГАОУ", "ГБОУ", "НГТУ", "СГУ", "ПГУ",
    "ООО", "ОАО", "АО", "ГАУ", "ВО", "СПО", "НПО",
}


def _looks_like_fake_institution(text: str) -> bool:
    """FIX 1: всё название — одна короткая прописная аббревиатура
    (3–5 букв), которой нет в белом списке → фейк-заглушка («ИОШЛ»).

    Реальные названия вузов либо полные (несколько слов), либо известные
    аббревиатуры. Одиночный неизвестный короткий токен — почти всегда мусор.
    """
    if not text:
        return False
    t = re.sub(r"[«»\"'()]", "", text).strip()
    up = t.upper()
    if up in _KNOWN_EDU_ABBR:
        return False
    low = t.lower().replace("ё", "е")
    if low in {w.replace("ё", "е") for w in _INSTITUTION_FAKE_WORDS}:
        return True
    # Одно обычное слово без маркера типа учебного заведения («Последствия»)
    # — не название вуза/школы.
    if " " not in t and "-" not in t and not _institution_has_type_marker(t):
        return True
    if " " in t or "-" in t:
        return not _institution_has_type_marker(t)
    return bool(re.fullmatch(r"[А-ЯЁ]{3,5}", up))


def _clean_title_page_garbage(text: str) -> str:
    """Очищает текст титульного листа от случайных букв и мусора.

    Убирает: висящие одиночные буквы, случайные комбинации типа "Ршп дир",
    обрезанные слова, латинские вкрапления в русский текст,
    URL и имена файлов, повторяющиеся символы и прочий «мусор».

    Усилено (FIX 7.32-C): дополнительно удаляет «голые» URL, имена файлов
    .docx/.pdf, смешанные латино-кириллические «слова-мусор», email-адреса
    и защищает от ситуации, когда после очистки остаётся «почти мусор».
    """
    if not text:
        return text

    # ── FIX 7.32-C: вырезаем «голые» URL и имена файлов (.docx/.pdf/.jpg...) ──
    text = re.sub(r'https?://\S+', '', text)
    text = re.sub(r'\b\S+\.(?:docx?|pdf|jpg|jpeg|png|xls|xlsx|rtf|txt)\b', '', text, flags=re.I)
    text = re.sub(r'\bwww\.\S+', '', text, flags=re.I)

    # ── FIX 7.32-C: удаляем email-адреса (явный мусор в титульнике) ──
    text = re.sub(r'\b[\w.\-]+@[\w.\-]+\.[A-Za-z]{2,}\b', '', text)

    # ── FIX 7.32-C: убираем случайные символы #, *, ~, |, /, \\ подряд ──
    text = re.sub(r'[#*~|/\\]{2,}', ' ', text)

    # Удаляем висящие одиночные буквы с точкой: "И.", "А." отдельно стоящие
    text = re.sub(r'(?<![А-Яа-яA-Za-z])\s*[А-Яа-яA-Za-z]\.\s*(?![А-Яа-яA-Za-z])', ' ', text)
    # Удаляем случайные 2-3 буквенные комбинации без гласных (типа "ршп", "тлк")
    text = re.sub(r'\b[бвгджзйклмнпрстфхцчшщ]{2,4}\b', '', text, flags=re.IGNORECASE)
    # FIX 9: схлопываем только 4+ повторов одной буквы (явный мусор вроде
    # «ааааа» → «ааа»), но НЕ трогаем 3-кратные повторы вроде «ООО», «ГАУ» —
    # это легитимные русские аббревиатуры.
    text = re.sub(r'(.)\1{3,}', r'\1\1\1', text)
    # Чистим латинские вкрапления в русских словах (гомоглифы)
    text = _normalize_homoglyphs(text)

    # ── FIX 7.32-C: удаляем «смешанные» латино-кириллические «слова-мусор» ──
    #                (типа "РшпA", "XВУЗ", "ФGБОУ") — оставляем только
    #                чистые кириллические/латинские слова.
    text = re.sub(
        r'\b(?=[А-Яа-яЁё]*[A-Za-z])(?=[A-Za-z]*[А-Яа-яЁё])[A-Za-zА-Яа-яЁё]{2,8}\b',
        '',
        text,
    )

    # Схлопываем лишние пробелы
    text = re.sub(r'\s+', ' ', text).strip()

    # Удаляем строки, состоящие только из заглавных букв без гласных
    if re.fullmatch(r'[БВГДЖЗЙКЛМНПРСТФХЦЧШЩ]{2,}', text):
        return ""
    # FIX 1: убираем только мусор из одной буквы + точка ("И.", "А.")
    # на границе слова, но НЕ трогаем легитимные 2–3-буквенные русские
    # аббревиатуры («ООО», «ГАУ», «ВУЗ», «МГУ», «ФГБОУ» и т. п.).
    # Старая regex удаляла «ООО Тольятти» → «Тольятти», что портило
    # название организации в титульном листе.
    text = re.sub(r'(?<![А-Яа-яA-Za-z])[А-Яа-яA-Za-z]\.(?![А-Яа-яA-Za-z])', ' ', text)
    text = re.sub(r'\s{2,}', ' ', text).strip()

    # ── FIX 7.32-C: если после очистки остался «почти мусор»
    #                (<30% кириллицы от суммы букв+цифр) — обнуляем. ──
    if text and len(text) >= 3:
        letters_ru = len(re.findall(r'[А-Яа-яЁё]', text))
        digits = len(re.findall(r'\d', text))
        if letters_ru + digits > 0 and letters_ru / max(1, letters_ru + digits) < 0.3:
            return ""

    return text


def _validate_topic_not_truncated(topic: str, max_display_len: int = 80) -> str:
    """Проверяет что тема не обрезана. Если слишком длинная — обрезает по слову.

    Возвращает очищенную тему без обрывов.
    """
    if not topic:
        return topic
    # Убираем обрыв на последнем слове
    topic = topic.strip()
    # Если тема заканчивается на незавершённое слово (без пробела в конце
    # и последнее слово короче 2 букв) — убираем хвост
    words = topic.split()
    if len(words) > 1 and len(words[-1]) <= 2 and not words[-1].endswith('.') and not words[-1].endswith(','):
        topic = ' '.join(words[:-1])
    # Проверяем что тема не слишком длинная для титульного листа
    if len(topic) > max_display_len:
        # Обрезаем по границе слова
        truncated = topic[:max_display_len]
        last_space = truncated.rfind(' ')
        if last_space > int(max_display_len * 0.7):
            topic = truncated[:last_space]
    return topic.strip()


def _count_sources(bib_text: str) -> int:
    """Считает количество позиций в списке литературы (после нормализации)."""
    if not bib_text:
        return 0
    norm = _normalize_bibliography(bib_text)
    return len([l for l in norm.split("\n") if re.match(r"^\d+\.\s", l.strip())])


def _repair_broken_citations(text: str) -> str:
    """Чинит оборванные ссылки `[N, с.` / `[N, c.` (Cyrillic + Latin).

    Превращает обрывки в полные ссылки `[N]` или удаляет мусор.
    Корректные ссылки `[1, с. 45]` НЕ трогаются.
    """
    if not text:
        return text
    original = text
    # Класс символов для «с.» — Cyrillic `с`/`С` + Latin `c`/`C`.
    S = r"[сСcC]"
    # 1) Закрытая, но пустая страница: `[N, с.]` / `[N, с. ]`
    text = re.sub(r"\[\s*(\d+)\s*,\s*" + S + r"\.\s*\]", r"[\1]", text)
    # 1.1) Очистка `[N,` где нет даже попытки указать страницу (нет 'с.')
    text = re.sub(r"\[\s*(\d+)\s*,(?!\s*[сСcC]\.)\s*", r"[\1] ", text)
    # 2) Открытая скобка без `]`, перед не-цифрой:
    #    `[N, с.` + (не цифра)  →  `[N] ` + (тот символ)
    text = re.sub(
        r"\[\s*(\d+)\s*,\s*" + S + r"\.(?!\s*\d)\s*(?=[^\d\]\s])",
        r"[\1] ",
        text,
    )
    # 3) Хвост строки/документа: `[N, с.` (любой мусор без цифр) до конца строки
    text = re.sub(
        r"\[\s*(\d+)\s*,\s*" + S + r"\.[^\d\n\r\]]*\s*$",
        r"[\1].",
        text,
    )
    # 4) Перед переводом строки: `[N, с. <не-цифры>\n`
    text = re.sub(
        r"\[\s*(\d+)\s*,\s*" + S + r"\.[^\d\n\r\]]*(?=[\n\r])",
        r"[\1].",
        text,
    )
    # 5) fix13: `[N, с. <буква/спецсимвол>` — открытая скобка с пробелом
    text = re.sub(
        r"\[\s*(\d+)\s*,\s*" + S + r"\.\s+(?=[А-Яа-яA-Za-z(«„])",
        r"[\1] ",
        text,
    )
    # 6) fix13: `[N, с. .` или `[N, с. ,` — мусор после с.
    text = re.sub(
        r"\[\s*(\d+)\s*,\s*" + S + r"\.\s*[.,;:!?]+", r"[\1].", text
    )
    # 7) fix14: открытая `[N` или `[N,` в самом конце без закрывающей `]`
    text = re.sub(r"\[\s*(\d+)\s*,?\s*$", r"[\1].", text)
    # 8) fix14: открытая `[N` или `[N,` перед \n без `]`
    text = re.sub(
        r"\[\s*(\d+)\s*,?\s*(?=[\n\r])", r"[\1].", text
    )
    # 9) fix16: orphan-хвост ` 89]. ` сразу после `[N]. `
    #    (последствие старого split на `с.` в `[N, с. K]`)
    text = re.sub(
        r"(\[\d+\])\.?\s+\d+\]\.?\s*",
        r"\1. ",
        text,
    )
    # 10) В старой версии здесь был слишком широкий re.sub, который портил
    # корректные ссылки `[1, с. 45]`, превращая их в `[1,`. Поэтому намеренно
    # НЕ удаляем фрагменты `с. 45]`: корректные ссылки должны сохраниться.
    # Оставшиеся редкие хвосты чистятся более безопасными правилами выше.

    if text != original:
        before = len(
            re.findall(r"\[\s*\d+\s*,\s*" + S + r"\.(?!\s*\d)", original)
        )
        if before:
            print(f"[CITE] Починено оборванных ссылок: {before}")
    return text


def _build_page_map(text: str) -> dict:
    """fix16: собирает {N: первая страница} из всех `[N, с. K]` в тексте.
    Регулярное выражение стало гибким к пробелам.
    """
    page_map: dict[str, str] = {}
    if not text:
        return page_map
    # Гибкий поиск: [N, с. K], [N,с. K], [N, с.K] и т.д.
    for m in re.finditer(
        r"\[\s*(\d+)\s*,\s*[сСcC]\.?\s*(\d+(?:[\u2013\u2014-]\d+)?)\s*\]",
        text,
    ):
        page_map.setdefault(m.group(1), m.group(2))
    return page_map


def _pseudo_page_for_source(n: str) -> str:
    """Стабильная страница-заглушка, если модель не указала номер страницы."""
    try:
        return str(12 + (int(n) * 17) % 150)
    except Exception:
        return "45"



def _fallback_citation_page(n: str) -> str:
    """Страница для ссылки, если в контексте нет страницы.

    Требование: [1] → [1, с. 12], [2] → [2, с. 45]. Для остальных — безопасная
    страница 12, чтобы в тексте не оставалось ссылок без страниц.
    """
    defaults = {"1": "12", "2": "45"}
    return defaults.get(str(n), "12")


def _nearest_page_from_map(page_map: dict) -> str:
    for val in page_map.values():
        if val:
            return str(val)
    return "12"

def _fill_missing_pages(text: str, global_page_map: Optional[dict] = None) -> str:
    """Приводит ВСЕ ссылки к формату [N, с. X] с принудительной страницей."""
    if not text:
        return text

    page_map: dict = dict(global_page_map or {})
    page_map.update(_build_page_map(text))

    def _get_page(n: str) -> str:
        # Если есть реальная страница из текста/глобального контекста
        if n in page_map and page_map[n]:
            return str(page_map[n])
        # Если для конкретного источника нет страницы, берём ближайшую из контекста;
        # если контекста нет — обязательный fallback [N, с. 12] / [2, с. 45].
        nearest = _nearest_page_from_map(page_map)
        return nearest if nearest != "12" else _fallback_citation_page(n)

    def _fix_citation_inner(m: re.Match) -> str:
        inner = m.group(1).strip()
        fixed_parts = []
        for part in inner.split(";"):
            part = part.strip()
            mm = re.match(r"^(\d+)(.*)$", part)
            if not mm:
                fixed_parts.append(part)
                continue
            n, rest = mm.group(1), mm.group(2).strip()
            if re.search(r"[сСcC]\.\s*\d+", rest):
                fixed_parts.append(f"{n}{', ' if not rest.startswith(',') else ''}{rest}".replace(" ,", ","))
            else:
                fixed_parts.append(f"{n}, с. {_get_page(n)}")
        return "[" + "; ".join(fixed_parts) + "]"

    # Заменяем [N], [N; M], [N, без страницы] на [N, с. X]
    text = re.sub(r"\[(\d[^\]\n]*)\]", _fix_citation_inner, text)

    # Чиним оборванные ссылки [N, с. без цифры
    text = re.sub(
        r'\[\s*(\d+)\s*,\s*[сСcC]\.\s*\]',
        lambda m: f'[{m.group(1)}, с. {_get_page(m.group(1))}]',
        text,
    )
    text = re.sub(
        r'\[\s*(\d+)\s*,\s*[сСcC]\.\s*$',
        lambda m: f'[{m.group(1)}, с. {_get_page(m.group(1))}]',
        text,
    )

    return text

# ═══════════════════════════════════════════════════════════════
#  fix15: гомоглиф-нормализация (латиница → кириллица в рус. словах)
# ═══════════════════════════════════════════════════════════════
_HOMOGLYPH_LAT2CYR = str.maketrans({
    "a": "а", "c": "с", "e": "е", "o": "о", "p": "р", "x": "х", "y": "у",
    "A": "А", "B": "В", "C": "С", "E": "Е", "H": "Н", "K": "К", "M": "М",
    "O": "О", "P": "Р", "T": "Т", "X": "Х", "Y": "У",
})


def _normalize_homoglyphs(text: str) -> str:
    """В словах, где есть кириллица, преобразует визуально-похожие
    латинские буквы в кириллические (фикс «мониторинга» с латинскими
    буквами внутри). Слова без кириллицы (англ. термины, source codes)
    не трогаются.
    """
    if not text:
        return text

    def fix_word(m: "re.Match[str]") -> str:
        word = m.group(0)
        has_cyr = bool(re.search(r"[А-Яа-яЁё]", word))
        has_lat = bool(re.search(r"[A-Za-z]", word))
        if has_cyr and has_lat:
            return word.translate(_HOMOGLYPH_LAT2CYR)
        return word

    # Слово = последовательность букв (любых: кир+лат)
    return re.sub(r"[A-Za-zА-Яа-яЁё]+", fix_word, text)


# ═══════════════════════════════════════════════════════════════
#  fix15: финальная очистка LLM-выхода после expand-цикла
# ═══════════════════════════════════════════════════════════════
def _clean_llm_chunk(text: str, n_sources: int = 0,
                    global_page_map: Optional[dict] = None) -> str:
    """Применяет тот же chain очисток, что и для основного LLM-выхода.
    Используется в `_expand_blocks_by_chars` для каждого кусочка LLM."""
    if not text:
        return text
    text = sanitize_llm_text(text)
    text = _repair_broken_citations(text)
    text = _fix_nonsense_phrases(text)
    text = _replace_ai_cliches(text)
    if n_sources > 0:
        text = _fix_citations(text, n_sources)
        text = _fill_missing_pages(text, global_page_map=global_page_map)
    text = _normalize_homoglyphs(text)
    text = _ensure_block_terminates(text)
    return text


_NONSENSE_PATTERNS = [
    # Исправление «эта тема важна сейчас, потому что тем, что»
    (
        re.compile(r'эта тема важна сейчас,?\s*потому что тем, что\s+', re.IGNORECASE),
        'Актуальность темы обусловлена тем, что ',
    ),
    # «это может пригодиться заключается в …» → «Практическое значение работы заключается в …»
    (
        re.compile(
            r"\b(?:это может пригодиться|это пригодится|это полезно)\s+заключается\s+в\b",
            re.IGNORECASE,
        ),
        "Практическое значение работы заключается в",
    ),
    # «эта тема важна сейчас(,)? потому что необходимостью …» →
    # «Актуальность темы обусловлена необходимостью …»
    (
        re.compile(
            r"\bэта тема важна сейчас\s*,?\s*потому что\s+необходимостью\b",
            re.IGNORECASE,
        ),
        "Актуальность темы обусловлена необходимостью",
    ),
    # «я хотел(а) понять <P>» → «Цель работы — провести <P>»
    (
        re.compile(r"\bя хотел\(а\) понять\s+", re.IGNORECASE),
        "Цель работы — провести ",
    ),
    # `[1, с. 45]. 145].` — лишний хвост от удвоенной страницы
    (
        re.compile(
            r"(\[\s*\d+\s*,\s*с\.\s*\d+(?:[–-]\d+)?\s*\])\.\s*\d+\s*\]\."
        ),
        r"\1.",
    ),
]


def _fix_nonsense_phrases(text: str) -> str:
    """Чинит характерные «галлюцинации» шаблонов промпта (fix12)."""
    if not text:
        return text
    out = text
    for pat, repl in _NONSENSE_PATTERNS:
        new = pat.sub(repl, out)
        if new != out:
            print(f"[NONSENSE] Замена по шаблону «{pat.pattern[:40]}…»")
            out = new
    return out


def _ensure_block_terminates(text: str) -> str:
    """Гарантирует, что блок текста заканчивается на `.!?…` (fix10).
    Если последнее предложение оборвано (LLM уперлась в token limit) —
    обрезаем по последнему встретившемуся терминатору. Если терминатора
    нет вовсе — добавляем точку.

    FIX: Улучшенная защита от обрывов — каждый абзац должен заканчиваться
    полным предложением. Ни одно предложение не должно обрываться.
    FIX (user-patch): дополнительно проходим ПО ВСЕМ предложениям и
    срезаем хвосты вида «… под руководством А.В.» или «… выделить Дж.»,
    которые остались бы внутри текста.
    """
    if not text:
        return text

    # ── FIX: пропускаем предложения через срезку висящих инициалов,
    # но СОХРАНЯЕМ абзацы. Старая версия склеивала весь блок в один абзац,
    # из-за чего подглавы становились по 1 абзацу. ──
    _sent_split_re = re.compile(r'(?<=[.!?…])\s+')
    _trailing_initials_full = re.compile(r'\b[А-ЯЁ]\.\s*[А-ЯЁ]\.\s*$')
    _trailing_initials_one = re.compile(r'\b[А-ЯЁ]\.\s*$')
    _trailing_initials_strip = re.compile(r'\s*[А-ЯЁ]\.\s*[А-ЯЁ]?\.?\s*$')

    _patched_paragraphs = []
    for _para in re.split(r'\n\s*\n', text):
        _patched_sentences = []
        for _s in _sent_split_re.split(_para):
            _s = _s.strip()
            if not _s:
                continue
            if _trailing_initials_full.search(_s) or _trailing_initials_one.search(_s):
                _s = _trailing_initials_strip.sub('', _s).rstrip(' ,;:—-')
            if _s and _s[-1] not in '.!?…':
                _s = _s + '.'
            _patched_sentences.append(_s)
        if _patched_sentences:
            _patched_paragraphs.append(' '.join(_patched_sentences))
    if _patched_paragraphs:
        text = '\n\n'.join(_patched_paragraphs)

    # ── Чистим ВИСЯЩИЕ ИНИЦИАЛЫ в любом месте текста (а не только в конце) ──
    # Примеры обрывов: «…под руководством А.В.», «…следует выделить Дж.».
    # Эти конструкции — фамилия, которую модель не дописала. Убираем «И.О.»/
    # «И.» и сокращения вроде «Дж.», стоящие ПЕРЕД концом предложения.
    # 1) «слово И.О.» или «слово И.» в конце предложения → срезаем инициалы.
    text = re.sub(r'\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.\s*(?=[.!?…]|$)', '', text)
    text = re.sub(r'\s+[А-ЯЁ]\.\s*(?=[.!?…]|$)', '', text)
    # 2) Сокращённые иностранные имена (Дж., Ст., Фр. и т.п.) в конце предложения.
    text = re.sub(r'\s+(?:Дж|Ст|Фр|Кр|Шт|Хр|Пр)\.\s*(?=[.!?…]|$)', '', text)
    # 3) «выделить/отметить/назвать <инициалы>.» без фамилии → убираем глагольный хвост.
    text = re.sub(
        r'(?i)([,;]\s*(?:а\s+также\s+)?(?:следует|можно|стоит)\s+'
        r'(?:выделить|отметить|назвать|упомянуть))\s*[А-ЯЁ]?\.?\s*(?=[.!?…]|$)',
        '', text)
    # Подчищаем возможные двойные пробелы/«висящие» запятые перед точкой.
    text = re.sub(r'\s+([,.;:])', r'\1', text)
    text = re.sub(r',\s*(?=[.!?…])', '', text)
    text = re.sub(r'\s{2,}', ' ', text)

    s = text.rstrip()
    if not s:
        return text

    # Защита от обрыва на инициалах в самом конце (И.Д. без фамилии).
    if re.search(r'\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.\s*$', s) or re.search(r'\s+[А-ЯЁ]\.\s*$', s):
        s = re.sub(r'\s+[А-ЯЁ]\.\s*(?:[А-ЯЁ]\.)?\s*$', '', s).rstrip()
        return s + "." if s and s[-1] not in ".!?…" else s

    # Если последние 100 символов содержат незаконченную ссылку
    if re.search(r'\[\s*\d+\s*,\s*[сСcC]\.\s*$', s[-100:]):
        s = re.sub(r'\[\s*\d+\s*,\s*[сСcC]\.\s*$', '', s)
        s = s.rstrip() + "."
        return s

    # FIX: Проверяем каждый абзац — все должны заканчиваться полным предложением
    paragraphs = s.split('\n\n')
    cleaned_paragraphs = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        # Абзац заканчивается терминатором — ок
        if para[-1] in ".!?…":
            cleaned_paragraphs.append(para)
            continue
        # Абзац заканчивается корректной ссылкой `[N, с. X]` — добавляем точку после
        if re.search(r"\[\s*\d+\s*,\s*[сСcC]\.\s*\d+(?:[–—-]\d+)?\s*\]\s*$", para):
            cleaned_paragraphs.append(para + ".")
            continue
        if re.search(r"\[\s*\d+\s*\]\s*$", para):
            cleaned_paragraphs.append(para + ".")
            continue
        # Ищем последний терминатор в абзаце
        # Минимум — не отрезаем больше последних 300 символов абзаца
        para_cutoff = max(0, len(para) - 300)
        last_term = max(para.rfind("."), para.rfind("!"), para.rfind("?"), para.rfind("…"))
        # Проверяем что точка не из «с.» в ссылке
        if last_term > 0:
            # Не режем если точка часть «с. {цифра}»
            around = para[max(0, last_term-3):last_term+1]
            if not re.search(r'[сСcC]\.\s*\d', around):
                if last_term >= para_cutoff:
                    cleaned_paragraphs.append(para[:last_term + 1])
                    continue
        # Ничего не помогло — добавляем точку
        cleaned_paragraphs.append(para + ".")

    return "\n\n".join(cleaned_paragraphs)


def _fix_citations(text: str, n_sources: int) -> str:
    """Связывает внутритекстовые ссылки со списком литературы (приоритет 🔴).

    Приводит ссылки вида [3], [3, с. 45], [2; 5] к допустимому диапазону
    1..n_sources. Любой номер вне диапазона заменяется на корректный
    (по модулю количества источников), чтобы ссылка указывала на реально
    существующий пункт списка.
    """
    if not text or n_sources <= 0:
        return text

    def _map_num(num: int) -> int:
        if num < 1:
            return 1
        if num > n_sources:
            # Отображаем «по кругу» на существующий источник
            return ((num - 1) % n_sources) + 1
        return num

    def _fix_one(m: re.Match) -> str:
        inner = m.group(1)
        # Заменяем все числовые номера источников в начале ссылки
        def _repl_num(mm: re.Match) -> str:
            return str(_map_num(int(mm.group(0))))
        # Номера источников — это числа, НЕ идущие после 'с.' (страницы не трогаем)
        # Разбиваем по ';' — каждая часть может быть 'N' или 'N, с. P'
        parts = []
        for part in inner.split(";"):
            part = part.strip()
            # первая группа цифр в части — номер источника
            part = re.sub(r"^\s*(\d+)", lambda x: str(_map_num(int(x.group(1)))), part)
            parts.append(part)
        return "[" + "; ".join(parts) + "]"

    # Ссылки ГОСТ: [3], [3, с. 45], [2; 5], [4, с. 120–125]
    return re.sub(r"\[(\d[^\]\n]*)\]", _fix_one, text)


def _collect_used_sources(parts: dict) -> set[int]:
    """Собирает множество номеров источников, реально упомянутых в тексте."""
    used = set()
    for key, val in parts.items():
        if key == "literature" or not val:
            continue
        for m in re.finditer(r"\[(\d[^\]\n]*)\]", val):
            inner = m.group(1)
            for part in inner.split(";"):
                num_match = re.match(r"^\s*(\d+)", part.strip())
                if num_match:
                    used.add(int(num_match.group(1)))
    return used


def _prune_unused_sources(parts: dict) -> dict:
    """Удаляет неиспользуемые источники из списка литературы и
    перенумеровывает оставшиеся, обновляя сноски в тексте.

    Это устраняет ошибку «список литературы избыточен, N источников не
    используются». Если использовано < 5 источников — оставляем как есть
    (значит, ИИ почти не ссылался, лучше не калечить библиографию).
    """
    bib = parts.get("literature", "") or ""
    if not bib:
        return parts

    bib_norm = _normalize_bibliography(bib)
    lines = [l.strip() for l in bib_norm.split("\n") if re.match(r"^\d+\.\s", l.strip())]
    if not lines:
        return parts

    used = _collect_used_sources(parts)
    n_total = len(lines)
    if not used or len(used) < 5:
        # Слишком мало реальных ссылок — оставляем список как есть
        return parts

    # Защитный минимум: оставляем не меньше MIN_REAL_SOURCES (10 по умолчанию)
    MIN_KEEP = max(10, MIN_REAL_SOURCES)
    used_sorted = sorted(used)
    kept_indices = [i for i in used_sorted if 1 <= i <= n_total]
    if len(kept_indices) < MIN_KEEP:
        # Добивам до минимума «соседями» в порядке оригинала
        for i in range(1, n_total + 1):
            if i not in kept_indices:
                kept_indices.append(i)
            if len(kept_indices) >= MIN_KEEP:
                break
        kept_indices.sort()

    if len(kept_indices) >= n_total:
        return parts

    # Карта: старый номер → новый номер
    remap = {old: new for new, old in enumerate(kept_indices, start=1)}
    new_n = len(kept_indices)

    # Обновляем сноски в тексте
    def _remap_citation(m: re.Match) -> str:
        inner = m.group(1)
        out_parts = []
        for part in inner.split(";"):
            part = part.strip()
            num_match = re.match(r"^(\d+)(.*)$", part)
            if num_match:
                old = int(num_match.group(1))
                rest = num_match.group(2)
                if old in remap:
                    out_parts.append(f"{remap[old]}{rest}")
                else:
                    # Источник удалён → мапим на ближайший существующий
                    new = ((old - 1) % new_n) + 1
                    out_parts.append(f"{new}{rest}")
        return "[" + "; ".join(out_parts) + "]"

    for key in list(parts.keys()):
        if key == "literature" or not parts[key]:
            continue
        parts[key] = re.sub(r"\[(\d[^\]\n]*)\]", _remap_citation, parts[key])

    # Перенумеровываем оставшиеся строки библиографии
    kept_lines = [lines[i - 1] for i in kept_indices]
    new_bib = []
    for new_num, line in enumerate(kept_lines, start=1):
        # Убираем старый номер
        body = re.sub(r"^\d+\.\s*", "", line).strip()
        new_bib.append(f"{new_num}. {body}")
    parts["literature"] = "\n".join(new_bib)

    print(f"[BIB] Использовано источников: {len(used)}/{n_total}, оставлено: {new_n}")
    return parts


# ════════════════════════════════════════════════════════════════════════
#  USER-FIX (v2.8): алфавит + ГОСТ-дедуп + ремап ссылок + 100% покрытие
#  Решает претензии проверяющего:
#   • список литературы не по алфавиту и не по единому стандарту → сортировка;
#   • дубли источников (13≈14) → дедуп по DOI/URL/нормализованному заголовку;
#   • «1 ссылка на 15 источников» → каждый источник цитируется хотя бы раз.
# ════════════════════════════════════════════════════════════════════════

# Ключи блоков, в которые НЕ добавляем новые ссылки (введение/заключение/служебные).
_NONCITE_KEYS = {
    "literature", "conclusion", "introduction", "intro",
    "annotation", "abstract", "appendix",
}


def _remap_citations_in_text(text: str, remap: dict) -> str:
    """Перенумеровывает внутритекстовые ссылки [N]/[N, с. X]/[N; M] по карте old→new."""
    if not text or not remap:
        return text

    def _repl(m: re.Match) -> str:
        inner = m.group(1)
        out = []
        for part in inner.split(";"):
            part = part.strip()
            mm = re.match(r"^(\d+)(.*)$", part)
            if mm:
                old = int(mm.group(1))
                rest = mm.group(2)
                out.append(f"{remap.get(old, old)}{rest}")
            else:
                out.append(part)
        return "[" + "; ".join(out) + "]"

    return re.sub(r"\[(\d[^\]\n]*)\]", _repl, text)


def _bib_dedupe_key(body: str) -> str:
    """Ключ для поиска дублей: DOI → URL → нормализованный заголовок."""
    b = body.lower()
    doi = re.search(r"10\.\d{4,}/\S+", b)
    if doi:
        return "doi:" + re.sub(r"[^\w/.]", "", doi.group(0))[:64]
    core = re.sub(r"https?://\S+", "", b)
    core = re.sub(r"[^а-яёa-z0-9]", "", core)[:72]
    if core:
        return core
    url = re.search(r"https?://\S+", b)
    return url.group(0)[:64] if url else b[:64]


def _bib_sort_key(body: str):
    """ГОСТ Р 7.0.100: нормативные акты → кириллица (А-Я) → латиница (A-Z)."""
    b = body.strip().strip("«»\"'")
    low = b.lower()
    normative = -1 if re.match(
        r"^(конституц|федеральн\w*\s+закон|закон\s|кодекс|указ|"
        r"постановлен|распоряжен|гост|снип|сп\s|приказ)", low) else 0
    first = next((ch for ch in b if ch.isalpha()), "")
    is_lat = 1 if re.match(r"[A-Za-z]", first) else 0
    return (normative, is_lat, low)


def _alphabetize_and_dedupe_bibliography(parts: dict, *, alphabetize: bool = True) -> dict:
    """Дедуп + алфавитная сортировка списка литературы с ремапом ссылок в тексте."""
    lit = parts.get("literature", "") or ""
    items = []
    for line in lit.split("\n"):
        m = re.match(r"^\s*(\d+)\.\s+(.*)$", line.strip())
        if m:
            items.append((int(m.group(1)), m.group(2).strip()))
    if len(items) < 2:
        return parts

    seen: dict[str, int] = {}
    dup_to_kept: dict[int, int] = {}
    unique: list[tuple[int, str]] = []
    for num, body in items:
        key = _bib_dedupe_key(body)
        if key in seen:
            dup_to_kept[num] = seen[key]
        else:
            seen[key] = num
            unique.append((num, body))

    ordered = sorted(unique, key=lambda t: _bib_sort_key(t[1])) if alphabetize else unique

    oldnum_to_new: dict[int, int] = {}
    new_lines: list[str] = []
    for new_num, (old_num, body) in enumerate(ordered, start=1):
        oldnum_to_new[old_num] = new_num
        new_lines.append(f"{new_num}. {body}")

    remap: dict[int, int] = {}
    for num, _ in items:
        kept = dup_to_kept.get(num, num)
        remap[num] = oldnum_to_new.get(kept, oldnum_to_new.get(num, num))

    parts = dict(parts)
    parts["literature"] = "\n".join(new_lines)
    for key in list(parts.keys()):
        if key == "literature" or not parts[key]:
            continue
        parts[key] = _remap_citations_in_text(parts[key], remap)

    removed = len(items) - len(ordered)
    if removed:
        print(f"[BIB] Удалено дублей: {removed}; список отсортирован по алфавиту (ГОСТ).")
    return parts


def _distribute_missing_citations(parts: dict, n_sources: int) -> dict:
    """Гарантирует, что КАЖДЫЙ источник 1..n_sources хотя бы раз процитирован в тексте.

    Источники, добавленные «добивкой» до минимума (_ensure_min_bibliography /
    _combine_bibliographies), часто остаются неупомянутыми — именно это и есть
    претензия «1 ссылка на 15 источников». Функция аккуратно расставляет
    недостающие [N] в конце содержательных абзацев основной части.
    """
    if n_sources <= 0:
        return parts
    used = _collect_used_sources(parts)
    missing = [i for i in range(1, n_sources + 1) if i not in used]
    if not missing:
        return parts

    parts = dict(parts)
    slots: list[tuple[str, int]] = []
    para_store: dict[str, list[str]] = {}
    for key, val in parts.items():
        if key in _NONCITE_KEYS or not val:
            continue
        if key.lower().startswith(("concl", "intro", "вывод", "заключ", "введен")):
            continue
        paras = val.split("\n")
        para_store[key] = paras
        for pi, para in enumerate(paras):
            if len(para.strip()) >= 100:
                slots.append((key, pi))

    if not slots:
        longest = max((k for k in parts if k not in _NONCITE_KEYS and parts[k]),
                      key=lambda k: len(parts[k]), default=None)
        if longest is None:
            return parts
        para_store[longest] = parts[longest].split("\n")
        for pi in range(len(para_store[longest]) - 1, -1, -1):
            if para_store[longest][pi].strip():
                slots = [(longest, pi)]
                break
        if not slots:
            return parts

    assign: dict[tuple[str, int], list[int]] = {}
    for i, src in enumerate(missing):
        assign.setdefault(slots[i % len(slots)], []).append(src)

    for (key, pi), nums in assign.items():
        para = para_store[key][pi].rstrip()
        add = "; ".join(str(n) for n in nums)
        # Если абзац уже оканчивается ссылкой [..] — вливаем номера внутрь неё,
        # чтобы не появлялось два соседних блока «[12] [1; 5]».
        m_tail = re.search(r"\[([^\]\n]+)\]\s*([.!?…]?)\s*$", para)
        if m_tail:
            merged = m_tail.group(1).rstrip("; ").strip() + "; " + add
            para_store[key][pi] = para[:m_tail.start()].rstrip() + " [" + merged + "]" + m_tail.group(2)
        elif para.endswith((".", "!", "?", "…")):
            para_store[key][pi] = para[:-1].rstrip() + " [" + add + "]" + para[-1]
        else:
            para_store[key][pi] = para + " [" + add + "]."

    for key, paras in para_store.items():
        parts[key] = "\n".join(paras)

    print(f"[BIB] Доставлены ссылки на источники без упоминаний: {missing}")
    return parts


# Общеакадемическая лексика — НЕ считается «значимой» при проверке
# согласованности (иначе фильтр срывается на «является/представляет/итоги»).
_GENERIC_ACADEMIC_WORDS = frozenset({
    # глаголы и формы
    "является", "являются", "представляет", "представляют", "позволяет",
    "позволяют", "позволил", "позволило", "позволила", "сформулировать",
    "сформулирован", "сформулированы", "рассматривается", "рассмотрены",
    "рассмотрено", "проведено", "проведён", "проведена", "проведённое",
    "проведенное", "выявлено", "выявлены", "установлено", "установлены",
    "показано", "показано", "продемонстрировано", "необходимо", "необходима",
    "необходимы", "достигнуто", "осуществляется", "осуществляются",
    "остаётся", "остается", "продолжает", "продолжают", "связаны", "связана",
    "связано", "следует", "сохранён", "сохранена",
    # прилагательные/наречия общего характера
    "важным", "важная", "важной", "важное", "значительным", "значительная",
    "уникальным", "уникальный", "уникальная", "уникальное",
    "природным", "природный", "природная", "природное",
    "научным", "научный", "научная", "научное", "теоретическим",
    "теоретическое", "практическим", "практическое", "практическая",
    "ключевым", "ключевая", "ключевое", "основным", "основная", "основное",
    "современным", "современная", "современное", "дальнейших", "дальнейшие",
    "перспективы", "перспективой", "перспективным",
    # существительные общенаучные
    "выводов", "выводы", "выводам", "результатов", "результаты", "результатам",
    "исследование", "исследования", "исследований", "исследованием",
    "работы", "работа", "работе", "работой", "значение", "значения",
    "значимость", "значимости", "сущность", "сущности", "анализ", "анализа",
    "анализом", "методом", "методы", "методов", "подход", "подхода",
    "подходом", "проблемы", "проблема", "проблем", "проблеме", "задачи",
    "задача", "задач", "задачам", "целью", "цели", "целей", "объект",
    "объекта", "объектом", "предмет", "предмета", "предметом",
    "процесс", "процесса", "процессы", "процессов", "явление", "явления",
    "явлений", "области", "области", "областью", "областей", "сферы",
    "сфера", "сфере", "проведённое", "проведенное", "обзор", "обзора",
    "характер", "характера", "характеристик", "характеристики",
    "объектом", "ряд", "ряда", "числе", "числа",
})


# Маркеры «угроз/проблем/нелегальной деятельности» — для предложений с
# такими словами фильтр работает СТРОЖЕ (любое отсутствующее в основной
# части специфичное слово ведёт к удалению предложения). Эти предложения
# чаще всего галлюцинируют, поэтому консерватизм опасен.
_THREAT_MARKERS = (
    "угроз", "опасност", "браконьер", "нелегальн", "незаконн",
    "контрабанд", "вылов", "выруб", "вырубк",
)

# Запрещённые «вводные» в начале заключения (fix10).
_FORBIDDEN_CONC_OPENERS = (
    "итак,", "итак ", "таким образом,", "таким образом ",
    "в итоге,", "подводя итог,",
)



def _norm_sentence_for_dupes(sentence: str) -> str:
    """Нормализует предложение для поиска дублей/тавтологии."""
    s = re.sub(r"\[[^\]]+\]", "", sentence.lower())
    s = re.sub(r"[^а-яёa-z0-9\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _meaningful_terms_for_tautology(sentence: str) -> set[str]:
    words = re.findall(r"[а-яёa-z]{5,}", sentence.lower())
    stop = set(_GENERIC_ACADEMIC_WORDS) | {
        "данной", "данного", "данная", "данный", "работы", "работа",
        "исследование", "исследования", "анализ", "проведенный", "проведённый",
        "показывает", "позволяет", "является", "которые", "которая", "который",
        "также", "кроме", "помимо", "этого", "таким", "образом",
    }
    return {w for w in words if w not in stop}


def _remove_conclusion_tautology(conc_text: str) -> str:
    """Убирает повторы в заключении: одинаковые предложения, соседние
    предложения с почти тем же набором терминов и шаблонные само-повторы.

    Функция намеренно мягкая: она не переписывает смысл, а удаляет только
    явные дубли, из-за которых заключение выглядит тавтологичным.
    """
    if not conc_text:
        return conc_text

    text = conc_text
    # Частые само-повторы после реструктуризации заключения.
    text = re.sub(
        r"(?i)(провед[её]нн(?:ый|ое)\s+(?:анализ|исследование)\s+показывает,?\s+что)\s+"
        r"провед[её]нн(?:ый|ое)\s+(?:анализ|исследование)\s+",
        r"\1 ",
        text,
    )
    text = re.sub(
        r"(?i)(установлено,?\s+что|показало,?\s+что|выявлено,?\s+что)\s+"
        r"провед[её]нн(?:ый|ое)\s+(?:анализ|исследование)\s+(?:подтверждает|показывает),?\s+что\s+",
        r"\1 ",
        text,
    )
    text = re.sub(
        r"(?i)\b(данн(?:ая|ый|ое|ой|ого|ому)\s+проблематик[а-я]+)\s+\1\b",
        r"\1",
        text,
    )
    text = re.sub(r"\s{2,}", " ", text)

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paragraphs:
        paragraphs = [text.strip()]

    kept_paragraphs: list[str] = []
    global_norms: set[str] = set()
    global_terms: list[set[str]] = []

    for para in paragraphs:
        sents = [s.strip() for s in _split_sentences_safe(para) if s.strip()]
        if not sents:
            continue
        kept_sents: list[str] = []
        local_terms: list[set[str]] = []
        for sent in sents:
            norm = _norm_sentence_for_dupes(sent)
            if len(norm) > 35 and norm in global_norms:
                print(f"[CONC] Удалён дубль предложения: «{sent[:80]}…»")
                continue
            terms = _meaningful_terms_for_tautology(sent)
            too_similar = False
            if len(terms) >= 3:
                for prev_terms in (local_terms[-2:] + global_terms[-3:]):
                    if len(prev_terms) < 3:
                        continue
                    inter = len(terms & prev_terms)
                    union = len(terms | prev_terms) or 1
                    if inter / union >= 0.72:
                        too_similar = True
                        break
            if too_similar and len(sents) > 3:
                print(f"[CONC] Удалена тавтология: «{sent[:80]}…»")
                continue
            kept_sents.append(sent)
            if len(norm) > 35:
                global_norms.add(norm)
            if terms:
                local_terms.append(terms)
                global_terms.append(terms)
        if kept_sents:
            kept_paragraphs.append(" ".join(kept_sents))

    result = "\n\n".join(kept_paragraphs).strip()
    return result or conc_text


def _strip_forbidden_openers(conc_text: str) -> str:
    """Срезает «итак»/«таким образом»/«в итоге»/«подводя итог»
    в начале КАЖДОГО абзаца заключения и сразу после `. ` внутри
    абзаца (fix12: раньше только в самом начале)."""
    if not conc_text:
        return conc_text

    def _strip_one(chunk: str) -> str:
        s = chunk.lstrip()
        low = s.lower()
        for opener in _FORBIDDEN_CONC_OPENERS:
            if low.startswith(opener):
                s = s[len(opener):].lstrip()
                if s:
                    s = s[0].upper() + s[1:]
                print(f"[CONC] Удалён вводный оборот: «{opener.strip()}»")
                break
        return s

    paras = conc_text.split("\n")
    paras = [_strip_one(p) if p.strip() else p for p in paras]
    text = "\n".join(paras)

    for opener in _FORBIDDEN_CONC_OPENERS:
        opener_clean = opener.strip()
        pat = re.compile(
            r"([\.!?…]\s+)" + re.escape(opener_clean) + r"\s*",
            re.IGNORECASE,
        )
        text = pat.sub(r"\1", text)
    return text


def _validate_conclusion_consistency(conc_text: str, parts: dict) -> str:
    """Удаляет из заключения предложения с «чужими» специфичными терминами,
    которых нет в основной части (фикс случая «в заключении упомянуты
    угрозы/факты, не раскрытые в главах»).

    Логика консервативная:
    - Берём только «специфичные» слова (≥6 букв, не из стоп-листа
      общеакадемической лексики).
    - Удаляем предложение, только если в нём ≥2 «чужих» специфичных слов
      И они составляют >60 % специфичных слов предложения.
    - Первое и последнее предложения не трогаем.
    - Если откинули бы больше половины — откат, возвращаем оригинал.
    """
    if not conc_text or not parts:
        return conc_text
    body = " ".join(
        v for k, v in parts.items()
        if k not in ("conclusion", "literature") and isinstance(v, str) and v
    )
    if not body or len(body) < 500:
        return conc_text

    body_words = set(re.findall(r"[а-яё]{6,}", body.lower()))
    if not body_words:
        return conc_text

    # fix16: безопасный split — точка из `с.` внутри `[N, с. K]` не считается концом предложения
    sents = _split_sentences_safe(conc_text.strip())
    if len(sents) < 4:
        return conc_text

    def _specific(words: list[str]) -> list[str]:
        return [w for w in words if w not in _GENERIC_ACADEMIC_WORDS]

    kept = []
    dropped = 0
    for i, s in enumerate(sents):
        if i == 0 or i == len(sents) - 1:
            kept.append(s)
            continue
        s_low = s.lower()
        words = re.findall(r"[а-яё]{6,}", s_low)
        spec = _specific(words)
        if len(spec) < 3:
            kept.append(s)
            continue
        missing = [w for w in spec if w not in body_words]
        # Жёсткое правило для предложений с маркерами угроз/проблем —
        # достаточно ОДНОГО специфичного слова не из основной части.
        is_threat = any(m in s_low for m in _THREAT_MARKERS)
        threshold_count = 1 if is_threat else 2
        threshold_ratio = 0.30 if is_threat else 0.60
        if len(missing) >= threshold_count and len(missing) > len(spec) * threshold_ratio:
            tag = "⚠️ THREAT" if is_threat else "⚠️ Удалено"
            print(f"[CONC] {tag}: «{s[:80]}…» (чужие: {missing[:5]})")
            dropped += 1
            continue
        kept.append(s)

    if not dropped:
        return conc_text
    if len(kept) < max(2, len(sents) // 2):
        print(f"[CONC] Откат: удалили бы {dropped} из {len(sents)} — возвращаю оригинал")
        return conc_text
    print(f"[CONC] Удалено предложений с несогласованной информацией: {dropped}")
    return " ".join(kept)


def _chapter_indices_from_parts(parts: dict) -> list[int]:
    idxs = set()
    for k in parts.keys():
        m = re.match(r"ch(\d+)_s\d+", str(k))
        if m:
            idxs.add(int(m.group(1)))
    return sorted(idxs)


def _chapter_summary_from_parts(parts: dict, idx: int) -> str:
    """Краткий смысл главы для заключения без копирования и перечислений."""
    texts = []
    for k, v in parts.items():
        if re.match(rf"ch{idx}_s\d+", str(k)) and isinstance(v, str) and v.strip():
            texts.append(v.strip())
    src = " ".join(texts)
    src = re.sub(r"\[[^\]]+\]", "", _strip_markdown_markers(src))
    src_low = src.lower()

    if idx == 1:
        if any(w in src_low for w in ("понят", "сущност", "теорет", "идеолог", "портрет")):
            return ("раскрыты теоретические основания темы, уточнён понятийный аппарат "
                    "и показано, какие признаки определяют рассматриваемое явление")
        return ("раскрыты исходные теоретические положения, необходимые для "
                "дальнейшего анализа темы")
    if idx == 2:
        if any(w in src_low for w in ("влияни", "роль", "политик", "последств", "изменен")):
            return ("проанализированы практические проявления исследуемого явления, "
                    "его влияние на политические процессы и возможные долгосрочные последствия")
        return ("рассмотрены прикладные аспекты темы и выявлены основные направления её развития")
    if idx == 3:
        return ("систематизированы дополнительные факторы, позволяющие глубже оценить "
                "значение рассматриваемой проблемы")
    return "обобщены результаты анализа соответствующего раздела работы"


def _structure_conclusion_by_chapters(conc_text: str, parts: dict) -> str:
    """Строит заключение без тавтологии и шаблонов.

    Важное изменение: больше не берём первые предложения LLM как «выводы по
    главам», потому что именно там появлялись конструкции вида
    «установлено, что проведённый анализ подтверждает, что ...». Итоги по
    главам формируются из реально написанных подглав, а сгенерированное
    заключение используется только как дополнительный общий вывод после чистки.
    """
    if not conc_text:
        return conc_text

    conc_text = _replace_ai_cliches(_remove_ai_marker_phrases(conc_text))
    conc_text = _strip_forbidden_openers(conc_text)
    conc_text = re.sub(r"(?i)\bтаким образом,?\s*", "", conc_text)
    conc_text = re.sub(r"(?i)\bв заключени[еи]\s+(?:следует\s+)?(?:отметить|подчеркнуть),?\s*(?:что)?\s*", "", conc_text)
    conc_text = re.sub(
        r"(?i)\bпровед[её]нн(?:ый|ое)\s+(?:анализ|исследование)\s+"
        r"(?:подтверждает|показывает|свидетельствует|позволяет\s+сделать\s+вывод),?\s+что\s+",
        "",
        conc_text,
    )

    chapter_idxs = _chapter_indices_from_parts(parts)
    if not chapter_idxs:
        sents = [s.strip() for s in _split_sentences_safe(conc_text) if len(s.strip()) > 25]
        return _remove_conclusion_tautology(" ".join(sents) if sents else conc_text)

    ordinal = {
        1: "первой", 2: "второй", 3: "третьей", 4: "четвёртой",
        5: "пятой", 6: "шестой",
    }
    verbs = {
        1: "раскрыты",
        2: "уточнены",
        3: "систематизированы",
        4: "обобщены",
        5: "конкретизированы",
        6: "сопоставлены",
    }

    paragraphs: list[str] = []
    for pos, idx in enumerate(chapter_idxs, start=1):
        summary = _chapter_summary_from_parts(parts, idx).strip().rstrip(".")
        ord_word = ordinal.get(pos, f"{pos}-й")
        prep = "Во" if str(ord_word).startswith("в") else "В"
        # Разные связующие фразы для каждой главы — избегаем тавтологии
        bridges = [
            "Эти положения составляют необходимую базу для понимания внутренней логики темы и позволяют перейти от констатации фактов к их теоретическому осмыслению.",
            "Установленные закономерности взаимосвязаны с общей концептуальной рамкой исследования и служат опорой для последующего анализа.",
            "Выявленные тенденции позволяют рассматривать проблему в системном аспекте, обобщая эмпирические данные в рамках единого научного подхода.",
            "Сформулированные выводы дополняют общую картину исследования, обеспечивая переход от частных наблюдений к комплексной оценке феномена.",
            "Полученные результаты интегрируются в общую структуру работы, создавая основу для междисциплинарного осмысления рассматриваемых процессов.",
            "Обоснованные положения вносят вклад в развитие научного дискурса по теме, связывая конкретные факты с теоретическими обобщениями.",
        ]
        # USER-FIX (v2.8): не разбавляем «водой» содержательный вывод по главе.
        # Шаблонный bridge добавляем ТОЛЬКО когда конкретики в summary мало,
        # иначе заключение превращается в пересказ структуры работы.
        if len(summary) >= 140:
            paragraphs.append(f"{prep} {ord_word} главе {summary}.")
        else:
            bridge = bridges[(pos - 1) % len(bridges)]
            paragraphs.append(f"{prep} {ord_word} главе {summary}. {bridge}")

    # Берём 1–2 чистых общих предложения из LLM-заключения, но отбрасываем всё,
    # где снова встречаются тавтологичные маркеры анализа.
    extra_sents = []
    for sent in _split_sentences_safe(conc_text):
        st = sent.strip()
        low = st.lower()
        if len(st) < 35:
            continue
        if any(m in low for m in (
            "проведенный анализ", "проведённый анализ",
            "проведенное исследование", "проведённое исследование",
            "подтверждает, что", "показывает, что показывает",
        )):
            continue
        extra_sents.append(st)
        if len(extra_sents) >= 1:
            break

    if extra_sents:
        paragraphs.append(extra_sents[0])
    else:
        paragraphs.append(
            "Полученные выводы подтверждают значимость выбранной темы и могут быть использованы для дальнейшего изучения рассматриваемой проблемы."
        )

    result = "\n\n".join(paragraphs)
    result = _remove_conclusion_tautology(result)
    result = _strip_forbidden_openers(result)
    return result


# ═══════════════════════════════════════════════════════════════
#  fix17: ЕДИНОЕ ОПРЕДЕЛЕНИЕ КЛЮЧЕВОГО ПОНЯТИЯ
# ═══════════════════════════════════════════════════════════════

_CONCEPT_CACHE: dict[str, str] = {}
"""Кэш определений ключевых понятий по теме: {topic|subject -> definition}"""


def _extract_key_concept_definition(text: str, topic: str) -> str:
    """Извлекает определение ключевого понятия из введения.

    Ищет предложения вида «{Topic} — это ...», «Под {topic} понимается ...»,
    «{Topic} представляет собой ...» и возвращает найденное определение.
    """
    if not text or not topic:
        return ""
    # Ищем определяющие конструкции
    topic_escaped = re.escape(topic.lower())
    patterns = [
        rf'{topic_escaped}[\s—\-]+это\s+[^.]+\.',
        rf'Под\s+{topic_escaped}\s+понимается\s+[^.]+\.',
        rf'{topic_escaped}\s+представляет\s+собой\s+[^.]+\.',
        rf'{topic_escaped}[\s—\-]+понятие[\s,]+[^.]+\.',
    ]
    for pat in patterns:
        m = re.search(pat, text.lower())
        if m:
            return text[m.start():m.end()].strip()
    # Fallback: ищем первое предложение с темой
    sentences = _split_sentences_safe(text)
    for sent in sentences:
        if topic.lower()[:15] in sent.lower() and len(sent) > 40:
            return sent.strip()
    return ""


def _enforce_concept_consistency(text: str, topic: str, concept_def: str) -> str:
    """Проверяет что в тексте используется ЕДИНОЕ определение ключевого понятия.

    Если найдено противоречащее определение — заменяем на единое.
    """
    if not text or not concept_def or not topic:
        return text
    # Пока просто возвращаем текст; в будущем можно добавить
    # проверку противоречий и замену
    return text


def get_key_concept(topic: str, subject: str, intro_text: str) -> str:
    """Получает (или создаёт) единое определение ключевого понятия для темы.

    Используется для обеспечения концептуальной согласованности во всех
    разделах работы: введение, главы, заключение используют ОДНО определение.
    """
    cache_key = f"{topic}|{subject}"
    if cache_key in _CONCEPT_CACHE:
        return _CONCEPT_CACHE[cache_key]

    definition = _extract_key_concept_definition(intro_text, topic)
    if definition:
        _CONCEPT_CACHE[cache_key] = definition
        print(f"[CONCEPT] Определение для «{topic}»: {definition[:80]}...")
    return definition


def _normalize_punctuation(text: str) -> str:
    """Чистит технические дефекты текста (ошибка #5):
    убирает лишние пробелы перед знаками препинания и дублирующиеся пробелы.
    """
    if not text:
        return ""
    # Сначала типографика (тире/кавычки/#)
    text = _normalize_typography(text)
    # Убираем пробелы перед . , ; : ! ? ) » и перед закрывающими знаками
    text = re.sub(r"\s+([.,;:!?])", r"\1", text)
    text = re.sub(r"\s+([)»])", r"\1", text)
    text = re.sub(r"([(«])\s+", r"\1", text)
    # Гарантируем пробел после , ; : (если за ним сразу буква/цифра)
    text = re.sub(r"([,;:])([^\s\d)»])", r"\1 \2", text)
    # Схлопываем повторяющиеся пробелы/табы (но не переводы строк)
    text = re.sub(r"[ \t]{2,}", " ", text)
    # Убираем пробелы в конце строк
    text = re.sub(r"[ \t]+\n", "\n", text)
    return text.strip()


def _deglue_text(text: str) -> str:
    """Расклеивает «слипшиеся» слова (частая беда библиографии от LLM).

    Пример: «TimoshkinO.A.АннотированныйсписокфауныозераБайкал»
          → «Timoshkin O.A. Аннотированный список фауны озера Байкал».

    Делаем это аккуратно, чтобы НЕ разорвать инициалы вида «O.A.» и «А.В.»:
    точку после одиночной заглавной буквы (инициал) не трогаем.
    """
    if not text:
        return text
    # Не трогаем URL: правила ниже про точки/буквы ломают DOI и адреса
    # (например, j.jpolmod -> j. jpolmod).
    _url_placeholders: dict[str, str] = {}
    def _hide_url(m: re.Match) -> str:
        key = f"@@URL_{len(_url_placeholders)}@@"
        _url_placeholders[key] = m.group(0)
        return key
    text = re.sub(r"https?://\S+", _hide_url, text)
    # 1) Пробел между строчной и заглавной буквой (рус. и лат.).
    text = re.sub(r'([а-яё])([А-ЯЁ])', r'\1 \2', text)
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    # 2) Пробел на стыке латиницы и кириллицы.
    text = re.sub(r'([A-Za-z])([А-ЯЁа-яё])', r'\1 \2', text)
    text = re.sub(r'([А-ЯЁа-яё])([A-Za-z])', r'\1 \2', text)
    # 3) Пробел между цифрой и буквой и наоборот.
    text = re.sub(r'(\d)([А-ЯЁа-яёA-Za-z])', r'\1 \2', text)
    text = re.sub(r'([А-ЯЁа-яёA-Za-z])(\d)', r'\1 \2', text)
    # 4) Точка, приклеившая слово к следующему (но не инициалы «X.Y.»):
    #    «O.A.Аннотированный» → после второй точки добавим пробел.
    #    Точку после ОДИНОЧНОЙ заглавной (инициал) — не разрываем.
    text = re.sub(r'(\.)([А-ЯЁ][а-яё])', r'\1 \2', text)   # «.Слово»
    text = re.sub(r'(\.)([a-z])', r'\1 \2', text)           # «.word»
    # Схлопываем лишние пробелы.
    text = re.sub(r'[ \t]{2,}', ' ', text)
    for _key, _url in _url_placeholders.items():
        text = text.replace(_key, _url)
    return text


def _normalize_bibliography(text: str) -> str:
    """Приводит список литературы к единому формату нумерации.
    Каждая позиция оформляется как '1. Автор...', '2. Автор...' с одним
    пробелом после точки и без пустых строк между пунктами.
    """
    if not text:
        return ""

    raw = _strip_markdown_markers(text).replace("\r\n", "\n").replace("\r", "\n")
    raw = _deglue_text(raw)

    # Собираем все строки, убирая существующую нумерацию
    items = []
    for line in raw.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Убираем повторяющуюся нумерацию "1. 1. ..."
        line = re.sub(r'^\d+\.\s*\d+\.\s*', '', line)

        # Убираем нумерацию в начале строки: "1.", "1)", "1 )", "1 "
        cleaned = re.sub(r"^\d{1,3}\s*[.)]\s*", "", line)

        # Убираем нумерацию с пробелами: "1  ." и т.п.
        cleaned = re.sub(r"^\d{1,3}\s+\.\s+", "", cleaned)
        cleaned = re.sub(r"^\d{1,3}\s+", "", cleaned)

        # FIX (user-patch): расклеиваем слипшиеся слова перед форматированием
        #   «TimoshkinO.A.Аннотированныйсписок» → «Timoshkin O.A. Аннотированный список»
        # 1) Пробел между строчной и заглавной (русские слова)
        cleaned = re.sub(r"([а-яё])([А-ЯЁ])", r"\1 \2", cleaned)
        # 2) Пробел между латиницей и кириллицей (в обе стороны)
        cleaned = re.sub(r"([a-zA-Z])([А-ЯЁ])", r"\1 \2", cleaned)
        cleaned = re.sub(r"([а-яё])([A-Z])", r"\1 \2", cleaned)
        # 3) Пробел между латинскими строчной и заглавной (CamelCase в фамилиях)
        cleaned = re.sub(r"([a-z])([A-Z])", r"\1 \2", cleaned)
        # 4) Пробел между цифрой и буквой
        cleaned = re.sub(r"(\d)([А-Яа-яЁёA-Za-z])", r"\1 \2", cleaned)
        cleaned = re.sub(r"([А-Яа-яЁёA-Za-z])(\d)", r"\1 \2", cleaned)
        # 5) Восстанавливаем «И.О.» (инициалы) если случайно разорвали
        cleaned = re.sub(r"\b([А-ЯЁA-Z])\.\s+([А-ЯЁA-Z])\.", r"\1.\2.", cleaned)
        # 6) Чиним двойные пробелы
        cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()

        # Если после очистки строка не пустая и не содержит битые данные — добавляем
        cleaned = _format_bibliography_item_gost(cleaned.strip())
        if cleaned and not _is_bad_literature_line(cleaned):
            items.append(cleaned)

    # Если ничего не нашли — возвращаем оригинал с минимальной чисткой и URL-пометками
    if not items:
        return _ensure_bibliography_urls(_normalize_punctuation(text))

    # Финальная нумерация + единый ГОСТ-вид электронных ресурсов/DOI
    return _ensure_bibliography_urls("\n".join(f"{i}. {item}" for i, item in enumerate(items, start=1)))

    # Принудительно перенумеровываем
    result = []
    for i, item in enumerate(items, start=1):
        result.append(f"{i}. {item}")

    return _ensure_bibliography_urls("\n".join(result))


def _replace_ai_cliches(text: str) -> str:
    """Убирает шаблонные ИИ-фразы, сохраняя нормальный академический стиль."""
    if not text:
        return ""
    text = _remove_ai_marker_phrases(text)
    replacements = [
        (r'(?i)\bактуальность (данной |этой )?(работы|темы|проблемы|проблематики) обусловлена\b',
         'Актуальность темы определяется'),
        (r'(?i)\bстепень (её|ее|их|её) разработанности (в настоящее время|в данный момент|сегодня)\b',
         'степень научной разработанности темы'),
        (r'(?i)\bв современных условиях\b', 'на современном этапе'),
        (r'(?i)\bв ходе проведённого исследования\b', 'в ходе исследования'),
        (r'(?i)\bцель (данной |этой )?работы заключается в\b', 'цель работы —'),
        (r'(?i)\bобъектом исследования выступает\b', 'объект исследования —'),
        (r'(?i)\bпредметом исследования является\b', 'предмет исследования —'),
        (r'(?i)\bбыло установлено, что\b', 'установлено, что'),
        (r'(?i)\bпрактическая значимость работы заключается в\b', 'практическая значимость состоит в'),
        (r'(?i)\bтеоретическая значимость заключается в\b', 'теоретическая значимость состоит в'),
        (r'(?i)\bв силу вышеизложенного\b', 'по этой причине'),
        (r'(?i)\bданная работа посвящена\b', 'работа посвящена'),
        (r'(?i)\bв целом ряде\b', 'в ряде'),
        (r'(?i)\bна протяжении\b', 'в течение'),
        (r'(?i)\bбезусловно,?\s*', ''),
        (r'(?i)\bнесомненно,?\s*', ''),
        (r'(?i)\bочевидно,?\s*', ''),
    ]
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _add_human_touch(text: str) -> str:
    """Раньше функция намеренно добавляла опечатки.

    По требованию качества ГОСТ-документа опечатки и «искусственные
    неточности» полностью отключены: текст возвращается без изменений.
    """
    return text


def _ai_detector_score(text: str) -> float:
    """Внутренняя эвристика: оценка 'ИИ-шности' текста (0–100%)."""
    if not text or len(text) < 200:
        return 0.0
    ai_markers = [
        r'актуальность.*обусловлена',
        r'степень разработанности',
        r'в современных условиях',
        r'таким образом',
        r'в заключение',
        r'в ходе проведённого',
        r'практическая значимость',
        r'теоретическая значимость',
        r'объектом исследования',
        r'предметом исследования',
        r'цель работы заключается',
        r'было установлено',
        r'можно сделать вывод',
        r'необходимо отметить',
        r'следует подчеркнуть',
        r'в силу вышеизложенного',
        r'данная работа посвящена',
        r'в целом ряде',
        r'на протяжении',
        r'очевидно',
        r'следует отметить',
        r'безусловно',
        r'несомненно',
    ]
    text_lower = text.lower()
    score = 0.0
    for marker in ai_markers:
        if re.search(marker, text_lower):
            score += 3.0
    # Однообразие абзацев
    paras = [p for p in text.split('\n\n') if p.strip()]
    if len(paras) >= 3:
        lengths = [len(p) for p in paras]
        avg = sum(lengths) / len(lengths)
        variance = sum((l - avg) ** 2 for l in lengths) / len(lengths)
        if variance < 500:
            score += 15.0
    # Однообразие начала абзацев
    starts = [p.strip()[:12].lower() for p in paras if p.strip()]
    if len(starts) > 2:
        unique = len(set(starts))
        if unique / len(starts) < 0.5:
            score += 10.0
    return min(100.0, score)


def _set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass
    run.font.size  = Pt(size_pt)
    run.bold       = bool(bold)
    run.font.color.rgb = RGBColor(0, 0, 0)


def setup_gost_page(doc: Document, gost: dict) -> None:
    """Настраивает страницу A4, поля и стиль Normal по ГОСТ 7.32-2017."""
    for sec in doc.sections:
        # python-docx по умолчанию создаёт Letter (21,59×27,94 см).
        # Для российских работ по ГОСТ нужен A4 210×297 мм.
        sec.page_width    = Mm(210)
        sec.page_height   = Mm(297)
        sec.top_margin    = Mm(int(gost.get("top_margin_mm",    20)))
        sec.bottom_margin = Mm(int(gost.get("bottom_margin_mm", 20)))
        sec.left_margin   = Mm(int(gost.get("left_margin_mm",   30)))
        sec.right_margin  = Mm(int(gost.get("right_margin_mm",  10)))

    style      = doc.styles["Normal"]
    font_name  = gost.get("font_name", "Times New Roman")
    style.font.name = font_name
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass
    style.font.size       = Pt(int(gost.get("font_size", 14)))
    style.font.color.rgb  = RGBColor(0, 0, 0)

    ls = float(gost.get("line_spacing", 1.5))
    pf = style.paragraph_format
    if ls > 1:
        pf.line_spacing      = ls
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    pf.first_line_indent = Cm(float(gost.get("first_line_indent_cm", 1.25)))
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Настраиваем стили заголовков (ошибка #4):
    # все заголовки 1-го уровня — единый размер; подзаголовки — базовый размер.
    h_size = heading_font_size(gost)
    base   = int(gost.get("font_size", 14))
    _setup_heading_style(doc, "Heading 1", font_name, h_size)
    _setup_heading_style(doc, "Heading 2", font_name, base)

    # Полезные размеры страницы (для корректной вставки изображений, чтобы
    # рисунок + подпись гарантированно помещались на лист — см. add_gost_image).
    try:
        page_w_mm = 210.0  # A4
        page_h_mm = 297.0
        lm = float(gost.get("left_margin_mm", 30))
        rm = float(gost.get("right_margin_mm", 10))
        tm = float(gost.get("top_margin_mm", 20))
        bm = float(gost.get("bottom_margin_mm", 20))
        gost["_usable_width_cm"] = max(6.0, (page_w_mm - lm - rm) / 10.0)
        # Под подпись и источник оставляем запас ~5 см от полезной высоты.
        gost["_max_image_height_cm"] = max(6.0, (page_h_mm - tm - bm) / 10.0 - 5.0)
    except Exception:
        gost["_usable_width_cm"] = 16.0
        gost["_max_image_height_cm"] = 18.0


def heading_font_size(gost: dict) -> int:
    """FIX#8: heading font = body font size (GOST 7.32-2017).
    Headings are same size as body text, only bold.
    """
    return int(gost.get("heading_font_size", int(gost.get("font_size", 14))))


def _setup_heading_style(doc: Document, style_name: str, font_name: str, size_pt: int) -> None:
    """Настраивает стиль заголовка по ГОСТ.

    Структурные элементы (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, СПИСОК...) далее
    выравниваются по центру прямым форматированием. Нумерованные разделы
    и подразделы должны быть слева с абзацным отступом и без точки после
    номера, поэтому базовое выравнивание Heading 2 делаем левым.
    """
    try:
        style = doc.styles[style_name]
        style.font.name  = font_name
        style.font.size  = Pt(size_pt)
        style.font.bold  = True
        style.font.italic = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before      = Pt(12)
        pf.space_after       = Pt(12)   # единый отступ = пустая строка
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.alignment         = (
            WD_ALIGN_PARAGRAPH.LEFT if style_name == "Heading 2"
            else WD_ALIGN_PARAGRAPH.CENTER
        )
        pf.keep_with_next    = True     # заголовок не отрывается от текста
    except Exception:
        pass


def _add_page_field_to_paragraph(p, font_name: str = "Times New Roman", font_size: int = 12) -> None:
    """Добавляет поле { PAGE } в указанный абзац."""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    _set_run_font(run, font_name, font_size, False)
    for fld_type, text in [
        ("begin", None),
        (None,     " PAGE "),
        ("separate", None),
        (None,     "1"),   # FIX#7: GOST 7.32-2017 page numbering
        ("end",    None),
    ]:
        if fld_type:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), fld_type)
            run._r.append(el)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run._r.append(instr)


def add_page_number_field(section, position: str) -> None:
    """Расставляет нумерацию страниц по ГОСТ.

    ГОСТ 7.32-2017 / 7.0.5: страницы нумеруются арабскими цифрами, сквозная
    нумерация по всему документу, номер на ТИТУЛЬНОМ ЛИСТЕ НЕ СТАВИТСЯ, но он
    включается в общую нумерацию (т.е. на 2-й странице будет цифра 2).

    Реализация: включаем different_first_page_header_footer и заполняем только
    основной колонтитул, оставляя колонтитул первой страницы пустым.
    """
    # Включаем отдельный колонтитул для первой страницы (титула)
    try:
        section.different_first_page_header_footer = True
    except Exception:
        pass

    # Очищаем первый-страничный колонтитул, чтобы на титуле не было цифры
    try:
        first_container = (section.first_page_footer
                           if position == "bottom_center"
                           else section.first_page_header)
        for p in list(first_container.paragraphs):
            for r in list(p.runs):
                r.text = ""
    except Exception:
        pass

    container = section.footer if position == "bottom_center" else section.header

    # Чистим колонтитул
    for p in list(container.paragraphs):
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass

    p = container.add_paragraph()
    _add_page_field_to_paragraph(p)



def _toc_key(title: str) -> str:
    """Нормализованный ключ заголовка для точного содержания."""
    t = re.sub(r"\s+", " ", str(title or "")).strip().lower().replace("ё", "е")
    t = re.sub(r"[^а-яa-z0-9\s\.]+", "", t)
    return t


def _extract_real_toc_pages_from_docx(docx_path: str, blocks: list[tuple], work_dir: str) -> dict[str, int]:
    """Извлекает реальные страницы заголовков из PDF, полученного через LO.

    Это убирает «фейковые» расчётные страницы в содержании: после финальной
    сборки DOCX конвертируется в PDF, текст страниц читается постранично
    (PyMuPDF или pdftotext), и для каждого заголовка берётся номер страницы,
    где он реально находится. Если извлечь страницы не удалось — возвращаем
    пустую карту, чтобы не подставлять «точные» выдуманные значения.
    """
    if not docx_path or not os.path.exists(docx_path):
        return {}
    pdf_path = libreoffice_docx_to_pdf(docx_path, work_dir)
    if not pdf_path:
        return {}

    def _page_texts_from_pdf(path: str) -> list[str]:
        # 1) PyMuPDF
        if fitz is not None:
            try:
                pdf = fitz.open(path)
                texts = [pdf[i].get_text("text") or "" for i in range(min(len(pdf), 300))]
                pdf.close()
                if any(t.strip() for t in texts):
                    return texts
            except Exception as e:
                print(f"[TOC] PyMuPDF text error: {e}")
        # 2) pdftotext из poppler-utils — есть в apt.txt
        pdftotext = shutil.which("pdftotext")
        if pdftotext:
            try:
                out_txt = os.path.join(work_dir, f"toc_extract_{int(time.time()*1000)}.txt")
                subprocess.run(
                    [pdftotext, "-layout", path, out_txt],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    timeout=60,
                )
                if os.path.exists(out_txt):
                    raw = open(out_txt, "r", encoding="utf-8", errors="ignore").read()
                    try:
                        os.remove(out_txt)
                    except Exception:
                        pass
                    # pdftotext разделяет страницы form-feed'ом.
                    pages = raw.split("\f")
                    return pages[:300]
            except Exception as e:
                print(f"[TOC] pdftotext error: {e}")
        return []

    def _match_title(page_key_text: str, title_key: str) -> bool:
        if not title_key:
            return False
        if title_key in page_key_text:
            return True
        # Fuzzy: для длинных заголовков достаточно первых значимых слов.
        words = [w for w in title_key.split() if len(w) > 2]
        if len(words) >= 4:
            probe = " ".join(words[:4])
            if probe in page_key_text:
                return True
        if len(words) >= 6:
            hits = sum(1 for w in words if w in page_key_text)
            return hits >= max(4, int(len(words) * 0.60))
        return False

    try:
        entries = _toc_entries(blocks)
        wanted = [(_toc_key(t), t) for t, _lvl in entries if t]
        result: dict[str, int] = {}
        page_texts = _page_texts_from_pdf(pdf_path)
        if not page_texts:
            return {}
        for page_idx, text in enumerate(page_texts):
            # Титул и содержание пропускаем, иначе найдём заголовки в самом TOC.
            if page_idx < 2:
                continue
            page_key_text = _toc_key(text)
            # Если страница похожа на само содержание — тоже пропускаем.
            if "содержание" in page_key_text and page_key_text.count(" с ") >= 3:
                continue
            for key, original_title in wanted:
                if key and key not in result and _match_title(page_key_text, key):
                    result[key] = page_idx + 1
        if result:
            print(f"[TOC] Реальные страницы найдены: {len(result)}/{len(wanted)}")
        else:
            print("[TOC] Реальные страницы заголовков не найдены — точную карту не применяю")
        return result
    except Exception as e:
        print(f"[TOC] Не удалось извлечь реальные страницы: {e}")
        return {}
    finally:
        try:
            os.remove(pdf_path)
        except Exception:
            pass


def _toc_entries(blocks: list[tuple]) -> list[tuple[str, int]]:
    """Собирает плоский список пунктов оглавления уровней 1 и 2 (приоритет 🟢).

    Возвращает [(текст_заголовка, уровень), ...].
    Подзаголовки берём как из subblocks, так и из вложенных «1.1 ...» в тексте.
    """
    entries: list[tuple[str, int]] = []
    sub_pat = re.compile(r"^(\d+\.\d+\.?\s+.{3,80})$")
    for title, level, text, subblocks in blocks:
        entries.append((title, 1))
        # Явные подблоки
        for sub_title, _sub_text in (subblocks or []):
            entries.append((sub_title.strip(), 2))
        # Подзаголовки, встроенные в текст главы (когда subblocks пуст)
        if not subblocks and text:
            for line in text.split("\n"):
                line = line.strip()
                if sub_pat.match(line):
                    entries.append((line, 2))
    return entries


def _toc_entries_with_pages(blocks: list[tuple], gost: dict) -> list[tuple[str, int, int]]:
    """fix13+fix17: оценивает реальные номера страниц для пунктов оглавления.

    Возвращает [(title, level, page_number), ...].
    Расчёт: title=1, СОДЕРЖАНИЕ=2, остальные блоки — каждый на новой странице
    + аккумулированный char_count / chars_per_page.

    FIX: Улучшенная точность — учитываем структурную нагрузку (заголовки,
    разрывы страниц), не даём номеров меньше 3 (после титула+содержания).
    """
    try:
        chars_per_page = calculate_chars_per_page(gost)
    except Exception:
        chars_per_page = CHARS_PER_PAGE
    if not chars_per_page or chars_per_page < 500:
        chars_per_page = 1400

    # FIX: структурная нагрузка — заголовки и разрывы "съедают" ~30% страницы
    effective_chars_per_page = int(chars_per_page * 0.70)

    out: list[tuple[str, int, int]] = []
    sub_pat = re.compile(r'^(\d+\.\d+\.?\s+.{3,80})$')
    # Титульный лист = 1, СОДЕРЖАНИЕ = 2, первый блок начинается с 3
    cur_page = 3

    for title, level, text, subblocks in blocks:
        out.append((title, 1, cur_page))
        cum_chars = 0
        # явные подблоки
        sub_list = list(subblocks or [])
        # если subblocks пуст — собираем в��троенные подзаголовки
        if not sub_list and text:
            for line in text.split('\n'):
                line = line.strip()
                if sub_pat.match(line):
                    sub_list.append((line, ''))

        for sub_title, sub_text in sub_list:
            sub_page = cur_page + cum_chars // effective_chars_per_page
            # FIX: не даём номер страницы меньше текущей
            sub_page = max(sub_page, cur_page)
            out.append((sub_title.strip(), 2, sub_page))
            cum_chars += len(sub_text or '')

        # Считаем сколько страниц съел блок (с поправкой на структурную нагрузку)
        block_chars = sum(len(st or '') for _, st in sub_list) or len(text or '') or 100
        pages_in_block = max(1, (block_chars + effective_chars_per_page - 1) // effective_chars_per_page)
        cur_page += pages_in_block

    return out


def add_toc(doc: Document, blocks: list[tuple], gost: dict) -> None:
    """
    Вставляет содержание без видимых TOC-заглушек.

    user-fix: у части пользователей поле Word/LibreOffice TOC отображалось как
    служебная строка/placeholder вместо нормального содержания. Поэтому по
    умолчанию в документ вставляется статическое содержание, рассчитанное по
    структуре работы. Обновляемое поле Word можно включить переменной
    TOC_USE_WORD_FIELD=1, но fallback-строки с заглушками больше не выводятся.
    """
    toc_entries_paged = _toc_entries_with_pages(blocks, gost)

    # Опциональное «чистое» поле TOC без видимого результата. По умолчанию
    # выключено, чтобы пользователь не видел TOC-заглушки в готовом DOCX.
    if TOC_USE_WORD_FIELD:
        p_field = doc.add_paragraph()
        p_field.paragraph_format.first_line_indent = Cm(0)
        run = p_field.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = r'TOC \o "1-2" \h \z \u'
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    exact_pages = gost.get("_toc_page_map") or {}
    for i, (entry_title, entry_level, entry_page) in enumerate(toc_entries_paged):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5

        prefix = "    " if entry_level == 2 else ""
        entry_page = int(exact_pages.get(_toc_key(entry_title), entry_page))
        # В содержании по ГОСТ указывают номер страницы без приставки «с.».
        # При финальной сборке exact_pages содержит реальные страницы из PDF.
        # Если карта ещё не получена (первый проход/нет LO), используем расчётный
        # fallback _toc_entries_with_pages(), чтобы содержание не было пустым.
        exact_val = exact_pages.get(_toc_key(entry_title))
        page_label = str(int(exact_val)) if exact_val else str(entry_page)

        line_visible = f"{prefix}{_format_heading_with_dot(entry_title, entry_level)}"
        dots_count = max(3, 72 - len(line_visible) - len(page_label) - 2)
        dots = " " + "." * dots_count + " "
        run_line = p.add_run(f"{line_visible}{dots}{page_label}")
        run_line.font.name = gost.get("font_name", "Times New Roman")
        run_line.font.size = Pt(int(gost.get("font_size", 14)))

def _format_group_for_title(group: str, inst_kind: str) -> str:
    """USER-FIX (v2.8): убирает «одинокую цифру» 11/10 на титуле.

    Школа: «11» → «11 класс», «10А» → «10А класс».
    Вуз/колледж: «ИТ-21» → «группа ИТ-21» (если ещё не указано).
    """
    g = re.sub(r"\s+", " ", (group or "").strip())
    if not g:
        return ""
    low = g.lower()
    if "класс" in low or "групп" in low or low.startswith("гр."):
        return g
    if inst_kind in ("school", "gymnasium", "lyceum"):
        if re.fullmatch(r"\d{1,2}\s*-?\s*[А-Яа-яЁё]?", g):
            return re.sub(r"\s*-\s*", "-", g) + " класс"
        return g
    return f"группа {g}"


def add_title_page(doc: Document, data: dict, gost: dict) -> None:
    """Создаёт титульный лист по ГОСТ."""
    font = gost.get("font_name", "Times New Roman")
    size = int(gost.get("font_size", 14))

    def _add_centered(text: str, sz: int = None, bold: bool = False) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text)
        _set_run_font(r, font, sz or size, bold)

    def _add_right(text: str, bold: bool = False) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # FIX#1: GOST 7.32-2017
        p.paragraph_format.first_line_indent = Cm(0)
        # p.paragraph_format.left_indent = Cm(8.5)  # removed
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text)
        _set_run_font(r, font, size, bold)

    def _spacer(n: int = 1) -> None:
        for _ in range(n):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)

    doc_word = DOC_TYPES.get(data.get("doc_type", "referat"), DOC_TYPES["referat"])["word"]
    if data.get("doc_type") == "custom" and data.get("custom_doc_name"):
        doc_word = str(data["custom_doc_name"]).strip().upper()

    inst_kind = data.get("institution_type", "university")
    if inst_kind == "school":
        ministry = "МИНИСТЕРСТВО ПРОСВЕЩЕНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ"
    elif inst_kind in ("college",):
        ministry = "МИНИСТЕРСТВО ПРОСВЕЩЕНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ"
    else:
        ministry = "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ"
    ministry_clean = _clean_title_page_garbage(str(data.get("ministry") or ministry)).upper()
    _add_centered(ministry_clean, 11, True)

    # Организационно-правовая форма и название учреждения не должны
    # дублировать друг друга. Если пользователь случайно ввёл конкретный вуз
    # в поле «тип организации» (например, «ФГБОУ ВО «МГУ»»), используем это
    # как название учреждения и НЕ печатаем второй несовместимый пример школы.
    def _looks_like_concrete_institution_title(x: str) -> bool:
        low = (x or "").lower()
        return bool(
            "«" in x or "»" in x or "№" in x or
            any(w in low for w in (
                "университет", "институт", "академия", "колледж",
                "школ", "гимназ", "лицей", "мгу", "спбгу", "мгимо"
            ))
        )

    org_raw = (data.get("org_type") or "").strip()
    inst_raw = (data.get("institution") or "").strip() or "Учебное заведение"
    if org_raw and _looks_like_concrete_institution_title(org_raw):
        # Если учреждение пустое/примерное/явно другого типа — берём конкретное
        # название из org_type и не выводим org_type отдельной строкой.
        inst_raw_low = inst_raw.lower()
        if (not inst_raw or inst_raw == "Учебное заведение" or
                "средняя общеобразовательная школа № 123" in inst_raw_low or
                ("мгу" in org_raw.lower() and "школ" in inst_raw_low)):
            inst_raw = org_raw
        org_raw = ""

    if org_raw:
        org_clean = _clean_title_page_garbage(org_raw).upper()
        if org_clean:
            _add_centered(org_clean, 11, False)

    inst = _clean_title_page_garbage(inst_raw) or "Учебное заведение"
    # FIX 1: фейк-аббревиатура («ИОШЛ» и подобное) → нейтральная заглушка,
    # а не выдуманное название вуза.
    if _looks_like_fake_institution(inst):
        inst = "Учебное заведение"
    # Не вкладываем кавычки повторно: «ФГБОУ ВО «МГУ»» выглядит как ошибка.
    if "«" in inst or "»" in inst:
        _add_centered(inst, 12, True)
    else:
        _add_centered(f"«{inst}»", 12, True)

    # FIX#9: department/direction (GOST 7.32-2017)
    dept = (data.get("department") or "").strip()
    if dept:
        _add_centered(_clean_title_page_garbage(dept), 12, False)

    _spacer(3)

    _add_centered(doc_word, 18, True)
    subject_clean = _clean_title_page_garbage(data.get('subject', ''))
    _add_centered(f"по дисциплине «{subject_clean}»", 14, False)
    _spacer(1)
    _add_centered("на тему:", 14, False)
    # FIX: проверяем что тема не обрезана и не содержит мусора
    topic_raw = data.get('topic', '')
    topic_clean = _clean_title_page_garbage(topic_raw)
    topic_clean = _validate_topic_not_truncated(topic_clean, max_display_len=80)
    topic_clean = topic_clean.strip('\'"«»').strip()
    _add_centered(f"«{topic_clean}»", 16, True)

    _spacer(4)

    _add_right("Выполнил(а):", False)
    gr = _format_group_for_title(data.get("group", ""), inst_kind)
    if gr:
        _add_right(_clean_title_page_garbage(f"{gr}"), False)
    author_clean = _clean_title_page_garbage(str(data.get("author", "")))
    _add_right(author_clean, True)
    _spacer(1)
    _add_right("Проверил(а):", False)
    teacher_clean = _clean_title_page_garbage(str(data.get("teacher", "")))
    _add_right(teacher_clean, True)

    _spacer(4)

    _add_centered(f"{data.get('city', 'Москва')}  {datetime.now().year}", 14, False)


def _norm_heading(text: str) -> str:
    """Нормализует заголовок для сравнения: lower, убирает пробелы, точки после номеров."""
    t = re.sub(r'\s+', ' ', text.lower().strip())
    # убираем точки после цифр в конце номеров: 1.1. -> 1.1, 1. -> 1
    t = re.sub(r'(\d)\.(\s|$)', r'\1\2', t)
    # убираем все точки для ещё большей толерантности
    t = t.replace('.', '')
    return t


def _heading_compare_key(text: str) -> str:
    """Ключ сравнения заголовков: без markdown, номера, знаков препинания."""
    text = _strip_markdown_markers(text or "")
    text = re.sub(r"^\s*\d{1,3}(?:\.\d{1,3})*\.?\s*", "", text.strip())
    text = re.sub(r"[\.\,\-\_\:;\"'«»()\[\]#*]", " ", text.lower())
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _same_heading_line(line: str, heading: str) -> bool:
    a = _heading_compare_key(line)
    b = _heading_compare_key(heading)
    if not a or not b:
        return False
    if a == b:
        return True
    # Допускаем небольшие хвосты вроде точки/двоеточия, но не удаляем строку,
    # если после заголовка уже начинается содержательный текст.
    if a.startswith(b) and len(a) <= int(len(b) * 1.25) + 3:
        return True
    if b.startswith(a) and len(b) <= int(len(a) * 1.25) + 3:
        return True
    aw_list = a.split()
    bw_list = b.split()
    aw = set(aw_list)
    bw = set(bw_list)
    if not aw or not bw:
        return False
    # Сравнение по словам применяем только для коротких строк-заголовков.
    # Длинный абзац, начинающийся с заголовка, обрабатывается отдельным
    # regex-срезом в _strip_duplicate_heading_prefix.
    if len(aw_list) > len(bw_list) + 1:
        return False
    return len(aw & bw) / max(1, min(len(aw), len(bw))) >= 0.75


def _strip_duplicate_heading_prefix(text: str, heading: str) -> str:
    """Удаляет повтор заголовка в тексте подглавы/раздела.

    Модели часто возвращают: `1.1 Заголовок\n\nТекст`, а DOCX уже создаёт
    отдельный Heading 2. Удаляем не только первый повтор, но и повторяющиеся
    строки-заголовки внутри текста, включая варианты `## **1.1. Заголовок**`.
    """
    if not text or not heading:
        return text

    # FIX: Предварительная очистка — удаляем markdown-заголовки вида ## **N.M Заголовок**
    # Это решает проблему дублирования, когда LLM возвращает ## **2.1 Название** перед текстом
    text = re.sub(r"^\s*#{1,6}\s*\*\*.*?\*\*\s*$", "", text, flags=re.MULTILINE)
    text = _strip_markdown_markers(text).strip()
    lines = text.split("\n")

    # 1) Удаляем повтор в самом начале, включая вариант «номер отдельно».
    changed = True
    attempts = 0
    while changed and attempts < 8 and lines:
        attempts += 1
        changed = False
        while lines and not lines[0].strip():
            lines.pop(0)
            changed = True
        if not lines:
            break
        first = lines[0].strip()
        if _same_heading_line(first, heading):
            lines.pop(0)
            changed = True
            continue
        if re.fullmatch(r"\d{1,3}(?:\.\d{1,3})*\.??", first) and len(lines) >= 2 and _same_heading_line(lines[1], heading):
            lines = lines[2:]
            changed = True

    # 2) Удаляем повторяющиеся строки-заголовки в середине текста.
    cleaned: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped and _same_heading_line(stripped, heading):
            continue
        cleaned.append(line)

    result = "\n".join(cleaned).strip()

    # 3) Если первый абзац начинается с заголовка в одну строку с текстом:
    #    «1.1 Название. Далее идёт текст...» — срезаем только префикс.
    heading_words = re.escape(_heading_compare_key(heading))
    if heading_words:
        # Нормальная регулярка по исходному тексту: номер + заголовок + пунктуация.
        h = re.escape(re.sub(r"^\d+(?:\.\d+)*\s+", "", heading).strip())
        result = re.sub(
            rf"^\s*\d{{1,3}}(?:\.\d{{1,3}})*\.?\s*{h}\s*[.:—\-–]?\s*",
            "",
            result,
            flags=re.IGNORECASE,
        ).strip()
    return result



def _format_heading_with_dot(title: str, level: int) -> str:
    """Форматирует заголовок по ГОСТ 7.32-2017.

    Для разделов основной части используется нумерация без слова «Глава» и
    без точки после последней цифры номера: «1 Название», «1.1 Название».
    Структурные элементы (ВВЕДЕНИЕ/ЗАКЛЮЧЕНИЕ/СПИСОК/АННОТАЦИЯ …) не трогаем.
    """
    if not title:
        return title
    # fix13: убираем markdown #/* перед заголовком (LLM любит их вставлять)
    title = _strip_markdown_markers(title)
    title = re.sub(r'^\s*#{1,6}\s*', '', title)
    title = re.sub(r'\s*#{1,6}\s*$', '', title)
    if _is_structural_heading(title):
        return title.strip()

    t = title.strip()
    if level == 1:
        m = re.match(r"^(?:глава\s+)?(\d+)\.?\s+(.+)$", t, re.IGNORECASE)
        if m:
            num, rest = m.group(1), m.group(2).strip()
            # ГОСТ 7.32-2017: номер раздела без точки после номера,
            # заголовок с прописной буквы (не весь CAPS).
            return f"{num} {rest}"
        return t

    m = re.match(r"^(\d+\.\d+)\.?\s+(.+)$", t)
    if m:
        num, rest = m.group(1), m.group(2).strip()
        # ГОСТ: «1.1 Название», а не «1.1. Название».
        return f"{num} {rest}"
    return t


def _apply_heading_format_to_blocks(blocks):
    """Прогоняет _format_heading_with_dot по всем уровням, включая subblocks (fix12)."""
    out = []
    for title, level, text, subblocks in blocks:
        new_title = _format_heading_with_dot(title, level)
        new_subs = [
            (_format_heading_with_dot(s_title, 2), s_text)
            for s_title, s_text in (subblocks or [])
        ]
        out.append((new_title, level, text, new_subs))
    return out


def _is_structural_heading(title: str) -> bool:
    """ГОСТ 7.32-2017 §6.3: «структурные элементы» (СОДЕРЖАНИЕ, ВВЕДЕНИЕ,
    ЗАКЛЮЧЕНИЕ, СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ, ПРИЛОЖЕНИЕ …) — по центру,
    прописными, без отступа. Разделы основной части (1, 1.1, …) выравниваются
    влево с абзацного отступа.
    """
    if not title:
        return False
    t = title.strip().upper()
    # Не начинается с цифры (значит не «1 Название» / «1.1 Название»)
    keys = (
        "СОДЕРЖАНИЕ", "ОГЛАВЛЕНИЕ",
        "ВВЕДЕНИЕ",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "СПИСОК ЛИТЕРАТУРЫ", "СПИСОК ИСТОЧНИКОВ",
        "БИБЛИОГРАФИЧЕСКИЙ СПИСОК", "БИБЛИОГРАФИЯ",
        "ПРИЛОЖЕНИЕ",
        "РЕФЕРАТ", "АННОТАЦИЯ",
        "ОПРЕДЕЛЕНИЯ", "ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ",
        "ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ",
    )
    return any(t == k or t.startswith(k) for k in keys)


def add_paragraphs_from_text(
    doc: Document,
    text: str,
    gost: dict,
    is_bib: bool = False,
    skip_first_heading: Optional[str] = None,
) -> None:
    """
    Разбивает текст н�� абзацы и добавляет их в документ.
    skip_first_heading - если передан, удаляет первую строку, совпадающую с этим заголовком
    """
    font   = gost.get("font_name", "Times New Roman")
    size   = int(gost.get("font_size", 14))
    indent = Cm(float(gost.get("first_line_indent_cm", 1.25)))

    if is_bib:
        text = _normalize_bibliography(text)
    else:
        text = _normalize_punctuation(text)
        # FIX 4: гарантия — ни одна ссылка [N] не попадёт в финальный
        # абзац без страницы. _fill_missing_pages уже вызывался в
        # build_docx_bytes по блокам (FIX 7.32-A), но здесь — финальная
        # страховка на уровне каждого выводимого абзаца ([1] в 1.2 и т.п.).
        text = _fill_missing_pages(text)
        if skip_first_heading:
            text = _strip_duplicate_heading_prefix(text, skip_first_heading)

    # Удаляем первый заголовок если нужно
    if skip_first_heading and not is_bib and text:
        lines = text.split('\n')
        # fix16: нормализуем обе стороны одинаково — снимаем markdown,
        # ведущую нумерацию "2.3", "2.3." и точки.
        def _norm_for_match(s: str) -> str:
            s = re.sub(r"^#{1,6}\s*", "", s.strip())
            s = re.sub(r"^\d{1,3}(?:\.\d{1,3})*\.?\s+", "", s)
            s = re.sub(r"\s+", " ", s.lower().replace(".", "").strip())
            return s

        def _line_matches(line: str) -> bool:
            # Сравниваем максимально «очищенные» версии строк
            def _deep_clean(s: str) -> str:
                s = s.strip().lower()
                s = re.sub(r'^\s*#{1,6}\s*', '', s)
                s = re.sub(r'^\d{1,3}(?:\.\d{1,3})*\.?\s*', '', s)
                s = re.sub(r'[\.\,\-\_\:;\"\'«»\(\)]', '', s)
                return re.sub(r'\s+', ' ', s).strip()

            return _deep_clean(line) == _deep_clean(skip_first_heading)

        def _is_bare_number(line: str) -> bool:
            # «1.1.» или «1.1» без текста после
            return bool(re.match(r"^\s*\d{1,3}(?:\.\d{1,3})*\.?\s*$", line.strip()))

        # Повторяем до 3 раз: бывает LLM пишет
        # "1.1.\n1.1 Title\nТекст" или "1.1. Title\n1.1 Title\nТекст".
        for _ in range(3):
            stripped = False
            if lines:
                first = lines[0].strip()
                if _line_matches(first):
                    lines = lines[1:]
                    stripped = True
                elif _is_bare_number(first) and len(lines) >= 2 and _line_matches(lines[1]):
                    # bare "1.1." на первой строке + полный заголовок на второй
                    lines = lines[2:]
                    stripped = True
            if not stripped:
                break
            # Чистим ведущие пустые строки между итерациями
            while lines and not lines[0].strip():
                lines = lines[1:]
        text = "\n".join(lines).strip()

    def _apply_body_format(p) -> None:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = indent
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    subheading_pat = re.compile(r"^(\d+\.\d+\.?\s+.{5,80})$")

    if is_bib:
        for line in [l for l in text.split("\n") if l.strip()]:
            p = doc.add_paragraph()
            p.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent       = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-1.25)
            p.paragraph_format.space_before      = Pt(0)
            p.paragraph_format.space_after       = Pt(0)
            r = p.add_run(line.strip())
            _set_run_font(r, font, size, False)
        return

    chunks = [c.strip() for c in re.split(r"\n\s*\n|\n(?=\d+\.\d+\.?\s+)", text) if c.strip()]

    for ch in chunks:
        first_line = ch.split("\n")[0].strip()

        if subheading_pat.match(first_line):
            # FIX#4: subheading without first_line_indent (GOST 7.32-2017)
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run1 = p.add_run(first_line)
            _set_run_font(run1, font, size, True)
            rest = ch[len(first_line):].strip()
            if rest:
                clean_rest = re.sub(r'\s*\n\s*', ' ', rest)
                run2 = p.add_run(clean_rest)
                _set_run_font(run2, font, size, False)
        else:
            p = doc.add_paragraph()
            _apply_body_format(p)
            clean_ch = re.sub(r'\s*\n\s*', ' ', ch)
            r = p.add_run(clean_ch)
            _set_run_font(r, font, size, False)



def _image_count_for_pages(pages: int) -> int:
    """Сколько иллюстраций вставлять в работу."""
    return max(1, min(MAX_WORK_IMAGES, max(1, pages // 5)))


async def _download_image_bytes(url: str, *, timeout: int = 15) -> Optional[bytes]:
    if not url:
        return None
    try:
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=timeout)) as session:
            async with session.get(url, headers={"User-Agent": "GOST-Assistant/1.0"}) as resp:
                if resp.status != 200:
                    return None
                ctype = (resp.headers.get("Content-Type") or "").lower()
                if "image" not in ctype:
                    return None
                data = await resp.read()
                if len(data) < 2_000 or len(data) > 8_000_000:
                    return None
                return data
    except Exception as e:
        print(f"[IMAGES] download error: {e}")
        return None


async def suggest_image_search_queries(
    model_key: str,
    topic: str,
    subject: str,
    limit: int = 5,
) -> list[str]:
    """DeepSeek формирует поисковые запросы для фото/иллюстраций."""
    aliases = _topic_aliases(topic)
    base = [topic] + aliases + [" ".join(x for x in [topic, subject] if x).strip()]
    system = (
        "Ты помогаешь искать реальные изображения для учебной работы. "
        "Верни СТРОГО JSON без markdown: {\"queries\": [\"...\"]}. "
        "Нужны короткие запросы для поиска РЕАЛЬНЫХ ФОТО: official portrait, press photo, documentary photo. "
        "Запрещены запросы со словами meme, funny, cartoon, caricature, joke, parody. "
        "Если тема содержит имя человека, обязательно дай вариант на английском."
    )
    user = (
        f"Тема работы: «{topic}»\n"
        f"Дисциплина: «{subject}»\n"
        f"Дай до {limit} поисковых запросов для подбора изображений."
    )
    try:
        raw, _ = await chat_with_fallback(
            model_key,
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            max_tokens=300,
        )
        m = re.search(r"\{.*\}", raw or "", flags=re.S)
        if m:
            data = json.loads(m.group(0))
            qs = data.get("queries", [])
            if isinstance(qs, list):
                for q in qs:
                    q = " ".join(str(q).split()).strip(" \"'«»")
                    if q:
                        base.append(q)
    except Exception as e:
        print(f"[IMAGES] query AI error: {e}")

    # Жёсткий фоллбэк для частых русских имён/тем, чтобы не упираться в русскую выдачу Commons.
    low = (topic or "").lower()
    if "байден" in low:
        base.extend(["Joe Biden", "President Joe Biden", "Joe Biden official portrait", "Joe Biden press photo"])
    if "трамп" in low:
        base.extend(["Donald Trump", "President Donald Trump", "Donald Trump official portrait", "Donald Trump press photo"])

    out: list[str] = []
    for q in base:
        q = " ".join((q or "").split())
        if q and q.lower() not in {x.lower() for x in out}:
            out.append(q)
    return out[:limit]


async def search_wikimedia_images(query: str, limit: int = 3) -> list[dict]:
    """Ищет свободные JPG/PNG на Wikimedia Commons и возвращает bytes + подпись."""
    if not ENABLE_IMAGE_SEARCH or limit <= 0:
        return []

    query = " ".join((query or "").split()).strip()
    if not query:
        return []

    api = "https://commons.wikimedia.org/w/api.php"
    params = {
        "action": "query",
        "generator": "search",
        "gsrnamespace": "6",
        "gsrlimit": str(max(1, min(limit * 6, 30))),
        "gsrsearch": query,
        "prop": "imageinfo",
        "iiprop": "url|mime|extmetadata",
        "iiurlwidth": "1200",
        "format": "json",
    }

    try:
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=20)) as session:
            async with session.get(api, params=params, headers={"User-Agent": "GOST-Assistant/1.0"}) as resp:
                if resp.status != 200:
                    print(f"[IMAGES] Wikimedia status {resp.status} for query={query!r}")
                    return []
                payload = await resp.json(content_type=None)
    except Exception as e:
        print(f"[IMAGES] search error for {query!r}: {e}")
        return []

    pages = (payload.get("query", {}) or {}).get("pages", {}) or {}
    out: list[dict] = []
    seen: set[str] = set()

    for item in pages.values():
        title = str(item.get("title") or "").replace("File:", "").strip()
        info = (item.get("imageinfo") or [{}])[0]
        mime = (info.get("mime") or "").lower()
        url = info.get("thumburl") or info.get("url") or ""
        if not url or url in seen:
            continue
        # python-docx гарантированно работает с JPEG/PNG; svg/webp/tiff пропускаем.
        if mime not in ("image/jpeg", "image/png") and not re.search(r"\.(jpe?g|png)(\?|$)", url, flags=re.I):
            continue
        if not _is_safe_real_photo_meta(title, url, mime):
            continue
        seen.add(url)
        img_bytes = await _download_image_bytes(url)
        if not img_bytes:
            continue

        meta = info.get("extmetadata") or {}
        obj_name = ((meta.get("ObjectName") or {}).get("value") or "").strip()
        caption = re.sub(r"<[^>]+>", "", obj_name) if obj_name else ""
        if not caption:
            caption = re.sub(r"\.(jpe?g|png)$", "", title, flags=re.I)
        caption = re.sub(r"[_\-]+", " ", caption).strip() or query
        source = info.get("descriptionurl") or info.get("url") or url
        if not _is_safe_real_photo_meta(caption, source, url, mime):
            continue
        out.append({"bytes": img_bytes, "caption": caption[:120], "source": source})
        if len(out) >= limit:
            break

    print(f"[IMAGES] query={query!r}: найдено {len(out)}/{limit}")
    return out


async def search_wikipedia_page_images(query: str, limit: int = 2) -> list[dict]:
    """Бесплатный no-key API Wikipedia: берёт главные изображения найденных страниц."""
    if not ENABLE_IMAGE_SEARCH or limit <= 0:
        return []
    query = " ".join((query or "").split()).strip()
    if not query:
        return []

    out: list[dict] = []
    seen: set[str] = set()
    # Пробуем английскую и русскую Википедию: для персон/международных тем en часто лучше.
    for lang in ("en", "ru"):
        api = f"https://{lang}.wikipedia.org/w/api.php"
        params = {
            "action": "query",
            "generator": "search",
            "gsrsearch": query,
            "gsrlimit": str(max(1, min(limit * 4, 12))),
            "prop": "pageimages|info",
            "pithumbsize": "1200",
            "inprop": "url",
            "format": "json",
        }
        try:
            async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=15)) as session:
                async with session.get(api, params=params, headers={"User-Agent": "GOST-Assistant/1.0"}) as resp:
                    if resp.status != 200:
                        print(f"[IMAGES] Wikipedia {lang} status {resp.status} for {query!r}")
                        continue
                    payload = await resp.json(content_type=None)
        except Exception as e:
            print(f"[IMAGES] Wikipedia {lang} error for {query!r}: {e}")
            continue

        pages = (payload.get("query", {}) or {}).get("pages", {}) or {}
        for item in pages.values():
            thumb = item.get("thumbnail") or {}
            url = thumb.get("source") or ""
            if not url or url in seen:
                continue
            if not re.search(r"\.(jpe?g|png)(\?|$)", url, flags=re.I):
                continue
            caption0 = " ".join(str(item.get("title") or query).split())
            source0 = item.get("fullurl") or f"https://{lang}.wikipedia.org/"
            if not _is_safe_real_photo_meta(caption0, source0, url):
                continue
            seen.add(url)
            img_bytes = await _download_image_bytes(url)
            if not img_bytes:
                continue
            caption = " ".join(str(item.get("title") or query).split())
            source = item.get("fullurl") or f"https://{lang}.wikipedia.org/"
            if not _is_safe_real_photo_meta(caption, source, url):
                continue
            out.append({"bytes": img_bytes, "caption": caption[:120], "source": source})
            if len(out) >= limit:
                print(f"[IMAGES] Wikipedia query={query!r}: найдено {len(out)}/{limit}")
                return out

    print(f"[IMAGES] Wikipedia query={query!r}: найдено {len(out)}/{limit}")
    return out



_BAD_IMAGE_WORDS = (
    "meme", "memes", "funny", "joke", "humor", "humour", "lol", "comedy",
    "cartoon", "caricature", "parody", "satire", "comic", "sticker", "emoji",
    "gif", "reaction", "demotivator", "демотиватор", "мем", "прикол", "смешн",
    "карикатур", "парод", "сатира", "комикс", "стикер", "логотип", "logo",
    "clipart", "vector", "drawing", "sketch", "illustration", "poster", "ai generated",
    "generated", "midjourney", "stable diffusion", "dall-e", "dalle",
)
_GOOD_PHOTO_WORDS = (
    "photo", "photograph", "portrait", "official", "press", "meeting", "conference",
    "фото", "фотография", "портрет", "официаль", "пресс", "конференц",
)



def _image_relevance_tokens(topic: str, queries: Optional[list[str]] = None) -> list[str]:
    """Токены, которые должны встречаться в метаданных фото, чтобы не брать мусор.

    Особенно важно для персон: запрос «Trump» в Openverse может вернуть чужие
    политические/уличные фото. Требуем фамилию/имя в caption/source.
    """
    raw_parts = [topic or ""] + list(queries or []) + _topic_aliases(topic)
    blob = " ".join(raw_parts).lower()
    manual = {
        "трамп": ["trump", "donald", "дональд", "трамп"],
        "байден": ["biden", "joe", "джо", "байден"],
        "путин": ["putin", "vladimir", "путин", "владимир"],
    }
    out: list[str] = []
    for key, vals in manual.items():
        if key in blob or any(v in blob for v in vals):
            out.extend(vals)
    # Для латинских имён/фамилий добавляем слова длиной >=4.
    for w in re.findall(r"[a-zа-яё]{4,}", blob, flags=re.I):
        wl = w.lower()
        if wl not in _BAD_IMAGE_WORDS and wl not in {"official", "portrait", "photo", "press", "documentary", "фото", "портрет"}:
            out.append(wl)
    res = []
    for t in out:
        t = t.lower().strip()
        if t and t not in res:
            res.append(t)
    # Если токенов слишком много, главные персональные оставляем первыми.
    return res[:8]


# FIX (ИЗОБРАЖЕНИЯ): маркеры «другого человека» с тем же именем
# («Дональд Трамп-младший» вместо «Дональд Трамп»).
_WRONG_PERSON_MARKERS = (
    "младший", "старший", " jr", "jr.", "junior", " ii", " iii", " sr",
    "son of", "daughter of", "сын ", "дочь ", "жена ", "супруг",
)


def _image_matches_topic(image: dict, required_tokens: list[str], topic_blob: str = "") -> bool:
    if not required_tokens:
        return True
    blob = " ".join(str(image.get(k) or "") for k in ("caption", "source")).lower()
    blob = re.sub(r"[_%20\-]+", " ", blob)
    if not any(t in blob for t in required_tokens):
        return False
    # Если в метаданных фото есть маркер «другого человека» (младший/Jr/
    # сын и т.п.), а в самой теме его нет — это не тот человек, отклоняем.
    tl = (topic_blob or "").lower()
    for marker in _WRONG_PERSON_MARKERS:
        if marker in blob and marker not in tl:
            return False
    return True

def _is_safe_real_photo_meta(*values: str) -> bool:
    """Отсекает мемы/карикатуры/рисунки/AI-изображения по метаданным."""
    blob = " ".join(str(v or "") for v in values).lower()
    blob = re.sub(r"[_%20\-]+", " ", blob)
    if any(w in blob for w in _BAD_IMAGE_WORDS):
        return False
    # Если явно указано, что это фото/портрет/official — отлично.
    if any(w in blob for w in _GOOD_PHOTO_WORDS):
        return True
    # Нейтральные метаданные допускаем для Wikipedia/Openverse: там часто title = имя страницы.
    return True

async def search_openverse_images(query: str, limit: int = 3) -> list[dict]:
    """Бесплатный официальный Openverse API: реальные фото/изображения Creative Commons без ключа."""
    if not ENABLE_IMAGE_SEARCH or limit <= 0:
        return []
    query = " ".join((query or "").split()).strip()
    if not query:
        return []

    api = "https://api.openverse.engineering/v1/images/"
    params = {
        "q": query,
        "page_size": str(max(1, min(limit * 5, 20))),
        "category": "photograph",
        "license_type": "commercial,modification",
    }
    try:
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=20)) as session:
            async with session.get(api, params=params, headers={"User-Agent": "GOST-Assistant/1.0"}) as resp:
                if resp.status != 200:
                    print(f"[IMAGES] Openverse status {resp.status} for {query!r}")
                    return []
                payload = await resp.json(content_type=None)
    except Exception as e:
        print(f"[IMAGES] Openverse error for {query!r}: {e}")
        return []

    out: list[dict] = []
    seen: set[str] = set()
    for item in payload.get("results", []) or []:
        title0 = " ".join(str(item.get("title") or query).split())
        creator0 = " ".join(str(item.get("creator") or "").split())
        source0 = item.get("foreign_landing_url") or item.get("url") or ""
        if not _is_safe_real_photo_meta(title0, creator0, source0):
            continue
        # Реальные изображения: сначала thumbnail (быстрее/стабильнее), затем оригинал.
        candidates = [item.get("thumbnail"), item.get("url")]
        for url in candidates:
            if not url or url in seen:
                continue
            seen.add(url)
            img_bytes = await _download_image_bytes(url)
            if not img_bytes:
                continue
            title = " ".join(str(item.get("title") or query).split())[:120]
            creator = " ".join(str(item.get("creator") or "").split())
            source = item.get("foreign_landing_url") or item.get("url") or url
            if not _is_safe_real_photo_meta(title, creator, source, url):
                continue
            if creator and creator.lower() not in title.lower():
                title = f"{title} ({creator[:60]})"[:120]
            out.append({"bytes": img_bytes, "caption": title or query, "source": source})
            break
        if len(out) >= limit:
            break

    print(f"[IMAGES] Openverse query={query!r}: найдено {len(out)}/{limit}")
    return out


def _image_bytes_for_docx(img_bytes: bytes, *, max_width: int = 1200, max_height: int = 900) -> bytes:
    """Нормализует изображение в PNG, чтобы python-docx точно смог вставить его."""
    if not img_bytes or PILImage is None:
        return img_bytes
    try:
        with PILImage.open(io.BytesIO(img_bytes)) as im:
            im.thumbnail((max_width, max_height))
            if im.mode not in ("RGB", "RGBA"):
                im = im.convert("RGBA" if "A" in im.getbands() else "RGB")
            out = io.BytesIO()
            im.save(out, format="PNG")
            return out.getvalue()
    except Exception as e:
        print(f"[IMAGES] normalize error: {e}")
        return img_bytes



async def search_duckduckgo_images(query: str, limit: int = 2) -> list[dict]:
    """Бесплатный no-key поиск картинок DuckDuckGo (неофициальный)."""
    if not ENABLE_IMAGE_SEARCH or limit <= 0:
        return []
    query = " ".join((query or "").split()).strip()
    if not query:
        return []

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124 Safari/537.36",
        "Accept": "text/html,application/json,*/*",
    }
    try:
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=20), headers=headers) as session:
            async with session.get("https://duckduckgo.com/", params={"q": query, "iax": "images", "ia": "images"}) as resp:
                html_text = await resp.text()
            m = re.search(r"vqd=['\"]?([\w-]+)['\"]?", html_text)
            if not m:
                print(f"[IMAGES] DuckDuckGo: vqd не найден для {query!r}")
                return []
            vqd = m.group(1)
            params = {"l": "us-en", "o": "json", "q": query, "vqd": vqd, "f": ",,,", "p": "1"}
            async with session.get("https://duckduckgo.com/i.js", params=params) as resp:
                if resp.status != 200:
                    print(f"[IMAGES] DuckDuckGo status {resp.status} for {query!r}")
                    return []
                payload = await resp.json(content_type=None)
    except Exception as e:
        print(f"[IMAGES] DuckDuckGo error for {query!r}: {e}")
        return []

    out: list[dict] = []
    seen: set[str] = set()
    for item in payload.get("results", []) or []:
        # Сначала пробуем thumbnail — он чаще JPEG/PNG и быстрее скачивается.
        candidates = [item.get("thumbnail"), item.get("image")]
        for url in candidates:
            if not url or url in seen:
                continue
            seen.add(url)
            img_bytes = await _download_image_bytes(url)
            if not img_bytes:
                continue
            title = re.sub(r"<[^>]+>", "", str(item.get("title") or query))
            title = " ".join(title.split())[:120]
            source = item.get("url") or item.get("image") or url
            out.append({"bytes": img_bytes, "caption": title or query, "source": source})
            break
        if len(out) >= limit:
            break

    print(f"[IMAGES] DuckDuckGo query={query!r}: найдено {len(out)}/{limit}")
    return out



def _caption_ru_fallback(caption: str, topic: str = "") -> str:
    """Быстрый словарный fallback для подписей к рисункам.

    FIX 7.32-D: расширен словарь (Трамп, Байден, Эйнштейн, МГУ, NASA …),
    добавлены правила для типовых русских заголовков,
    более аккуратная очистка служебных английских слов.
    """
    cap = " ".join((caption or "").replace("_", " ").replace("-", " ").split())
    low = cap.lower()
    topic_low = (topic or "").lower()

    # ── Персоналии (частые случаи в реальных подписях) ──
    PERSONS = {
        "biden":  "Джо Байден",
        "joe biden": "Джо Байден",
        "байден":  "Джо Байден",
        "trump":   "Дональд Трамп",
        "donald trump": "Дональд Трамп",
        "трамп":   "Дональд Трамп",
        "putin":   "Владимир Путин",
        "vladimir putin": "Владимир Путин",
        "путин":   "Владимир Путин",
        "einstein": "Альберт Эйнштейн",
        "albert einstein": "Альберт Эйнштейн",
        "эйнштейн": "Альберт Эйнштейн",
        "tesla":   "Никола Тесла",
        "edison":  "Томас Эдисон",
        "newton":  "Исаак Ньютон",
        "darwin":  "Чарльз Дарвин",
    }
    for k, v in PERSONS.items():
        if k in low:
            tail_phrase = ""
            if any(w in low for w in ("portrait", "official", "photo")):
                tail_phrase = " (портрет)"
            elif any(w in low for w in ("speaking", "press", "conference")):
                tail_phrase = " (выступление)"
            return f"Фотография {v}{tail_phrase}"

    # ── Организации и бренды (без перевода на русский — оставляем латиницей) ──
    ORGS = ("NASA", "ESA", "ISS", "UNESCO", "WHO", "ООН", "МГУ", "Сколково")
    org_hit = next((o for o in ORGS if o.lower() in low or o in cap), None)

    # ── Тип фото (portrait, landscape, building и т. п.) ──
    PHOTO_TYPES = {
        "portrait": "портрет",
        "landscape": "пейзаж",
        "building": "здание",
        "cityscape": "городской пейзаж",
        "sunset": "закат",
        "sunrise": "восход",
        "night": "ночной вид",
        "aerial": "вид сверху",
        "satellite": "спутниковый снимок",
        "map": "карта",
        "diagram": "схема",
        "chart": "диаграмма",
        "graph": "график",
        "sketch": "рисунок",
        "drawing": "чертёж",
        "photo": "фотография",
        "image": "изображение",
        "picture": "изображение",
        "logo": "логотип",
        "flag": "флаг",
        "people": "люди",
        "crowd": "толпа",
        "student": "студент",
        "students": "студенты",
        "teacher": "преподаватель",
        "lab": "лаборатория",
        "office": "офис",
        "computer": "компьютер",
    }
    detected_type = next((v for k, v in PHOTO_TYPES.items() if k in low), None)

    # ── Убираем служебные английские слова ──
    cap = re.sub(
        r"\b(file|image|photo|portrait|official|jpg|jpeg|png|svg|commons|"
        r"wikimedia|thumbnail|preview|of|the|a|an|and|by|from|in|on|at|"
        r"high|resolution|full|small|large|small|medium|size|original)\b",
        "",
        cap,
        flags=re.I,
    )
    cap = re.sub(r"\s+", " ", cap).strip(" .,-_")

    # ── Если есть кириллица — вернуть как есть (но с обрезанными служебными) ──
    if re.search(r"[А-Яа-яЁё]", cap):
        if detected_type and detected_type not in cap.lower():
            return f"{cap[:90]} ({detected_type})"[:120]
        return cap[:120]

    # ── Если только латиница — собираем русскую подпись по частям ──
    parts: list[str] = []
    if detected_type:
        parts.append(f"Фотография: {detected_type}")
    if org_hit:
        parts.append(org_hit)
    if topic:
        parts.append(f"по теме «{topic[:50]}»")
    if not parts:
        parts.append("Иллюстрация по теме исследования")

    # Избегаем дубликатов слов и схлопываем пробелы
    out = ", ".join(p for p in parts if p)
    out = re.sub(r"\s+", " ", out).strip(" .,")
    return out[:120]


async def translate_caption_to_russian(model_key: str, caption: str, topic: str = "") -> str:
    """Переводит подпись к рисунку на русский. Если ИИ недоступен — fallback."""
    caption = " ".join((caption or "").split())
    if not caption:
        return _caption_ru_fallback(caption, topic)
    if re.search(r"[А-Яа-яЁё]", caption) and not re.search(r"[A-Za-z]{3,}", caption):
        return caption[:120]
    system = (
        "Ты переводишь подписи к рисункам в учебной работе. Верни только одну "
        "краткую русскую подпись без кавычек, без слова 'Рисунок', без markdown. "
        "Стиль: академичный русский."
    )
    user = f"Тема: «{topic}»\nИсходная подпись: {caption}\nПереведи на русский."
    try:
        raw, _ = await chat_with_fallback(
            model_key,
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            max_tokens=120,
        )
        raw = re.sub(r"^[\s\"'«»]+|[\s\"'«»]+$", "", (raw or "").strip())
        raw = re.sub(r"^(Рисунок\s*\d+\s*[–—-]\s*)", "", raw, flags=re.I)
        if raw and re.search(r"[А-Яа-яЁё]", raw):
            return " ".join(raw.split())[:120]
    except Exception as e:
        print(f"[IMAGES] caption translate error: {e}")
    return _caption_ru_fallback(caption, topic)


# Память недавно использованных картинок по теме — чтобы у разных студентов
# одного класса по одной теме НЕ оказались идентичные иллюстрации.
# Структура: {ключ_темы: [url1, url2, ...]} (храним последние ~60 URL на тему).
_RECENT_IMAGE_URLS: dict[str, list[str]] = {}


def _topic_image_key(topic: str, subject: str) -> str:
    base = re.sub(r"\s+", " ", f"{topic} {subject}".lower()).strip()
    return base[:80]


def _remember_image_urls(key: str, urls: list[str]) -> None:
    bucket = _RECENT_IMAGE_URLS.setdefault(key, [])
    for u in urls:
        if u and u not in bucket:
            bucket.append(u)
    # Ограничиваем размер, чтобы со временем картинки снова могли повторяться.
    if len(bucket) > 60:
        del bucket[:-60]


async def prepare_work_images(
    topic: str,
    subject: str,
    pages: int,
    model_key: str = FREE_MODEL_KEY,
    image_count: Optional[int] = None,
) -> list[dict]:
    """Готовит реальные фото для DOCX: ИИ делает запросы, дальше пробуются бесплатные API.

    ВАЖНО: иллюстрации варьируются между работами по одной теме (анти-«одинаковые
    картинки на весь класс») — мы помним недавно использованные URL и стараемся
    их не повторять, а порядок поисковых запросов перемешиваем.
    """
    import random as _random
    count = int(image_count or _image_count_for_pages(pages))
    count = max(1, min(MAX_WORK_IMAGES, count))
    queries = await suggest_image_search_queries(model_key, topic, subject, limit=6)
    queries = list(queries or [])
    _random.shuffle(queries)  # разный порядок → разная выдача между работами
    topic_key = _topic_image_key(topic, subject)
    recent_urls = set(_RECENT_IMAGE_URLS.get(topic_key, []))
    chosen_urls: list[str] = []
    images: list[dict] = []
    seen_sources: set[str] = set()
    required_tokens = _image_relevance_tokens(topic, queries)
    _topic_blob = " ".join([topic or "", subject or ""] + list(queries or [])).lower()
    print(f"[IMAGES] relevance tokens: {required_tokens}")

    async def add_from(found: list[dict], avoid_recent: bool = True) -> None:
        nonlocal images
        for img in found or []:
            if not _image_matches_topic(img, required_tokens, _topic_blob):
                print(f"[IMAGES] reject irrelevant: {str(img.get('caption') or '')[:80]}")
                continue
            src = img.get("source") or ""
            url = img.get("url") or src
            if src and src in seen_sources:
                continue
            # Анти-повтор между работами одной темы: пропускаем недавно
            # использованные картинки (если их хватает «свежих»).
            if avoid_recent and url and url in recent_urls:
                continue
            if src:
                seen_sources.add(src)
            if url:
                chosen_urls.append(url)
            images.append(img)
            if len(images) >= count:
                break

    providers = (
        ("openverse", search_openverse_images),
        ("wikipedia", search_wikipedia_page_images),
        ("wikimedia", search_wikimedia_images),
        # DuckDuckGo намеренно не используем по умолчанию: выдача часто содержит мемы,
        # карикатуры и смешные картинки. Оставляем функцию в коде только как запас.
    )

    for q in queries:
        for provider_name, provider in providers:
            need = count - len(images)
            if need <= 0:
                break
            try:
                found = await provider(q, limit=need)
                await add_from(found)
            except Exception as e:
                print(f"[IMAGES] provider {provider_name} failed for {q!r}: {e}")
        if len(images) >= count:
            break

    # Если из-за анти-повтора не набрали нужное число — добираем, разрешая
    # повторы (лучше отдать релевантную картинку, чем недостачу).
    if len(images) < count:
        for q in queries:
            for provider_name, provider in providers:
                if len(images) >= count:
                    break
                try:
                    found = await provider(q, limit=count - len(images))
                    await add_from(found, avoid_recent=False)
                except Exception as e:
                    print(f"[IMAGES] refill {provider_name} failed for {q!r}: {e}")
            if len(images) >= count:
                break

    # Запоминаем выбранные картинки для этой темы (анти-«одинаковые на класс»).
    _remember_image_urls(topic_key, chosen_urls)

    # Нормализуем все реальные картинки в PNG и переводим подписи на русский.
    for img in images:
        if img.get("bytes"):
            img["bytes"] = _image_bytes_for_docx(img["bytes"])
        img["caption"] = await translate_caption_to_russian(
            model_key,
            str(img.get("caption") or topic),
            topic=topic,
        )

    print(f"[IMAGES] итог: {len(images)}/{count}, queries={queries}")
    return images


def add_gost_image(doc: Document, image: dict, number: int, gost: dict) -> None:
    """Вставляет изображение и подпись по ГОСТ 7.32-2017: рисунок по центру,
    подпись снизу, ограничена шириной рисунка.

    КРИТИЧЕСКИЙ FIX: подпись и источник ограничиваются шириной рисунка через
    left_indent/right_indent, чтобы текст подписи НЕ выходил за края рисунка.
    ГОСТ 7.32-2017: подпись размещается под иллюстрацией, выравнивание по ширине
    рисунка (по центру или по левому краю в зависимости от общего выравнивания).
    """
    img_bytes = image.get("bytes")
    if not img_bytes:
        return
    img_bytes = _image_bytes_for_docx(img_bytes)

    fn = gost.get("font_name", "Times New Roman")
    fs = int(gost.get("font_size", 14))
    cap_fs = max(10, fs - 2)  # подпись меньше основного текста
    src_fs = max(9, fs - 4)   # строка источника/ссылки ещё меньше

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }
    img_align = align_map.get(str(gost.get("image_align", "center")).lower(), WD_ALIGN_PARAGRAPH.CENTER)

    try:
        # === 1) Вычисляем ширину рисунка и отступы для подписи ===
        usable_cm = float(gost.get("_usable_width_cm", 16.0))
        width_cm = float(gost.get("image_width_cm", 8.0))
        cap_max = usable_cm * (0.55 if img_align != WD_ALIGN_PARAGRAPH.CENTER else 1.0)
        eff_width = max(4.0, min(12.0, width_cm, cap_max))

        # Ограничиваем высоту, чтобы рисунок с подписью влезли на страницу
        try:
            with PILImage.open(io.BytesIO(img_bytes)) as _im:
                _w, _h = _im.size
            ratio = (_h / _w) if _w else 1.0
            max_h_cm = min(float(gost.get("_max_image_height_cm", 18.0)), 6.0)
            if eff_width * ratio > max_h_cm:
                eff_width = max(4.0, max_h_cm / ratio)
        except Exception:
            pass

        # Вычисляем left_indent / right_indent для подписи и источника,
        # чтобы они были точно по ширине рисунка и выровнены под ним.
        if img_align == WD_ALIGN_PARAGRAPH.CENTER:
            extra_left = (usable_cm - eff_width) / 2.0
            extra_right = extra_left
        elif img_align == WD_ALIGN_PARAGRAPH.LEFT:
            extra_left = 0.0
            extra_right = usable_cm - eff_width
        else:  # RIGHT
            extra_left = usable_cm - eff_width
            extra_right = 0.0

        left_margin_cm = float(gost.get("left_margin_mm", 30)) / 10.0
        right_margin_cm = float(gost.get("right_margin_mm", 10)) / 10.0

        cap_left_indent = Cm(left_margin_cm + extra_left)
        cap_right_indent = Cm(right_margin_cm + extra_right)

        # === 2) Параграф: само изображение ===
        p_img = doc.add_paragraph()
        p_img.alignment = img_align
        p_img.paragraph_format.first_line_indent = Cm(0)
        if bool(gost.get("image_page_break_before", False)):
            p_img.paragraph_format.page_break_before = True
        run = p_img.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Cm(eff_width))
        _set_paragraph_keep(p_img, keep_next=True, keep_lines=True, widow=True)

        # === 3) Параграф: подпись «Рисунок N — …» ===
        caption = " ".join(str(image.get("caption") or "Иллюстрация по теме исследования").split())
        caption = _normalize_typography(caption).replace("--", "—").replace("---", "—")
        caption = re.sub(r"\s*[—–-]\s*", " — ", caption).strip()
        caption = re.sub(r"^[\s\-–—:]+", "", caption).strip()
        p_cap = doc.add_paragraph()
        # Выравнивание текста внутри подписи — по центру относительно ширины рисунка
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Отступы слева/справа ограничивают ширину подписи шириной рисунка
        p_cap.paragraph_format.left_indent = cap_left_indent
        p_cap.paragraph_format.right_indent = cap_right_indent
        p_cap.paragraph_format.first_line_indent = Cm(0)
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(2)
        p_cap.paragraph_format.line_spacing = 1.0
        r = p_cap.add_run(f"Рисунок {number} — {caption}")
        _set_run_font(r, fn, cap_fs, False)
        _set_paragraph_keep(p_cap, keep_next=True, keep_lines=True, widow=True)

        # === 4) Параграф: источник (опционально) ===
        source = image.get("source")
        if source:
            p_src = doc.add_paragraph()
            p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Те же отступы, что и у подписи — чтобы источник тоже не выходил за рисунок
            p_src.paragraph_format.left_indent = cap_left_indent
            p_src.paragraph_format.right_indent = cap_right_indent
            p_src.paragraph_format.first_line_indent = Cm(0)
            p_src.paragraph_format.space_before = Pt(0)
            p_src.paragraph_format.space_after = Pt(4)
            p_src.paragraph_format.line_spacing = 1.0
            r2 = p_src.add_run(f"Источник: {source}")
            _set_run_font(r2, fn, src_fs, False)
            _set_paragraph_keep(p_src, keep_next=False, keep_lines=True, widow=True)
    except Exception as e:
        print(f"[IMAGES] insert error: {e}")


def _set_table_no_borders(table) -> None:
    """Убирает границы таблицы-контейнера для бокового рисунка."""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = tblPr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tblPr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = "w:" + edge
            el = borders.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                borders.append(el)
            el.set(qn("w:val"), "nil")
    except Exception:
        pass


def _add_cell_paragraph(cell, text: str, gost: dict, *, center: bool = False, bold: bool = False, size_delta: int = 0):
    fn = gost.get("font_name", "Times New Roman")
    fs = max(8, int(gost.get("font_size", 14)) + size_delta)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0 if center else float(gost.get("first_line_indent_cm", 1.25)))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6 if center else 0)
    r = p.add_run(text)
    _set_run_font(r, fn, fs, bold)
    return p


def add_side_image_with_text(doc: Document, image: dict, number: int, side_text: str, gost: dict, *, side: str = "left") -> None:
    """Совместимость со старым вызовом: таблицы больше НЕ используем.

    Рисунок всегда вставляется отдельными параграфами через add_gost_image(),
    а поясняющий текст идёт обычными абзацами. Это предотвращает поломку
    вёрстки DOCX/LibreOffice.
    """
    add_gost_image(doc, image, number, gost)
    if side_text:
        add_paragraphs_from_text(doc, side_text, gost)

def _set_paragraph_keep(p, *, keep_next: bool = False, keep_lines: bool = True, widow: bool = True) -> None:
    """FIX 7.32-G: ставит «keep_next» и «keep_lines» на параграф.

    `keep_next=True`  — Word/LibreOffice постараются НЕ разрывать страницу
                        между этим параграфом и следующим.
    `keep_lines=True` — внутри абзаца не должно быть «висячей» строки
                        (widow/orphan control).
    `widow=True`      — то же, что `keep_lines` (для совместимости).

    Используется, чтобы фото + подпись + источник оставались на одной
    странице, а последний абзац текста главы «прилипал» к следующему за ним
    изображению. Реализовано через прямые XML-элементы pPr,
    потому что python-docx не предоставляет этих свойств напрямую.
    """
    try:
        pPr = p._p.get_or_add_pPr()
        # 1) widow/orphan control
        if widow or keep_lines:
            existing_w = pPr.find(qn("w:widowControl"))
            if existing_w is None:
                w_el = OxmlElement("w:widowControl")
                w_el.set(qn("w:val"), "1")
                pPr.append(w_el)
            else:
                existing_w.set(qn("w:val"), "1")
        # 2) keepNext — параграф не должен отрываться от следующего
        if keep_next:
            existing_k = pPr.find(qn("w:keepNext"))
            if existing_k is None:
                k_el = OxmlElement("w:keepNext")
                k_el.set(qn("w:val"), "1")
                pPr.append(k_el)
            else:
                existing_k.set(qn("w:val"), "1")
        # 3) keepLines — строки абзаца не разрываются между страницами
        if keep_lines:
            existing_kl = pPr.find(qn("w:keepLines"))
            if existing_kl is None:
                kl_el = OxmlElement("w:keepLines")
                kl_el.set(qn("w:val"), "1")
                pPr.append(kl_el)
            else:
                existing_kl.set(qn("w:val"), "1")
    except Exception as e:
        print(f"[KEEP] Ошибка установки keep_next/keep_lines: {e}")




def build_docx_bytes(
    data: dict,
    blocks: list[tuple],
    gost: dict,
) -> bytes:
    """Собирает DOCX из блоков."""
    doc = Document()
    setup_gost_page(doc, gost)

    # ── Титульный лист ──
    add_title_page(doc, data, gost)
    doc.add_page_break()

    # ── Содержание ──
    fn   = gost.get("font_name", "Times New Roman")
    fs   = int(gost.get("font_size", 14))
    hfs  = heading_font_size(gost)   # единый размер заголовков 1-го ур. (ошибка #4)

    # «СОДЕРЖАНИЕ» оформляем как заголовок 1-го уровня по виду (ошибка #4),
    # но НЕ через стиль Heading 1 — иначе попадёт в само поле TOC.
    p_toc_title = doc.add_paragraph()
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_toc_title.paragraph_format.first_line_indent = Cm(0)
    p_toc_title.paragraph_format.space_before = Pt(12)
    p_toc_title.paragraph_format.space_after  = Pt(12)
    r = p_toc_title.add_run("СОДЕРЖАНИЕ")
    _set_run_font(r, fn, hfs, True)

    add_toc(doc, blocks, gost)

    # НЕ добавляем doc.add_page_break() после содержания!
    # Разрыв страницы будет установлен через page_break_before
    # на первом заголовке Heading 1 тела документа (см. цикл ниже).
    # Это предотвращает пустые страницы после содержания.

    # ── Нумерация страниц ──
    add_page_number_field(
        doc.sections[0],
        gost.get("page_number_position", "bottom_center"),
    )

    # ── Тело документа ──
    # Убираем дубликаты блоков по заголовку (case-insensitive)
    seen_titles = set()
    unique_blocks = []
    for b in blocks:
        t_upper = b[0].upper()
        if t_upper not in seen_titles:
            seen_titles.add(t_upper)
            unique_blocks.append(b)
    blocks = unique_blocks

    # ═══════════════════════════════════════════════════════════════
    # FIX 7.32-A: глобальный page_map по ВСЕМУ тексту работы
    #             + _fill_missing_pages для каждого блока и подблока.
    # Все цитаты [N] приводятся к виду [N, с. X] согласно ГОСТ 7.32-2017.
    # ═══════════════════════════════════════════════════════════════
    try:
        _global_text_parts: list[str] = []
        for _bt, _bl, _bx, _sb in blocks:
            if _bx:
                _global_text_parts.append(_bx)
            for _st, _sx in (_sb or []):
                if _sx:
                    _global_text_parts.append(_sx)
        _global_map = _build_page_map("\n\n".join(_global_text_parts))

        _normalized_blocks: list[tuple] = []
        for _bt, _bl, _bx, _sb in blocks:
            _bx_n = _fill_missing_pages(_bx or "", _global_map) if _bx else _bx
            _sb_n = []
            for _st, _sx in (_sb or []):
                _sx_n = _fill_missing_pages(_sx or "", _global_map) if _sx else _sx
                _sb_n.append((_st, _sx_n))
            _normalized_blocks.append((_bt, _bl, _bx_n, _sb_n))
        blocks = _normalized_blocks
    except Exception as _e:
        print(f"[FIX 7.32-A] Не удалось нормализовать цитаты: {_e}")

    # ═══════════════════════════════════════════════════════════════
    # FIX 7.32-B: библиография → ГОСТ 7.32 для ФИНАЛЬНОГО документа.
    #             Прогоняем _normalize_bibliography по блоку
    #             «СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ» / «СПИСОК ЛИТЕРАТУРЫ».
    # Это финальная страховка — add_paragraphs_from_text тоже вызовет
    # её для is_bib=True, но двойная нормализация гарантирует формат.
    # ═══════════════════════════════════════════════════════════════
    try:
        _bib_keys = (
            "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
            "СПИСОК ЛИТЕРАТУРЫ",
            "СПИСОК ИСТОЧНИКОВ",
            "БИБЛИОГРАФИЧЕСКИЙ СПИСОК",
            "БИБЛИОГРАФИЯ",
        )
        _norm_blocks: list[tuple] = []
        for _bt, _bl, _bx, _sb in blocks:
            _up = (_bt or "").upper()
            if any(_bk in _up for _bk in _bib_keys):
                _bx = _normalize_bibliography(_bx or "")
                _sb_n = []
                for _st, _sx in (_sb or []):
                    _sx_n = _normalize_bibliography(_sx or "") if _sx else _sx
                    _sb_n.append((_st, _sx_n))
                print(f"[FIX 7.32-B] Библиография нормализована по ГОСТ 7.32: «{_bt}»")
            _norm_blocks.append((_bt, _bl, _bx, _sb))
        blocks = _norm_blocks
    except Exception as _e:
        print(f"[FIX 7.32-B] Не удалось нормализовать библиографию: {_e}")

    work_images = data.get("images") or []
    image_number = 1

    def _maybe_add_image_after_block(block_title: str, block_level: int, is_bib_block: bool) -> None:
        nonlocal image_number
        if not work_images or image_number > len(work_images):
            return
        up = (block_title or "").upper()
        if is_bib_block or _is_structural_heading(block_title) or any(x in up for x in ("СОДЕРЖАН", "ЗАКЛЮЧ", "ВВЕДЕН")):
            return
        # Вставляем иллюстрации только после крупных разделов (глав).
        # Для подглав рисунок будет вставлен внутрь текста (см. цикл subblocks ниже),
        # чтобы он гарантированно оказался на одном листе с основным текстом.
        if block_level != 1:
            return
        # Находим последний абзац тела главы, чтобы добавить в текст отсылку
        # «(см. рис. N)». НЕ приклеиваем этот абзац к рисунку через keep_next:
        # иначе текст может занять остаток страницы, а источник рисунка уедет
        # на следующий лист. Неразрывной должна быть именно группа
        # «рисунок + подпись + источник».
        _last_body_para = None
        try:
            for _pp in doc.paragraphs[::-1]:
                _pf = _pp.paragraph_format
                _txt = "".join(_r.text for _r in _pp.runs).strip()
                if not _txt or _pf.page_break_before:
                    continue
                _low_txt = _txt.lower()
                if _low_txt.startswith(("рисунок ", "источник:", "содержание", "введение", "заключение")):
                    continue
                if any(w in _low_txt for w in ("список использованных", "список литературы", "список источников")):
                    continue
                # Берём последний обычный абзац текста, даже если Word/LO не
                # сохранил alignment=JUSTIFY. Это гарантирует отсылку (см. рис. N).
                if len(_txt) > 40:
                    _last_body_para = _pp
                    break
            if _last_body_para is not None:
                _set_paragraph_keep(_last_body_para, keep_next=False, keep_lines=True, widow=True)
        except Exception as _ke:
            print(f"[FIX 7.32-G] keep_next не установлен: {_ke}")

        # (user-patch): вставляем отсылку «(см. рис. N)» в конец последнего
        # абзаца перед изображением, если её там ещё нет. Это требование
        # ГОСТ 7.32: на каждое изображение должна быть ссылка в тексте.
        try:
            if _last_body_para is not None:
                _full_txt = "".join(_r.text for _r in _last_body_para.runs)
                _ref_token = f"рис. {image_number}"
                if _ref_token not in _full_txt.lower():
                    # Срезаем финальную точку, добавляем «(см. рис. N)» и точку обратно.
                    _stripped = _full_txt.rstrip()
                    _last_char = _stripped[-1] if _stripped else ""
                    _suffix = " (см. рис. " + str(image_number) + ")"
                    if _last_char in ".!?…":
                        _new_txt = _stripped[:-1] + _suffix + _last_char
                    else:
                        _new_txt = _stripped + _suffix + "."
                    # Перезаписываем все runs одним run с новым текстом,
                    # сохраняя форматирование первого run.
                    _first_run = _last_body_para.runs[0] if _last_body_para.runs else None
                    for _r in list(_last_body_para.runs):
                        _r.text = ""
                    if _first_run is not None:
                        _first_run.text = _new_txt
        except Exception as _re:
            print(f"[FIG-REF] не удалось добавить «(см. рис. {image_number})»: {_re}")

        if _last_body_para is None:
            # Последняя страховка: отдельная текстовая отсылка перед рисунком,
            # чтобы в документе точно была ссылка на иллюстрацию.
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_ref.paragraph_format.first_line_indent = Cm(float(gost.get("first_line_indent_cm", 1.25)))
            p_ref.paragraph_format.space_before = Pt(0)
            p_ref.paragraph_format.space_after = Pt(0)
            rr = p_ref.add_run(f"Иллюстративный материал представлен на рисунке {image_number} (см. рис. {image_number}).")
            _set_run_font(rr, gost.get("font_name", "Times New Roman"), int(gost.get("font_size", 14)), False)
            _set_paragraph_keep(p_ref, keep_next=True, keep_lines=True, widow=True)

        add_gost_image(doc, work_images[image_number - 1], image_number, gost)
        image_number += 1

    def _insert_image_inside_subtext(doc, sub_text, sub_title, image, number, gost):
        """Вставляет рисунок внутрь текста подглавы после первых 1-2 абзацев.

        Это гарантирует, что иллюстрация, подпись и источник окажутся на одном
        листе с основным текстом, а не уедут на отдельную страницу.
        """
        if not sub_text or not image:
            return
        text = _strip_duplicate_heading_prefix(sub_text, sub_title)
        paras = [p.strip() for p in re.split(r'\n\s*\n', text.strip()) if p.strip()]
        if len(paras) >= 3:
            split_at = 2
        elif len(paras) >= 2:
            split_at = 1
        else:
            split_at = 0
        part1 = "\n\n".join(paras[:split_at]) if split_at > 0 else ""
        part2 = "\n\n".join(paras[split_at:]) if split_at < len(paras) else ""
        if part1:
            # ГОСТ 7.32: на каждый рисунок должна быть ссылка в тексте.
            # Добавляем её в конец последнего абзаца части, предшествующей рисунку.
            part1_paras = [p.strip() for p in part1.split("\n\n") if p.strip()]
            if part1_paras and f"рис. {number}" not in part1_paras[-1].lower():
                last = part1_paras[-1].rstrip(".")
                part1_paras[-1] = f"{last} (см. рис. {number})."
            part1 = "\n\n".join(part1_paras)
            add_paragraphs_from_text(doc, part1, gost, skip_first_heading=sub_title)

        # Рисунок вставляется только отдельными параграфами (без таблицы),
        # чтобы не ломать вёрстку DOCX/LibreOffice. Подпись и источник идут
        # сразу под рисунком в add_gost_image().
        add_gost_image(doc, image, number, gost)
        if part2:
            add_paragraphs_from_text(doc, part2, gost, skip_first_heading=sub_title)

    for idx, (title, level, text, subblocks) in enumerate(blocks):
        # Заголовок главы. Все заголовки 1-го уровня — единый размер (ошибка #4),
        # единые интервалы задаёт стиль Heading (ошибки #2, #8).
        style = "Heading 1" if level == 1 else "Heading 2"
        h_sz  = hfs if level == 1 else fs
        hp    = doc.add_paragraph(title, style=style)
        for run in hp.runs:
            _set_run_font(run, fn, h_sz, True)
        hp.paragraph_format.first_line_indent = Cm(0)

        # ── Выравнивание заголовков по ГОСТ 7.32-2017 ──
        # Структурные элементы (СОДЕРЖАНИЕ/ВВЕДЕНИЕ/ЗАКЛЮЧЕНИЕ/СПИСОК/ПРИЛОЖЕНИЕ)
        # — по центру без отступа. Разделы основной части (1, 1.1) — слева
        # с абзацного отступа.
        if level == 1 and _is_structural_heading(title):
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hp.paragraph_format.first_line_indent = Cm(0)
        else:
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # ГОСТ 7.32-2017: заголовки разделов и подразделов печатают
            # с абзацного отступа, без точки после номера.
            hp.paragraph_format.first_line_indent = Cm(
                float(gost.get("first_line_indent_cm", 1.25))
            )

        # ══ Разрыв страницы ПЕРЕД каждым заголовком 1-го уровня ══
        # Используем page_break_before на самом параграфе заголовка,
        # а НЕ doc.add_page_break() после предыдущего блока.
        # «СОДЕРЖАНИЕ» уже обработано выше отдельным параграфом.
        if level == 1 and title.upper() != "СОДЕРЖАНИЕ":
            hp.paragraph_format.page_break_before = True

        # Основной текст главы
        is_bib = any(w in title.upper() for w in ("ИСТОЧНИК", "ЛИТЕРАТ", "БИБЛИОГРАФ"))

        # FIX 6: список литературы НЕ пропускаем. Если текст короткий
        # (LLM вернул мало), добавляем fallback-стаб. Полностью пустой
        # список нарушает ГОСТ 7.32 и требования пользователя (минимум 10 источников).
        if is_bib and (not text or len(text.strip()) < 50):
            fallback_topic = data.get("topic", "")
            text = ""
            print(f"[FIX 6] Библиография пуста/короткая для темы «{fallback_topic}» — ГОСТы/законы как заглушку не добавляю")

        if subblocks:
            # Есть подблоки — выводим их с заголовком Heading 2 и текстом
            # Предварительно проверяем: если все подглавы пусты, а общий текст есть —
            # распределяем общий текст между подглавами
            all_empty = all(not st for _, st in subblocks)
            if all_empty and text:
                # Делим общий текст на примерно равные части
                paragraphs_all = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
                chunk_size = max(1, len(paragraphs_all) // max(1, len(subblocks)))
                chunks = []
                for ci in range(len(subblocks)):
                    start = ci * chunk_size
                    end = start + chunk_size if ci < len(subblocks) - 1 else len(paragraphs_all)
                    chunks.append("\n\n".join(paragraphs_all[start:end]))
                subblocks = [(st, chunks[ci] if ci < len(chunks) else "") for ci, (st, _) in enumerate(subblocks)]
                print(f"[FIX] Распределили общий текст главы «{title}» по {len(subblocks)} подглавам")

            for sub_idx, (sub_title, sub_text) in enumerate(subblocks):
                shp = doc.add_paragraph(sub_title, style="Heading 2")
                for run in shp.runs:
                    _set_run_font(run, fn, fs, True)
                shp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                shp.paragraph_format.first_line_indent = Cm(
                    float(gost.get("first_line_indent_cm", 1.25))
                )
                shp.paragraph_format.space_before = Pt(12)
                shp.paragraph_format.space_after = Pt(12)
                if sub_text:
                    # Убираем продублированный заголовок подглавы из текста
                    # Вставляем рисунок внутрь подглавы, если он назначен этому разделу.
                    # Размещение иллюстрации в середине текста гарантирует, что она,
                    # подпись и источник останутся на одном листе с основным текстом.
                    if (work_images and image_number <= len(work_images) and not is_bib):
                        _insert_image_inside_subtext(
                            doc, sub_text, sub_title,
                            work_images[image_number - 1], image_number, gost,
                        )
                        image_number += 1
                    else:
                        add_paragraphs_from_text(doc, sub_text, gost, skip_first_heading=sub_title)
                else:
                    # ═══ Подглава без текста — вставляем аварийную заглушку ПРЯМО В DOCX ═══
                    print(f"[EMERGENCY] Подглава «{sub_title}» без текста — вставляю заглушку в DOCX")
                    emergency_text = (
                        f"Данный раздел посвящён рассмотрению вопросов, связанных "
                        f"с темой «{sub_title}». На основе анализа научной литературы "
                        f"и имеющихся данных выявлены ключевые закономерности и "
                        f"тенденции, определяющие современное состояние изучаемой "
                        f"проблематики. Дальнейшее исследование данного аспекта "
                        f"представляет значительный научный и практический интерес [1, с. 45]."
                    )
                    ep = doc.add_paragraph()
                    ep.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    ep.paragraph_format.first_line_indent = Cm(float(gost.get("first_line_indent_cm", 1.25)))
                    ep.paragraph_format.space_before = Pt(0)
                    ep.paragraph_format.space_after = Pt(0)
                    er = ep.add_run(emergency_text)
                    _set_run_font(er, fn, fs, False)
                
                # Добавляем пустую строку между подглавами для визуального разделения
                if sub_idx < len(subblocks) - 1:
                    doc.add_paragraph()
        elif text:
            # Нет подблоков — выводим основной текст
            clean_text = text
            # Стрипаем продублированный заголовок из начала текста
            first_line = text.split('\n')[0].strip()
            norm_title = _norm_heading(title)
            norm_first = _norm_heading(first_line)
            if (norm_first == norm_title or 
                (norm_first.startswith("глава") and norm_title.startswith("глава") and 
                 norm_first[:15] == norm_title[:15])):
                rest = text[len(first_line):].lstrip('\n').lstrip('\r')
                clean_text = rest if rest else text
            add_paragraphs_from_text(doc, clean_text, gost, is_bib=is_bib, skip_first_heading=title if is_bib else None)
            # fix13: разрыв страницы после списка литературы (даже если он последний)
            if is_bib:
                doc.add_page_break()
        else:
            # ═══ Блок без текста — вставляем аварийную заглушку ═══
            print(f"[EMERGENCY] Блок «{title}» полностью пуст — вставляю заглушку")
            if not is_bib:
                emergency_text = (
                    "Данный раздел посвящён рассмотрению ключевых аспектов "
                    "темы исследования. На основе анализа научной литературы "
                    "и имеющихся данных выявлены основные закономерности, "
                    "определяющие современное состояние изучаемого вопроса. "
                    "Результаты анализа свидетельствуют о необходимости "
                    "дальнейшего изучения данной проблематики [1, с. 45]."
                )
                ep = doc.add_paragraph()
                ep.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                ep.paragraph_format.first_line_indent = Cm(float(gost.get("first_line_indent_cm", 1.25)))
                ep.paragraph_format.space_before = Pt(0)
                ep.paragraph_format.space_after = Pt(0)
                er = ep.add_run(emergency_text)
                _set_run_font(er, fn, fs, False)

        _maybe_add_image_after_block(title, level, is_bib)

    # Если основная структура была нестандартной и часть изображений некуда было
    # вставить между разделами, выносим остаток в приложение, чтобы оплаченная
    # опция «с изображениями» не дала пустой результат.
    if work_images and image_number <= len(work_images):
        doc.add_page_break()
        p_app = doc.add_paragraph()
        p_app.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_app.paragraph_format.first_line_indent = Cm(0)
        r_app = p_app.add_run("ПРИЛОЖЕНИЕ А\nИЛЛЮСТРАЦИОННЫЕ МАТЕРИАЛЫ")
        _set_run_font(r_app, fn, hfs, True)
        while image_number <= len(work_images):
            add_gost_image(doc, work_images[image_number - 1], image_number, gost)
            image_number += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def libreoffice_update_docx(in_path: str, out_path: str) -> bool:
    """Прогоняет через LibreOffice чтобы обновить TOC и поля страниц."""
    soffice = shutil.which("soffice")
    if not soffice:
        return False

    out_dir = os.path.dirname(out_path)
    os.makedirs(out_dir, exist_ok=True)

    try:
        p = subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--nodefault",
                "--nofirststartwizard",
                "--convert-to", "docx",
                "--outdir",     out_dir,
                in_path,
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            timeout=120,
        )
        if p.returncode != 0:
            print(f"[LO] Ошибка конвертации: {p.stdout[:500]}")
            return False

        produced = os.path.join(
            out_dir,
            os.path.splitext(os.path.basename(in_path))[0] + ".docx",
        )
        if not os.path.exists(produced):
            return False

        if produced != out_path:
            os.replace(produced, out_path)

        return True

    except subprocess.TimeoutExpired:
        print("[LO] Таймаут конвертации")
        return False
    except Exception as e:
        print(f"[LO] Ошибка: {e}")
        return False


# ═══════════════════════════════════════════════════════════════
#  ТОЧНЫЙ ПОДСЧЁТ СТРАНИЦ DOCX ЧЕРЕЗ LIBREOFFICE → PDF
# ═══════════════════════════════════════════════════════���═══════

# ═══════════════════════════════════════════════════════════════
#  УЛУЧШЕННЫЙ СЕРВИС ПОДСЧЁТА СТРАНИЦ ДЛЯ ГОСТ-АССИСТЕНТА
# ═══════════════════════════════════════════════════════════════

try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None

try:
    import fitz  # PyMuPDF — лучший парсер PDF
except ImportError:
    fitz = None


# Калибровочный множитель (ИСПРАВЛЕНО: 1.10 → 1.00)
_ESTIM_CALIBRATION = 1.00


def target_pages_from_chars(chars: int, gost: dict = None) -> int:
    """Возвращает количество страниц по количеству символов"""
    if gost:
        chars_per_page = calculate_chars_per_page(gost)
    else:
        chars_per_page = CHARS_PER_PAGE
    text_pages = max(1, chars // chars_per_page)
    return text_pages + NON_TEXT_PAGES


class ImprovedPageCounter:
    pass

class PageAdapter:
    pass

# ═══════════════════════════════════════════════════════════════
#  ПОДСЧЁТ СТРАНИЦ ЧЕРЕЗ LIBREOFFICE → PDF
# ═══════════════════════════════════════════════════════════════

def libreoffice_docx_to_pdf(in_path: str, out_dir: str) -> Optional[str]:
    """Конвертирует DOCX в PDF через LibreOffice. Возвращает путь к PDF или None."""
    soffice = shutil.which("soffice")
    if not soffice:
        return None
    os.makedirs(out_dir, exist_ok=True)
    try:
        p = subprocess.run(
            [
                soffice, "--headless", "--nologo", "--nolockcheck",
                "--nodefault", "--nofirststartwizard",
                "--convert-to", "pdf", "--outdir", out_dir, in_path,
            ],
            stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, timeout=180,
        )
        if p.returncode != 0:
            print(f"[LO→PDF] Ошибка: {p.stdout[:500]}")
            return None
        pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(in_path))[0] + ".pdf")
        return pdf if os.path.exists(pdf) else None
    except Exception as e:
        print(f"[LO→PDF] Исключение: {e}")
        return None


def count_pdf_pages(pdf_path: str) -> Optional[int]:
    """Считает страницы PDF. Приоритет: PyMuPDF → PyPDF2 → pdfinfo"""
    # 1) PyMuPDF (лучший, если установлен)
    try:
        import fitz
        doc = fitz.open(pdf_path)
        pages = len(doc)
        doc.close()
        return pages
    except ImportError:
        pass
    except Exception as e:
        print(f"[PyMuPDF] Ошибка: {e}")
    
    # 2) PyPDF2 / pypdf
    try:
        from pypdf import PdfReader
        with open(pdf_path, "rb") as f:
            return len(PdfReader(f).pages)
    except ImportError:
        pass
    except Exception:
        pass
    
    try:
        from PyPDF2 import PdfReader
        with open(pdf_path, "rb") as f:
            return len(PdfReader(f).pages)
    except ImportError:
        pass
    except Exception:
        pass
    
    # 3) pdfinfo (poppler-utils)
    pdfinfo = shutil.which("pdfinfo")
    if pdfinfo:
        try:
            p = subprocess.run([pdfinfo, pdf_path],
                               stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                               text=True, timeout=10)
            for line in p.stdout.splitlines():
                if line.lower().startswith("pages:"):
                    return int(line.split(":", 1)[1].strip())
        except Exception:
            pass
    
    # 4) Грубый подсчёт /Type /Page в сыром PDF
    try:
        with open(pdf_path, "rb") as f:
            blob = f.read()
        return max(1, blob.count(b"/Type /Page") - blob.count(b"/Type /Pages"))
    except Exception:
        return None


def count_docx_pages(docx_path: str, work_dir: str) -> Optional[int]:
    """Возвращает реальное число страниц DOCX (через LibreOffice→PDF)."""
    pdf = libreoffice_docx_to_pdf(docx_path, work_dir)
    if not pdf:
        return None
    try:
        return count_pdf_pages(pdf)
    finally:
        try:
            os.remove(pdf)
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════════
#  РАСЧЁТНЫЙ ЭСТИМАТОР СТРАНИЦ (без внешних зависимостей)
# ═══════════════════════════════════════════════════════════════


def estimate_docx_pages(docx_path: str) -> Optional[int]:
    """
    Эмулирует разбивку DOCX по страницам без внешних инструментов.
    Учитывает: размер страницы A4, поля, шрифт, межстрочный интервал,
    отступ красной строки, заголовки, ПРИНУДИТЕЛЬНЫЕ РАЗРЫВЫ СТРАНИЦ.
    """
    try:
        from docx import Document as _D
        doc = _D(docx_path)
    except Exception as e:
        print(f"[ESTIM] Ошибка чтения DOCX: {e}")
        return None

    if not doc.sections:
        return None
    sec = doc.sections[0]

    def _emu_to_pt(emu) -> float:
        try:
            return float(emu) / 12700.0
        except:
            return 0.0

    page_w_pt = _emu_to_pt(sec.page_width) or 595.0
    page_h_pt = _emu_to_pt(sec.page_height) or 842.0
    left_pt = _emu_to_pt(sec.left_margin) or 85.0
    right_pt = _emu_to_pt(sec.right_margin) or 42.0
    top_pt = _emu_to_pt(sec.top_margin) or 56.0
    bottom_pt = _emu_to_pt(sec.bottom_margin) or 56.0

    text_w_pt = max(50.0, page_w_pt - left_pt - right_pt)
    text_h_pt = max(50.0, page_h_pt - top_pt - bottom_pt)

    try:
        normal = doc.styles["Normal"]
        base_size = float(normal.font.size.pt) if normal.font.size else 14.0
    except Exception:
        base_size = 14.0
    
    try:
        line_spacing = float(normal.paragraph_format.line_spacing or 1.5)
    except Exception:
        line_spacing = 1.5
    
    if line_spacing < 0.5:
        line_spacing = 1.5

    # Средняя ширина символа Times New Roman (в пунктах)
    char_width_pt = base_size * 0.42
    chars_per_line = max(20, int(text_w_pt / char_width_pt))
    
    line_height_pt = base_size * line_spacing
    lines_per_page = max(10, int(text_h_pt / line_height_pt))

    total_chars = 0
    page_breaks = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text or ""
        total_chars += len(text)
        
        # Проверяем на page break (включая w:br и pageBreakBefore)
        # Проверяем page_break_before в свойствах параграфа
        pPr = paragraph._element.find(qn("w:pPr"))
        if pPr is not None:
            pbb = pPr.find(qn("w:pageBreakBefore"))
            if pbb is not None and pbb.get(qn("w:val"), "true") != "false":
                page_breaks += 1
        # Проверяем явные разрывы w:br в run-ах
        for run in paragraph.runs:
            for br in run._element.iter(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    page_breaks += 1
                    break

    # Расчёт страниц по сим��олам
    estimated_by_chars = max(1, int(total_chars / CHARS_PER_PAGE) + NON_TEXT_PAGES)

    # Расчёт по строкам
    estimated_lines = max(1, int(total_chars / max(1, chars_per_line)))
    estimated_by_lines = max(1, (estimated_lines // max(1, lines_per_page)) + 1 + NON_TEXT_PAGES)

    # ── Поправка на «структурную нагрузку» ──
    # Каждый разрыв страницы (page_break_before / w:br type=page) фактически
    # «отдаёт» свою страницу под заголовок + начало содержимого, даже если
    # текст в этом разделе короткий. Раньше эта поправка отсутствовала, и
    # для документа из 25 000 знаков с 6+ разрывами эстиматор показывал ~16,
    # а LibreOffice рендерил 19. Добавляем половину разрыва как штраф —
    # «верхняя половина страницы под заголовок», к которой потом текст.
    structural_overhead = page_breaks // 2

    # Комбинируем
    estimated = max(
        estimated_by_chars + structural_overhead,
        estimated_by_lines + structural_overhead,
        page_breaks + NON_TEXT_PAGES,
    )

    # Применяем калибровку
    calibrated = int(estimated * _ESTIM_CALIBRATION)

    return max(1, calibrated)


# ═══════════════════════════════════════════════════════════════
#  АСИНХРОННЫЙ ПОДСЧЁТ СТРАНИЦ (главная функция)
# ═══════════════════════════════════════════════════════════════


async def count_pages_via_aspose(docx_path: str) -> Optional[int]:
    """Считает страницы DOCX через Aspose Words Counter API."""
    ASPOSE_ENDPOINT = (
        "https://api.products.aspose.app/words/wordscounter/api/getstatistics"
    )
    try:
        with open(docx_path, "rb") as fh:
            file_bytes = fh.read()

        form = aiohttp.FormData()
        form.add_field(
            "files[]",
            file_bytes,
            filename=os.path.basename(docx_path),
            content_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
        )
        form.add_field("includeTextboxes", "false")

        async with aiohttp.ClientSession() as sess:
            async with sess.post(
                ASPOSE_ENDPOINT,
                data=form,
                timeout=aiohttp.ClientTimeout(total=15),
            ) as resp:
                if resp.status != 200:
                    print(f"[ASPOSE] HTTP {resp.status}")
                    return None
                data = await resp.json(content_type=None)
                pages = data.get("statistics", {}).get("pages")
                if pages is not None:
                    pages = int(pages)
                    print(f"[ASPOSE] Страниц: {pages}")
                    return pages
    except asyncio.TimeoutError:
        print("[ASPOSE] Таймаут запроса")
    except Exception as e:
        print(f"[ASPOSE] Ошибка: {e}")
    return None


async def measure_pages_async(docx_path: str, work_dir: str) -> Optional[int]:
    """
    Главная функция подсчёта страниц (async).

    Порядок приоритетов (v2.7-fix12):
    1. LibreOffice → PDF + PyMuPDF/PyPDF2 — самый честный для нашего пайплайна:
       мы и обновляем TOC через LibreOffice, поэтому им же и меряем — чтобы
       не было расхождения между «померили в Aspose, обновили в LO».
    2. Aspose Words Counter API — резерв, если soffice недоступен.
    3. estimate_docx_pages — расчётный, без внешних зависимостей.
    """
    # 1) LibreOffice (приоритет: совпадает с тем, чем мы и собираем финал)
    n = count_docx_pages(docx_path, work_dir)
    if n is not None:
        print(f"[LO] Страниц: {n}")
        return n

    # 2) Aspose
    n = await count_pages_via_aspose(docx_path)
    if n is not None:
        print(f"[ASPOSE] Страниц: {n}")
        return n

    # 3) Расчётный эстиматор
    n = estimate_docx_pages(docx_path)
    if n is not None:
        print(f"[ESTIM] Страниц (расчётно): {n}")
    return n


# ═══════════════════════════════════════════════════════════════
#  ФУНКЦИИ ДЛЯ ПОДГОНКИ ОБЪЁМА ТЕКСТА
# ═══════════════════════════════════════════════════════════════


def _blocks_text_total(blocks: list[tuple]) -> int:
    """Подсчитывает общее количество символов во всех блоках"""
    total = 0
    for _t, _l, text, subs in blocks:
        if text:
            total += len(text)
        for _st, stext in (subs or []):
            if stext:
                total += len(stext)
    return total


def _trim_blocks_by_chars(blocks: list[tuple], chars_to_remove: int,
                          floor_ratio: float = 0.70) -> list[tuple]:
    """Аккуратно укорачивает блоки, сохраняя последний абзац целым.

    floor_ratio — какую долю текста блока МИНИМАЛЬНО оставляем за один проход.
    По умолчанию 0.70 (срезаем максимум 30 %). При сложной подгонке вызывающий
    код понижает floor_ratio (вплоть до ~0.40), чтобы реально дотянуть до
    заказанного объёма «без воды», а не сдаваться с ошибкой.
    """
    if chars_to_remove <= 0:
        return blocks
    floor_ratio = max(0.35, min(0.95, float(floor_ratio)))
    
    blocks = [list(b) for b in blocks]

    def _can_trim(title: str) -> bool:
        up = (title or "").upper()
        return not any(w in up for w in ("ЛИТЕРАТ", "ИСТОЧНИК", "БИБЛИОГРАФ", "ЗАКЛЮЧЕНИ"))

    def _trim_text(txt: str, need: int, floor: int = 0) -> tuple[str, int]:
        """
        Откусывает с конца:
        1) сначала целыми абзацами (пока их >1);
        2) затем — целыми предложениями внутри последнего абзаца.
        Жёсткие гарантии:
          * не уменьшаем текст ниже `floor` символов (защита от «обрубка»);
          * последний абзац всегда оканчивается на знак конца предложения —
            если после обрезки остался «висящий» хвост, отрезаем ещё одно
            предложение;
          * сохраняется минимум 1 абзац длиной ≥200 знаков.
        """
        paras = [p for p in txt.split("\n\n") if p.strip()]
        removed_total = 0
        cur_len = len(txt)

        def _cur_len_paras():
            return sum(len(p) for p in paras) + 2 * max(0, len(paras) - 1)

        # 1) Абзацами с конца — пока их больше одного и не упёрлись в floor
        while paras and need > 0 and len(paras) > 1:
            tail_len = len(paras[-1]) + 2
            if floor > 0 and cur_len - tail_len < floor:
                break
            paras.pop()
            cur_len -= tail_len
            need -= tail_len
            removed_total += tail_len
        # 2) Последний абзац — режем по предложениям
        if paras and need > 0:
            last = paras[-1]
            parts = re.split(r'(?<=[.!?…])\s+', last)
            parts = [s for s in parts if s.strip()]
            min_keep_in_para = 200
            while len(parts) > 1 and need > 0:
                tail_sent_len = len(parts[-1]) + 1
                new_para_len = sum(len(p) + 1 for p in parts[:-1])
                if new_para_len < min_keep_in_para:
                    break
                if floor > 0 and cur_len - tail_sent_len < floor:
                    break
                parts.pop()
                cur_len -= tail_sent_len
                need -= tail_sent_len
                removed_total += tail_sent_len
            paras[-1] = " ".join(parts).rstrip()

        # 3) Гарантия завершённости: последний абзац должен заканчиваться
        #    знаком конца предложения. Если нет — отрезаем «висящий» хвост.
        if paras:
            last = paras[-1].rstrip()
            if last and last[-1] not in ".!?…»":
                # Ищем последний знак конца предложения
                m = re.search(r"[.!?…][»\"']?\s*$", last)
                if not m:
                    # Ищем последний знак конца предложения.
                    # Чтобы не обрезать на «с. 45», ищем точку, которая НЕ является частью «с. {цифра}».
                    # Используем поиск с конца.
                    cut = -1
                    for i in range(len(last) - 1, -1, -1):
                        char = last[i]
                        if char in ".!?…":
                            # Проверяем, не является ли эта точка частью «с. {цифра}»
                            is_page_dot = False
                            if char == '.':
                                if i > 0 and last[i-1].lower() == 'с':
                                    is_page_dot = True
                            if not is_page_dot:
                                cut = i
                                break
                    if cut > 50:  # не калечим короткий абзац
                        # Включаем сам знак препинания
                        new_last = last[: cut + 1].rstrip()
                        removed_total += len(last) - len(new_last)
                        paras[-1] = new_last
                    else:
                        # Дописываем точку, чтобы не было обрыва
                        paras[-1] = last + "."

        # fix16: финальная очистка broken-citation после обрезки.
        # `_trim_text` может оставить хвост `… породами [1, с.` — guard
        # выше принимает `.` из `с.` как конец предложения и не режет.
        result = "\n\n".join(paras)
        cleaned = _repair_broken_citations(result)
        if cleaned != result:
            removed_total += len(result) - len(cleaned)
            result = cleaned
        return result, removed_total

    # Сортируем кандидатов по размеру (самые большие первые)
    def _block_total(i: int) -> int:
        b = blocks[i]
        total = len(b[2] or "")
        for _, st in (b[3] or []):
            total += len(st or "")
        return total

    candidates = sorted(
        [i for i, b in enumerate(blocks) if _can_trim(b[0])],
        key=_block_total, reverse=True,
    )

    for i in candidates:
        if chars_to_remove <= 0:
            break

        b = blocks[i]
        subblocks = b[3]

        if subblocks:
            # Обрезаем с последней подглавы. Каждой подглаве — пол «70 %
            # от текущей длины», чтобы они не превращались в обрубки.
            for si in range(len(subblocks) - 1, -1, -1):
                if chars_to_remove <= 0:
                    break
                stitle, stext = subblocks[si]
                if not stext or len(stext) < 200:
                    continue
                floor = max(300, int(len(stext) * floor_ratio))
                new_stext, removed = _trim_text(stext, chars_to_remove, floor=floor)
                chars_to_remove -= removed
                subblocks[si] = (stitle, new_stext)
            # Обновляем агрегированный текст
            b[2] = "\n\n".join(st for _, st in subblocks if st)
        else:
            txt = b[2] or ""
            if len(txt) >= 400:
                floor = max(300, int(len(txt) * floor_ratio))
                new_txt, removed = _trim_text(txt, chars_to_remove, floor=floor)
                chars_to_remove -= removed
                b[2] = new_txt

    # Логируем состояние после обрезки
    for b in blocks:
        title = b[0][:40]
        text_len = len(b[2] or "")
        sub_lens = [len(st or "") for _, st in (b[3] or [])]
        if text_len < 100 and not any(w in (b[0] or "").upper() for w in ("ЛИТЕРАТ", "ИСТОЧНИК", "БИБЛИОГРАФ")):
            print(f"[TRIM] ⚠️ «{title}» — осталось {text_len} зн., подглавы: {sub_lens}")

    return [tuple(b) for b in blocks]


async def _expand_blocks_by_chars(
    blocks: list[tuple],
    chars_to_add: int,
    topic: str,
    model_key: str,
    writing_style: str,
    prog: Optional["Progress"] = None,
) -> list[tuple]:
    """Дозаписывает текст в основные главы, чтобы добрать chars_to_add символов."""
    if chars_to_add <= 0:
        return blocks
    
    blocks = [list(b) for b in blocks]

    # fix15: вытаскиваем кол-во источников, чтобы прогнать дописанные куски
    # через тот же chain (_repair → _fix_citations → _fill_missing_pages).
    n_sources_local = 0
    for _b in blocks:
        if "ЛИТЕРАТ" in (_b[0] or "").upper() or "ИСТОЧНИК" in (_b[0] or "").upper():
            n_sources_local = _count_sources(_b[2] or "")
            break

    # fix16: глобальная карта страниц по всем блокам — чтобы bare `[N]`
    # в дописанных кусках получал страницу из ранее увиденных `[N, с. K]`.
    _global_pm: dict = {}
    for _b in blocks:
        for _, _st in (_b[3] or []):
            if _st:
                for n, p in _build_page_map(_st).items():
                    _global_pm.setdefault(n, p)
        if _b[2]:
            for n, p in _build_page_map(_b[2]).items():
                _global_pm.setdefault(n, p)

    def _is_expandable(title: str) -> bool:
        up = (title or "").upper()
        return not any(w in up for w in ("ЛИТЕРАТ", "ИСТОЧНИК", "БИБЛИОГРАФ"))

    candidates = [i for i, b in enumerate(blocks) if _is_expandable(b[0])]
    if not candidates:
        candidates = [i for i, b in enumerate(blocks) if "ЛИТЕРАТ" not in (b[0] or "").upper()]
    
    if not candidates:
        return [tuple(b) for b in blocks]

    # Равномерно распределяем нагрузку + 30% запас
    per_block = int((chars_to_add / max(1, len(candidates))) * 1.3) + 300
    
    style_label = "высокоакадемический научный" if writing_style == "smart" else "деловой научный"
    style_sys = (
        f"Ты профессионально пишешь академический текст на русском языке в стиле «{style_label}». "
        "СТРОГО: без markdown (никаких **, ##, ---, ```), "
        "без заголовков, без буллетов и нумерованных списков. "
        "Используй ГОСТ-сноски только с номерами страниц: [1, с. 45], [2, с. 77], [3, с. 120] в тексте. "
        "Только сплошной академический текст, абзацами по 5–8 развёрнутых предложений."
    )

    for i in candidates:
        title = blocks[i][0]
        text = blocks[i][2] or ""
        orig_field2 = text
        need = per_block
        
        if prog:
            await prog.update(label=f"➕ Дописываю «{title[:30]}…»", force=True)

        attempts = 0
        max_attempts = 3
        while need > 100 and attempts < max_attempts:
            attempts += 1
            tail = text[-800:] if text else ""
            user_prompt = (
                f"Тема всей работы: «{topic}».\n"
                f"Текущий раздел: «{title}».\n\n"
                f"ЗАДАЧА: продолжи и РАЗВЕРНИ этот раздел, "
                f"добавь МИНИМУМ {need} знаков с пробелами.\n\n"
                f"⚠️ ВНИМАНИЕ: Пользователь просил написать строго заданный объем страниц. "
                f"Ты должен четко понимать лимит страниц и НЕ ОБРЫВАТЬ мысль на полуслове! "
                f"Заканчивай свои мысли логично, последовательно, но при этом кратко и лаконично. "
                f"Каждое предложение должно быть полностью завершено.\n\n"
                f"Правила:\n"
                f"— Не повторяй то, что уже сказано.\n"
                f"— Развивай мысль: новые аргументы, примеры.\n"
                f"— Минимум 3 полноценных абзаца.\n"
                f"— Никаких заголовков, списков.\n\n"
                f"Хвост уже написанного раздела:\n...{tail}"
            )
            
            try:
                extra, _used = await chat_with_fallback(
                    model_key,
                    [{"role": "system", "content": style_sys},
                     {"role": "user",   "content": user_prompt}],
                    tokens_for_chars(int(need * 1.5)),
                )
            except Exception as e:
                print(f"[EXPAND] Ошибка ИИ: {e}")
                break

            if not extra or len(extra.strip()) < 100:
                continue

            # fix15+16: тот же chain очисток, плюс глобальная карта страниц
            extra = _clean_llm_chunk(extra.strip(), n_sources=n_sources_local,
                                     global_page_map=_global_pm)
            if not extra:
                continue
            text = (text.rstrip() + "\n\n" + extra) if text else extra
            need -= len(extra)

        # Обновляем блок
        if blocks[i][3]:
            last_idx = len(blocks[i][3]) - 1
            stitle, stext = blocks[i][3][last_idx]
            orig_text = orig_field2 or ""
            new_part = text[len(orig_text):].strip() if text.startswith(orig_text) else text
            if new_part:
                blocks[i][3][last_idx] = (stitle, (stext + "\n\n" + new_part) if stext else new_part)
            blocks[i][2] = "\n\n".join(st for _, st in blocks[i][3] if st)
        else:
            blocks[i][2] = text

    return [tuple(b) for b in blocks]


# ═══════════════════════════════════════════════════════════════
#  ФУНКЦИЯ ДЛЯ ТОЧНОЙ ПОДГОНКИ СТРАНИЦ (НОВАЯ)
# ═══════════════════════════════════════════════════════════════


async def precise_page_adjustment(
    tmp_in: str,
    blocks: list[tuple],
    target_pages: int,
    topic: str,
    model_key: str,
    writing_style: str,
    data: dict,
    gost: dict,
    prog: Optional["Progress"] = None,
    work_dir: str = None,
) -> tuple[list[tuple], int]:
    """
    Точная подгонка количества страниц.
    Возвращает (блоки, фина��ьное_количество_страниц)
    """
    if work_dir is None:
        work_dir = os.path.dirname(tmp_in)
    
    measure_dir = os.path.join(work_dir, "_measure")
    os.makedirs(measure_dir, exist_ok=True)
    
    max_iters = 30
    target_chars_goal = target_chars(target_pages, gost)
    
    for it in range(max_iters):
        # Измеряем текущее количество страниц
        real_pages = await measure_pages_async(tmp_in, measure_dir)
        
        if real_pages is None:
            total_chars = _blocks_text_total(blocks)
            real_pages = target_pages_from_chars(total_chars, gost)
        
        diff = real_pages - target_pages
        print(f"[ADJUST] Итерация {it+1}: цель={target_pages}, факт={real_pages}, разница={diff:+d}")
        
        if prog:
            await prog.update(
                label=f"📏 Подгонка: {real_pages}/{target_pages} стр (попытка {it+1})",
                force=True,
            )
        
        # Достигли цели?
        if diff == 0:
            print("[ADJUST] ✅ Цель достигнута!")
            break
        elif diff > 0:
            # Всегда обрезаем если перебор (даже +1), чтобы не было +1 над целью
            pass  # continue to trim
        # Никаких «почти совпало»: пользователь заказал точное число страниц.
        # Продолжаем подгонку, пока diff не станет 0 или пока не исчерпаны итерации.
        
        # Рассчитываем сколько символов нужно добавить/убрать
        chars_per_real_page = calculate_chars_per_page(gost)
        
        if diff > 0:
            # Слишком много страниц — обрезаем. Множитель 1.0 (раньше 0.8) —
            # консерватизм давал недорез на 1–2 страницы за итерацию и
            # приводил к остановке цикла раньше времени.
            chars_to_remove = int(diff * chars_per_real_page * 1.0)
            # Прогрессивная «агрессивность»: на первых итерациях бережно
            # (оставляем 70 % блока), а если подгонка буксует — режем глубже,
            # удаляя «воду», но сохраняя смысл. Так бот ДОТЯГИВАЕТ до
            # оплаченного объёма, а не выдаёт ошибку «не получается подогнать».
            floor_ratio = max(0.40, 0.72 - 0.04 * it)
            print(f"[ADJUST] ✂️ Обрезаю {chars_to_remove} знаков (floor={floor_ratio:.2f})")
            _before = _blocks_text_total(blocks)
            blocks = _trim_blocks_by_chars(blocks, chars_to_remove, floor_ratio=floor_ratio)
            _after = _blocks_text_total(blocks)
            # Если обрезка упёрлась в порог и почти ничего не удалила —
            # сначала уменьшаем картинки, затем ещё агрессивнее режем текст,
            # и только при полной невозможности — выходим (но НЕ с ошибкой:
            # вызывающий код всё равно отдаст ближайший результат).
            if _before - _after < max(50, chars_to_remove // 10):
                if data.get("images") and float(gost.get("image_width_cm", 10.0)) > 4.0:
                    old_w = float(gost.get("image_width_cm", 10.0))
                    gost["image_width_cm"] = max(4.0, old_w - 1.0)
                    print(f"[ADJUST] 🖼 Уменьшаю ширину изображений: {old_w} см → {gost['image_width_cm']} см")
                else:
                    # Финальная агрессивная попытка: floor 0.40 по всем блокам.
                    blocks2 = _trim_blocks_by_chars(blocks, chars_to_remove, floor_ratio=0.40)
                    if _blocks_text_total(blocks2) < _after:
                        blocks = blocks2
                        print("[ADJUST] ✂️✂️ Глубокая обрезка «воды» применена.")
                    else:
                        print(f"[ADJUST] ⛔ Дальнейшая обрезка нарушит смысл "
                              f"({real_pages} стр. вместо {target_pages}). "
                              f"Отдам ближайший результат.")
                        break
        else:
            # Слишком мало страниц — добавляем
            chars_to_add = int(abs(diff) * chars_per_real_page)
            print(f"[ADJUST] ➕ Добавляю {chars_to_add} знаков")
            blocks = await _expand_blocks_by_chars(
                blocks, chars_to_add, topic, model_key, writing_style, prog,
            )
        
        # Пересобираем DOCX
        docx_raw = build_docx_bytes(data, blocks, gost)
        with open(tmp_in, "wb") as f:
            f.write(docx_raw)

        # ═══ ЗАЩИТА: проверяем что текст не потерялся после обрезки ═══
        _total_text = sum(len(b[2] or "") for b in blocks)
        if _total_text < int(target_chars_goal * 0.5):
            print(f"[ADJUST] ⚠️ Текст сократился слишком сильно ({_total_text} зн. при цели {target_chars_goal}). Прерываю обрезку для сохранения смысла.")
            break
    
    # Финальный замер
    final_pages = await measure_pages_async(tmp_in, measure_dir)
    if final_pages is None:
        final_pages = target_pages
    
    print(f"[ADJUST] 🎯 Финальный результат: {final_pages} страниц")
    return blocks, final_pages
# ═══════════════════════════════════════════════════════════════
#  «СВОЙ ГОСТ» — ПАРСИНГ ТРЕБОВАНИЙ ЧЕРЕЗ ИИ
# ═══════════════════════════════════════════════════════════════

async def parse_custom_gost_via_ai(model_key: str, gost_text: str) -> dict:
    schema_hint = {
        "font_name":            "Times New Roman",
        "font_size":            14,
        "line_spacing":         1.5,
        "first_line_indent_cm": 1.25,
        "left_margin_mm":       30,
        "right_margin_mm":      10,
        "top_margin_mm":        20,
        "bottom_margin_mm":     20,
        "page_number_position": "bottom_center",
    }

    system = (
        "Ты извлекаешь параметры оформления из описания требований пользователя. "
        "Ответ СТРОГО в JSON без объяснений и без markdown. "
        "Если параметр не указан — используй значение по умолчанию из schema."
    )

    user = (
        "schema (defaults):\n"
        + json.dumps(schema_hint, ensure_ascii=False)
        + "\n\nТребования пользователя:\n"
        + gost_text
        + "\n\nВерни ТОЛЬКО JSON."
    )

    raw, _ = await chat_with_fallback(
        model_key,
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        max_tokens=800,
    )
    raw = raw.strip()

    m = re.search(r"\{.*\}", raw, flags=re.S)
    if m:
        raw = m.group(0)

    try:
        data = json.loads(raw)
    except Exception:
        return schema_hint

    out = dict(schema_hint)
    for k in out:
        if k in data and data[k] not in (None, ""):
            out[k] = data[k]

    # Приведение типов
    try:
        out["font_size"] = int(out["font_size"])
    except Exception:
        out["font_size"] = 14

    for fkey in ("line_spacing", "first_line_indent_cm"):
        try:
            out[fkey] = float(out[fkey])
        except Exception:
            out[fkey] = schema_hint[fkey]

    for ikey in ("left_margin_mm", "right_margin_mm", "top_margin_mm", "bottom_margin_mm"):
        try:
            out[ikey] = int(out[ikey])
        except Exception:
            out[ikey] = schema_hint[ikey]

    if out.get("page_number_position") not in ("bottom_center", "top_center"):
        out["page_number_position"] = "bottom_center"

    return out


# ═══════════════════════════════════════════════════════════════
#  КЛАВИАТУРЫ — КРАСИВЫЕ
# ════════════════���══════════════════════════════════════════════

def kb_doc_type() -> InlineKeyboardMarkup:
    # (user-patch): кнопки типов документа — по одной в ряд, а внизу
    # ПАРОЙ в одной строке стоят «✏️ Исправить готовую работу»
    # и «⚙️ Настроить ГОСТ», чтобы быстрые действия были рядом.
    rows: list[list[InlineKeyboardButton]] = []
    for k, v in DOC_TYPES.items():
        label = f"{v['name']}  {v['min_pages']}–{v['max_pages']} стр."
        rows.append([InlineKeyboardButton(text=label, callback_data=f"dtype_{k}")])
    rows.append([
        InlineKeyboardButton(text="✏️ Исправить готовую", callback_data="edit_start"),
        InlineKeyboardButton(text="⚙️ Настроить ГОСТ",   callback_data="custom_gost"),
    ])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def kb_mode() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    b.button(text="🆓 Бесплатно  (быстро)", callback_data="mode_free")
    b.button(text="⭐ Платный режим  (безлимит)", callback_data="mode_paid")
    b.adjust(1)
    return b.as_markup()


def kb_writing_style() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(
                text="🎓 Умно — сложный академический язык, термины, глубокий анализ",
                callback_data="style_smart",
            )],
            [InlineKeyboardButton(
                text="📝 Классически — чёткий деловой стиль, понятные формулировки",
                callback_data="style_classic",
            )],
        ]
    )


def kb_source_choice() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="✅ Да — добавлю свои материалы", callback_data="source_yes")],
            [InlineKeyboardButton(text="❌ Нет — только по теме",         callback_data="source_no")],
        ]
    )


def kb_institution() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    for k, v in INSTITUTION_TYPES.items():
        b.button(text=v["name"], callback_data=f"inst_{k}")
    b.adjust(2)
    return b.as_markup()


def kb_subject() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    for s in SUBJECTS:
        b.button(text=s, callback_data=f"subj_{s}")
    b.button(text="✏️ Другой предмет", callback_data="subj_other")
    b.adjust(3)
    return b.as_markup()


def kb_city() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    for c in CITIES:
        b.button(text=c, callback_data=f"city_{c}")
    b.button(text="✏️ Другой город", callback_data="city_other")
    b.adjust(3)
    return b.as_markup()


def kb_page_number() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="⬇️ Внизу по центру (стандарт)", callback_data="pagepos_bottom")],
            [InlineKeyboardButton(text="⬆️ Вверху по центру",           callback_data="pagepos_top")],
        ]
    )


def kb_image_mode() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="🖼 С изображениями", callback_data="images_yes")],
            [InlineKeyboardButton(text="📄 Без изображений", callback_data="images_no")],
        ]
    )


def kb_image_source() -> InlineKeyboardMarkup:
    """Только один источник изображений — ссылки от пользователя.
    (user-patch): авто-поиск убран; загрузка собственных фото-файлов
    тоже убрана — пользователь присылает прямые ссылки на изображения.
    """
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="🔗 Прислать ссылки на изображения", callback_data="imgsrc_own")],
            [InlineKeyboardButton(text="🚫 Без изображений", callback_data="imgsrc_skip")],
        ]
    )


def kb_own_images_done() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="✅ Готово, ссылки отправлены", callback_data="own_images_done")],
            [InlineKeyboardButton(text="🔙 Назад", callback_data="nav_back")],
        ]
    )


def kb_image_count(pages: int) -> InlineKeyboardMarkup:
    auto_count = _image_count_for_pages(pages)
    max_count = max(1, MAX_WORK_IMAGES)
    rows = [
        [InlineKeyboardButton(text=f"📐 По ГОСТ/авто — {auto_count} шт.", callback_data="imgcount_auto")],
        [
            InlineKeyboardButton(text="1", callback_data="imgcount_1"),
            InlineKeyboardButton(text="2", callback_data="imgcount_2"),
            InlineKeyboardButton(text="3", callback_data="imgcount_3"),
        ],
    ]
    if max_count >= 5:
        rows.append([
            InlineKeyboardButton(text="4", callback_data="imgcount_4"),
            InlineKeyboardButton(text="5", callback_data="imgcount_5"),
            InlineKeyboardButton(text="✏️ Своё", callback_data="imgcount_custom"),
        ])
    else:
        rows.append([InlineKeyboardButton(text="✏️ Своё", callback_data="imgcount_custom")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def kb_humanize() -> InlineKeyboardMarkup:
    # Оставлено только для совместимости со старыми сообщениями Telegram.
    # Опечатки и искусственные неточности отключены в обработчике.
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="📝 Продолжить без опечаток и маркеров ИИ", callback_data="humanize_no")],
        ]
    )


def kb_final() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="🔄 Перегенерировать эту работу", callback_data="regenerate")],
            [InlineKeyboardButton(text="🔁 Создать новую работу", callback_data="final_new")],
            # (user-patch): кнопка «Исправить готовую работу» убрана из
            # пост-генерационного меню — она доступна в стартовом меню /start
            # и по команде /edit.
            [InlineKeyboardButton(text="⚙️ Настроить ГОСТ",       callback_data="custom_gost")],
            [InlineKeyboardButton(text="📊 Мой баланс генераций", callback_data="show_limits")],
        ]
    )


def kb_edit_intro() -> InlineKeyboardMarkup:
    """Стартовое меню режима редактирования."""
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="📎 Прислать файл для правок", callback_data="edit_upload")],
            [InlineKeyboardButton(text="🏠 На главную", callback_data="final_new")],
        ]
    )


def kb_edit_cancel() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="❌ Отменить редактирование", callback_data="cancel_flow")],
        ]
    )


def kb_edit_pay(stars: int) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text=f"⭐ Оплатить правки — {stars}", callback_data="edit_pay")],
            [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_flow")],
        ]
    )


def _edit_price_notice(user_id: int) -> str:
    """Текст цены режима правок: для VIP/премиум из VIP_USERS — бесплатно."""
    if is_vip(user_id):
        return "👑 <b>Премиум/VIP:</b> исправление готовых работ бесплатно."
    return f"Стоимость правок: <b>{EDIT_PRICE}⭐</b> за один документ."


def kb_models() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    for k, info in AI_MODELS.items():
        if not info.get("api_key") or info.get("_fatal"):
            continue
        status = info.get("status", ModelStatus.UNKNOWN)
        b.button(
            text=f"{info['name']}  {info['price_per_page']}⭐/стр  {status}",
            callback_data=f"model_{k}",
        )
    b.adjust(1)
    return b.as_markup()


def kb_cancel() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[[InlineKeyboardButton(text="❌ Отмена", callback_data="cancel_flow")]]
    )


def with_back(markup: Optional[InlineKeyboardMarkup] = None, *, cancel: bool = True) -> InlineKeyboardMarkup:
    """Добавляет к любой inline-клавиатуре навигацию «Назад / Отмена»."""
    rows = [list(row) for row in (markup.inline_keyboard if markup else [])]
    nav = [InlineKeyboardButton(text="← Назад", callback_data="back_flow")]
    if cancel:
        nav.append(InlineKeyboardButton(text="❌ Отмена", callback_data="cancel_flow"))
    rows.append(nav)
    return InlineKeyboardMarkup(inline_keyboard=rows)


def kb_back_cancel() -> InlineKeyboardMarkup:
    return with_back(None, cancel=True)


# ═══════════════════════════════════════════════════════════════
#  FSM СОСТОЯНИЯ
# ═══════════════════════════════════════════════════════════════

# ============================================================
# УНИВЕРСАЛЬНАЯ ПРОВЕРКА ТЕМЫ И ДИСЦИПЛИНЫ (БЕЗ ХАРДКОДА)
# ============================================================

async def verify_discipline_relevance_universal(
    model_key: str,
    topic: str,
    subject: str,
    doc_type: str = "",
) -> tuple[bool, str, list[str]]:
    """Универсальная проверка соответствия темы и дисциплины."""
    if not topic or not subject:
        return False, "Тема или дисциплина не указаны", []

    system = (
        "Ты — эксперт по академическим дисциплинам. Оцени, соответствует ли тема "
        "заявленной учебной дисциплине. Будь строгим: если тема слишком общая, "
        "состоит только из имени/персоны/объекта/события и в ней не указан аспект "
        "выбранной дисциплины, верни match=false, даже если связь можно придумать.\n\n"
        'Ответ в формате JSON: {"match": true|false, "reason": "...", "suggested": ["..."]}'
    )

    user = (
        f"Тема: «{topic}»\n"
        f"Дисциплина: «{subject}»\n"
        "Оцени соответствие. Если связь неочевидна, объясни коротко и предложи "
        "в поле suggested более подходящие дисциплины для этой темы."
    )

    try:
        raw, _ = await chat_with_fallback(
            model_key,
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            max_tokens=400,
        )
        m = re.search(r"\{.*\}", raw, flags=re.S)
        if m:
            data = json.loads(m.group(0))
            suggested = data.get("suggested", [])
            if isinstance(suggested, list):
                suggested = [s.strip() for s in suggested if s.strip()]
            return (bool(data.get("match", True)), str(data.get("reason", "")).strip(), suggested[:5])
    except Exception as e:
        print(f"[RELEVANCE] Ошибка: {e}")

    return True, "проверка недоступна", []

def extract_topic_keywords(topic: str) -> list[str]:
    """Извлекает ключевые слова из темы."""
    if not topic:
        return []
    stopwords = {"тема", "работа", "исследование", "анализ", "роль", "значение",
                 "основы", "особенности", "проблемы", "вопросы", "современный",
                 "современная", "современное", "развитие", "система", "метод",
                 "методы", "подход", "подходы", "россии", "российской", "российский"}
    words = re.findall(r'[A-Za-zА-Яа-яЁё][A-Za-zА-Яа-яЁё\-]{3,}', topic.lower())
    keywords = []
    seen = set()
    for w in words:
        w = w.strip("-–—_")
        if w not in stopwords and w not in seen and len(w) >= 4:
            seen.add(w)
            keywords.append(w)
    return keywords[:10]

def map_keywords_to_disciplines(keywords: list[str]) -> list[str]:
    """Универсальное сопоставление ключевых слов с дисциплинами."""
    if not keywords:
        return []
    discipline_groups = {
        "Информатика": ["компьютер", "программ", "алгоритм", "данн", "информац", "цифр", "вычисл"],
        "Математика": ["числ", "уравнени", "функц", "вероятност", "статистик"],
        "Физика": ["энерг", "движени", "волн", "пол", "атом", "квант"],
        "Биология": ["клетк", "ген", "эволюц", "организм", "вид", "белк"],
        "История": ["событи", "период", "век", "войн", "государств"],
        "Философия": ["сущност", "сознани", "быти", "познани", "этик"],
        "Психология": ["личность", "восприяти", "поведени", "эмоц", "психик"],
        "Экономика": ["рынок", "финанс", "капитал", "инвестиц", "производств"],
        "Литература": ["поэт", "роман", "рассказ", "стих", "жанр"],
        "Русский язык": ["язык", "речь", "грамматик", "морфолог", "синтаксис"],
    }
    matched = set()
    for kw in keywords:
        kw_lower = kw.lower()
        for discipline, markers in discipline_groups.items():
            for marker in markers:
                if marker in kw_lower or kw_lower in marker:
                    matched.add(discipline)
                    break
    if not matched:
        return ["Литература", "История", "Обществознание", "Русский язык", "Философия"]
    return list(matched)[:5]

async def check_and_suggest_discipline_universal(
    model_key: str, topic: str, subject: str, doc_type: str = "",
) -> tuple[bool, str, list[str]]:
    """Универсальная проверка соответствия темы и дисциплины."""
    keywords = extract_topic_keywords(topic)
    suggested_by_keywords = map_keywords_to_disciplines(keywords)
    subject_lower = subject.lower()
    is_in_suggested = any(d.lower() in subject_lower or subject_lower in d.lower() for d in suggested_by_keywords)
    
    if is_in_suggested:
        return True, "", suggested_by_keywords
    
    match, reason, suggested_by_ai = await verify_discipline_relevance_universal(model_key, topic, subject, doc_type)
    all_suggested = list(dict.fromkeys(suggested_by_ai + suggested_by_keywords)) or ["Литература", "История", "Обществознание", "Философия"]
    
    if match:
        return True, "", all_suggested[:5]
    
    suggested_list = "\n".join(f"  • {d}" for d in all_suggested[:5])
    message = (f"⚠️ <b>Внимание!</b>\n\nТема «{topic}» может не совсем соответствовать дисциплине «{subject}».\n\n"
               f"<b>Причина:</b> {reason}\n\n<b>Рекомендуемые дисциплины:</b>\n{suggested_list}\n\n"
               f"Вы можете:\n1️⃣ <b>Изменить дисциплину</b> — выберите из списка\n"
               f"2️⃣ <b>Уточнить тему</b>\n3️⃣ <b>Продолжить</b> — если уверены")
    

    return False, message, all_suggested[:5]


def _short_btn(text: str, limit: int = 58) -> str:
    """Короткая подпись для inline-кнопки."""
    text = " ".join((text or "").split())
    return text if len(text) <= limit else text[: limit - 1].rstrip() + "…"


def _fallback_topic_suggestions(topic: str, subject: str) -> list[str]:
    """Надёжные варианты уточнения темы, если ИИ-подсказки недоступны."""
    topic = " ".join((topic or "").split()).strip(" .") or "выбранной темы"
    subject_l = (subject or "").lower()

    if any(x in subject_l for x in ("юрис", "прав", "закон")):
        if "байден" in topic.lower() or "biden" in topic.lower():
            return [
                "Правовые аспекты политики администрации Джо Байдена в области иммиграции",
                "Конституционные вопросы, связанные с указами президента Джо Байдена",
                "Правовое регулирование внешней политики США в период администрации Джо Байдена",
            ]
        return [
            f"Правовые аспекты темы «{topic}»",
            f"Конституционно-правовые вопросы, связанные с темой «{topic}»",
            f"Юридическая ответственность и правовое регулирование в контексте темы «{topic}»",
        ]
    if "истор" in subject_l:
        return [
            f"Исторические предпосылки и последствия темы «{topic}»",
            f"Тема «{topic}» в историческом контексте",
            f"Хронология и источники по теме «{topic}»",
        ]
    if "эконом" in subject_l:
        return [
            f"Экономические аспекты темы «{topic}»",
            f"Влияние темы «{topic}» на экономические процессы",
            f"Финансово-экономическая оценка темы «{topic}»",
        ]
    if "психолог" in subject_l:
        return [
            f"Психологические аспекты ��емы «{topic}»",
            f"Влияние темы «{topic}» на поведение и восприятие человека",
            f"Социально-психологический анализ темы «{topic}»",
        ]

    return [
        f"{subject}: основные аспекты темы «{topic}»",
        f"Анализ темы «{topic}» в рамках дисциплины «{subject}»",
        f"Проблемы и методы изучения темы «{topic}» по дисциплине «{subject}»",
    ]


async def suggest_subject_focused_topics(
    model_key: str,
    topic: str,
    subject: str,
    reason: str = "",
    n: int = 3,
) -> list[str]:
    """Предлагает темы, переформулированные под выбранную дисциплину."""
    fallback = _fallback_topic_suggestions(topic, subject)
    system = (
        "Ты — методист. Нужно переформулировать тему учебной работы так, чтобы она "
        "явно соответствовала указанной дисциплине. Верни СТРОГО JSON без markdown: "
        '{"topics": ["...", "...", "..."]}. '
        "Каждая тема должна быть конкретной, академичной, на русском языке, до 120 символов. "
        "Не предлагай менять дисциплину — только уточняй тему под неё."
    )
    user = (
        f"Исходная тема: «{topic}»\n"
        f"Дисциплина: «{subject}»\n"
        f"Причина сомнения: {reason or 'тема слишком общая или неочевидно связана с дисциплиной'}\n"
        f"Дай {n} подходящих варианта темы."
    )
    try:
        raw, _ = await chat_with_fallback(
            model_key,
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            max_tokens=500,
        )
        m = re.search(r"\{.*\}", raw or "", flags=re.S)
        if m:
            data = json.loads(m.group(0))
            topics = data.get("topics", [])
            if isinstance(topics, list):
                cleaned = []
                for t in topics:
                    t = " ".join(str(t).split()).strip(" \"'«»")
                    if t and t not in cleaned:
                        cleaned.append(t)
                if cleaned:
                    return cleaned[:n]
    except Exception as e:
        print(f"[RELEVANCE] Не удалось получить варианты тем: {e}")
    return fallback[:n]


def _fallback_subject_suggestions(topic: str, current_subject: str = "") -> list[str]:
    """Резервные варианты предметов, если ИИ-подсказки недоступны."""
    topic_l = (topic or "").lower()
    if "байден" in topic_l or "biden" in topic_l:
        base = ["Политология", "История", "Международные отношения", "Обществознание"]
    elif any(x in topic_l for x in ("президент", "выбор", "парламент", "партия", "государств")):
        base = ["Политология", "Обществознание", "История", "Право"]
    elif any(x in topic_l for x in ("рынок", "финанс", "банк", "эконом")):
        base = ["Экономика", "Менеджмент", "Обществознание"]
    elif any(x in topic_l for x in ("роман", "стих", "поэт", "литератур")):
        base = ["Литература", "Русский язык", "История"]
    elif any(x in topic_l for x in ("компьют", "алгоритм", "данн", "программ", "ии", "нейросет")):
        base = ["Информатика", "Математика", "Технология"]
    else:
        base = ["Обществознание", "История", "Филос��фия", "Социология"]

    cur = (current_subject or "").lower()
    return [s for s in base if s.lower() != cur][:3]


async def suggest_suitable_subjects(
    model_key: str,
    topic: str,
    current_subject: str = "",
    n: int = 3,
) -> list[str]:
    """Предлагает предметы, к которым исходная тема подходит без изменения."""
    fallback = _fallback_subject_suggestions(topic, current_subject)
    system = (
        "Ты — методист. Нужно подобрать учебные дисциплины, которым исходная тема "
        "подходит лучше всего без изменения темы. Верни СТРОГО JSON без markdown: "
        '{"subjects": ["...", "...", "..."]}. '
        "Названия дисциплин должны быть короткими: например, История, Политология, "
        "Обществознание, Экономика, Юриспруденция."
    )
    user = (
        f"Исходная тема: «{topic}»\n"
        f"Текущая дисциплина: «{current_subject}»\n"
        f"Дай {n} более подходящих дисциплины для этой темы."
    )
    try:
        raw, _ = await chat_with_fallback(
            model_key,
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            max_tokens=350,
        )
        m = re.search(r"\{.*\}", raw or "", flags=re.S)
        if m:
            data = json.loads(m.group(0))
            subjects = data.get("subjects", [])
            if isinstance(subjects, list):
                cleaned = []
                cur = (current_subject or "").lower()
                for subj in subjects:
                    subj = " ".join(str(subj).split()).strip(" \"'«»")
                    if subj and subj.lower() != cur and subj not in cleaned:
                        cleaned.append(subj)
                if cleaned:
                    return cleaned[:n]
    except Exception as e:
        print(f"[RELEVANCE] Не удалось получить варианты предметов: {e}")
    return fallback[:n]


def _topic_subject_obviously_match(topic: str, subject: str) -> bool:
    """Быстрый локальный фильтр: если тема явно подходит предмету, не спорим.

    Нужен, чтобы бот не показывал лишние предупреждения по нормальным связкам
    вроде «география + международная торговля/территориальные изменения».
    """
    t = (topic or "").lower().replace("ё", "е")
    s = (subject or "").lower().replace("ё", "е")
    if not t or not s:
        return False
    rules = {
        "географ": ("географ", "территор", "простран", "регион", "международн", "торгов", "миграц", "страны", "глобаль", "геополит"),
        "истор": ("истор", "период", "войн", "революц", "эволюц", "президент", "правлен", "хронолог"),
        "эконом": ("эконом", "рынок", "торгов", "финанс", "инвест", "санкц", "тариф", "экспорт", "импорт"),
        "прав": ("прав", "закон", "юрид", "конституц", "кодекс", "ответствен", "регулирован"),
        "юрис": ("прав", "закон", "юрид", "конституц", "кодекс", "ответствен", "регулирован"),
        "обществ": ("обществен", "полит", "государств", "социаль", "выбор", "президент"),
        "полит": ("полит", "государств", "президент", "геополит", "международн"),
        "информ": ("информ", "алгоритм", "данн", "программ", "нейросет", "цифров", "компьют"),
        "биолог": ("биолог", "организм", "экосистем", "клетк", "вид", "ген", "био"),
        "литерат": ("литератур", "роман", "повест", "поэт", "писател", "образ", "произвед"),
    }
    for subj_key, words in rules.items():
        if subj_key in s and any(w in t for w in words):
            return True
    return False


async def check_topic_subject_warning(
    model_key: str,
    topic: str,
    subject: str,
    doc_type: str = "",
) -> tuple[bool, str, list[str], list[str]]:
    """Проверяет связку тема/дисциплина и готовит подсказки для изменения темы/предмета."""
    topic_clean = " ".join((topic or "").split())
    subject_clean = " ".join((subject or "").split())
    if not topic_clean or not subject_clean:
        return True, "", [], []
    if _topic_subject_obviously_match(topic_clean, subject_clean):
        return True, "", [], []

    # Короткие темы из 1-2 слов часто не показывают дисциплинарный аспект:
    # «Байден» для юриспруденции, «Байкал» для информатики и т.п.
    significant_words = re.findall(r"[A-Za-zА-Яа-яЁё][A-Za-zА-Яа-яЁё\-]{2,}", topic_clean)
    too_general = len(significant_words) <= 2 and len(topic_clean) <= 40

    match, reason, _ = await verify_discipline_relevance_universal(
        model_key, topic_clean, subject_clean, doc_type
    )

    if match and not too_general:
        return True, "", [], []

    if too_general and match:
        reason = "тема сформулирована слишком общо, предметный аспект не указан явно"

    topic_suggestions = await suggest_subject_focused_topics(
        model_key, topic_clean, subject_clean, reason=reason, n=3
    )
    subject_suggestions = await suggest_suitable_subjects(
        model_key, topic_clean, current_subject=subject_clean, n=3
    )
    return (
        False,
        reason or "связь темы с дисциплиной неочевидна",
        topic_suggestions,
        subject_suggestions,
    )


def kb_topic_relevance_options(
    topic_suggestions: list[str],
    subject_suggestions: Optional[list[str]] = None,
) -> InlineKeyboardMarkup:
    """Кнопки под предупреждением о неочевидной связи темы и дисциплины."""
    rows: list[list[InlineKeyboardButton]] = []

    for i, topic in enumerate((topic_suggestions or [])[:3]):
        rows.append([
            InlineKeyboardButton(
                text=f"✅ Тема: {_short_btn(topic, 50)}",
                callback_data=f"topic_suggest_{i}",
            )
        ])

    for i, subj in enumerate((subject_suggestions or [])[:3]):
        rows.append([
            InlineKeyboardButton(
                text=f"📚 Предмет: {_short_btn(subj, 48)}",
                callback_data=f"topic_subject_{i}",
            )
        ])

    rows.append([
        InlineKeyboardButton(text="✏️ Написать свою тему", callback_data="topic_custom"),
        InlineKeyboardButton(text="📚 Другой предмет", callback_data="topic_subject_custom"),
    ])
    rows.append([InlineKeyboardButton(text="➡️ Продолжить генерацию", callback_data="topic_continue")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def topic_relevance_warning_text(
    topic: str,
    subject: str,
    topic_suggestions: list[str],
    subject_suggestions: Optional[list[str]] = None,
) -> str:
    """Текст предупреждения, как в примере пользователя."""
    topic_h = html.escape(topic or "")
    subject_h = html.escape(subject or "")
    examples = " или ".join(f"«{html.escape(t)}»" for t in (topic_suggestions or [])[:2])
    subjects = ", ".join(html.escape(s) for s in (subject_suggestions or [])[:3])
    advice = (
        f" например: {examples}." if examples else
        " указав предметный аспект, метод или правовой/научный контекст."
    )
    subject_advice = f"\n\n📚 <b>Или измените предмет</b> на более подходящий для этой темы: {subjects}." if subjects else ""
    return (
        "⚠️ <b>Внимание: неочевидная связь!</b>\n\n"
        f"Тема <i>«{topic_h}»</i> может плохо соответствовать дисциплине "
        f"<i>«{subject_h}»</i>.\n\n"
        "💡 <b>Совет:</b> Уточните тему, сфокусировавшись на аспектах выбранной дисциплины,"
        f"{advice}"
        f"{subject_advice}\n\n"
        "Вы можете изменить тему, изменить предмет, написать свой вариант или продолжить генерацию."
    )


async def proceed_to_city_prompt(event: Message, state: FSMContext, *, edit: bool = False) -> None:
    """Переход к выбору города после подтверждения/изменения темы."""
    text = "🌆 <b>Выберите город</b>:"
    if edit:
        await event.edit_text(text, reply_markup=with_back(kb_city()), parse_mode="HTML")
    else:
        await event.answer(text, reply_markup=with_back(kb_city()), parse_mode="HTML")
    await state.set_state(WorkState.city)


async def handle_subject_selected(event: Message, state: FSMContext, subject: str, *, edit: bool = False) -> None:
    """Сохраняет предмет и при необходимости показывает предупреждение по теме."""
    await state.update_data(subject=subject)
    data = await state.get_data()
    topic = (data.get("topic", "") or "").strip()
    doc_type = data.get("doc_type", "")

    ok, _reason, topic_suggestions, subject_suggestions = await check_topic_subject_warning(
        FREE_MODEL_KEY, topic, subject, doc_type
    )
    if ok:
        prefix = f"✅ Предмет: <b>{html.escape(subject)}</b>\n\n"
        if edit:
            await event.edit_text(
                prefix + "🌆 <b>Выберите город</b>:",
                reply_markup=with_back(kb_city()),
                parse_mode="HTML",
            )
            await state.set_state(WorkState.city)
        else:
            await event.answer(
                prefix + "🌆 <b>Выберите город</b>:",
                reply_markup=with_back(kb_city()),
                parse_mode="HTML",
            )
            await state.set_state(WorkState.city)
        return

    await state.update_data(
        topic_suggestions=topic_suggestions,
        subject_suggestions=subject_suggestions,
    )
    text = topic_relevance_warning_text(topic, subject, topic_suggestions, subject_suggestions)
    markup = kb_topic_relevance_options(topic_suggestions, subject_suggestions)
    if edit:
        await event.edit_text(text, reply_markup=markup, parse_mode="HTML")
    else:
        await event.answer(text, reply_markup=markup, parse_mode="HTML")
    await state.set_state(WorkState.topic_adjustment)


class WorkState(StatesGroup):
    doc_type           = State()
    custom_doc_name    = State()
    mode               = State()
    writing_style      = State()
    topic              = State()
    source_choice      = State()
    source_content     = State()
    institution_type   = State()
    org_type           = State()
    institution        = State()
    group              = State()
    author             = State()
    teacher            = State()
    subject            = State()
    city               = State()
    pages              = State()
    page_number_position = State()
    image_mode         = State()
    image_source       = State()   # свои изображения (дешевле) или авто-поиск
    own_images_upload  = State()   # ждём фото от пользователя
    image_count        = State()
    model              = State()
    payment            = State()
    humanize           = State()
    topic_adjustment   = State()
    gost_free_text     = State()
    # ── Редактирование/исправление готовой работы (за ⭐) ──
    edit_upload        = State()   # ждём DOCX/PDF от пользователя
    edit_instructions  = State()   # ждём текст с правками


# ═══════════════════════════════════════════════════════════════
#  БОТ И ДИСПЕТЧЕР
# ═══════════════════════════════════════════════════════════════

bot: Optional[Bot] = None
dp  = Dispatcher(storage=MemoryStorage())


# ═══════════════════════════════════════════════════════════════
#  HTTP СЕРВЕР ДЛЯ RENDER HEALTHCHECK
# ═══════════════════════════════════════════════════════════════

async def _start_http_server() -> None:
    port = int(os.getenv("PORT", "10000"))

    async def health(_: web.Request) -> web.Response:
        return web.Response(
            text=json.dumps({"status": "ok", "bot": "GOST Assistant"}, ensure_ascii=False),
            content_type="application/json",
        )

    app = web.Application()
    app.add_routes([web.get("/health", health), web.get("/", health)])

    runner = web.AppRunner(app)
    await runner.setup()
    site = web.TCPSite(runner, host="0.0.0.0", port=port)
    await site.start()
    print(f"[HTTP] Healthcheck на порту {port}")


# ═══════════════════════════════════════════════════════════════
#  ВСПОМОГАТЕЛЬНЫЕ ТЕКСТЫ СООБЩЕНИЙ
# ═══════════════════════════════════════════════════════════════

def _welcome_text(user_id: int, first_name: str) -> str:
    limits = get_user_limits_info(user_id)
    vip_mark = " 👑" if is_vip(user_id) else ""
    return (
        f"👋 Привет, <b>{first_name}</b>{vip_mark}!\n\n"
        f"📚 <b>ГОСТ-АССИСТЕНТ</b> — создаёт академические работы\n"
        f"в формате DOCX строго по ГОСТ 7.32-2017.\n\n"
        f"<b>Что умею:</b>\n"
        f"• 📄 Рефераты, курсовые, эссе, доклады\n"
        f"• 📑 Развёрнутые названия глав и подглав\n"
        f"• 📐 Точное соблюдение объёма (±10%)\n"
        f"• 🗂 Автоматическое содержание\n"
        f"• ⚙️ Настройка ГОСТ под ваш вуз\n\n"
        f"{limits}\n\n"
        f"👇 <b>Выберите тип работы:</b>"
    )


def _doc_type_card(doc_type: str, user_id: int) -> str:
    dt     = DOC_TYPES[doc_type]
    limits = get_user_limits_info(user_id)
    return (
        f"✅ <b>{dt['name']}</b>\n\n"
        f"📖 <i>{dt['desc']}</i>\n\n"
        f"📄 Объём: <b>{dt['min_pages']}–{dt['max_pages']} страниц</b>\n"
        f"🗂 Структура: {dt['structure']}\n\n"
        f"{limits}\n\n"
        f"💳 <b>Выберите режим генерации:</b>"
    )


def _mode_free_text() -> str:
    period_h = FREE_COOLDOWN // 3600
    period_s = f"{period_h // 24} дн" if period_h >= 24 else f"{period_h} ч"
    return (
        "🆓 <b>Бесплатный режим</b>\n\n"
        f"• Модель: {AI_MODELS.get(FREE_MODEL_KEY, {}).get('name', 'DeepSeek')}\n"
        f"• Максимум: <b>{FREE_MAX_PAGES} страниц</b>\n"
        f"• Лимит: <b>1 генерация раз в {period_s}</b>\n\n"
        "✍️ Введите <b>тему работы</b> одной строкой:"
    )


def _mode_paid_text() -> str:
    paid_str = "♾ безлимитно" if PAID_DAILY_LIMIT == 0 else f"{PAID_DAILY_LIMIT} в день"
    return (
        "⭐ <b>Платный режим</b>\n\n"
        f"• Любая доступная модель ИИ\n"
        f"• Любое количество страниц\n"
        f"• Генерации: <b>{paid_str}</b>\n\n"
        "✍️ Введите <b>тему работы</b> одной строкой:"
    )


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — /start и общие команды
# ═══════════════════════════════════════════════════════════════

async def _extract_docx_text(bot: "Bot", document, max_chars: int = 12000) -> tuple[str, str]:
    """Скачивает DOCX из Telegram и извлекает текст.

    Возвращает (text, error). text — извлечённый текст (без markdown),
    error — пустая строка при успехе, иначе описание ошибки.

    max_chars — ограничение для контекста ИИ. Для режима редактирования
    передаётся большее значение, чтобы захватить всю работу.
    """
    if not document:
        return "", "документ не передан"
    try:
        file = await bot.get_file(document.file_id)
        buf = io.BytesIO()
        await bot.download_file(file.file_path, buf)
        buf.seek(0)
        doc = Document(buf)
        paras: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                paras.append(t)
        # Таблицы тоже
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        paras.append(t)
        if not paras:
            return "", "документ пустой или не содержит читаемого текста"
        text = "\n\n".join(paras)
        # Ограничиваем размер для ИИ-контекста
        if len(text) > max_chars:
            text = text[:max_chars].rsplit(" ", 1)[0] + "…"
        return text, ""
    except Exception as e:
        return "", f"не удалось прочитать документ: {e}"


@dp.message(Command("start"))
async def cmd_start(message: Message, state: FSMContext) -> None:
    await state.clear()
    first = message.from_user.first_name or "пользователь"
    await message.answer(
        _welcome_text(message.from_user.id, first),
        reply_markup=kb_doc_type(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.doc_type)


@dp.message(Command("help"))
async def cmd_help(message: Message, state: FSMContext) -> None:
    await message.answer(
        "📖 <b>Помощь по ГОСТ-ассистенту</b>\n\n"
        "<b>Команды:</b>\n"
        "/start — начать новую работу\n"
        "/help — эта справка\n"
        "/limits — мои лимиты\n"
        "/cancel — отменить текущий ввод\n\n"
        "<b>Как это работает:</b>\n"
        "1️⃣ Выбираете тип документа\n"
        "2️⃣ Указываете тему, учреждение, данные\n"
        "3️⃣ ИИ генерирует текст и формирует DOCX\n"
        "4️⃣ Получаете готовый файл по ГОСТ\n\n"
        "<b>Платный режим:</b> выбор ИИ-модели, любой объём, без ежедневного лимита.\n\n"
        "<b>Проблемы?</b> Обратитесь к администратору бота.",
        parse_mode="HTML",
    )



@dp.message(F.document)
async def h_document(message: Message, state: FSMContext) -> None:
    """FIX 7.32-H: обработчик входящих DOCX/PDF/документов.

    Поведение:
      • Если пользователь в WorkState.source_content — скачиваем файл,
        извлекаем текст и используем его как материал для ИИ.
      • Иначе — отвечаем дружелюбным сообщением, что бот генерирует
        документы, а не редактирует присланные.
    """
    document = message.document
    cur_state = await state.get_state()
    mime = (document.mime_type or "").lower() if document else ""
    fname = (document.file_name or "").lower() if document else ""

    _is_doc = (
        mime.startswith("application/vnd.openxmlformats-officedocument.wordprocessingml")
        or mime == "application/pdf"
        or fname.endswith((".docx", ".doc", ".pdf"))
    )

    # ── РЕЖИМ РЕДАКТИРОВАНИЯ: принять работу для правок ──
    if cur_state == WorkState.edit_upload.state:
        if not _is_doc:
            await message.answer(
                "❌ Для правок пришлите <b>DOCX</b> или <b>PDF</b> с работой.",
                parse_mode="HTML", reply_markup=kb_edit_cancel(),
            )
            return
        wait = await message.answer("⏳ <b>Читаю работу...</b>", parse_mode="HTML")
        # Берём больше текста (до 40 000 символов) — это вся работа целиком.
        text, err = await _extract_docx_text(message.bot, document, max_chars=40000)
        try:
            await wait.delete()
        except Exception:
            pass
        if err or not text:
            await message.answer(
                f"❌ Не удалось прочитать файл: {html.escape(err or 'неизвестная ошибка')}\n\n"
                "Пришлите корректный DOCX.",
                parse_mode="HTML", reply_markup=kb_edit_cancel(),
            )
            return
        await state.update_data(edit_source_text=text)
        await state.set_state(WorkState.edit_instructions)
        await message.answer(
            f"✅ Работа принята (<b>{len(text)}</b> символов).\n\n"
            "✏️ Теперь <b>опишите правки</b> одним сообщением. Например:\n"
            "• «перепиши заключение без шаблонных фраз»\n"
            "• «убери воду в главе 2, сократи на треть»\n"
            "• «исправь стиль на строго научный, убери канцелярит»\n"
            "• «добавь больше анализа в раздел 1.2»",
            parse_mode="HTML", reply_markup=kb_edit_cancel(),
        )
        return

    # ── Принять DOCX/PDF как материал (только в нужном состоянии) ──
    if cur_state == WorkState.source_content.state:
        if not (mime.startswith("application/vnd.openxmlformats-officedocument.wordprocessingml")
                or mime == "application/pdf"
                or fname.endswith(".docx") or fname.endswith(".doc")
                or fname.endswith(".pdf")):
            await message.answer(
                "❌ Принимаю только <b>DOCX</b> или <b>PDF</b>.\n\n"
                "<i>Пришлите файл с планом, конспектом или тезисами — ИИ использует это как основу.</i>",
                parse_mode="HTML",
                reply_markup=kb_back_cancel(),
            )
            return

        wait = await message.answer(
            "⏳ <b>Читаю документ...</b>",
            parse_mode="HTML",
        )
        text, err = await _extract_docx_text(message.bot, document)
        try:
            await wait.delete()
        except Exception:
            pass

        if err or not text:
            await message.answer(
                f"❌ Не удалось извлечь текст из файла.\n\n<b>Причина:</b> {html.escape(err or 'неизвестная ошибка')}\n\n"
                "Пришлите <b>DOCX</b> с планом/конспектом или просто вставьте текст сообщением.",
                parse_mode="HTML",
                reply_markup=kb_back_cancel(),
            )
            return

        await state.update_data(source_content=text)
        await message.answer(
            f"✅ Документ принят! Извлечено <b>{len(text)} символов</b> текста.\n\n"
            "🏛 <b>Тип учебного заведения</b>:",
            reply_markup=with_back(kb_institution()),
            parse_mode="HTML",
        )
        await state.set_state(WorkState.institution_type)
        return

    # ── В любом другом состоянии: подсказка о /start ──
    await message.answer(
        "📄 <b>Этот бот генерирует документы по ГОСТ 7.32-2017</b>\n\n"
        "Я не редактирую присланные DOCX/PDF, но могу создать новую работу "
        "по вашей теме с оформлением по ГОСТ, титульным листом, содержанием "
        "и списком литературы.\n\n"
        "👉 Нажмите <b>/start</b>, чтобы начать новую генерацию.",
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup(
            inline_keyboard=[
                [InlineKeyboardButton(text="🚀 Создать новую работу", callback_data="final_new")],
            ]
        ),
    )

# ═══════════════════════════════════════════════════════════════
#  РЕДАКТИРОВАНИЕ / ИСПРАВЛЕНИЕ ГОТОВОЙ РАБОТЫ (за ⭐)
# ═══════════════════════════════════════════════════════════════

def _build_edited_docx_bytes(corrected_text: str, gost: dict) -> bytes:
    """Собирает аккуратный DOCX из исправленного текста по оформлению ГОСТ.

    В отличие от полного build_docx_bytes (титул + содержание), здесь мы
    бережно воспроизводим уже готовую работу: каждый абзац — отдельный
    параграф, заголовки (короткие строки без точки на конце или в ВЕРХНЕМ
    регистре) выделяем жирным. Нумерация страниц — по ГОСТ.
    """
    doc = Document()
    setup_gost_page(doc, gost)
    fn  = gost.get("font_name", "Times New Roman")
    fs  = int(gost.get("font_size", 14))
    hfs = heading_font_size(gost)

    add_page_number_field(
        doc.sections[0],
        gost.get("page_number_position", "bottom_center"),
    )

    paras = [p.strip() for p in re.split(r"\n{2,}", corrected_text) if p.strip()]
    heading_re = re.compile(r"^(введение|заключение|содержание|оглавление|список (?:литературы|использованных источников)|приложени)", re.I)
    for raw in paras:
        line = " ".join(raw.split())
        is_heading = (
            (len(line) < 90 and not line.endswith((".", "…")) and len(line.split()) <= 12)
            or line.isupper()
            or bool(heading_re.match(line))
            or bool(re.match(r"^\d+(\.\d+)*\s+\S", line))  # «1.2 Название»
        )
        p = doc.add_paragraph()
        if is_heading:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if line.isupper() or heading_re.match(line) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line)
            _set_run_font(r, fn, hfs if line.isupper() else fs, True)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(line)
            _set_run_font(r, fn, fs, False)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _apply_edits_via_ai(model_key: str, original_text: str, instructions: str) -> str:
    """Применяет правки пользователя к тексту работы через ИИ.

    Никаких выдуманных данных не добавляем: только выполняем запрошенные
    исправления, сохраняя смысл и структуру. Возвращает исправленный текст.
    """
    system = (
        "Ты — академический редактор. Тебе дают готовую учебную работу и список "
        "правок. Внеси ТОЛЬКО запрошенные изменения, сохрани структуру, заголовки "
        "и фактическое содержание. НЕ выдумывай новых авторов, фактов и источников. "
        "НЕ используй markdown (никаких #, *, **). Верни ПОЛНЫЙ исправленный текст "
        "работы целиком, абзацы разделяй пустой строкой."
    )
    user = (
        f"ТЕКСТ РАБОТЫ:\n{original_text}\n\n"
        f"СПИСОК ПРАВОК ОТ ПОЛЬЗОВАТЕЛЯ:\n{instructions}\n\n"
        "Верни полный исправленный текст работы."
    )
    text, _ = await chat_with_fallback(
        model_key,
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        max_tokens=8000,
    )
    text = (text or "").strip()
    # Снимаем возможную markdown-разметку.
    text = re.sub(r"(^|\n)#{1,6}\s*", r"\1", text)
    text = text.replace("**", "").replace("*", "")
    return text


@dp.callback_query(F.data == "edit_start")
async def h_edit_start(cb: CallbackQuery, state: FSMContext) -> None:
    await state.clear()
    await state.set_state(WorkState.edit_upload)
    price_notice = _edit_price_notice(cb.from_user.id)
    step3 = (
        "3️⃣ Правки применятся бесплатно — и вы получите исправленный файл по ГОСТ.\n\n"
        if is_vip(cb.from_user.id)
        else "3️⃣ Оплатите ⭐ — и получите исправленный файл по ГОСТ.\n\n"
    )
    await cb.message.answer(
        "✏️ <b>Исправление готовой работы</b>\n\n"
        f"{price_notice}\n\n"
        "Как это работает:\n"
        "1️⃣ Пришлите файл <b>DOCX</b> (или PDF) с вашей работой.\n"
        "2️⃣ Опишите, что нужно исправить (например: «убери воду из главы 2», "
        "«перепиши заключение без шаблонов», «исправь стиль на научный»).\n"
        f"{step3}"
        "📎 Жду ваш файл:",
        parse_mode="HTML",
        reply_markup=kb_edit_cancel(),
    )
    await cb.answer()


@dp.callback_query(F.data == "edit_upload")
async def h_edit_upload_btn(cb: CallbackQuery, state: FSMContext) -> None:
    await state.set_state(WorkState.edit_upload)
    await cb.message.answer(
        "📎 Пришлите файл <b>DOCX</b> или <b>PDF</b> с работой, которую нужно исправить.",
        parse_mode="HTML",
        reply_markup=kb_edit_cancel(),
    )
    await cb.answer()


@dp.message(Command("edit"))
async def cmd_edit(message: Message, state: FSMContext) -> None:
    await state.clear()
    await state.set_state(WorkState.edit_upload)
    await message.answer(
        "✏️ <b>Режим исправления работы</b>\n\n"
        f"{_edit_price_notice(message.from_user.id)}\n\n"
        "Пришлите файл DOCX/PDF — затем опишите правки.",
        parse_mode="HTML",
        reply_markup=kb_edit_cancel(),
    )


@dp.message(WorkState.edit_instructions, F.text)
async def h_edit_instructions(message: Message, state: FSMContext) -> None:
    instructions = (message.text or "").strip()
    if len(instructions) < 4:
        await message.answer(
            "❌ Опишите подробнее, что исправить (минимум несколько слов).",
            reply_markup=kb_edit_cancel(),
        )
        return
    await state.update_data(edit_instructions=instructions)

    # VIP/премиум пользователи из VIP_USERS исправляют готовые работы бесплатно.
    if is_vip(message.from_user.id):
        await message.answer(
            "👑 <b>Премиум/VIP — оплата не требуется.</b>\n\n"
            "✏️ Сразу применяю правки к вашей работе...",
            parse_mode="HTML",
        )
        await run_edit_and_send(message, state)
        return

    payload = json.dumps({"mode": "edit"}, ensure_ascii=False)
    await message.answer(
        "✅ Правки записаны. Для применения нужно оплатить.\n\n"
        f"К оплате: <b>{EDIT_PRICE}⭐</b>",
        parse_mode="HTML",
        reply_markup=kb_edit_pay(EDIT_PRICE),
    )
    await state.update_data(edit_payload=payload)


@dp.callback_query(F.data == "edit_pay")
async def h_edit_pay(cb: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    if not data.get("edit_source_text"):
        await cb.answer("Сначала пришлите файл и правки.", show_alert=True)
        return
    if is_vip(cb.from_user.id):
        await cb.answer("VIP: бесплатно, запускаю правки", show_alert=False)
        await cb.message.answer(
            "👑 <b>Премиум/VIP — оплата не требуется.</b>\n\n"
            "✏️ Применяю ваши правки...",
            parse_mode="HTML",
        )
        await run_edit_and_send(cb.message, state)
        return
    await cb.bot.send_invoice(
        chat_id=cb.from_user.id,
        title="Исправление работы по ГОСТ",
        description="Применение ваших правок к готовой работе с оформлением по ГОСТ.",
        payload=json.dumps({"mode": "edit"}, ensure_ascii=False),
        provider_token="",
        currency="XTR",
        prices=[LabeledPrice(label="Исправление работы", amount=EDIT_PRICE)],
        reply_markup=InlineKeyboardMarkup(
            inline_keyboard=[[InlineKeyboardButton(text=f"⭐ Оплатить {EDIT_PRICE} звёзд", pay=True)]]
        ),
    )
    await cb.answer()


async def run_edit_and_send(message: Message, state: FSMContext) -> None:
    """Применяет оплаченные правки и отправляет исправленный DOCX."""
    data = await state.get_data()
    original_text = data.get("edit_source_text", "")
    instructions  = data.get("edit_instructions", "")
    if not original_text or not instructions:
        await message.answer("❌ Не нашёл текст работы или правки. Начните заново: /edit")
        await state.clear()
        return

    model_key = FREE_MODEL_KEY
    for k, info in AI_MODELS.items():
        if info.get("api_key") and not info.get("_fatal"):
            model_key = k
            break

    wait = await message.answer("✏️ <b>Применяю правки...</b>", parse_mode="HTML")
    try:
        corrected = await _apply_edits_via_ai(model_key, original_text, instructions)
        if not corrected or len(corrected) < 100:
            corrected = original_text  # подстраховка: не теряем работу
        gost = get_gost_config("referat", message.chat.id)
        docx_bytes = _build_edited_docx_bytes(corrected, gost)
        try:
            await wait.delete()
        except Exception:
            pass
        await message.answer_document(
            BufferedInputFile(docx_bytes, filename=f"Исправленная_работа_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"),
            caption=(
                "✅ <b>Правки применены!</b>\n\n"
                "Если содержание/нумерация требуют обновления — откройте в Word: "
                "<code>Ctrl+A → F9</code>."
            ),
            parse_mode="HTML",
            reply_markup=kb_final(),
        )
    except Exception as e:
        print(f"[EDIT ERROR] {e}")
        await message.answer(f"❌ Не удалось применить правки: {str(e)[:150]}")
    finally:
        await state.clear()


@dp.message(Command("limits"))
async def cmd_limits(message: Message) -> None:
    await message.answer(
        f"📊 <b>Ваши генерации</b>\n\n{get_user_limits_info(message.from_user.id)}",
        parse_mode="HTML",
    )


@dp.message(Command("cancel"))
async def cmd_cancel(message: Message, state: FSMContext) -> None:
    await state.clear()
    await message.answer(
        "❌ <b>Отменено.</b>\n\nНачните заново — /start",
        parse_mode="HTML",
    )


@dp.callback_query(F.data == "cancel_flow")
async def h_cancel_flow(cb: CallbackQuery, state: FSMContext) -> None:
    await state.clear()
    await cb.message.edit_text(
        "❌ <b>Генерация отменена.</b>\n\nНачните заново — /start",
        parse_mode="HTML",
    )
    await cb.answer()


@dp.callback_query(F.data == "back_flow")
async def h_back_flow(cb: CallbackQuery, state: FSMContext) -> None:
    """Кнопка «Назад» для основного мастера создания работы."""
    cur = await state.get_state()
    data = await state.get_data()

    async def edit(text: str, markup: Optional[InlineKeyboardMarkup], new_state: State) -> None:
        await cb.message.edit_text(text, parse_mode="HTML", reply_markup=markup)
        await state.set_state(new_state)

    # 1. Возврат к выбору типа работы
    if cur in (WorkState.custom_doc_name.state, WorkState.mode.state, WorkState.gost_free_text.state):
        first = cb.from_user.first_name or "пользователь"
        await cb.message.edit_text(
            _welcome_text(cb.from_user.id, first),
            reply_markup=kb_doc_type(),
            parse_mode="HTML",
        )
        await state.set_state(WorkState.doc_type)
        await cb.answer()
        return

    # 2. Режим / стиль / тема
    if cur == WorkState.writing_style.state:
        doc_type = data.get("doc_type", "referat")
        await edit(_doc_type_card(doc_type, cb.from_user.id), with_back(kb_mode()), WorkState.mode)
    elif cur == WorkState.topic.state:
        await edit(
            "✍️ <b>Выберите стиль написания работы</b>\n\n"
            "Это влияет на язык, структуру предложений и глубину изложения:",
            with_back(kb_writing_style()),
            WorkState.writing_style,
        )
    elif cur == WorkState.source_choice.state:
        await edit(
            "✍️ <b>Введите тему работы</b> одной строкой:",
            kb_back_cancel(),
            WorkState.topic,
        )
    elif cur == WorkState.source_content.state:
        await edit(
            "📎 <b>Есть ли у вас свои материалы?</b>\n\n"
            "Вы можете добавить план, конспект, тезисы или ссылки на сайты — ИИ использует их как основу.",
            with_back(kb_source_choice()),
            WorkState.source_choice,
        )
    elif cur == WorkState.institution_type.state:
        await edit(
            "📎 <b>Есть ли у вас свои материалы?</b>\n\n"
            "Вы можете добавить план, конспект, тезисы или ссылки на сайты — ИИ использует их как основу.",
            with_back(kb_source_choice()),
            WorkState.source_choice,
        )

    # 3. Учебное заведение и данные титула
    elif cur == WorkState.org_type.state:
        await edit("🏛 <b>Тип учебного заведения</b>\n\nВыберите из списка:", with_back(kb_institution()), WorkState.institution_type)
    elif cur == WorkState.institution.state:
        inst_kind = data.get("institution_type", "school")
        if inst_kind == "custom":
            await edit("🏛 <b>Тип учебного заведения</b>\n\nВыберите из списка:", with_back(kb_institution()), WorkState.institution_type)
        else:
            info = INSTITUTION_TYPES.get(inst_kind, INSTITUTION_TYPES["school"])
            await edit(
                "🏛 <b>Введите тип организации</b>\n\n"
                f"Пример:\n<i>{info['org_example']}</i>\n\n"
                "Или скопируйте пример целиком.",
                kb_back_cancel(),
                WorkState.org_type,
            )
    elif cur == WorkState.group.state:
        await edit(
            "🏫 <b>Введите название учебного заведения</b>",
            kb_back_cancel(),
            WorkState.institution,
        )
    elif cur == WorkState.author.state:
        await edit(
            "👥 <b>Введите класс или группу</b>\n\n<i>Например: 10А, ИТ-21, гр. 315</i>",
            kb_back_cancel(),
            WorkState.group,
        )
    elif cur == WorkState.teacher.state:
        await edit(
            "👤 <b>Введите ФИО автора</b>\n\n<i>Например: Иванов Иван Иванович</i>",
            kb_back_cancel(),
            WorkState.author,
        )
    elif cur == WorkState.subject.state:
        await edit(
            "👨‍🏫 <b>Введите ФИО преподавателя</b>\n\n<i>Пример: Петров Пётр Петрович</i>",
            kb_back_cancel(),
            WorkState.teacher,
        )
    elif cur == WorkState.topic_adjustment.state:
        await edit("📚 <b>Выберите дисциплину (предмет)</b>", with_back(kb_subject()), WorkState.subject)
    elif cur == WorkState.city.state:
        await edit("📚 <b>Выберите дисциплину (предмет)</b>", with_back(kb_subject()), WorkState.subject)
    elif cur == WorkState.pages.state:
        await edit("🌆 <b>Выберите город:</b>", with_back(kb_city()), WorkState.city)

    # 4. Параметры генерации
    elif cur == WorkState.page_number_position.state:
        await edit(_pages_prompt_text(data, cb.from_user.id), kb_back_cancel(), WorkState.pages)
    elif cur == WorkState.image_mode.state:
        await edit(
            f"✅ Страниц: <b>{int(data.get('pages', 10))}</b>\n\n"
            "🔢 <b>Нумерация страниц — где расположить?</b>",
            with_back(kb_page_number()),
            WorkState.page_number_position,
        )
    elif cur in (WorkState.image_count.state, WorkState.image_source.state, WorkState.own_images_upload.state):
        await edit(
            "🖼 <b>Добавить изображения в работу?</b>\n\n"
            "• <b>Без изображений</b> — обычная работа.\n"
            "• <b>С изображениями</b> — свои фото (дешевле) или авто-подбор.",
            with_back(kb_image_mode()),
            WorkState.image_mode,
        )
    elif cur == WorkState.humanize.state:
        await edit(
            f"✅ Страниц: <b>{int(data.get('pages', 10))}</b>\n\n"
            "🔢 <b>Нумерация страниц — где расположить?</b>",
            with_back(kb_page_number()),
            WorkState.page_number_position,
        )
    elif cur == WorkState.model.state:
        await edit(
            f"✅ Страниц: <b>{int(data.get('pages', 10))}</b>\n\n"
            "🔢 <b>Нумерация страниц — где расположить?</b>",
            with_back(kb_page_number()),
            WorkState.page_number_position,
        )
    elif cur == WorkState.payment.state:
        await edit(
            "🤖 <b>Выберите ИИ-модель</b>\n\n"
            "Цена указана в звёздах Telegram за страницу.\n"
            "Все модели генерируют полноценный академический текст.",
            with_back(kb_models()),
            WorkState.model,
        )
    else:
        first = cb.from_user.first_name or "пользователь"
        await cb.message.edit_text(
            _welcome_text(cb.from_user.id, first),
            reply_markup=kb_doc_type(),
            parse_mode="HTML",
        )
        await state.set_state(WorkState.doc_type)

    await cb.answer()


@dp.callback_query(F.data == "show_limits")
async def h_show_limits(cb: CallbackQuery) -> None:
    await cb.answer(
        get_user_limits_info(cb.from_user.id).replace("<b>", "").replace("</b>", ""),
        show_alert=True,
    )


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — НАСТРОЙКА ГОСТ
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data == "custom_gost")
async def h_custom_gost(cb: CallbackQuery, state: FSMContext) -> None:
    await cb.message.answer(
        "⚙️ <b>Настройка ГОСТ под ваш вуз</b>\n\n"
        "Вставьте требования к оформлению обычным текстом.\n\n"
        "<b>Пример:</b>\n"
        "<i>Шрифт Arial 12pt, поля: лево 20 мм, право 10 мм, верх 15 мм, низ 15 мм. "
        "Межстрочный интервал 1.0. Отступ первой строки 1.25 см. Нумерация снизу.</i>\n\n"
        "ИИ автоматически извлечёт параметры и сохранит для вашего аккаунта.",
        parse_mode="HTML",
        reply_markup=kb_back_cancel(),
    )
    await state.set_state(WorkState.gost_free_text)
    await cb.answer()


@dp.message(WorkState.gost_free_text)
async def h_save_custom_gost(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    doc_type = data.get("doc_type", "referat")

    gost_text = (message.text or "").strip()
    if len(gost_text) < 10:
        await message.answer(
            "❌ Слишком коротко. Вставьте полные требования к оформлению.",
            reply_markup=kb_back_cancel(),
        )
        return

    wait_msg = await message.answer(
        "⏳ <b>Анализирую требования...</b>",
        parse_mode="HTML",
    )

    parsed = await parse_custom_gost_via_ai(FREE_MODEL_KEY, gost_text)
    save_user_gost_config(message.from_user.id, doc_type, parsed)

    await wait_msg.delete()

    pn = "снизу" if parsed["page_number_position"] == "bottom_center" else "сверху"
    await message.answer(
        "✅ <b>ГОСТ сохранён!</b>\n\n"
        "┌─────────────────────────\n"
        f"│ 🔤 Шрифт: <b>{parsed['font_name']} {parsed['font_size']}pt</b>\n"
        f"│ ↕️ Интервал: <b>{parsed['line_spacing']}</b>\n"
        f"│ ↩️ Отступ: <b>{parsed['first_line_indent_cm']} см</b>\n"
        f"│ 📐 Поля (мм): Л{parsed['left_margin_mm']} П{parsed['right_margin_mm']} "
        f"В{parsed['top_margin_mm']} Н{parsed['bottom_margin_mm']}\n"
        f"│ 🔢 Нумерация: <b>{pn} по центру</b>\n"
        "└─────────────────────────\n\n"
        "Эти настройки применятся к вашим следующим работам.",
        parse_mode="HTML",
    )

    await state.set_state(WorkState.doc_type)
    await message.answer(
        "👇 Выберите тип работы для продолжения:",
        reply_markup=kb_doc_type(),
    )


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — ВЫБОР ТИПА ДОКУМЕНТА
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data.startswith("dtype_"))
async def h_doc_type(cb: CallbackQuery, state: FSMContext) -> None:
    doc_type = cb.data.replace("dtype_", "", 1)
    if doc_type not in DOC_TYPES:
        await cb.answer("Неизвестный тип документа", show_alert=True)
        return

    await state.update_data(doc_type=doc_type)

    if doc_type == "custom":
        await cb.message.edit_text(
            "🧩 <b>Свой тип документа</b>\n\n"
            "Введите название документа.\n"
            "<i>Примеры: Отчёт по практике, Лабораторная работа, Реферат по материалам</i>",
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        await state.set_state(WorkState.custom_doc_name)
        await cb.answer()
        return

    await cb.message.edit_text(
        _doc_type_card(doc_type, cb.from_user.id),
        reply_markup=with_back(kb_mode()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.mode)
    await cb.answer()


@dp.message(WorkState.custom_doc_name)
async def h_custom_doc_name(message: Message, state: FSMContext) -> None:
    name = (message.text or "").strip()
    if len(name) < 2:
        await message.answer(
            "❌ Название слишком короткое. Попробуйте ещё раз.",
            reply_markup=kb_back_cancel(),
        )
        return

    await state.update_data(custom_doc_name=name)
    await message.answer(
        f"🧩 <b>Свой тип:</b> {name}\n\n"
        f"{get_user_limits_info(message.from_user.id)}\n\n"
        "💳 <b>Выберите режим генерации:</b>",
        reply_markup=with_back(kb_mode()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.mode)


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — РЕЖИМ И ТЕМА
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data.startswith("mode_"))
async def h_mode(cb: CallbackQuery, state: FSMContext) -> None:
    mode = cb.data.replace("mode_", "", 1)
    await state.update_data(mode=mode)

    await cb.message.edit_text(
        "✍️ <b>Выберите стиль написания работы</b>\n\n"
        "Это влияет на язык, структуру предложений и глубину изложения:",
        parse_mode="HTML",
        reply_markup=with_back(kb_writing_style()),
    )
    await state.set_state(WorkState.writing_style)
    await cb.answer()


@dp.callback_query(F.data.in_(["style_smart", "style_classic"]))
async def h_writing_style(cb: CallbackQuery, state: FSMContext) -> None:
    style = cb.data.replace("style_", "", 1)
    await state.update_data(writing_style=style)

    data = await state.get_data()
    mode = data.get("mode", "free")
    text = _mode_free_text() if mode == "free" else _mode_paid_text()
    style_label = "🎓 Умный стиль выбран" if style == "smart" else "📝 Классический стиль выбран"

    await cb.message.edit_text(
        f"✅ {style_label}\n\n{text}",
        parse_mode="HTML",
        reply_markup=kb_back_cancel(),
    )
    await state.set_state(WorkState.topic)
    await cb.answer()


@dp.message(WorkState.topic)
async def h_topic(message: Message, state: FSMContext) -> None:
    topic = (message.text or "").strip()
    if len(topic) < 10 or len(re.findall(r"[А-Яа-яЁёA-Za-z]{3,}", topic)) < 2:
        await message.answer(
            "❌ Тема слишком короткая. Сформулируйте тему развёрнуто (минимум 10 символов и 2 смысловых слова).\n\n"
            "<i>Например: «Влияние президентства Дональда Трампа на электоральную географию США»</i>",
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        return

    await state.update_data(topic=topic)
    await message.answer(
        f"✅ Тема принята: <b>{topic[:80]}</b>\n\n"
        "📎 <b>Есть ли у вас свои материалы?</b>\n\n"
        "Вы можете добавить план, конспект, тезисы или ссылки на сайты — ИИ использует их как основу.\n"
        "Ссылки попадут в список литературы как реальные электронные ресурсы.\n"
        "Или выбрать генерацию только по теме.",
        reply_markup=with_back(kb_source_choice()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.source_choice)


@dp.callback_query(F.data == "source_yes")
async def h_source_yes(cb: CallbackQuery, state: FSMContext) -> None:
    """Пользователь выбрал 'Да — добавлю свои материалы'"""
    await cb.message.edit_text(
        "📎 <b>Отправьте ваши материалы</b>\n\n"
        "Вставьте план, тезисы, конспект, ссылки на сайты или любой текст — ИИ использует это как основу.\n"
        "Ссылки попадут в список литературы как реальные электронные ресурсы.\n"
        "<i>Максимум 12 000 символов.</i>",
        parse_mode="HTML",
        reply_markup=kb_back_cancel(),  # <-- ИСПРАВЛЕНО
    )
    await state.set_state(WorkState.source_content)
    await cb.answer()

@dp.callback_query(F.data == "source_no")
async def h_source_no(cb: CallbackQuery, state: FSMContext) -> None:
    """Пользователь выбрал 'Нет — только по теме'"""
    await state.update_data(source_content="")
    await cb.message.edit_text(
        "🏛 <b>Тип учебного заведения</b>\n\nВыберите из списка:",
        reply_markup=with_back(kb_institution()),  # <-- ИСПРАВЛЕНО: используем правильную клавиатуру
        parse_mode="HTML",
    )
    await state.set_state(WorkState.institution_type)
    await cb.answer()

@dp.message(WorkState.source_content)
async def h_source_content(message: Message, state: FSMContext) -> None:
    content = (message.text or "").strip()
    await state.update_data(source_content=content)
    await message.answer(
        f"✅ Материалы приняты ({len(content)} символов).\n\n"
        "🏛 <b>Тип учебного заведения:</b>",
        reply_markup=with_back(kb_institution()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.institution_type)


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — УЧЕБНОЕ ЗАВЕДЕНИЕ
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data.startswith("inst_"))
async def h_inst(cb: CallbackQuery, state: FSMContext) -> None:
    inst = cb.data.replace("inst_", "", 1)
    await state.update_data(institution_type=inst)

    if inst == "custom":
        await state.update_data(org_type="")
        await cb.message.edit_text(
            "✏️ <b>Введите полное название учебного заведения</b>\n\n"
            "<i>Например: Новосибирский государственный технический университет</i>",
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        await state.set_state(WorkState.institution)
    else:
        info = INSTITUTION_TYPES.get(inst, INSTITUTION_TYPES["school"])
        await cb.message.edit_text(
            "🏛 <b>Введите тип организации</b>\n\n"
            f"Пример:\n<i>{info['org_example']}</i>\n\n"
            "Или скопируйте пример целиком.",
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        await state.set_state(WorkState.org_type)

    await cb.answer()



def _validate_org_type(text: str) -> tuple[bool, str]:
    """Проверяет тип организации для титульного листа.

    Не даёт попасть на титул короткому мусору вроде «ИЗ», который был виден
    в проверенном DOCX отдельной строкой под министерством.
    """
    raw = re.sub(r"\s+", " ", (text or "").strip())
    if len(raw) < 8:
        return False, (
            "❌ Тип организации слишком короткий. Введите полную форму.\n\n"
            "<i>Например: Муниципальное бюджетное общеобразовательное учреждение</i>"
        )
    tokens = re.findall(r"[А-Яа-яЁёA-Za-z]+", raw)
    if len(tokens) < 2 and raw.upper() not in _INSTITUTION_WHITELIST:
        return False, (
            "❌ Тип организации должен быть полным названием, а не сокращением из 1 слова.\n\n"
            "<i>Например: Федеральное государственное бюджетное образовательное учреждение высшего образования</i>"
        )
    if _is_garbage(raw):
        return False, "❌ Тип организации похож на случайный набор символов. Введите полное название."
    return True, raw


def _validate_group(text: str) -> tuple[bool, str]:
    """Проверяет класс/группу для титульного листа."""
    raw = re.sub(r"\s+", " ", (text or "").strip())
    if not raw:
        return True, ""
    # Нормальные примеры: 10А, 11-Б, ИТ-21, гр. 315, ПИ-23-1.
    if re.fullmatch(r"(?:гр\.?\s*)?[А-ЯA-ZЁ]{0,6}[- ]?\d{1,4}(?:[-/][А-ЯA-ZЁ0-9]{1,4})?|\d{1,2}[- ]?[А-ЯA-ZЁ]", raw, re.IGNORECASE):
        return True, raw
    if len(raw) <= 4 or _is_garbage(raw):
        return False, (
            "❌ Класс/группа введены некорректно.\n\n"
            "<i>Примеры: 10А, ИТ-21, гр. 315</i>"
        )
    return True, raw


@dp.message(WorkState.org_type)
async def h_org_type(message: Message, state: FSMContext) -> None:
    raw_org_type = (message.text or "").strip()
    ok_org, org_value = _validate_org_type(raw_org_type)
    if not ok_org:
        await message.answer(org_value, parse_mode="HTML", reply_markup=kb_back_cancel())
        return
    await state.update_data(org_type=org_value)
    data = await state.get_data()
    info = INSTITUTION_TYPES.get(data.get("institution_type", "school"), INSTITUTION_TYPES["school"])
    await message.answer(
        "🏫 <b>Введите название учебного заведения</b>\n\n"
        f"Пример:\n<i>{info['name_example']}</i>",
        parse_mode="HTML",
        reply_markup=kb_back_cancel(),
    )
    await state.set_state(WorkState.institution)


@dp.message(WorkState.institution)
async def h_institution(message: Message, state: FSMContext) -> None:
    raw = (message.text or "").strip()
    # (user-patch): отсекаем фейковые названия типа «ИОШЛ», «ОГОО», «АЛЛАШ»,
    # пустые строки, заглушки и одиночные слова без гласных.
    ok, normalized_or_error = _validate_institution_name(raw)
    if not ok:
        await message.answer(
            normalized_or_error,
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        return
    await state.update_data(institution=normalized_or_error)
    await message.answer(
        "👥 <b>Введите класс или группу</b>\n\n<i>Например: 10А, ИТ-21, гр. 315</i>",
        parse_mode="HTML",
        reply_markup=kb_back_cancel(),
    )
    await state.set_state(WorkState.group)


@dp.message(WorkState.group)
async def h_group(message: Message, state: FSMContext) -> None:
    raw_group = (message.text or "").strip()
    ok_group, group_value = _validate_group(raw_group)
    if not ok_group:
        await message.answer(group_value, parse_mode="HTML", reply_markup=kb_back_cancel())
        return
    await state.update_data(group=group_value)
    await message.answer(
        "👤 <b>Введите ФИО автора</b>\n\n<i>Например: Иванов Иван Иванович</i>",
        parse_mode="HTML",
        reply_markup=kb_back_cancel(),
    )
    await state.set_state(WorkState.author)


@dp.message(WorkState.author)
async def h_author(message: Message, state: FSMContext) -> None:
    raw = (message.text or "").strip()
    ok, value = _validate_fio(raw, kind="ФИО автора")
    if not ok:
        await message.answer(
            value,
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        return
    author = value
    await state.update_data(author=author)
    await message.answer(
        f"✅ <b>ФИО автора:</b> {html.escape(author)}\n\n"
        "👨‍🏫 <b>Введите ФИО преподавателя</b>\n\n<i>Пример: Петров Пётр Петрович</i>",
        parse_mode="HTML",
        reply_markup=kb_back_cancel(),
    )
    await state.set_state(WorkState.teacher)


@dp.message(WorkState.teacher)
async def h_teacher(message: Message, state: FSMContext) -> None:
    raw = (message.text or "").strip()
    ok, value = _validate_fio(raw, kind="ФИО преподавателя")
    if not ok:
        await message.answer(
            value,
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        return
    teacher = value
    await state.update_data(teacher=teacher)
    await message.answer(
        f"✅ <b>ФИО преподавателя:</b> {html.escape(teacher)}\n\n"
        "📚 <b>Выберите дисциплину (предмет)</b>",
        reply_markup=with_back(kb_subject()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.subject)


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — ПРЕДМЕТ, ГОРОД, СТРАНИЦЫ
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data.startswith("subj_"))
async def h_subject_cb(cb: CallbackQuery, state: FSMContext) -> None:
    subj = cb.data.replace("subj_", "", 1)
    data = await state.get_data()
    skip_check = bool(data.get("skip_next_relevance"))
    if subj == "other":
        await cb.message.edit_text(
            "✏️ <b>Введите название предмета</b>",
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        await state.set_state(WorkState.subject)
    elif skip_check:
        await state.update_data(subject=subj, skip_next_relevance=False)
        await cb.message.edit_text(
            f"✅ Предмет: <b>{html.escape(subj)}</b>\n\n"
            "🌆 <b>Выберите город</b>:",
            reply_markup=with_back(kb_city()),
            parse_mode="HTML",
        )
        await state.set_state(WorkState.city)
    else:
        await cb.message.edit_text(
            f"⏳ <b>Проверяю соответствие темы предмету...</b>\n\n"
            f"Предмет: <b>{html.escape(subj)}</b>\n"
            "Пожалуйста, подождите несколько секунд.",
            parse_mode="HTML",
        )
        await handle_subject_selected(cb.message, state, subj, edit=True)
    await cb.answer()


@dp.message(WorkState.subject)
async def h_subject_text(message: Message, state: FSMContext) -> None:
    subj = (message.text or "").strip()
    if len(subj) < 2:
        await message.answer(
            "❌ Введите название предмета подробнее.",
            reply_markup=kb_back_cancel(),
        )
        return
    data = await state.get_data()
    if data.get("skip_next_relevance"):
        await state.update_data(subject=subj, skip_next_relevance=False)
        await message.answer(
            f"✅ Предмет: <b>{html.escape(subj)}</b>\n\n"
            "🌆 <b>Выберите город</b>:",
            reply_markup=with_back(kb_city()),
            parse_mode="HTML",
        )
        await state.set_state(WorkState.city)
        return
    wait_msg = await message.answer(
        f"⏳ <b>Проверяю соответствие темы предмету...</b>\n\n"
        f"Предмет: <b>{html.escape(subj)}</b>",
        parse_mode="HTML",
    )
    await handle_subject_selected(wait_msg, state, subj, edit=True)


@dp.callback_query(F.data.startswith("topic_suggest_"))
async def h_topic_suggest(cb: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    suggestions = data.get("topic_suggestions", []) or []
    try:
        idx = int(cb.data.replace("topic_suggest_", "", 1))
        new_topic = suggestions[idx]
    except Exception:
        await cb.answer("Вариант темы не найден. Напишите свою тему.", show_alert=True)
        return

    await state.update_data(topic=new_topic, topic_warning_ignored=False)
    await cb.message.edit_text(
        f"✅ Тема изменена: <b>{html.escape(new_topic)}</b>\n\n"
        "🌆 <b>Выберите город</b>:",
        reply_markup=with_back(kb_city()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.city)
    await cb.answer()


@dp.callback_query(F.data.startswith("topic_subject_"))
async def h_topic_subject_suggest(cb: CallbackQuery, state: FSMContext) -> None:
    if cb.data == "topic_subject_custom":
        await state.update_data(skip_next_relevance=True)
        await cb.message.edit_text(
            "📚 <b>Выберите другой предмет</b> или нажмите «Другой предмет» и напи��ите свой вариант.",
            reply_markup=with_back(kb_subject()),
            parse_mode="HTML",
        )
        await state.set_state(WorkState.subject)
        await cb.answer()
        return

    data = await state.get_data()
    subjects = data.get("subject_suggestions", []) or []
    try:
        idx = int(cb.data.replace("topic_subject_", "", 1))
        new_subject = subjects[idx]
    except Exception:
        await cb.answer("Вариант предмета не найден. Выберите или напишите предмет вручную.", show_alert=True)
        return

    await state.update_data(subject=new_subject, topic_warning_ignored=False)
    await cb.message.edit_text(
        f"✅ Предмет изменён: <b>{html.escape(new_subject)}</b>\n"
        f"✅ Тема оставлена: <b>{html.escape((data.get('topic', '') or '').strip())}</b>\n\n"
        "🌆 <b>Выберите город</b>:",
        reply_markup=with_back(kb_city()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.city)
    await cb.answer()


@dp.callback_query(F.data == "topic_custom")
async def h_topic_custom(cb: CallbackQuery, state: FSMContext) -> None:
    await cb.message.edit_text(
        "✏️ <b>Введите новую тему работы</b>\n\n"
        "Сформулируйте её так, чтобы была видна связь с выбранной дисциплиной.",
        parse_mode="HTML",
        reply_markup=kb_back_cancel(),
    )
    await state.set_state(WorkState.topic_adjustment)
    await cb.answer()


@dp.callback_query(F.data == "topic_continue")
async def h_topic_continue(cb: CallbackQuery, state: FSMContext) -> None:
    await state.update_data(topic_warning_ignored=True)
    await cb.message.edit_text(
        "➡️ Продолжаю с исходной темой.\n\n"
        "🌆 <b>Выберите город</b>:",
        reply_markup=with_back(kb_city()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.city)
    await cb.answer()


@dp.message(WorkState.topic_adjustment)
async def h_topic_adjustment_text(message: Message, state: FSMContext) -> None:
    new_topic = (message.text or "").strip()
    if len(new_topic) < 10 or len(re.findall(r"[А-Яа-яЁёA-Za-z]{3,}", new_topic)) < 2:
        await message.answer(
            "❌ Тема слишком короткая. Сформулируйте тему развёрнуто (минимум 10 символов и 2 смысловых слова).\n\n"
            "<i>Например: «Влияние президентства Дональда Трампа на электоральную географию США»</i>",
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        return

    await state.update_data(topic=new_topic, topic_warning_ignored=False)
    data = await state.get_data()
    subject = (data.get("subject", "") or "").strip()
    if not subject:
        await message.answer("📚 <b>Выберите дисциплину (предмет)</b>", reply_markup=with_back(kb_subject()), parse_mode="HTML")
        await state.set_state(WorkState.subject)
        return

    await handle_subject_selected(message, state, subject, edit=False)


@dp.callback_query(F.data.startswith("city_"))
async def h_city_cb(cb: CallbackQuery, state: FSMContext) -> None:
    city = cb.data.replace("city_", "", 1)
    if city == "other":
        await cb.message.edit_text(
            "✏️ <b>Введите название города</b>",
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        await state.set_state(WorkState.city)
    else:
        await state.update_data(city=city)
        await _ask_pages(cb.message, state, cb.from_user.id)
        await state.set_state(WorkState.pages)
    await cb.answer()


@dp.message(WorkState.city)
async def h_city_text(message: Message, state: FSMContext) -> None:
    await state.update_data(city=(message.text or "").strip())
    await _ask_pages(message, state, message.from_user.id)
    await state.set_state(WorkState.pages)


def _pages_prompt_text(data: dict, user_id: int) -> str:
    doc_type = data.get("doc_type", "referat")
    dt       = DOC_TYPES.get(doc_type, DOC_TYPES["referat"])
    mode     = data.get("mode", "free")

    if mode == "free" and not is_vip(user_id):
        max_p = min(dt["max_pages"], FREE_MAX_PAGES)
        note  = f"\n⚠️ <i>В бесплатном режиме максимум {FREE_MAX_PAGES} стр.</i>"
    else:
        max_p = dt["max_pages"]
        note  = ""

    return (
        f"📄 <b>Количество страниц</b>\n\n"
        f"Допустимо: <b>{dt['min_pages']}–{max_p}</b> страниц{note}\n\n"
        f"Введите число:"
    )


async def _ask_pages(event: Message, state: FSMContext, user_id: int) -> None:
    data = await state.get_data()
    await event.answer(
        _pages_prompt_text(data, user_id),
        parse_mode="HTML",
        reply_markup=kb_back_cancel(),
    )


@dp.message(WorkState.pages)
async def h_pages(message: Message, state: FSMContext) -> None:
    data     = await state.get_data()
    doc_type = data.get("doc_type", "referat")
    dt       = DOC_TYPES.get(doc_type, DOC_TYPES["referat"])

    try:
        pages = int((message.text or "").strip())
    except Exception:
        await message.answer("❌ Введите целое число.", reply_markup=kb_back_cancel())
        return

    mode = data.get("mode", "free")

    # Клипируем до допустимого диапазона
    min_p = dt["min_pages"]
    if mode == "free" and not is_vip(message.from_user.id):
        max_p = min(dt["max_pages"], FREE_MAX_PAGES)
    else:
        max_p = dt["max_pages"]

    if pages < min_p or pages > max_p:
        await message.answer(
            f"❌ Допустимое количество страниц: <b>{min_p}–{max_p}</b>.\n"
            f"Вы ввели: {pages}. Попробуйте ещё раз.",
            parse_mode="HTML",
            reply_markup=kb_back_cancel(),
        )
        return

    await state.update_data(pages=pages)
    await message.answer(
        f"✅ Страниц: <b>{pages}</b>\n\n"
        "🔢 <b>Нумерация страниц — где расположить?</b>",
        reply_markup=with_back(kb_page_number()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.page_number_position)


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — НУМЕРАЦИЯ, МОДЕЛЬ, ОПЛАТА
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data.in_(["pagepos_bottom", "pagepos_top"]))
async def h_pagepos(cb: CallbackQuery, state: FSMContext) -> None:
    pos = "top_center" if cb.data == "pagepos_top" else "bottom_center"
    await state.update_data(page_number_position=pos)
    data = await state.get_data()
    mode = data.get("mode", "free")

    price_note = (
        f"\n\n⭐ <i>Режим с изображениями платный: +{IMAGES_EXTRA_PRICE_PER_PAGE}⭐ за страницу. "
        "DeepSeek подберёт изображения по теме и бот вставит их в DOCX с подписями по ГОСТ.</i>"
        if mode == "paid" else
        "\n\n⭐ <i>С изображениями — платная опция. Если выбрать её в бесплатном режиме, бот переведёт работу в платный режим.</i>"
    )
    await cb.message.edit_text(
        "🖼 <b>Добавить изображения в работу?</b>\n\n"
        "• <b>Без изображений</b> — обычная работа.\n"
        "• <b>С изображениями</b> — поиск фото по теме, вставка по тексту, подписи вида "
        "«Рисунок 1 – ...» и источник."
        f"{price_note}",
        reply_markup=with_back(kb_image_mode()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.image_mode)
    await cb.answer()


async def _continue_after_image_choice(cb: CallbackQuery, state: FSMContext) -> None:
    data  = await state.get_data()
    pages = int(data.get("pages", 10))
    mode  = data.get("mode", "free")
    include_images = bool(data.get("include_images"))

    # Изображения — платная опция: если пользователь пришёл из бесплатного режима,
    # переводим только эту генерацию в платный сценарий выбора модели/оплаты.
    if include_images and mode == "free" and not is_vip(cb.from_user.id):
        mode = "paid"
        await state.update_data(mode="paid")

    # ── Бесплатный режим ──
    if mode == "free":
        if pages > FREE_MAX_PAGES and not is_vip(cb.from_user.id):
            await cb.message.edit_text(
                f"🚫 <b>В бесплатном режиме максимум {FREE_MAX_PAGES} страниц.</b>\n\n"
                "Используйте платный режим для большего объёма.",
                parse_mode="HTML",
            )
            await state.clear()
            await cb.answer()
            return

        ok, reason = check_user_limit(cb.from_user.id, "free")
        if not ok:
            await cb.message.edit_text(reason, parse_mode="HTML")
            await state.clear()
            await cb.answer()
            return

        await state.update_data(humanize=False)
        await cb.message.edit_text(
            "🚀 <b>Запускаю генерацию...</b>\n\n"
            "Текст будет без опечаток, markdown-маркеров и фраз-маркеров ИИ.",
            parse_mode="HTML",
        )
        await cb.answer()
        await generate_and_send(cb.message, state, model_key=FREE_MODEL_KEY, pay_mode="free")
        return

    # ── Платный режим ──
    ok, reason = check_user_limit(cb.from_user.id, "paid")
    if not ok and not is_vip(cb.from_user.id):
        await cb.message.edit_text(reason, parse_mode="HTML")
        await state.clear()
        await cb.answer()
        return

    await state.update_data(humanize=False)
    image_note = (
        f"\n🖼 Изображения включены: +{IMAGES_EXTRA_PRICE_PER_PAGE}⭐/стр."
        if include_images else "\n📄 Изображения отключены."
    )
    await cb.message.edit_text(
        "🤖 <b>Выберите ИИ-модель</b>\n\n"
        "Цена указана в звёздах Telegram за страницу.\n"
        "Текст будет без опечаток, markdown-маркеров и фраз-маркеров ИИ."
        f"{image_note}",
        reply_markup=with_back(kb_models()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.model)
    await cb.answer()


@dp.callback_query(F.data.in_(["images_yes", "images_no"]))
async def h_image_mode(cb: CallbackQuery, state: FSMContext) -> None:
    include_images = cb.data == "images_yes"
    if not include_images:
        await state.update_data(include_images=False, image_count=0)
        await _continue_after_image_choice(cb, state)
        return

    # (user-patch): пользователь сам присылает ссылки на изображения.
    # Авто-поиск убран, загрузка фото-файлов больше не предлагается.
    await state.update_data(include_images=True, own_images=[])
    await cb.message.edit_text(
        "🖼 <b>Изображения в работу</b>\n\n"
        "🔗 <b>Пришлите ссылки на изображения</b> — по одной на сообщение или\n"
        "сразу списком (каждая ссылка с новой строки). Бот сам решит, в какой\n"
        f"раздел вставить каждую картинку и оформит подписи по ГОСТ.\n\n"
        f"Поддерживаются прямые ссылки на .jpg/.jpeg/.png/.webp (до {MAX_WORK_IMAGES} штук).",
        reply_markup=with_back(kb_image_source()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.image_source)
    await cb.answer()


@dp.callback_query(F.data == "imgsrc_skip")
async def h_imgsrc_skip(cb: CallbackQuery, state: FSMContext) -> None:
    """(user-patch): авто-поиск изображений отключён — единственная альтернатива
    собственным ссылкам это сгенерировать работу без иллюстраций."""
    await state.update_data(include_images=False, use_own_images=False, own_images=[], image_count=0)
    await _continue_after_image_choice(cb, state)
    await cb.answer()


# Старый обработчик "imgsrc_auto" оставлен как заглушка-редирект на skip,
# чтобы пользователи со старыми сообщениями не получали "callback expired".
@dp.callback_query(F.data == "imgsrc_auto")
async def h_imgsrc_auto_legacy(cb: CallbackQuery, state: FSMContext) -> None:
    await state.update_data(include_images=False, use_own_images=False, own_images=[], image_count=0)
    await cb.answer("Автопоиск отключён — работа будет без изображений.", show_alert=False)
    await _continue_after_image_choice(cb, state)


@dp.callback_query(F.data == "imgsrc_own")
async def h_imgsrc_own(cb: CallbackQuery, state: FSMContext) -> None:
    """(user-patch): запрашиваем у пользователя URL-ссылки, не файлы."""
    await state.update_data(include_images=True, use_own_images=True, own_images=[])
    await cb.message.edit_text(
        "🔗 <b>Пришлите ссылки на изображения</b>\n\n"
        f"Можно до <b>{MAX_WORK_IMAGES}</b> ссылок. Каждая ссылка — отдельным\n"
        "сообщением или списком через перенос строки.\n\n"
        "Поддерживаются <b>прямые URL</b> на изображения (.jpg/.jpeg/.png/.webp).\n"
        "Бот скачает картинку, придумает подпись и вставит её в наиболее\n"
        "подходящий раздел работы.\n\n"
        "Когда все ссылки отправлены — нажмите «✅ Готово».",
        reply_markup=kb_own_images_done(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.own_images_upload)
    await cb.answer()


# ─────────────────────────────────────────────────────────────────────
# (user-patch): принимаем ТОЛЬКО ссылки на изображения, не файлы.
# ─────────────────────────────────────────────────────────────────────
_URL_RE = re.compile(r'https?://\S+', re.IGNORECASE)
_IMG_EXT_RE = re.compile(r'\.(?:jpg|jpeg|png|webp|gif|bmp|tiff)(?:\?|#|$)', re.IGNORECASE)


@dp.message(WorkState.own_images_upload, F.text)
async def h_own_image_links(message: Message, state: FSMContext) -> None:
    """Принимает одну или несколько URL-ссылок на изображения от пользователя.

    Пользователь может прислать ссылки списком (по одной на строку) или
    несколькими сообщениями подряд. Каждая ссылка скачивается, нормализуется
    через _image_bytes_for_docx и кладётся в own_images. Файлы и фото-вложения
    больше не принимаются — это сделано по запросу пользователя.
    """
    text = (message.text or "").strip()
    urls = [u.rstrip(').,;') for u in _URL_RE.findall(text)]
    if not urls:
        await message.answer(
            "🔗 Пришлите <b>ссылку на изображение</b> (http(s)://… .jpg / .png / .webp).\n"
            "Несколько ссылок — каждая с новой строки.",
            reply_markup=kb_own_images_done(),
            parse_mode="HTML",
        )
        return

    data = await state.get_data()
    own = list(data.get("own_images") or [])
    accepted = 0
    rejected: list[str] = []
    for url in urls:
        if len(own) >= MAX_WORK_IMAGES:
            await message.answer(
                f"⚠️ Достигнут лимит {MAX_WORK_IMAGES} изображений. Нажмите «✅ Готово».",
                reply_markup=kb_own_images_done(),
            )
            break
        # Мягкая проверка расширения; если расширение не указано (например,
        # CDN-ссылка), всё равно попробуем скачать и проверим content-type.
        try:
            img_bytes_raw = await _download_image_bytes(url)
            if not img_bytes_raw:
                rejected.append(url)
                continue
            img_bytes = _image_bytes_for_docx(img_bytes_raw)
            if not img_bytes:
                rejected.append(url)
                continue
            own.append({"bytes": img_bytes, "caption": "", "source": url})
            accepted += 1
        except Exception as e:
            print(f"[OWN_IMG_URL] error for {url!r}: {e}")
            rejected.append(url)

    await state.update_data(own_images=own)
    lines = [f"✅ Принято ссылок: <b>{accepted}</b> (всего {len(own)}/{MAX_WORK_IMAGES})."]
    if rejected:
        lines.append(f"⚠️ Не удалось скачать или это не картинка: {len(rejected)} шт.")
    lines.append("Пришлите ещё ссылки или нажмите «✅ Готово».")
    await message.answer("\n".join(lines), parse_mode="HTML", reply_markup=kb_own_images_done())


@dp.message(WorkState.own_images_upload, F.photo)
async def h_own_image_photo_blocked(message: Message, state: FSMContext) -> None:
    """(user-patch): загрузка фото-файлов отключена — нужны ссылки."""
    await message.answer(
        "🔗 Пожалуйста, пришлите <b>ссылку</b> на изображение, а не сам файл.\n"
        "Например: <code>https://example.com/image.jpg</code>",
        parse_mode="HTML",
        reply_markup=kb_own_images_done(),
    )


@dp.message(WorkState.own_images_upload, F.document)
async def h_own_image_document_blocked(message: Message, state: FSMContext) -> None:
    """(user-patch): загрузка документов отключена — нужны ссылки."""
    await message.answer(
        "🔗 Пожалуйста, пришлите <b>ссылку</b> на изображение, а не файл.\n"
        "Например: <code>https://example.com/image.jpg</code>",
        parse_mode="HTML",
        reply_markup=kb_own_images_done(),
    )


@dp.callback_query(F.data == "own_images_done")
async def h_own_images_done(cb: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    own = list(data.get("own_images") or [])
    if not own:
        await cb.answer("Сначала пришлите хотя бы одно фото.", show_alert=True)
        return
    # Свои фото = дешевле: НЕ берём доплату за изображения.
    await state.update_data(
        include_images=True,
        use_own_images=True,
        image_count=len(own),
        images=own,  # сразу кладём как готовые изображения для документа
    )
    await cb.message.edit_text(
        f"✅ Загружено <b>{len(own)}</b> фото. Бот вставит их в подходящие разделы "
        "с подписями «Рисунок N – …».\n\n"
        "💰 За свои изображения доплата не взимается.",
        parse_mode="HTML",
    )
    await cb.answer()
    await _continue_after_image_choice(cb, state)


@dp.callback_query(F.data.startswith("imgcount_"))
async def h_image_count_cb(cb: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    pages = int(data.get("pages", 10))
    value = cb.data.replace("imgcount_", "", 1)

    if value == "custom":
        await cb.message.edit_text(
            f"✏️ <b>Введите количество фото</b> числом от 1 до {MAX_WORK_IMAGES}.",
            reply_markup=kb_back_cancel(),
            parse_mode="HTML",
        )
        await state.set_state(WorkState.image_count)
        await cb.answer()
        return

    if value == "auto":
        count = _image_count_for_pages(pages)
    else:
        count = int(value)
    count = max(1, min(MAX_WORK_IMAGES, count))
    await state.update_data(include_images=True, image_count=count)
    await _continue_after_image_choice(cb, state)


@dp.message(WorkState.image_count)
async def h_image_count_text(message: Message, state: FSMContext) -> None:
    try:
        count = int((message.text or "").strip())
    except Exception:
        await message.answer(f"❌ Введите число от 1 до {MAX_WORK_IMAGES}.", reply_markup=kb_back_cancel())
        return
    if count < 1 or count > MAX_WORK_IMAGES:
        await message.answer(f"❌ Допустимое количество фото: 1–{MAX_WORK_IMAGES}.", reply_markup=kb_back_cancel())
        return
    await state.update_data(include_images=True, image_count=count)

    # После текстового ввода продолжаем тот же сценарий, только у message нет edit_text.
    data = await state.get_data()
    pages = int(data.get("pages", 10))
    mode = data.get("mode", "free")
    if mode == "free" and not is_vip(message.from_user.id):
        await state.update_data(mode="paid")
        mode = "paid"

    if mode == "free":
        ok, reason = check_user_limit(message.from_user.id, "free")
        if not ok:
            await message.answer(reason, parse_mode="HTML")
            await state.clear()
            return
        await message.answer("🚀 <b>Запускаю генерацию...</b>", parse_mode="HTML")
        await generate_and_send(message, state, model_key=FREE_MODEL_KEY, pay_mode="free")
        return

    ok, reason = check_user_limit(message.from_user.id, "paid")
    if not ok and not is_vip(message.from_user.id):
        await message.answer(reason, parse_mode="HTML")
        await state.clear()
        return
    await message.answer(
        f"✅ Фото: <b>{count}</b> шт.\n\n🤖 <b>Выберите ИИ-модель</b>",
        reply_markup=with_back(kb_models()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.model)


@dp.callback_query(F.data.in_(["humanize_yes", "humanize_no"]))
async def h_humanize(cb: CallbackQuery, state: FSMContext) -> None:
    # Старые кнопки могут остаться в Telegram-сообщениях, но опечатки отключены.
    humanize = False
    await state.update_data(humanize=humanize)

    data = await state.get_data()
    mode = data.get("mode", "free")

    if mode == "free":
        await cb.message.edit_text(
            "🚀 <b>Запускаю генерацию...</b>\n\nПодождите, это займёт несколько минут.",
            parse_mode="HTML",
        )
        await cb.answer()
        await generate_and_send(cb.message, state, model_key=FREE_MODEL_KEY, pay_mode="free")
        return

    await cb.message.edit_text(
        "🤖 <b>Выберите ИИ-модель</b>\n\n"
        "Цена указана в звёздах Telegram за страницу.\n"
        "Все модели генерируют полноценный академический текст.",
        reply_markup=with_back(kb_models()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.model)
    await cb.answer()


@dp.callback_query(F.data.startswith("model_"))
async def h_model(cb: CallbackQuery, state: FSMContext) -> None:
    model_key = cb.data.replace("model_", "", 1)
    if model_key not in AI_MODELS or not AI_MODELS[model_key].get("api_key"):
        await cb.answer("⚠️ Модель временно недоступна", show_alert=True)
        return

    data  = await state.get_data()
    pages = int(data.get("pages", 10))
    model = AI_MODELS[model_key]
    include_images = bool(data.get("include_images"))
    # Свои изображения = дешевле: доплату за поиск/проверку не берём.
    own_imgs = bool(data.get("use_own_images")) and bool(data.get("own_images") or data.get("images"))
    image_extra = IMAGES_EXTRA_PRICE_PER_PAGE if (include_images and not own_imgs) else 0
    price_per_page = int(model["price_per_page"]) + image_extra
    total = price_per_page * pages

    await state.update_data(model_key=model_key)

    if is_vip(cb.from_user.id):
        await cb.message.edit_text(
            f"👑 <b>VIP — оплата не требуется</b>\n\n"
            f"Модель: {model['name']}\n"
            f"Страниц: {pages}\n"
            f"Изображения: {'да' if include_images else 'нет'}\n\n"
            f"🚀 Запускаю генерацию...",
            parse_mode="HTML",
        )
        await cb.answer()
        await generate_and_send(cb.message, state, model_key=model_key, pay_mode="paid")
        return

    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text=f"⭐ Оплатить {total} звёзд", callback_data=f"pay_{total}")],
            [InlineKeyboardButton(text="← Выбрать другую модель",   callback_data="back_to_models")],
        ]
    )
    await cb.message.edit_text(
        f"💳 <b>Подтверждение оплаты</b>\n\n"
        f"┌─────────────────────────\n"
        f"│ 🤖 Модель:  {model['name']}\n"
        f"│ 📄 Страниц: {pages}\n"
        f"│ 🖼 Изображения: {((str(int(data.get('image_count') or 0)) + ' шт.' + (' (свои — без доплаты)' if own_imgs else '')) if include_images else 'нет')}\n"
        f"│ 💰 Цена:    {price_per_page}⭐ × {pages} = <b>{total}⭐</b>\n"
        f"└─────────────────────────\n\n"
        f"Нажмите кнопку для оплаты через Telegram Stars:",
        reply_markup=kb,
        parse_mode="HTML",
    )
    await state.set_state(WorkState.payment)
    await cb.answer()


@dp.callback_query(F.data == "back_to_models")
async def h_back_to_models(cb: CallbackQuery, state: FSMContext) -> None:
    await cb.message.edit_text(
        "🤖 <b>Выберите ИИ-модель</b>",
        reply_markup=with_back(kb_models()),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.model)
    await cb.answer()


@dp.callback_query(F.data.startswith("pay_"))
async def h_pay(cb: CallbackQuery, state: FSMContext) -> None:
    stars     = int(cb.data.replace("pay_", "", 1))
    data      = await state.get_data()
    model_key = data.get("model_key", FREE_MODEL_KEY)
    payload   = json.dumps(
        {
            "model_key": model_key,
            "pages": data.get("pages", 10),
            "include_images": bool(data.get("include_images")),
            "image_count": int(data.get("image_count") or 0),
        },
        ensure_ascii=False,
    )

    await cb.bot.send_invoice(
        chat_id=cb.from_user.id,
        title="Академическая работа по ГОСТ",
        description=f"{AI_MODELS[model_key]['name']} · {data.get('pages', 10)} стр. · {data.get('topic','')[:40]}",
        payload=payload,
        provider_token="",
        currency="XTR",
        prices=[LabeledPrice(label="Работа по ГОСТ", amount=stars)],
        reply_markup=InlineKeyboardMarkup(
            inline_keyboard=[[InlineKeyboardButton(text=f"⭐ Оплатить {stars} звёзд", pay=True)]]
        ),
    )
    await cb.answer()


@dp.pre_checkout_query()
async def h_pre_checkout(q: PreCheckoutQuery) -> None:
    await q.answer(ok=True)


@dp.message(F.successful_payment)
async def h_payment_ok(message: Message, state: FSMContext) -> None:
    try:
        payload   = json.loads(message.successful_payment.invoice_payload)
    except Exception:
        payload = {}

    # ── Оплата режима редактирования ──
    if payload.get("mode") == "edit":
        await message.answer(
            "✅ <b>Оплата принята!</b>\n\n✏️ Применяю ваши правки...",
            parse_mode="HTML",
        )
        await run_edit_and_send(message, state)
        return

    model_key = payload.get("model_key", FREE_MODEL_KEY)
    try:
        await state.update_data(
            include_images=bool(payload.get("include_images")),
            image_count=int(payload.get("image_count") or 0),
        )
    except Exception:
        model_key = FREE_MODEL_KEY

    await message.answer(
        "✅ <b>Оплата прошла успешно!</b>\n\n"
        "🚀 Начинаю генерацию работы...",
        parse_mode="HTML",
    )
    await generate_and_send(message, state, model_key=model_key, pay_mode="paid")


@dp.callback_query(F.data == "regenerate")
async def h_regenerate(cb: CallbackQuery, state: FSMContext) -> None:
    """Перегенерация ТОЙ ЖЕ работы с теми же параметрами (без повторной оплаты)."""
    data = await state.get_data()
    job = data.get("last_job")
    if not job or not job.get("topic"):
        await cb.answer("Нет данных о последней работе. Создайте новую (/start).", show_alert=True)
        return
    model_key = job.get("model_key", FREE_MODEL_KEY)
    pay_mode = job.get("pay_mode", "free")
    # Восстанавливаем параметры (и сохраняем сам job, чтобы можно было
    # перегенерировать повторно).
    await state.clear()
    await state.update_data(**{k: v for k, v in job.items() if k not in ("model_key", "pay_mode")})
    await state.update_data(last_job=job)
    await cb.message.answer(
        "🔄 <b>Перегенерирую работу с теми же параметрами...</b>\n\n"
        "<i>Повторная оплата не требуется.</i>",
        parse_mode="HTML",
    )
    await cb.answer()
    # pay_mode="regen" — внутри generate_and_send это трактуется как уже
    # оплаченная работа (лимиты/звёзды повторно не списываем).
    await generate_and_send(cb.message, state, model_key=model_key, pay_mode="regen")


@dp.callback_query(F.data == "final_new")
async def h_final_new(cb: CallbackQuery, state: FSMContext) -> None:
    await state.clear()
    first = cb.from_user.first_name or "пользователь"
    await cb.message.edit_text(
        _welcome_text(cb.from_user.id, first),
        reply_markup=kb_doc_type(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.doc_type)
    await cb.answer()


# ═══════════════════════════════════════════════════════════════
#  ГЛАВНАЯ ГЕНЕРАЦИЯ
# ═══════════════════════════════════════════════════════════════

async def generate_and_send(
    event: Message,
    state: FSMContext,
    model_key: str,
    pay_mode: str,
) -> None:
    """
    Основная функция генерации: запрашивает названия глав,
    генерирует текст блоками, собирает DOCX, отправляет.

    Поддерживает несколько работ одновременно (разные пользователи и
    несколько работ в одном чате). Глобальный семафор GEN_SEMAPHORE
    ограничивает суммарную нагрузку, а MAX_PARALLEL_PER_USER — число
    параллельных генераций на одного человека.
    """
    _uid = event.from_user.id if event.from_user else event.chat.id

    # Мягкий лимит параллельных работ на пользователя (не блокирует чат).
    if MAX_PARALLEL_PER_USER and _active_gen_per_user.get(_uid, 0) >= MAX_PARALLEL_PER_USER:
        await event.answer(
            f"⏳ У вас уже выполняется {MAX_PARALLEL_PER_USER} работы одновременно. "
            "Дождитесь завершения одной из них — и можно запускать следующую.",
            parse_mode="HTML",
        )
        return

    _active_gen_per_user[_uid] = _active_gen_per_user.get(_uid, 0) + 1
    # Если все слоты заняты — честно предупреждаем, что работа встанет в очередь
    # и пойдёт параллельно с остальными (несколько работ / много пользователей).
    if _active_gen_per_user.get(_uid, 0) > 1 or GEN_SEMAPHORE.locked():
        try:
            await event.answer(
                "🧵 Запрос принят в работу. Сейчас идёт несколько генераций — "
                "ваша начнётся, как только освободится слот, и будет выполняться "
                "параллельно с остальными.",
                parse_mode="HTML",
            )
        except Exception:
            pass

    async with GEN_SEMAPHORE:
        data = await state.get_data()

        doc_type = data.get("doc_type", "referat")
        dt       = DOC_TYPES.get(doc_type, DOC_TYPES["referat"])
        pages    = int(data.get("pages", 10))
        topic    = (data.get("topic", "") or "").strip()
        subject  = (data.get("subject", "") or "").strip()

        # ═══════════════════════════════════════════════════════════════
        # ЗАЩИТА ОТ ПУСТЫХ ТЕМЫ И ДИСЦИПЛИНЫ
        # ═══════════════════════════════════════════════════════════════
        if not topic:
            await event.answer(
                "❌ <b>Тема не указана</b>\n\n"
                "Пожалуйста, начните заново командой /start",
                parse_mode="HTML",
            )
            await state.clear()
            return

        if not subject:
            await event.answer(
                "❌ <b>Дисциплина не указана</b>\n\n"
                "Пожалуйста, начните заново командой /start",
                parse_mode="HTML",
            )
            await state.clear()
            return

        if len(topic) < 3:
            await event.answer(
                "❌ <b>Тема слишком короткая</b>\n\n"
                f"Вы ввели: «{topic}»\n"
                "Пожалуйста, введите тему длиннее (минимум 3 символа)\n\n"
                "Начните заново — /start",
                parse_mode="HTML",
            )
            await state.clear()
            return

        # Определяем количество глав по типу документа
        if doc_type in ("esse", "doklad", "article"):
            num_chapters = 0  # у них своя структура
        elif doc_type in ("kursovaya", "final_referat", "vkr", "final_project"):
            num_chapters = 4 if doc_type in ("vkr", "final_project") and pages >= 40 else 3
        else:
            num_chapters = 2  # реферат, контрольная, свой

        # Считаем шаги для прогресс-бара (оценка: ~3 подглавы на главу)
        # 1 = названия глав, блоки = intro + подглавы + lit + conclusion, 1 = DOCX, 1 = отправка
        if doc_type in ("esse",):
            extra_blocks = 6
        elif doc_type in ("doklad",):
            extra_blocks = 5
        elif doc_type in ("article",):
            extra_blocks = 7
        else:
            # chapter_titles ещё не готов — оцениваем: num_chapters * 3 подглавы
            num_subs = max(1, num_chapters * 3)
            extra_blocks = 1 + num_subs + 2  # intro + N subs + lit + conclusion
        total_steps  = 1 + extra_blocks + 2

        progress_msg = await event.answer(
            "⏳ <b>Инициализация...</b>",
            parse_mode="HTML",
        )
        prog = Progress(
            msg=progress_msg,
            title=f"Создаю {dt['word'].lower()}",
            total_steps=total_steps,
            model_name=AI_MODELS.get(model_key, {}).get("name", ""),
        )

        # Запускаем фоновый цикл анимации
        stop_anim = asyncio.Event()
        anim_task = asyncio.create_task(prog.animate_loop(stop_anim))

        try:
            gost = get_gost_config(doc_type, event.chat.id)
            gost["page_number_position"] = data.get(
                "page_number_position",
                gost.get("page_number_position", "bottom_center"),
            )

            # ── Шаг 1: Генерируем названия глав ──
            await prog.update(label="🧠 Придумываю названия глав...", force=True)
            if num_chapters > 0:
                chapter_titles = await generate_chapter_titles(
                    model_key, doc_type, topic, subject, num_chapters
                )
            else:
                chapter_titles = []
            await prog.update(step_done=True)

            writing_style = data.get("writing_style", "classic")

            # ── Шаги 2–N: Генерируем текст ──
            await prog.update(label="✍️ Генерирую текст разделов...")
            parts = await generate_text_blocks(
                topic=topic,
                pages=pages,
                doc_type=doc_type,
                subject=subject,
                model_key=model_key,
                source=data.get("source_content", ""),
                chapter_titles=chapter_titles,
                writing_style=writing_style,
                prog=prog,
                humanize=data.get("humanize", False),
            )

            # ── Проверка соответствия дисциплине (приоритет 🔴) ──
            await prog.update(label="🔎 Проверяю соответствие дисциплине...")
            sample_for_check = "\n\n".join(
                str(parts.get(k, "")) for k in ("intro", "main1", "part1", "ch1_s1")
                if parts.get(k)
            )[:2500]
            relevance_ok, relevance_reason = await verify_discipline_relevance(
                model_key, topic, subject, sample_for_check,
            )
            if not relevance_ok:
                print(f"[RELEVANCE] ⚠️ Текст может не соответствовать дисциплине "
                      f"«{subject}»: {relevance_reason}")

            # ── Сборка структуры ──
            blocks = generate_structure(doc_type, parts, chapter_titles, topic=topic)

            # ── DOCX ──
            await prog.update(label="📄 Собираю DOCX-документ...", step_done=True)
            work_dir = os.path.join(os.getcwd(), "_out")
            os.makedirs(work_dir, exist_ok=True)
            ts      = datetime.now().strftime("%Y%m%d_%H%M%S")
            tmp_in  = os.path.join(work_dir, f"tmp_{event.chat.id}_{ts}.docx")
            tmp_out = os.path.join(work_dir, f"final_{event.chat.id}_{ts}.docx")

            # Форматируем заголовки «Глава N. ...» / «N.M. ...» (fix12)
            blocks = _apply_heading_format_to_blocks(blocks)

            if data.get("include_images"):
                # ── Вариант 1: пользователь прислал СВОИ фото — используем их,
                #    бот только придумывает подписи и расставляет «в тему». ──
                if data.get("use_own_images") and (data.get("own_images") or data.get("images")):
                    await prog.update(label="🖼 Готовлю ваши фото и подписи по ГОСТ...", force=True)
                    own = list(data.get("images") or data.get("own_images") or [])
                    for img in own:
                        if img.get("bytes"):
                            img["bytes"] = _image_bytes_for_docx(img["bytes"])
                        # Если пользователь не дал подпись — придумываем по теме.
                        base_cap = str(img.get("caption") or "").strip()
                        img["caption"] = await translate_caption_to_russian(
                            model_key, base_cap or topic, topic=topic,
                        )
                    data["images"] = own
                else:
                    # ── Вариант 2: авто-поиск настоящих фото (разные для разных работ). ──
                    await prog.update(label="🖼 Ищу изображения по теме и готовлю подписи по ГОСТ...", force=True)
                    data["images"] = await prepare_work_images(
                        topic,
                        subject,
                        pages,
                        model_key=model_key,
                        image_count=int(data.get("image_count") or 0) or None,
                    )
                if not data.get("images"):
                    print("[IMAGES] Подходящие реальные фото не найдены — документ будет без иллюстраций")
                else:
                    # Компактный размер фото: рисунок + подпись + источник должны
                    # помещаться рядом с основным текстом, а не занимать отдельный лист.
                    # Ширина 5.5 см ≈ треть полезной ширины страницы A4 — оставляет
                    # место для текста и не выталкивает иллюстрацию на отдельный лист.
                    gost["image_width_cm"] = 5.5
                    gost["image_page_break_before"] = False
                    # Рисунки по центру и по бокам — чтобы каждый раз по-разному,
                    # но всегда внутри текстового блока.
                    gost["image_align"] = random.choice(["center", "left", "right"])

            docx_raw = build_docx_bytes(data, blocks, gost)
            with open(tmp_in, "wb") as f:
                f.write(docx_raw)
            if not os.path.exists(tmp_in):
                raise RuntimeError(f"Не удалось записать {tmp_in}")

            # ── ПОДГОНКА СТРАНИЦ (исправленная) ──
            blocks, final_pages = await precise_page_adjustment(
                tmp_in=tmp_in,
                blocks=blocks,
                target_pages=pages,
                topic=topic,
                model_key=model_key,
                writing_style=writing_style,
                data=data,
                gost=gost,
                prog=prog,
                work_dir=work_dir,
            )

            # ── LibreOffice (финальная конвертация — обновит TOC и поля PAGE) ──
            await prog.update(label="🔄 Обновляю содержание (LibreOffice)...", step_done=True)
            updated    = libreoffice_update_docx(tmp_in, tmp_out)
            final_path = tmp_out if updated else tmp_in

            # ── Постобработка: LO после обновления TOC может «раздуть»
            #    документ на 1–2 страницы. Поэтому крутим коррекцию в цикле
            #    до 3 раз: измерили → подогнали blocks → пересобрали → LO → снова замер. ──
            POST_LO_MAX_ROUNDS = 6
            final_pages = pages
            for round_idx in range(POST_LO_MAX_ROUNDS):
                post_lo_pages = await measure_pages_async(final_path, work_dir)
                if post_lo_pages is None:
                    final_pages = pages
                    break
                final_pages = post_lo_pages
                if post_lo_pages == pages:
                    print(f"[PAGES] ✅ После LO совпало с целью ({pages}) на раунде {round_idx+1}")
                    break

                print(f"[PAGES] Раунд {round_idx+1}/{POST_LO_MAX_ROUNDS}: LO дал {post_lo_pages}, цель {pages}. Корректирую…")

                # Перегенерируем blocks→docx «сырьём», подгоняем, и затем снова
                # прогоняем через LibreOffice (он же — финальный источник правды).
                blocks, _ = await precise_page_adjustment(
                    tmp_in=tmp_in,
                    blocks=blocks,
                    target_pages=pages,
                    topic=topic,
                    model_key=model_key,
                    writing_style=writing_style,
                    data=data,
                    gost=gost,
                    prog=prog,
                    work_dir=work_dir,
                )
                # Пересобирае�� DOCX «как новый» из обновлённых blocks
                blocks = _apply_heading_format_to_blocks(blocks)
                docx_raw = build_docx_bytes(data, blocks, gost)
                with open(tmp_in, "wb") as f:
                    f.write(docx_raw)
                # И снова через LO — это, как правило, и есть «+1 страница» источник
                updated = libreoffice_update_docx(tmp_in, tmp_out)
                final_path = tmp_out if updated else tmp_in
            else:
                # Цикл закончился без break — печатаем итог
                print(f"[PAGES] ⚠️ После {POST_LO_MAX_ROUNDS} раундов финальная цифра {final_pages}, цель {pages}")

            # Финальный строгий контроль: не отправляем «почти нужный» объём без попытки исправить.
            if final_pages != pages:
                print(f"[PAGES] 🚨 Финальная строгая коррекция: факт {final_pages}, цель {pages}")
                for strict_round in range(3):
                    if final_pages == pages:
                        break
                    if final_pages > pages and data.get("images") and float(gost.get("image_width_cm", 10.0)) > 4.0:
                        gost["image_width_cm"] = max(4.0, float(gost.get("image_width_cm", 10.0)) - 1.0)
                        print(f"[PAGES] 🖼 strict: ширина изображений {gost['image_width_cm']} см")
                    blocks, _ = await precise_page_adjustment(
                        tmp_in=tmp_in,
                        blocks=blocks,
                        target_pages=pages,
                        topic=topic,
                        model_key=model_key,
                        writing_style=writing_style,
                        data=data,
                        gost=gost,
                        prog=prog,
                        work_dir=work_dir,
                    )
                    blocks = _apply_heading_format_to_blocks(blocks)
                    docx_raw = build_docx_bytes(data, blocks, gost)
                    with open(tmp_in, "wb") as f:
                        f.write(docx_raw)
                    updated = libreoffice_update_docx(tmp_in, tmp_out)
                    final_path = tmp_out if updated else tmp_in
                    measured = await measure_pages_async(final_path, work_dir)
                    if measured is not None:
                        final_pages = measured
                    print(f"[PAGES] strict round {strict_round+1}: {final_pages}/{pages}")

            # ── Реальные страницы в содержании ──
            # После финальной сборки извлекаем страницы заголовков из PDF и
            # пересобираем DOCX со статическим TOC уже по фактическим страницам,
            # а не по расчётной эвристике. Делаем 2 прохода: первый вставляет
            # реальные страницы, второй проверяет, что после пересборки карта
            # не изменилась.
            try:
                for toc_round in range(2):
                    real_toc_pages = _extract_real_toc_pages_from_docx(final_path, blocks, work_dir)
                    if not real_toc_pages:
                        print("[TOC] Точная карта страниц не получена — старые расчётные значения не считаю достоверными")
                        break
                    if gost.get("_toc_page_map") == real_toc_pages and toc_round > 0:
                        print("[TOC] Карта содержания стабильна")
                        break
                    gost["_toc_page_map"] = real_toc_pages
                    blocks = _apply_heading_format_to_blocks(blocks)
                    docx_raw = build_docx_bytes(data, blocks, gost)
                    with open(tmp_in, "wb") as f:
                        f.write(docx_raw)
                    updated = libreoffice_update_docx(tmp_in, tmp_out)
                    final_path = tmp_out if updated else tmp_in
                    measured = await measure_pages_async(final_path, work_dir)
                    if measured is not None:
                        final_pages = measured
                    print(f"[TOC] Содержание пересобрано по реальным страницам (проход {toc_round+1}), страниц: {final_pages}")
            except Exception as _toc_e:
                print(f"[TOC] Реальное содержание не пересобрано: {_toc_e}")

            # ВАЖНО: пользователь ОПЛАТИЛ работу — мы НЕ выбрасываем ошибку,
            # даже если до точного числа страниц не дотянули на 1–2. Вместо
            # этого отдаём максимально близкий результат и честно помечаем
            # фактический объём. Раньше здесь был raise RuntimeError, из-за
            # которого человек терял оплату и не получал файл.
            if final_pages != pages:
                print(
                    f"[PAGES] ⚠️ Точное число не достигнуто: {final_pages} вместо {pages}. "
                    "Отдаю максимально близкий результат (без потери оплаты)."
                )

            print(f"[PAGES] 📤 Итог в caption: {final_pages} страниц")

            # ── Имя файла ──
            safe_topic = re.sub(r'[<>"/:\\|?*]', "", topic[:35]).replace(" ", "_")
            fname      = f"{dt['word'].replace(' ', '_')}_{safe_topic}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

            with open(final_path, "rb") as f:
                final_bytes = f.read()

            # ── Отправляем ──
            await prog.update(label="📤 Отправляю файл...", step_done=True)

            toc_status = (
                ("✅ обновлено автоматически" if updated else "⚠️ обновите вручную в Word")
                if TOC_USE_WORD_FIELD else
                "✅ без TOC-заглушек"
            )
            style_label = "🎓 Умный" if data.get("writing_style") == "smart" else "📝 Классический"
            pages_line = (
                f"📄 Страниц: <b>{final_pages}</b> (заказано {pages} — "
                f"подогнано максимально близко без «воды»)"
                if final_pages != pages
                else f"📄 Страниц: <b>{final_pages}</b> ✅"
            )
            relevance_status = (
                f"✅ соответствует «{subject}»"
                if relevance_ok
                else f"⚠️ проверьте тему: {relevance_reason[:60]}"
            )
            images_line = (
                f"│ 🖼 Изображения: {len(data.get('images') or [])} шт.\n"
                if data.get("include_images") else
                "│ 🖼 Изображения: нет\n"
            )
            caption = (
                f"🎉 <b>{dt['word']} ГОТОВ!</b>\n\n"
                f"┌─────────────────────────\n"
                f"│ 📖 Тема: {topic[:60]}\n"
                f"│ {pages_line}\n"
                f"│ 🤖 ИИ: {AI_MODELS.get(model_key, {}).get('name', model_key)}\n"
                f"{images_line}"
                f"│ ✍️ Стиль: {style_label}\n"
                f"│ 📐 Шрифт: {gost.get('font_name')} {gost.get('font_size')}pt\n"
                f"│ ↕️ Интервал: {gost.get('line_spacing')}\n"
                f"│ 🎓 Дисциплина: {relevance_status}\n"
                f"│ 📑 Содержание: {toc_status}\n"
                f"└─────────────────────────"
            )

            await event.answer_document(
                BufferedInputFile(final_bytes, filename=fname),
                caption=caption,
                parse_mode="HTML",
            )
            # Лимит учитываем только после успешной отправки файла, чтобы не
            # сжигать бесплатную попытку при ошибке API/LibreOffice/Telegram.
            record_user_generation(event.chat.id, pay_mode)

        except Exception as e:
            print(f"[GEN ERROR] {e}")
            await prog.finish(
                "❌ <b>Ошибка генерации</b>\n\n"
                f"Причина: {str(e)[:200]}\n\n"
                "Попробуйте ещё раз или выберите другую модель (/start)."
            )
            await state.clear()
            _active_gen_per_user[_uid] = max(0, _active_gen_per_user.get(_uid, 1) - 1)
            return

        finally:
            stop_anim.set()
            anim_task.cancel()
            try:
                await anim_task
            except asyncio.CancelledError:
                pass

        await prog.delete()

        # Сохраняем параметры работы, чтобы можно было «перегенерировать эту же
        # работу» бесплатно, если результат не устроил.
        _last_job = {
            **{k: data.get(k) for k in (
                "doc_type", "custom_doc_name", "topic", "subject", "pages",
                "writing_style", "institution", "institution_type", "org_type",
                "group", "author", "teacher", "city", "page_number_position",
                "source_content", "include_images", "image_count", "use_own_images",
                "images", "own_images",
            )},
            "model_key": model_key,
            "pay_mode": pay_mode,
        }
        await state.clear()
        await state.update_data(last_job=_last_job)

        await event.answer(
            "✅ <b>Работа готова!</b>\n\n"
            "Если содержание не обновилось — откройте файл в Word и нажмите:\n"
            "<code>Ctrl+A → F9 → Обновить всё поле</code>\n\n"
            "Если результат не тот — можно <b>перегенерировать эту же работу</b> "
            "с теми же параметрами.",
            reply_markup=kb_final(),
            parse_mode="HTML",
        )

        # Чистим временные файлы
        for tmp in (tmp_in, tmp_out):
            try:
                if os.path.exists(tmp):
                    os.remove(tmp)
            except Exception:
                pass

        # Освобождаем слот пользователя (несколько работ в одном чате).
        _active_gen_per_user[_uid] = max(0, _active_gen_per_user.get(_uid, 1) - 1)


# ═══════════════════════════════════════════════════════════════
#  ЗАПУСК
# ═══════════════════════════════════════════════════════════════

async def main() -> None:
    global bot
    if not BOT_TOKEN:
        raise SystemExit(
            "❌ ОШИБКА: не задан BOT_TOKEN. Укажите его в переменной окружения "
            "BOT_TOKEN или в bot_config.json."
        )
    bot = Bot(token=BOT_TOKEN)

    print("═" * 62)
    print("  🤖  ГОСТ-АССИСТЕНТ v3.0-gost")
    print("═" * 62)
    print(f"  LibreOffice : {shutil.which('soffice') or '❌ не найден'}")
    print(f"  DeepSeek    : {'✅' if DEEPSEEK_KEY else '❌ нет ключа'}")
    print(f"  OpenRouter  : {'✅' if OPENROUTER_KEY else '❌ нет ключа'}")
    print(f"  Groq        : {'✅' if GROQ_KEY else '❌ нет ключа'}")
    print(f"  VIP users   : {VIP_USERS or 'нет'}")
    _fcd_days = FREE_COOLDOWN / 86400
    print(f"  Free limit  : 1 ген / {_fcd_days:.1f} дн, max {FREE_MAX_PAGES} стр.")
    print(f"  Paid limit  : {'безлимит' if PAID_DAILY_LIMIT == 0 else str(PAID_DAILY_LIMIT)+'/день'}")
    print("═" * 62)

    await _start_http_server()

    while True:
        try:
            print("[BOT] Запускаю polling (конкурентная обработка апдейтов)...")
            # handle_as_tasks=True — каждый апдейт обрабатывается отдельной
            # задачей, поэтому несколько пользователей (и несколько работ в
            # одном чате) обслуживаются параллельно, а не по очереди.
            await dp.start_polling(bot, handle_as_tasks=True)
        except Exception as e:
            print(f"[BOT] Ошибка поллинга: {e}")
            await asyncio.sleep(5)


if __name__ == "__main__":
    asyncio.run(main())

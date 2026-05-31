# -*- coding: utf-8 -*-
from __future__ import annotations

"""ГОСТ-АССИСТЕНТ v2.0 — ПОЛНАЯ ПЕРЕПИСЬ

Что исправлено и улучшено:
────────────────────────────────────────────────────────────────
1. СТРАНИЦЫ: точный подсчёт символов. Генератор теперь запрашивает
   строго нужный объём и дополняет текст если он короче цели.
   max_tokens рассчитывается корректно (~750 токен/страница).

2. ГЛАВЫ: названия больше не "Глава 1" — ИИ генерирует развёрнутые
   заголовки («Теоретические основы ...», «Анализ современного
   состояния ...»), а также подглавы (1.1, 1.2, 1.3).

3. ЛИМИТЫ: платные пользователи (mode=paid) — БЕЗЛИМИТНЫЕ генерации.
   Кулдаун для платных убран до 0. FREE_DAILY_LIMIT остаётся.

4. АНИМАЦИЯ: Progress показывает:
   - красивый анимированный спиннер (меняется каждую секунду)
   - прогресс-бар с процентом
   - текущий шаг и его название
   - прошедшее время и примерное время до конца (ETA)

5. СООБЩЕНИЯ: все reply-сообщения переработаны — красивое
   оформление через HTML, эмодзи-разделители, форматированные
   карточки для каждого шага.

6. СТРУКТУРА DOCX: содержание теперь действительно заполнено
   (Heading 1 + Heading 2 для подглав), заголовки глав развёрнуты.

Примечание по "точным страницам":
Точность ±10% достигается за счёт корректного расчёта max_tokens
и дозаполнения коротких секций. LibreOffice дополнительно обновляет TOC.
"""

import asyncio
import io
import json
import os
import re
import shutil
import subprocess
import time
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

import aiohttp
from aiohttp import web
from aiogram import Bot, Dispatcher, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    BufferedInputFile,
    CallbackQuery,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    LabeledPrice,
    Message,
    PreCheckoutQuery,
)
from aiogram.utils.keyboard import InlineKeyboardBuilder

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


# ═══════════════════════════════════════════════════════════════
#  КРАСИВАЯ АНИМАЦИЯ ПРОГРЕССА С ETA
# ═══════════════════════════════════════════════════════════════

# Рамки спиннера — меняются каждую секунду для иллюзии движения
_SPINNER_FRAMES = ["⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏"]
_BAR_FULL  = "█"
_BAR_HALF  = "▓"
_BAR_EMPTY = "░"


def _spinner() -> str:
    """Возвращает текущий кадр спиннера на основе текущего времени."""
    idx = int(time.monotonic() * 4) % len(_SPINNER_FRAMES)
    return _SPINNER_FRAMES[idx]


def _progress_bar(done: int, total: int, width: int = 16) -> str:
    """Рисует красивый прогресс-бар с половинным блоком на границе."""
    total = max(1, int(total))
    done  = max(0, min(int(done), total))
    ratio = done / total
    filled_f = width * ratio
    filled_i = int(filled_f)
    half     = filled_f - filled_i >= 0.5
    bar      = _BAR_FULL * filled_i
    if half and filled_i < width:
        bar += _BAR_HALF
        bar += _BAR_EMPTY * (width - filled_i - 1)
    else:
        bar += _BAR_EMPTY * (width - filled_i)
    return bar


def _fmt_time(seconds: float) -> str:
    """Форматирует секунды в mm:ss."""
    s = max(0, int(seconds))
    m, s = divmod(s, 60)
    if m:
        return f"{m}м {s:02d}с"
    return f"{s}с"


@dataclass
class Progress:
    """Красивый прогресс с анимацией, ETA и шагами."""
    msg: Message
    title: str
    total_steps: int
    model_name: str = ""
    done: int = 0
    label: str = "Подготовка..."
    _last_text: str = field(default="", repr=False)
    _last_ts: float = field(default=0.0, repr=False)
    _start_ts: float = field(default_factory=time.monotonic, repr=False)
    _step_labels: list = field(default_factory=list, repr=False)

    def _elapsed(self) -> float:
        return time.monotonic() - self._start_ts

    def _eta(self) -> str:
        elapsed = self._elapsed()
        if self.done == 0:
            return "считаю..."
        rate = self.done / elapsed  # шагов в секунду
        remaining_steps = self.total_steps - self.done
        if rate <= 0:
            return "..."
        eta_sec = remaining_steps / rate
        return _fmt_time(eta_sec)

    def render(self) -> str:
        pct      = int(self.done * 100 / max(1, self.total_steps))
        bar      = _progress_bar(self.done, self.total_steps)
        spin     = _spinner()
        elapsed  = _fmt_time(self._elapsed())
        eta      = self._eta()
        model    = f"\n🤖 <b>Модель:</b> <i>{self.model_name}</i>" if self.model_name else ""
        step_num = min(self.done + 1, self.total_steps)

        return (
            f"{spin} <b>{self.title}</b>\n"
            f"━━━━━━━━━━━━━━━━━━━━━━\n"
            f"<code>[{bar}] {pct}%</code>\n"
            f"━━━━━━━━━━━━━━━━━━━━━━\n"
            f"📌 <b>Шаг {step_num}/{self.total_steps}:</b> {self.label}{model}\n"
            f"⏱ Прошло: <code>{elapsed}</code> · До конца: <code>{eta}</code>"
        )

    async def update(
        self,
        *,
        label: Optional[str] = None,
        step_done: bool = False,
        model_name: Optional[str] = None,
        force: bool = False,
        min_interval: float = 1.0,
    ) -> None:
        if label is not None:
            self.label = label
        if model_name is not None:
            self.model_name = model_name
        if step_done:
            self.done = min(self.total_steps, self.done + 1)

        text = self.render()
        now  = time.monotonic()

        # Не спамим обновлениями чаще min_interval
        if not force and text == self._last_text:
            return
        if not force and (now - self._last_ts) < min_interval:
            return

        try:
            await self.msg.edit_text(text, parse_mode="HTML")
            self._last_text = text
            self._last_ts   = now
        except Exception:
            pass

    async def finish(self, text: str) -> None:
        try:
            await self.msg.edit_text(text, parse_mode="HTML")
        except Exception:
            pass

    async def delete(self) -> None:
        try:
            await self.msg.delete()
        except Exception:
            pass

    async def animate_loop(self, stop_event: asyncio.Event) -> None:
        """Фоновый цикл: обновляет спиннер и ETA каждую секунду."""
        while not stop_event.is_set():
            try:
                text = self.render()
                now  = time.monotonic()
                if text != self._last_text and (now - self._last_ts) >= 1.0:
                    await self.msg.edit_text(text, parse_mode="HTML")
                    self._last_text = text
                    self._last_ts   = now
            except Exception:
                pass
            await asyncio.sleep(1.2)


# ═══════════════════════════════════════════════════════════════
#  КОНФИГ
# ═══════════════════════════════════════════════════════════════

TOKENS = {
    "BOT_TOKEN": "",
    "OPENROUTER_KEY": "",
    "DEEPSEEK_KEY": "",
    "GROQ_KEY": "",
    "VIP_USERS": "5291613279",
}

CONFIG_FILE      = "bot_config.json"
GOST_CONFIG_FILE = "gost_configs.json"
USAGE_FILE       = "usage_limits.json"


def _load_json(path: str, default):
    if not os.path.exists(path):
        return default
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default


def _save_json(path: str, data) -> None:
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"[WARN] Ошибка сохранения {path}: {e}")


CONFIG      = _load_json(CONFIG_FILE, {})
GOST_CONFIGS = _load_json(GOST_CONFIG_FILE, {})


def cfg(name: str, default: str = "") -> str:
    return (TOKENS.get(name) or os.getenv(name) or CONFIG.get(name) or default).strip()


BOT_TOKEN = cfg("BOT_TOKEN")
if not BOT_TOKEN:
    raise SystemExit("❌ ОШИБКА: не вставлен BOT_TOKEN (в TOKENS, .env или bot_config.json)")

OPENROUTER_KEY = cfg("OPENROUTER_KEY")
DEEPSEEK_KEY   = cfg("DEEPSEEK_KEY")
GROQ_KEY       = cfg("GROQ_KEY")

OPENROUTER_BASE_URL = cfg("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1").rstrip("/")
DEEPSEEK_BASE_URL   = cfg("DEEPSEEK_BASE_URL",   "https://api.deepseek.com/v1").rstrip("/")
GROQ_BASE_URL       = cfg("GROQ_BASE_URL",       "https://api.groq.com/openai/v1").rstrip("/")

DEEPSEEK_MODEL         = cfg("DEEPSEEK_MODEL",         "deepseek-chat")
OPENROUTER_R1_MODEL    = cfg("OPENROUTER_R1_MODEL",    "deepseek/deepseek-r1-0528:free")
OPENROUTER_GEMINI_MODEL = cfg("OPENROUTER_GEMINI_MODEL", "google/gemini-2.5-flash")
GROQ_MODEL             = cfg("GROQ_MODEL",             "llama-3.3-70b-versatile")

FREE_MODEL_KEY = cfg("FREE_MODEL_KEY", "deepseek")

# Лимиты для БЕСПЛАТНОГО режима
FREE_MAX_PAGES    = int(cfg("FREE_MAX_PAGES",         "15"))
FREE_DAILY_LIMIT  = int(cfg("FREE_DAILY_LIMIT",       "1"))
FREE_COOLDOWN     = int(cfg("FREE_COOLDOWN_SECONDS",  "1800"))

# Лимиты для ПЛАТНОГО режима — платят деньги, получают безлимит
# PAID_DAILY_LIMIT = 0 означает "без лимита"
PAID_DAILY_LIMIT  = int(cfg("PAID_DAILY_LIMIT",       "0"))   # 0 = безлимит
PAID_COOLDOWN     = int(cfg("PAID_COOLDOWN_SECONDS",  "0"))   # 0 = нет кулдауна

# Символов на страницу (ГОСТ: ~1800 знаков с пробелами на стр A4 14pt 1.5 интервал)
CHARS_PER_PAGE = int(cfg("CHARS_PER_PAGE", "1800"))
# Страниц без текста (титул + содержание)
NON_TEXT_PAGES = int(cfg("NON_TEXT_PAGES", "2"))

DEEPSEEK_PRICE  = int(cfg("DEEPSEEK_PRICE",           "5"))
GROQ_PRICE      = int(cfg("GROQ_PRICE",               "3"))
OR_GEMINI_PRICE = int(cfg("OPENROUTER_GEMINI_PRICE",  "7"))

MAX_PARALLEL  = int(cfg("MAX_PARALLEL", "1"))
GEN_SEMAPHORE = asyncio.Semaphore(MAX_PARALLEL)

VIP_USERS = {int(x) for x in cfg("VIP_USERS", "").replace(" ", "").split(",") if x.isdigit()}


# ═══════════════════════════════════════════════════════════════
#  ТИПЫ ДОКУМЕНТОВ
# ═══════════════════════════════════════════════════════════════

DOC_TYPES = {
    "referat": {
        "name": "📄 Реферат",
        "word": "РЕФЕРАТ",
        "min_pages": 10,
        "max_pages": 30,
        "structure": "Введение · 2 главы с подглавами · Заключение · Список литературы",
        "desc": "Обзор литературы по теме, изложение основных концепций и выводы",
    },
    "kursovaya": {
        "name": "📚 Курсовая работа",
        "word": "КУРСОВАЯ РАБОТА",
        "min_pages": 25,
        "max_pages": 50,
        "structure": "Введение · 3 главы с подглавами (§) · Заключение · Библиография",
        "desc": "Самостоятельное исследование с анализом, практической частью и выводами",
    },
    "doklad": {
        "name": "🎤 Доклад",
        "word": "ДОКЛАД",
        "min_pages": 5,
        "max_pages": 15,
        "structure": "Введение · Основная часть · Заключение",
        "desc": "Краткий обзор темы для устного выступления",
    },
    "esse": {
        "name": "✍️ Эссе",
        "word": "ЭССЕ",
        "min_pages": 3,
        "max_pages": 10,
        "structure": "Вступление · Аргументы · Авторская позиция",
        "desc": "Авторский взгляд на проблему с аргументацией и личными выводами",
    },
    "kontrolnaya": {
        "name": "📝 Контрольная работа",
        "word": "КОНТРОЛЬНАЯ РАБОТА",
        "min_pages": 10,
        "max_pages": 25,
        "structure": "Теоретическая часть · Практическая часть · Выводы",
        "desc": "Проверочная работа по теории и практике дисциплины",
    },
    "final_project": {
        "name": "📦 Итоговый проект",
        "word": "ИТОГОВЫЙ ПРОЕКТ",
        "min_pages": 35,
        "max_pages": 80,
        "structure": "Введение · 4 главы · Заключение · Приложения · Источники",
        "desc": "Комплексный проект с теорией, практикой и проектной частью",
    },
    "custom": {
        "name": "🧩 Свой тип",
        "word": "РАБОТА",
        "min_pages": 3,
        "max_pages": 100,
        "structure": "Пользовательская структура",
        "desc": "Любой тип документа с вашими требованиями к структуре",
    },
}

INSTITUTION_TYPES = {
    "school": {
        "name": "🏫 Школа",
        "org_example": "Муниципальное бюджетное общеобразовательное учреждение",
        "name_example": "Средняя общеобразовательная школа № 123",
    },
    "college": {
        "name": "🏛 Колледж / СПО",
        "org_example": "Государственное бюджетное профессиональное образовательное учреждение",
        "name_example": "Колледж информационных технологий № 42",
    },
    "university": {
        "name": "🎓 ВУЗ / Университет",
        "org_example": "Федеральное государственное бюджетное образовательное учреждение высшего образования",
        "name_example": "Московский государственный университет имени М.В. Ломоносова",
    },
    "custom": {
        "name": "✏️ Свой вариант",
        "org_example": "",
        "name_example": "",
    },
}

SUBJECTS = [
    "История", "Обществознание", "Литература", "Русский язык",
    "Математика", "Физика", "Информатика", "Биология",
    "Химия", "География", "Экономика", "Право",
    "Философия", "Психология", "Социология", "Менеджмент",
    "Педагогика", "Медицина", "Архитектура", "Юриспруденция",
]

CITIES = [
    "Москва", "Санкт-Петербург", "Новосибирск", "Екатеринбург",
    "Казань", "Нижний Новгород", "Красноярск", "Челябинск",
    "Омск", "Ростов-на-Дону", "Уфа", "Волгоград",
]


# ═══════════════════════════════════════════════════════════════
#  ГОСТ — дефолты и пользовательские настройки
# ═══════════════════════════════════════════════════════════════

DEFAULT_GOST_CONFIGS: dict = {
    "_base": {
        "font_name":              "Times New Roman",
        "font_size":              14,
        "line_spacing":           1.5,
        "first_line_indent_cm":   1.25,
        "left_margin_mm":         30,
        "right_margin_mm":        15,
        "top_margin_mm":          20,
        "bottom_margin_mm":       20,
        "alignment":              "justify",
        "page_number_position":   "bottom_center",
    }
}

for _t in DOC_TYPES.keys():
    DEFAULT_GOST_CONFIGS.setdefault(_t, dict(DEFAULT_GOST_CONFIGS["_base"]))

if not isinstance(GOST_CONFIGS, dict):
    GOST_CONFIGS = {}

for _t in DOC_TYPES.keys():
    GOST_CONFIGS.setdefault(_t, {})
    for _k, _v in DEFAULT_GOST_CONFIGS[_t].items():
        GOST_CONFIGS[_t].setdefault(_k, _v)

_save_json(GOST_CONFIG_FILE, GOST_CONFIGS)


def get_gost_config(doc_type: str, user_id: Optional[int] = None) -> dict:
    doc_type = doc_type if doc_type in DOC_TYPES else "referat"
    config = dict(GOST_CONFIGS.get(doc_type, DEFAULT_GOST_CONFIGS[doc_type]))
    if user_id:
        user_file = f"user_gost_{user_id}.json"
        u = _load_json(user_file, {})
        if isinstance(u, dict) and doc_type in u and isinstance(u[doc_type], dict):
            config.update(u[doc_type])
    return config


def save_user_gost_config(user_id: int, doc_type: str, config: dict) -> None:
    user_file = f"user_gost_{user_id}.json"
    u = _load_json(user_file, {})
    if not isinstance(u, dict):
        u = {}
    u[doc_type] = config
    _save_json(user_file, u)


# ═══════════════════════════════════════════════════════════════
#  ЛИМИТЫ / VIP
# ═══════════════════════════════════════════════════════════════

def is_vip(user_id: int) -> bool:
    return int(user_id) in VIP_USERS


def today_key() -> str:
    return datetime.now().strftime("%Y-%m-%d")


def fmt_seconds(seconds: int) -> str:
    m, s = divmod(max(0, int(seconds)), 60)
    if m:
        return f"{m} мин {s:02d} сек"
    return f"{s} сек"


def load_usage() -> dict:
    return _load_json(USAGE_FILE, {})


def save_usage(d: dict) -> None:
    _save_json(USAGE_FILE, d)


def check_user_limit(user_id: int, mode: str) -> tuple[bool, str]:
    """Проверяет ограничения.

    Платные пользователи (mode='paid') — БЕЗЛИМИТНЫ:
    PAID_DAILY_LIMIT=0 означает "без лимита", PAID_COOLDOWN=0 — без кулдауна.
    """
    if is_vip(user_id):
        return True, ""

    is_free = mode == "free"

    # Платный режим — без ограничений если PAID_DAILY_LIMIT == 0
    if not is_free and PAID_DAILY_LIMIT == 0 and PAID_COOLDOWN == 0:
        return True, ""

    data   = load_usage()
    uid    = str(user_id)
    today  = today_key()
    rec    = data.get(uid, {})
    if rec.get("date") != today:
        rec = {"date": today, "free": 0, "paid": 0, "last_free_ts": 0, "last_paid_ts": 0}

    now = int(datetime.now().timestamp())

    # Кулдаун
    cooldown = FREE_COOLDOWN if is_free else PAID_COOLDOWN
    ts_key   = "last_free_ts" if is_free else "last_paid_ts"
    last_ts  = int(rec.get(ts_key, 0))
    if cooldown > 0 and last_ts and (now - last_ts) < cooldown:
        wait = cooldown - (now - last_ts)
        return False, (
            f"⏳ <b>Подождите перед следующей генерацией</b>\n\n"
            f"Осталось: <code>{fmt_seconds(wait)}</code>"
        )

    # Дневной лимит
    limit = FREE_DAILY_LIMIT if is_free else PAID_DAILY_LIMIT
    used  = int(rec.get("free" if is_free else "paid", 0))
    if limit > 0 and used >= limit:
        kind = "бесплатных" if is_free else "платных"
        return False, (
            f"🚫 <b>Дневной лимит {kind} генераций исчерпан</b>\n\n"
            f"Использовано: <b>{used}/{limit}</b>\n"
            f"Лимит обновится в полночь 🕛"
        )

    return True, ""


def record_user_generation(user_id: int, mode: str) -> None:
    if is_vip(user_id):
        return

    data  = load_usage()
    uid   = str(user_id)
    today = today_key()

    rec = data.get(uid, {})
    if rec.get("date") != today:
        rec = {"date": today, "free": 0, "paid": 0, "last_free_ts": 0, "last_paid_ts": 0}

    key    = "free" if mode == "free" else "paid"
    ts_key = "last_free_ts" if mode == "free" else "last_paid_ts"
    rec[key]    = int(rec.get(key, 0)) + 1
    rec[ts_key] = int(datetime.now().timestamp())
    rec.setdefault("date", today)

    data[uid] = rec
    save_usage(data)


def get_user_limits_info(user_id: int) -> str:
    """Возвращает красивую карточку с лимитами пользователя."""
    if is_vip(user_id):
        return (
            "┌─────────────────────────\n"
            "│ 👑 <b>Статус: VIP</b>\n"
            "│ ♾ Генерации — без лимитов\n"
            "└─────────────────────────"
        )

    data  = load_usage()
    uid   = str(user_id)
    today = today_key()
    rec   = data.get(uid, {})
    if rec.get("date") != today:
        rec = {"date": today, "free": 0, "paid": 0}

    free_used = int(rec.get("free", 0))
    paid_used = int(rec.get("paid", 0))

    paid_str = (
        "♾ Безлимитно" if PAID_DAILY_LIMIT == 0
        else f"{paid_used}/{PAID_DAILY_LIMIT}"
    )

    return (
        "┌─────────────────────────\n"
        f"│ 🆓 Бесплатно сегодня: <b>{free_used}/{FREE_DAILY_LIMIT}</b>\n"
        f"│ ⭐ Платные сегодня:    <b>{paid_str}</b>\n"
        "└─────────────────────────"
    )


# ═══════════════════════════════════════════════════════════════
#  AI МОДЕЛИ
# ═══════════════════════════════════════════════════════════════

class ModelStatus:
    AVAILABLE = "✅"
    LIMIT     = "❌ лимит"
    UNKNOWN   = "❓"
    FATAL     = "🔴 ошибка"


AI_MODELS: dict = {
    "deepseek": {
        "name":           "🐋 DeepSeek Chat",
        "base_url":       DEEPSEEK_BASE_URL,
        "api_key":        DEEPSEEK_KEY,
        "model":          DEEPSEEK_MODEL,
        "price_per_page": DEEPSEEK_PRICE,
        "status":         ModelStatus.UNKNOWN,
        "_fatal":         False,
    },
    "deepseek_r1": {
        "name":           "🧠 DeepSeek R1 (OpenRouter)",
        "base_url":       OPENROUTER_BASE_URL,
        "api_key":        OPENROUTER_KEY,
        "model":          OPENROUTER_R1_MODEL,
        "price_per_page": DEEPSEEK_PRICE,
        "status":         ModelStatus.UNKNOWN,
        "_fatal":         False,
    },
    "gemini_or": {
        "name":           "🌟 Gemini 2.5 Flash (OpenRouter)",
        "base_url":       OPENROUTER_BASE_URL,
        "api_key":        OPENROUTER_KEY,
        "model":          OPENROUTER_GEMINI_MODEL,
        "price_per_page": OR_GEMINI_PRICE,
        "status":         ModelStatus.UNKNOWN,
        "_fatal":         False,
    },
    "groq": {
        "name":           "⚡ Groq (LLaMA 3.3 70B)",
        "base_url":       GROQ_BASE_URL,
        "api_key":        GROQ_KEY,
        "model":          GROQ_MODEL,
        "price_per_page": GROQ_PRICE,
        "status":         ModelStatus.UNKNOWN,
        "_fatal":         False,
    },
}

if FREE_MODEL_KEY not in AI_MODELS:
    FREE_MODEL_KEY = "deepseek"


def sanitize_llm_text(raw: str) -> str:
    """Чистит мусор от LLM: markdown-разметку, тройные переводы строк."""
    if not raw:
        return ""
    text = raw.strip()
    # убираем ```код``` блоки
    text = re.sub(r"```[^\n]*\n?", "", text)
    # убираем **жирный** и *курсив* markdown
    text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
    # убираем # заголовки markdown
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # убираем тройные+ переводы строк
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


async def call_openai_compat(
    info: dict,
    messages: list[dict],
    max_tokens: int = 4096,
    timeout: int = 300,
) -> str:
    """Вызов OpenAI-совместимого API."""
    if info.get("_fatal") or not info.get("api_key"):
        return ""

    base    = info["base_url"].rstrip("/")
    headers = {
        "Content-Type":  "application/json",
        "Authorization": f"Bearer {info['api_key']}",
    }

    if "openrouter.ai" in base:
        headers["HTTP-Referer"] = "https://t.me/gost_assistant_bot"
        headers["X-Title"]      = "GOST Assistant Bot"

    payload = {
        "model":       info["model"],
        "messages":    messages,
        "temperature": 0.7,
        "max_tokens":  max_tokens,
    }

    try:
        async with aiohttp.ClientSession() as sess:
            async with sess.post(
                f"{base}/chat/completions",
                headers=headers,
                json=payload,
                timeout=aiohttp.ClientTimeout(total=timeout),
            ) as r:
                txt = await r.text()
                if r.status == 200:
                    data = json.loads(txt)
                    return data["choices"][0]["message"]["content"]
                if r.status in (401, 402, 403):
                    info["_fatal"] = True
                    info["status"] = ModelStatus.FATAL
                    print(f"[FATAL] {info['name']} — auth error {r.status}")
                elif r.status == 429:
                    info["status"] = ModelStatus.LIMIT
                    print(f"[LIMIT] {info['name']} — rate limit")
                else:
                    print(f"[ERROR] {info['name']} — HTTP {r.status}: {txt[:200]}")
    except asyncio.TimeoutError:
        print(f"[TIMEOUT] {info['name']}")
    except Exception as e:
        print(f"[ERR] {info['name']}: {e}")

    return ""


async def chat_with_model(info: dict, messages: list[dict], max_tokens: int = 4096) -> str:
    raw = await call_openai_compat(info, messages, max_tokens=max_tokens)
    return sanitize_llm_text(raw)


def fallback_chain(primary: str) -> list[str]:
    """Цепочка фоллбэков: сначала primary, потом остальные доступные."""
    preferred = [primary, "deepseek", "deepseek_r1", "gemini_or", "groq"]
    out: list[str] = []
    for k in preferred:
        if k not in AI_MODELS or k in out:
            continue
        info = AI_MODELS[k]
        if info.get("_fatal"):
            continue
        if info.get("api_key"):
            out.append(k)
    return out


async def chat_with_fallback(
    primary: str,
    messages: list[dict],
    max_tokens: int,
) -> tuple[str, str]:
    """Пробует модели по цепочке, возвращает (текст, ключ_модели)."""
    for k in fallback_chain(primary):
        info = AI_MODELS[k]
        text = await chat_with_model(info, messages, max_tokens=max_tokens)
        if text and len(text.strip()) > 50:
            info["status"] = ModelStatus.AVAILABLE
            return text, k
        info["status"] = ModelStatus.LIMIT
    return "", primary


# ═══════════════════════════════════════════════════════════════
#  ГЕНЕРАЦИЯ СТРУКТУРЫ И НАЗВАНИЙ ГЛАВ ЧЕРЕЗ ИИ
# ═══════════════════════════════════════════════════════════════

async def generate_chapter_titles(
    model_key: str,
    doc_type: str,
    topic: str,
    subject: str,
    num_chapters: int,
) -> list[dict]:
    """
    Просит ИИ придумать развёрнутые названия глав и подглав.

    Возвращает список словарей:
    [
      {"title": "Глава 1. Теоретические основы ...", "subs": ["1.1. ...", "1.2. ..."]},
      ...
    ]
    """
    system = (
        "Ты помогаешь составлять структуру академических работ. "
        "Отвечай СТРОГО в формате JSON-массива без пояснений и без markdown. "
        "Каждый элемент: {\"title\": \"...\", \"subs\": [\"...\", \"...\"]}. "
        "Название главы должно быть развёрнутым и конкретным — не 'Глава 1', "
        "а 'Глава 1. Теоретические основы изучения ...'."
    )

    user = (
        f"Тема работы: «{topic}». "
        f"Дисциплина: {subject}. "
        f"Тип документа: {DOC_TYPES.get(doc_type, DOC_TYPES['referat'])['word']}. "
        f"Количество глав: {num_chapters}. "
        f"Каждая глава — 2–3 подглавы. "
        f"Язык: русский. "
        f"Верни ТОЛЬКО JSON-массив."
    )

    messages = [
        {"role": "system", "content": system},
        {"role": "user",   "content": user},
    ]

    raw, _ = await chat_with_fallback(model_key, messages, max_tokens=1500)
    raw    = raw.strip()

    # Вытаскиваем JSON даже если модель добавила мусор
    m = re.search(r"\[.*\]", raw, flags=re.S)
    if m:
        raw = m.group(0)

    try:
        result = json.loads(raw)
        if isinstance(result, list) and all("title" in r for r in result):
            return result
    except Exception:
        pass

    # Фоллбэк: базовые названия
    return _default_chapter_titles(doc_type, topic, num_chapters)


def _default_chapter_titles(doc_type: str, topic: str, num_chapters: int) -> list[dict]:
    """Запасные названия глав если ИИ не ответил."""
    defaults = [
        {
            "title": f"Глава 1. Теоретические основы исследования темы «{topic[:40]}»",
            "subs":  [
                f"1.1. Понятие и сущность исследуемой проблематики",
                f"1.2. Исторические предпосылки и современное состояние",
                f"1.3. Нормативно-правовая база и научные подходы",
            ],
        },
        {
            "title": f"Глава 2. Анализ и современное состояние проблемы",
            "subs":  [
                f"2.1. Характеристика основных факторов и условий",
                f"2.2. Сравнительный анализ подходов и методов",
                f"2.3. Проблемы и противоречия в изучаемой области",
            ],
        },
        {
            "title": f"Глава 3. Практические аспекты и пути решения",
            "subs":  [
                f"3.1. Практика применения теоретических положений",
                f"3.2. Рекомендации по совершенствованию",
                f"3.3. Перспективы развития исследуемой сферы",
            ],
        },
        {
            "title": f"Глава 4. Оценка результатов и выводы",
            "subs":  [
                f"4.1. Обобщение результатов исследования",
                f"4.2. Практическая значимость выводов",
            ],
        },
    ]
    return defaults[:num_chapters]


# ═══════════════════════════════════════════════════════════════
#  РАСЧЁТ ОБЪЁМА ТЕКСТА — ИСПРАВЛЕННЫЙ
# ═══════════════════════════════════════════════════════════════

def target_chars(pages: int) -> int:
    """
    Целевое количество символов с пробелами для основного текста.
    Вычитаем NON_TEXT_PAGES (титул + содержание), остаток умножаем на CHARS_PER_PAGE.
    """
    text_pages = max(1, pages - NON_TEXT_PAGES)
    return text_pages * CHARS_PER_PAGE


def tokens_for_chars(chars: int) -> int:
    """
    Примерно 1 токен = 3.0 символа (для русского, с запасом).
    Добавляем 30% запас чтобы не обрезало.
    """
    return max(1000, min(16000, int(chars / 3.0 * 1.3)))


def _style_instruction(writing_style: str) -> str:
    """Возвращает инструкцию по стилю для системного промпта."""
    if writing_style == "smart":
        return (
            "Используй сложный академический язык: многосоставные предложения, "
            "научную терминологию дисциплины, ссылки на теории и концепции, "
            "глубокий аналитический разбор. Язык должен быть характерен для "
            "диссертаций и монографий."
        )
    return (
        "Используй чёткий деловой научный стиль: ясные формулировки, "
        "логичная структура, конкретные утверждения без излишней сложности. "
        "Язык должен быть понятен образованному читателю без специальной подготовки."
    )


def strict_prompt(task: str, chars: int, writing_style: str = "classic") -> str:
    """Системная инструкция для строго академического текста нужного объёма."""
    style_instr = _style_instruction(writing_style)
    return (
        f"{task}\n\n"
        f"Требуемый объём: МИНИМУМ {chars} знаков с пробелами (это строго обязательно — "
        f"пиши развёрнуто до достижения нужного объёма).\n"
        f"Стиль: {style_instr}\n"
        "Каждый абзац — 5–8 предложений, без маркированных списков, без нумерации абзацев.\n"
        "ЗАПРЕЩЕНО: markdown-разметка, символы **, ##, [], скобки со цифрами типа [1] [2], "
        "фразы «как ИИ», «давайте рассмотрим», «следует отметить», «таким образом» в начале абзацев. "
        "Не повторяй одинаковые слова в одном предложении. Не начинай два абзаца подряд одним словом."
    )


# ═══════════════════════════════════════════════════════════════
#  ПОСТРОЕНИЕ ПРОМПТОВ ПО ТИПАМ ДОКУМЕНТОВ
# ═══════════════════════════════════════════════════════════════

def build_prompts(
    doc_type: str,
    topic: str,
    subject: str,
    pages: int,
    source: str,
    chapter_titles: list[dict],
    writing_style: str = "classic",
) -> dict[str, str]:
    """
    Возвращает словарь {ключ_блока: промпт}.
    chapter_titles — список из generate_chapter_titles().
    """
    total = target_chars(pages)
    s     = (source or "").strip()[:12000]
    ctx   = f"\n\nИсходные материалы для использования:\n{s}\n" if s else ""

    if doc_type == "esse":
        return {
            "intro":      strict_prompt(
                f"Напиши вступление эссе на тему «{topic}», предмет «{subject}».{ctx}"
                f"Обозначь проблему, её актуальность, цель и подход автора.",
                int(total * 0.20),
                writing_style,
            ),
            "main1":      strict_prompt(
                f"Напиши первый аргумент в эссе «{topic}»."
                f"Приведи конкретные факты, мнения учёных и доказательства.",
                int(total * 0.28),
                writing_style,
            ),
            "main2":      strict_prompt(
                f"Напиши второй аргумент в эссе «{topic}» с контраргументом."
                f"Рассмотри противоположную точку зрения и опровергни её.",
                int(total * 0.28),
                writing_style,
            ),
            "conclusion": strict_prompt(
                f"Напиши заключение эссе «{topic}»."
                f"Подведи итог, выскажи авторскую позицию, дай прогноз.",
                int(total * 0.12),
                writing_style,
            ),
            "literature": (
                f"Составь список из 10–14 источников по теме «{topic}» "
                f"строго по ГОСТ Р 7.0.5-2008. "
                f"Формат: 1. Автор А.А. Название. — М.: Издательство, год. — N с.\n"
                f"Только список, без заголовков и пояснений."
            ),
        }

    if doc_type == "doklad":
        return {
            "intro":      strict_prompt(
                f"Введение доклада «{topic}», предмет «{subject}».{ctx}"
                f"Актуальность, цель, задачи.",
                int(total * 0.14),
                writing_style,
            ),
            "part1":      strict_prompt(
                f"Раздел 1 доклада «{topic}»: ключевые понятия и определения."
                f"Раскрой теоретическую базу.",
                int(total * 0.28),
                writing_style,
            ),
            "part2":      strict_prompt(
                f"Раздел 2 доклада «{topic}»: факты, статистика, примеры из практики.",
                int(total * 0.32),
                writing_style,
            ),
            "conclusion": strict_prompt(
                f"Заключение доклада «{topic}»: выводы, практическое значение.",
                int(total * 0.14),
                writing_style,
            ),
            "literature": (
                f"Составь список из 8–12 источников по теме «{topic}» по ГОСТ. "
                f"Только нумерованный список."
            ),
        }

    # Реферат, курсовая, контрольная, итоговый проект, свой тип
    # Используем chapter_titles для развёрнутых названий
    num_ch  = len(chapter_titles)
    prompts = {}

    # Введение — 10% от текста
    prompts["intro"] = strict_prompt(
        f"Напиши введение для {DOC_TYPES.get(doc_type, DOC_TYPES['referat'])['word'].lower()}а "
        f"на тему «{topic}», предмет «{subject}».{ctx}"
        f"Раскрой: актуальность темы, степень разработанности, цель, задачи (3–5 задач), "
        f"объект и предмет исследования, методы, структуру работы.",
        int(total * 0.10),
        writing_style,
    )

    # Главы — равномерно делим ~75% объёма
    chapter_share = 0.75 / max(1, num_ch)
    for i, ch in enumerate(chapter_titles, start=1):
        key      = f"ch{i}"
        ch_chars = int(total * chapter_share)
        subs_str = "; ".join(ch.get("subs", []))
        prompts[key] = strict_prompt(
            f"Напиши полный текст для главы: «{ch['title']}».\n"
            f"Тема работы: «{topic}».\n"
            f"Подглавы: {subs_str}.\n"
            f"Пиши подглавы последовательно: сначала {ch['subs'][0] if ch.get('subs') else ''}, "
            f"затем остальные. Каждая подглава — отдельный абзац с подзаголовком подглавы "
            f"(написать как обычный текст перед абзацем, например: '1.1. Название'). ",
            ch_chars,
            writing_style,
        )

    # Заключение — 8%
    prompts["conclusion"] = strict_prompt(
        f"Напиши заключение для работы «{topic}».\n"
        f"Подведи итог по каждой главе, сформулируй общие выводы, "
        f"практическую значимость и перспективы дальнейшего исследования.",
        int(total * 0.08),
        writing_style,
    )

    # Библиография
    num_sources = 12 if pages <= 20 else 20
    prompts["literature"] = (
        f"Составь список из {num_sources}–{num_sources + 5} источников по теме «{topic}» "
        f"строго по ГОСТ Р 7.0.5-2008. Включи: монографии, учебники, статьи из журналов, "
        f"нормативные документы (если уместно), интернет-ресурсы. "
        f"Формат: 1. Автор А.А. Название / А.А. Автор. — М.: Изд-во, год. — N с.\n"
        f"Только нумерованный список без заголовков."
    )

    return prompts


# ═══════════════════════════════════════════════════════════════
#  СТРУКТУРА DOCX-БЛОКОВ
# ═══════════════════════════════════════════════════════════════

def generate_structure(
    doc_type: str,
    parts: dict[str, str],
    chapter_titles: list[dict],
) -> list[tuple]:
    """
    Возвращает список блоков для DOCX:
    (заголовок, уровень_heading, текст_абзаца, список_подглав)

    Уровень: 1 = Heading 1, 2 = Heading 2
    список_подглав = [] или [(название_подглавы, текст), ...]
    """
    if doc_type == "esse":
        return [
            ("ВВЕДЕНИЕ", 1, parts.get("intro", ""), []),
            ("ОСНОВНАЯ ЧАСТЬ", 1, "", [
                ("Аргумент 1", parts.get("main1", "")),
                ("Аргумент 2. Контраргумент и опровержение", parts.get("main2", "")),
            ]),
            ("ЗАКЛЮЧЕНИЕ", 1, parts.get("conclusion", ""), []),
            ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1, parts.get("literature", ""), []),
        ]

    if doc_type == "doklad":
        return [
            ("ВВЕДЕНИЕ", 1, parts.get("intro", ""), []),
            ("1. Теоретические основы и ключевые понятия", 1, parts.get("part1", ""), []),
            ("2. Факты, статистика и практические примеры", 1, parts.get("part2", ""), []),
            ("ЗАКЛЮЧЕНИЕ", 1, parts.get("conclusion", ""), []),
            ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1, parts.get("literature", ""), []),
        ]

    # Универсальная структура (реферат / курсовая / контрольная / итоговый / свой)
    blocks = [("ВВЕДЕНИЕ", 1, parts.get("intro", ""), [])]

    for i, ch in enumerate(chapter_titles, start=1):
        key  = f"ch{i}"
        text = parts.get(key, "")
        subs = ch.get("subs", [])
        # Подглавы извлекаем из текста (ИИ уже вставил подзаголовки внутри текста)
        # Возвращаем пустой список — текст уже содержит подзаголовки
        blocks.append((ch["title"].upper() if ch["title"].upper() == ch["title"] else ch["title"], 1, text, []))

    blocks.append(("ЗАКЛЮЧЕНИЕ", 1, parts.get("conclusion", ""), []))
    blocks.append(("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1, parts.get("literature", ""), []))

    return blocks


# ═══════════════════════════════════════════════════════════════
#  ГЕНЕРАЦИЯ ТЕКСТА С ДОЗАПОЛНЕНИЕМ
# ═══════════════════════════════════════════════════════════════

def _clean_ai_artifacts(text: str) -> str:
    """Удаляет признаки ИИ: [1], [2], ***, ## и повторяющиеся слова в предложениях."""
    if not text:
        return ""
    # Убираем сноски [1], [2], [3] и т.п.
    text = re.sub(r"\s*\[\d+\]", "", text)
    # Убираем **жирный** markdown
    text = re.sub(r"\*{1,3}([^*\n]+)\*{1,3}", r"\1", text)
    # Убираем # заголовки markdown
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Убираем горизонтальные линии ***
    text = re.sub(r"^\*{3,}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^-{3,}\s*$", "", text, flags=re.MULTILINE)
    # Убираем тройные+ переводы строк
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Убираем фразы-маркеры ИИ
    ai_phrases = [
        r"(?i)как (языковая )?модель[,\s]",
        r"(?i)я (являюсь|являюсь )?ИИ[,\s]",
        r"(?i)в заключение следует отметить,?\s+что\s+",
        r"(?i)таким образом,?\s+можно сделать вывод",
    ]
    for p in ai_phrases:
        text = re.sub(p, "", text)
    return text.strip()


async def generate_text_blocks(
    topic: str,
    pages: int,
    doc_type: str,
    subject: str,
    model_key: str,
    source: str,
    chapter_titles: list[dict],
    writing_style: str = "classic",
    prog: Optional["Progress"] = None,
) -> dict[str, str]:
    """
    Генерирует все текстовые блоки.
    Если текст получился короче цели — дозаполняет до нужного объёма.
    """
    prompts    = build_prompts(doc_type, topic, subject, pages, source, chapter_titles, writing_style)
    total_chars = target_chars(pages)
    num_blocks = max(1, len(prompts))

    parts: dict[str, str] = {}
    step = 0

    style_sys = (
        "Ты профессионально пишешь академические тексты на русском языке строго для учебных работ. "
        "НЕ используй markdown-разметку (никаких **, ##, ---). "
        "НЕ используй сноски вида [1], [2] — вместо этого упоминай авторов в тексте: «по мнению И.И. Иванова». "
        "Не повторяй одно слово несколько раз в одном предложении. "
        "Не начинай два абзаца подряд одним и тем же словом. "
    )
    if writing_style == "smart":
        style_sys += (
            "Стиль — высокоакадемический: сложные синтаксические конструкции, "
            "специализированная терминология, ссылки на научные концепции и теории."
        )
    else:
        style_sys += (
            "Стиль — чёткий деловой научный: ясные формулировки, "
            "логичная структура, конкретные утверждения."
        )

    for key, prompt in prompts.items():
        step += 1

        # Рассчитываем целевой объём для блока
        if key == "intro":
            block_share = 0.10
        elif key == "conclusion":
            block_share = 0.08
        elif key == "literature":
            block_share = 0.0
        else:
            num_ch = max(1, len([k for k in prompts if k.startswith("ch")]))
            block_share = 0.75 / num_ch

        block_chars = int(total_chars * block_share) if block_share > 0 else 0
        max_tok     = tokens_for_chars(block_chars) if block_chars > 0 else 1500

        if prog:
            block_name = {
                "intro":       "Введение",
                "conclusion":  "Заключение",
                "literature":  "Список литературы",
                "main1":       "Основная часть (аргумент 1)",
                "main2":       "Основная часть (аргумент 2)",
                "part1":       "Раздел 1",
                "part2":       "Раздел 2",
            }.get(key, f"Глава {key.replace('ch', '')}")
            await prog.update(label=f"✍️ Пишу: {block_name}")

        messages = [
            {"role": "system", "content": style_sys},
            {"role": "user",   "content": prompt},
        ]

        text, used_model = await chat_with_fallback(model_key, messages, max_tok)

        if prog and used_model and used_model != model_key:
            await prog.update(model_name=AI_MODELS.get(used_model, {}).get("name", used_model))

        if not text or len(text) < 80:
            text = _stub_text(key, topic)

        text = _clean_ai_artifacts(text)

        # Дозаполняем если текст короче цели на 25%+ (кроме литературы)
        if key not in ("literature",) and block_chars > 0 and len(text) < int(block_chars * 0.75):
            max_refills = 3
            refill_count = 0
            while len(text) < int(block_chars * 0.9) and refill_count < max_refills:
                extra_chars = block_chars - len(text)
                ext_tok     = tokens_for_chars(extra_chars)
                extra_prompt = (
                    f"Продолжи и дополни следующий академический текст по теме «{topic}». "
                    f"Добавь ещё минимум {extra_chars} знаков. "
                    f"Пиши плавно, без заголовков, без [1] [2], без **, без ##:\n\n"
                    f"{text[-600:]}"
                )
                ext_messages = [
                    {"role": "system", "content": style_sys},
                    {"role": "user",   "content": extra_prompt},
                ]
                extra, _ = await chat_with_fallback(model_key, ext_messages, ext_tok)
                if extra and len(extra.strip()) > 50:
                    extra = _clean_ai_artifacts(extra)
                    text = text.rstrip() + "\n\n" + extra.strip()
                else:
                    break
                refill_count += 1

        parts[key] = text

        if prog:
            await prog.update(step_done=True)

    return parts


def _stub_text(key: str, topic: str) -> str:
    """Заглушка если ИИ не ответил."""
    stubs = {
        "intro":       f"Данная работа посвящена исследованию темы «{topic}». В современных условиях данная проблематика приобретает особую актуальность и практическую значимость для науки и общества.",
        "conclusion":  f"Проведённое исследование по теме «{topic}» позволило сформулировать следующие выводы: изученная проблематика имеет важное теоретическое и практическое значение.",
        "literature":  f"1. Иванов А.А. {topic} / А.А. Иванов. — М.: Наука, 2023. — 256 с.\n2. Петров Б.Б. Основы исследования. — СПб.: Питер, 2022. — 312 с.",
    }
    return stubs.get(key, f"Текст раздела по теме «{topic}» временно недоступен. Повторите генерацию.")


# ═══════════════════════════════════════════════════════════════
#  ГОСТ-DOCX: СТИЛИ, TOC, НУМЕРАЦИЯ, ПОДГЛАВЫ
# ═══════════════════════════════════════════════════════════════

def _set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass
    run.font.size  = Pt(size_pt)
    run.bold       = bool(bold)
    run.font.color.rgb = RGBColor(0, 0, 0)


def setup_gost_page(doc: Document, gost: dict) -> None:
    """Настраивает поля страницы и стиль Normal по ГОСТ."""
    for sec in doc.sections:
        sec.top_margin    = Mm(int(gost.get("top_margin_mm",    20)))
        sec.bottom_margin = Mm(int(gost.get("bottom_margin_mm", 20)))
        sec.left_margin   = Mm(int(gost.get("left_margin_mm",   30)))
        sec.right_margin  = Mm(int(gost.get("right_margin_mm",  15)))

    style      = doc.styles["Normal"]
    font_name  = gost.get("font_name", "Times New Roman")
    style.font.name = font_name
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass
    style.font.size       = Pt(int(gost.get("font_size", 14)))
    style.font.color.rgb  = RGBColor(0, 0, 0)

    ls = float(gost.get("line_spacing", 1.5))
    pf = style.paragraph_format
    if ls > 1:
        pf.line_spacing      = ls
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    pf.first_line_indent = Cm(float(gost.get("first_line_indent_cm", 1.25)))
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Настраиваем стили заголовков
    _setup_heading_style(doc, "Heading 1", font_name, int(gost.get("font_size", 14)))
    _setup_heading_style(doc, "Heading 2", font_name, int(gost.get("font_size", 14)))


def _setup_heading_style(doc: Document, style_name: str, font_name: str, size_pt: int) -> None:
    """Настраивает стиль заголовка по ГОСТ (полужирный, без отступа, по центру)."""
    try:
        style = doc.styles[style_name]
        style.font.name  = font_name
        style.font.size  = Pt(size_pt)
        style.font.bold  = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before      = Pt(12)
        pf.space_after       = Pt(6)
        pf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


def add_page_number_field(section, position: str) -> None:
    """Добавляет поле PAGE в колонтитул."""
    container = section.footer if position == "bottom_center" else section.header

    for p in list(container.paragraphs):
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass

    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)

    run = p.add_run()
    _set_run_font(run, "Times New Roman", 12, False)

    for fld_type, text in [
        ("begin", None),
        (None,     " PAGE "),
        ("separate", None),
        ("end",    None),
    ]:
        if fld_type:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), fld_type)
            run._r.append(el)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run._r.append(instr)


def add_toc(doc: Document, blocks: list[tuple], gost: dict) -> None:
    """
    Вставляет содержание с реальными заголовками.
    Сначала вставляет поле TOC (обновится в Word/LibreOffice),
    затем добавляет текстовую копию для немедленной читаемости.
    """
    fn = gost.get("font_name", "Times New Roman")
    fs = int(gost.get("font_size", 14))

    # Поле TOC для автообновления
    p_field = doc.add_paragraph()
    p_field.paragraph_format.first_line_indent = Cm(0)
    run = p_field.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)

    # Текстовое содержание (отображается сразу)
    page_num = 3  # Начинаем с 3-й страницы (1=титул, 2=содержание)
    for title, level, text, subblocks in blocks:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)

        # Заполнитель между названием и номером страницы
        tab_stops = p.paragraph_format.tab_stops
        from docx.oxml.ns import qn as _qn
        pPr = p._p.get_or_add_pPr()
        tabs_el = OxmlElement("w:tabs")
        tab_el  = OxmlElement("w:tab")
        tab_el.set(_qn("w:val"),    "right")
        tab_el.set(_qn("w:leader"), "dot")
        tab_el.set(_qn("w:pos"),    "8500")
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

        indent = Cm(0) if level == 1 else Cm(1.0)
        p.paragraph_format.left_indent = indent

        run_t = p.add_run(title)
        _set_run_font(run_t, fn, fs if level == 1 else fs - 1, level == 1)
        run_pg = p.add_run(f"\t{page_num}")
        _set_run_font(run_pg, fn, fs if level == 1 else fs - 1, False)

        # Грубая оценка страниц: считаем знаки текста
        text_len = len(text) if text else 0
        for _, st in subblocks:
            text_len += len(st) if st else 0
        pages_for_block = max(1, round(text_len / CHARS_PER_PAGE))
        page_num += pages_for_block


def add_title_page(doc: Document, data: dict, gost: dict) -> None:
    """Создаёт титульный лист по ГОСТ."""
    font = gost.get("font_name", "Times New Roman")
    size = int(gost.get("font_size", 14))

    def _add_centered(text: str, sz: int = None, bold: bool = False) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text)
        _set_run_font(r, font, sz or size, bold)

    def _add_right(text: str, bold: bool = False) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(8.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text)
        _set_run_font(r, font, size, bold)

    def _spacer(n: int = 1) -> None:
        for _ in range(n):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)

    doc_word = DOC_TYPES.get(data.get("doc_type", "referat"), DOC_TYPES["referat"])["word"]
    if data.get("doc_type") == "custom" and data.get("custom_doc_name"):
        doc_word = str(data["custom_doc_name"]).strip().upper()

    _add_centered("МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ РОССИЙСКОЙ ФЕДЕРАЦИИ", 11, True)

    if data.get("org_type"):
        _add_centered(str(data["org_type"]).upper(), 11, False)

    inst = data.get("institution") or "Учебное заведение"
    _add_centered(f"«{inst}»", 12, True)

    _spacer(4)

    _add_centered(doc_word, 18, True)
    _add_centered(f"по дисциплине «{data.get('subject', '')}»", 14, False)
    _spacer(1)
    _add_centered("на тему:", 14, False)
    _add_centered(f"«{data.get('topic', '')}»", 16, True)

    _spacer(4)

    _add_right("Выполнил(а):", False)
    gr = data.get("group", "")
    if gr:
        _add_right(f"{gr}", False)
    _add_right(str(data.get("author", "")), True)
    _spacer(1)
    _add_right("Проверил(а):", False)
    _add_right(str(data.get("teacher", "")), True)

    _spacer(4)

    _add_centered(f"{data.get('city', 'Москва')}  {datetime.now().year}", 14, False)


def add_paragraphs_from_text(
    doc: Document,
    text: str,
    gost: dict,
    is_bib: bool = False,
) -> None:
    """
    Разбивает текст на абзацы и добавляет их в документ.
    Строки вида '1.1. Название подглавы' оформляются как Heading 2.
    """
    font = gost.get("font_name", "Times New Roman")
    size = int(gost.get("font_size", 14))

    # Паттерн для подглав: "1.1. Текст" или "2.3. Текст"
    subheading_pat = re.compile(r"^(\d+\.\d+\.?\s+.{5,80})$")

    chunks = [c.strip() for c in re.split(r"\n\s*\n|\n(?=\d+\.\d+\.)", text) if c.strip()]

    for ch in chunks:
        first_line = ch.split("\n")[0].strip()

        if subheading_pat.match(first_line) and not is_bib:
            # Это подзаголовок — Heading 2
            hp = doc.add_paragraph(first_line, style="Heading 2")
            for run in hp.runs:
                _set_run_font(run, font, size, True)
            hp.paragraph_format.first_line_indent = Cm(0)
            # Остаток как обычный текст
            rest = ch[len(first_line):].strip()
            if rest:
                p = doc.add_paragraph()
                r = p.add_run(rest)
                _set_run_font(r, font, size, False)
                if is_bib:
                    p.paragraph_format.left_indent       = Cm(1.25)
                    p.paragraph_format.first_line_indent = Cm(-1.25)
        else:
            p = doc.add_paragraph()
            if is_bib:
                p.paragraph_format.left_indent       = Cm(1.25)
                p.paragraph_format.first_line_indent = Cm(-1.25)
            r = p.add_run(ch)
            _set_run_font(r, font, size, False)


def build_docx_bytes(
    data: dict,
    blocks: list[tuple],
    gost: dict,
) -> bytes:
    """Собирает DOCX из блоков."""
    doc = Document()
    setup_gost_page(doc, gost)

    # ── Титульный лист ──
    add_title_page(doc, data, gost)
    doc.add_page_break()

    # ── Содержание ──
    fn   = gost.get("font_name", "Times New Roman")
    fs   = int(gost.get("font_size", 14))

    p_toc_title = doc.add_paragraph()
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_toc_title.paragraph_format.first_line_indent = Cm(0)
    r = p_toc_title.add_run("СОДЕРЖАНИЕ")
    _set_run_font(r, fn, 16, True)

    add_toc(doc, blocks, gost)

    doc.add_page_break()

    # ── Нумерация страниц ──
    add_page_number_field(
        doc.sections[0],
        gost.get("page_number_position", "bottom_center"),
    )

    # ── Тело документа ──
    for title, level, text, subblocks in blocks:
        # Заголовок главы
        style = "Heading 1" if level == 1 else "Heading 2"
        hp    = doc.add_paragraph(title, style=style)
        for run in hp.runs:
            _set_run_font(run, fn, fs, True)
        hp.paragraph_format.first_line_indent = Cm(0)

        # Основной текст главы
        is_bib = any(w in title.upper() for w in ("ИСТОЧНИК", "ЛИТЕРАТ", "БИБЛИОГРАФ"))
        if text:
            add_paragraphs_from_text(doc, text, gost, is_bib=is_bib)

        # Подблоки (для эссе и т.п.)
        for sub_title, sub_text in subblocks:
            shp = doc.add_paragraph(sub_title, style="Heading 2")
            for run in shp.runs:
                _set_run_font(run, fn, fs, True)
            shp.paragraph_format.first_line_indent = Cm(0)
            if sub_text:
                add_paragraphs_from_text(doc, sub_text, gost)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def libreoffice_update_docx(in_path: str, out_path: str) -> bool:
    """Прогоняет через LibreOffice чтобы обновить TOC и поля страниц."""
    soffice = shutil.which("soffice")
    if not soffice:
        return False

    out_dir = os.path.dirname(out_path)
    os.makedirs(out_dir, exist_ok=True)

    try:
        p = subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--nodefault",
                "--nofirststartwizard",
                "--convert-to", "docx",
                "--outdir",     out_dir,
                in_path,
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            timeout=120,
        )
        if p.returncode != 0:
            print(f"[LO] Ошибка конвертации: {p.stdout[:500]}")
            return False

        produced = os.path.join(
            out_dir,
            os.path.splitext(os.path.basename(in_path))[0] + ".docx",
        )
        if not os.path.exists(produced):
            return False

        if produced != out_path:
            os.replace(produced, out_path)

        return True

    except subprocess.TimeoutExpired:
        print("[LO] Таймаут конвертации")
        return False
    except Exception as e:
        print(f"[LO] Ошибка: {e}")
        return False


# ═══════════════════════════════════════════════════════════════
#  «СВОЙ ГОСТ» — ПАРСИНГ ТРЕБОВАНИЙ ЧЕРЕЗ ИИ
# ═══════════════════════════════════════════════════════════════

async def parse_custom_gost_via_ai(model_key: str, gost_text: str) -> dict:
    schema_hint = {
        "font_name":            "Times New Roman",
        "font_size":            14,
        "line_spacing":         1.5,
        "first_line_indent_cm": 1.25,
        "left_margin_mm":       30,
        "right_margin_mm":      15,
        "top_margin_mm":        20,
        "bottom_margin_mm":     20,
        "page_number_position": "bottom_center",
    }

    system = (
        "Ты извлекаешь параметры оформления из описания требований пользователя. "
        "Ответ СТРОГО в JSON без объяснений и без markdown. "
        "Если параметр не указан — используй значение по умолчанию из schema."
    )

    user = (
        "schema (defaults):\n"
        + json.dumps(schema_hint, ensure_ascii=False)
        + "\n\nТребования пользователя:\n"
        + gost_text
        + "\n\nВерни ТОЛЬКО JSON."
    )

    raw, _ = await chat_with_fallback(
        model_key,
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        max_tokens=800,
    )
    raw = raw.strip()

    m = re.search(r"\{.*\}", raw, flags=re.S)
    if m:
        raw = m.group(0)

    try:
        data = json.loads(raw)
    except Exception:
        return schema_hint

    out = dict(schema_hint)
    for k in out:
        if k in data and data[k] not in (None, ""):
            out[k] = data[k]

    # Приведение типов
    try:
        out["font_size"] = int(out["font_size"])
    except Exception:
        out["font_size"] = 14

    for fkey in ("line_spacing", "first_line_indent_cm"):
        try:
            out[fkey] = float(out[fkey])
        except Exception:
            out[fkey] = schema_hint[fkey]

    for ikey in ("left_margin_mm", "right_margin_mm", "top_margin_mm", "bottom_margin_mm"):
        try:
            out[ikey] = int(out[ikey])
        except Exception:
            out[ikey] = schema_hint[ikey]

    if out.get("page_number_position") not in ("bottom_center", "top_center"):
        out["page_number_position"] = "bottom_center"

    return out


# ═══════════════════════════════════════════════════════════════
#  КЛАВИАТУРЫ — КРАСИВЫЕ
# ═══════════════════════════════════════════════════════════════

def kb_doc_type() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    for k, v in DOC_TYPES.items():
        label = f"{v['name']}  {v['min_pages']}–{v['max_pages']} стр."
        b.button(text=label, callback_data=f"dtype_{k}")
    b.adjust(1)
    return b.as_markup()


def kb_mode() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    b.button(text="🆓 Бесплатно  (быстро)", callback_data="mode_free")
    b.button(text="⭐ Платный режим  (безлимит)", callback_data="mode_paid")
    b.adjust(1)
    return b.as_markup()


def kb_writing_style() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(
                text="🎓 Умно — сложный академический язык, термины, глубокий анализ",
                callback_data="style_smart",
            )],
            [InlineKeyboardButton(
                text="📝 Классически — чёткий деловой стиль, понятные формулировки",
                callback_data="style_classic",
            )],
        ]
    )


def kb_source_choice() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="✅ Да — добавлю свои материалы", callback_data="source_yes")],
            [InlineKeyboardButton(text="❌ Нет — только по теме",         callback_data="source_no")],
        ]
    )


def kb_institution() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    for k, v in INSTITUTION_TYPES.items():
        b.button(text=v["name"], callback_data=f"inst_{k}")
    b.adjust(2)
    return b.as_markup()


def kb_subject() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    for s in SUBJECTS:
        b.button(text=s, callback_data=f"subj_{s}")
    b.button(text="✏️ Другой предмет", callback_data="subj_other")
    b.adjust(3)
    return b.as_markup()


def kb_city() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    for c in CITIES:
        b.button(text=c, callback_data=f"city_{c}")
    b.button(text="✏️ Другой город", callback_data="city_other")
    b.adjust(3)
    return b.as_markup()


def kb_page_number() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="⬇️ Внизу по центру (стандарт)", callback_data="pagepos_bottom")],
            [InlineKeyboardButton(text="⬆️ Вверху по центру",           callback_data="pagepos_top")],
        ]
    )


def kb_final() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="🔁 Создать новую работу", callback_data="final_new")],
            [InlineKeyboardButton(text="⚙️ Настроить ГОСТ",       callback_data="custom_gost")],
            [InlineKeyboardButton(text="📊 Мой баланс генераций", callback_data="show_limits")],
        ]
    )


def kb_models() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    for k, info in AI_MODELS.items():
        if not info.get("api_key") or info.get("_fatal"):
            continue
        status = info.get("status", ModelStatus.UNKNOWN)
        b.button(
            text=f"{info['name']}  {info['price_per_page']}⭐/стр  {status}",
            callback_data=f"model_{k}",
        )
    b.adjust(1)
    return b.as_markup()


def kb_cancel() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[[InlineKeyboardButton(text="❌ Отмена", callback_data="cancel_flow")]]
    )


# ═══════════════════════════════════════════════════════════════
#  FSM СОСТОЯНИЯ
# ═══════════════════════════════════════════════════════════════

class WorkState(StatesGroup):
    doc_type           = State()
    custom_doc_name    = State()
    mode               = State()
    writing_style      = State()
    topic              = State()
    source_choice      = State()
    source_content     = State()
    institution_type   = State()
    org_type           = State()
    institution        = State()
    group              = State()
    author             = State()
    teacher            = State()
    subject            = State()
    city               = State()
    pages              = State()
    page_number_position = State()
    model              = State()
    payment            = State()
    gost_free_text     = State()


# ═══════════════════════════════════════════════════════════════
#  БОТ И ДИСПЕТЧЕР
# ═══════════════════════════════════════════════════════════════

bot = Bot(token=BOT_TOKEN)
dp  = Dispatcher(storage=MemoryStorage())


# ═══════════════════════════════════════════════════════════════
#  HTTP СЕРВЕР ДЛЯ RENDER HEALTHCHECK
# ═══════════════════════════════════════════════════════════════

async def _start_http_server() -> None:
    port = int(os.getenv("PORT", "10000"))

    async def health(_: web.Request) -> web.Response:
        return web.Response(
            text=json.dumps({"status": "ok", "bot": "GOST Assistant"}, ensure_ascii=False),
            content_type="application/json",
        )

    app = web.Application()
    app.add_routes([web.get("/health", health), web.get("/", health)])

    runner = web.AppRunner(app)
    await runner.setup()
    site = web.TCPSite(runner, host="0.0.0.0", port=port)
    await site.start()
    print(f"[HTTP] Healthcheck на порту {port}")


# ═══════════════════════════════════════════════════════════════
#  ВСПОМОГАТЕЛЬНЫЕ ТЕКСТЫ СООБЩЕНИЙ
# ═══════════════════════════════════════════════════════════════

def _welcome_text(user_id: int, first_name: str) -> str:
    limits = get_user_limits_info(user_id)
    vip_mark = " 👑" if is_vip(user_id) else ""
    return (
        f"👋 Привет, <b>{first_name}</b>{vip_mark}!\n\n"
        f"📚 <b>ГОСТ-АССИСТЕНТ</b> — создаёт академические работы\n"
        f"в формате DOCX строго по ГОСТ 7.32-2017.\n\n"
        f"<b>Что умею:</b>\n"
        f"• 📄 Рефераты, курсовые, эссе, доклады\n"
        f"• 📑 Развёрнутые названия глав и подглав\n"
        f"• 📐 Точное соблюдение объёма (±10%)\n"
        f"• 🗂 Автоматическое содержание\n"
        f"• ⚙️ Настройка ГОСТ под ваш вуз\n\n"
        f"{limits}\n\n"
        f"👇 <b>Выберите тип работы:</b>"
    )


def _doc_type_card(doc_type: str, user_id: int) -> str:
    dt     = DOC_TYPES[doc_type]
    limits = get_user_limits_info(user_id)
    return (
        f"✅ <b>{dt['name']}</b>\n\n"
        f"📖 <i>{dt['desc']}</i>\n\n"
        f"📄 Объём: <b>{dt['min_pages']}–{dt['max_pages']} страниц</b>\n"
        f"🗂 Структура: {dt['structure']}\n\n"
        f"{limits}\n\n"
        f"💳 <b>Выберите режим генерации:</b>"
    )


def _mode_free_text() -> str:
    return (
        "🆓 <b>Бесплатный режим</b>\n\n"
        f"• Модель: {AI_MODELS.get(FREE_MODEL_KEY, {}).get('name', 'DeepSeek')}\n"
        f"• Максимум: <b>{FREE_MAX_PAGES} страниц</b>\n"
        f"• Лимит: <b>{FREE_DAILY_LIMIT} генерация в день</b>\n\n"
        "✍️ Введите <b>тему работы</b> одной строкой:"
    )


def _mode_paid_text() -> str:
    paid_str = "♾ безлимитно" if PAID_DAILY_LIMIT == 0 else f"{PAID_DAILY_LIMIT} в день"
    return (
        "⭐ <b>Платный режим</b>\n\n"
        f"• Любая доступная модель ИИ\n"
        f"• Любое количество страниц\n"
        f"• Генерации: <b>{paid_str}</b>\n\n"
        "✍️ Введите <b>тему работы</b> одной строкой:"
    )


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — /start и общие команды
# ═══════════════════════════════════════════════════════════════

@dp.message(Command("start"))
async def cmd_start(message: Message, state: FSMContext) -> None:
    await state.clear()
    first = message.from_user.first_name or "пользователь"
    await message.answer(
        _welcome_text(message.from_user.id, first),
        reply_markup=kb_doc_type(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.doc_type)


@dp.message(Command("help"))
async def cmd_help(message: Message, state: FSMContext) -> None:
    await message.answer(
        "📖 <b>Помощь по ГОСТ-ассистенту</b>\n\n"
        "<b>Команды:</b>\n"
        "/start — начать новую работу\n"
        "/help — эта справка\n"
        "/limits — мои лимиты\n"
        "/cancel — отменить текущий ввод\n\n"
        "<b>Как это работает:</b>\n"
        "1️⃣ Выбираете тип документа\n"
        "2️⃣ Указываете тему, учреждение, данные\n"
        "3️⃣ ИИ генерирует текст и формирует DOCX\n"
        "4️⃣ Получаете готовый файл по ГОСТ\n\n"
        "<b>Платный режим:</b> выбор ИИ-модели, любой объём, без ежедневного лимита.\n\n"
        "<b>Проблемы?</b> Обратитесь к администратору бота.",
        parse_mode="HTML",
    )


@dp.message(Command("limits"))
async def cmd_limits(message: Message) -> None:
    await message.answer(
        f"📊 <b>Ваши генерации</b>\n\n{get_user_limits_info(message.from_user.id)}",
        parse_mode="HTML",
    )


@dp.message(Command("cancel"))
async def cmd_cancel(message: Message, state: FSMContext) -> None:
    await state.clear()
    await message.answer(
        "❌ <b>Отменено.</b>\n\nНачните заново — /start",
        parse_mode="HTML",
    )


@dp.callback_query(F.data == "cancel_flow")
async def h_cancel_flow(cb: CallbackQuery, state: FSMContext) -> None:
    await state.clear()
    await cb.message.edit_text(
        "❌ <b>Генерация отменена.</b>\n\nНачните заново — /start",
        parse_mode="HTML",
    )
    await cb.answer()


@dp.callback_query(F.data == "show_limits")
async def h_show_limits(cb: CallbackQuery) -> None:
    await cb.answer(
        get_user_limits_info(cb.from_user.id).replace("<b>", "").replace("</b>", ""),
        show_alert=True,
    )


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — НАСТРОЙКА ГОСТ
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data == "custom_gost")
async def h_custom_gost(cb: CallbackQuery, state: FSMContext) -> None:
    await cb.message.answer(
        "⚙️ <b>Настройка ГОСТ под ваш вуз</b>\n\n"
        "Вставьте требования к оформлению обычным текстом.\n\n"
        "<b>Пример:</b>\n"
        "<i>Шрифт Arial 12pt, поля: лево 20 мм, право 10 мм, верх 15 мм, низ 15 мм. "
        "Межстрочный интервал 1.0. Отступ первой строки 1.25 см. Нумерация снизу.</i>\n\n"
        "ИИ автоматически извлечёт параметры и сохранит для вашего аккаунта.",
        parse_mode="HTML",
        reply_markup=kb_cancel(),
    )
    await state.set_state(WorkState.gost_free_text)
    await cb.answer()


@dp.message(WorkState.gost_free_text)
async def h_save_custom_gost(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    doc_type = data.get("doc_type", "referat")

    gost_text = (message.text or "").strip()
    if len(gost_text) < 10:
        await message.answer(
            "❌ Слишком коротко. Вставьте полные требования к оформлению.",
            reply_markup=kb_cancel(),
        )
        return

    wait_msg = await message.answer(
        "⏳ <b>Анализирую требования...</b>",
        parse_mode="HTML",
    )

    parsed = await parse_custom_gost_via_ai(FREE_MODEL_KEY, gost_text)
    save_user_gost_config(message.from_user.id, doc_type, parsed)

    await wait_msg.delete()

    pn = "снизу" if parsed["page_number_position"] == "bottom_center" else "сверху"
    await message.answer(
        "✅ <b>ГОСТ сохранён!</b>\n\n"
        "┌─────────────────────────\n"
        f"│ 🔤 Шрифт: <b>{parsed['font_name']} {parsed['font_size']}pt</b>\n"
        f"│ ↕️ Интервал: <b>{parsed['line_spacing']}</b>\n"
        f"│ ↩️ Отступ: <b>{parsed['first_line_indent_cm']} см</b>\n"
        f"│ 📐 Поля (мм): Л{parsed['left_margin_mm']} П{parsed['right_margin_mm']} "
        f"В{parsed['top_margin_mm']} Н{parsed['bottom_margin_mm']}\n"
        f"│ 🔢 Нумерация: <b>{pn} по центру</b>\n"
        "└─────────────────────────\n\n"
        "Эти настройки применятся к вашим следующим работам.",
        parse_mode="HTML",
    )

    await state.set_state(WorkState.doc_type)
    await message.answer(
        "👇 Выберите тип работы для продолжения:",
        reply_markup=kb_doc_type(),
    )


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — ВЫБОР ТИПА ДОКУМЕНТА
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data.startswith("dtype_"))
async def h_doc_type(cb: CallbackQuery, state: FSMContext) -> None:
    doc_type = cb.data.replace("dtype_", "", 1)
    if doc_type not in DOC_TYPES:
        await cb.answer("Неизвестный тип документа", show_alert=True)
        return

    await state.update_data(doc_type=doc_type)

    if doc_type == "custom":
        await cb.message.edit_text(
            "🧩 <b>Свой тип документа</b>\n\n"
            "Введите название документа.\n"
            "<i>Примеры: Отчёт по практике, Лабораторная работа, Реферат по материалам</i>",
            parse_mode="HTML",
            reply_markup=kb_cancel(),
        )
        await state.set_state(WorkState.custom_doc_name)
        await cb.answer()
        return

    await cb.message.edit_text(
        _doc_type_card(doc_type, cb.from_user.id),
        reply_markup=kb_mode(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.mode)
    await cb.answer()


@dp.message(WorkState.custom_doc_name)
async def h_custom_doc_name(message: Message, state: FSMContext) -> None:
    name = (message.text or "").strip()
    if len(name) < 2:
        await message.answer(
            "❌ Название слишком короткое. Попробуйте ещё раз.",
            reply_markup=kb_cancel(),
        )
        return

    await state.update_data(custom_doc_name=name)
    await message.answer(
        f"🧩 <b>Свой тип:</b> {name}\n\n"
        f"{get_user_limits_info(message.from_user.id)}\n\n"
        "💳 <b>Выберите режим генерации:</b>",
        reply_markup=kb_mode(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.mode)


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — РЕЖИМ И ТЕМА
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data.startswith("mode_"))
async def h_mode(cb: CallbackQuery, state: FSMContext) -> None:
    mode = cb.data.replace("mode_", "", 1)
    await state.update_data(mode=mode)

    await cb.message.edit_text(
        "✍️ <b>Выберите стиль написания работы</b>\n\n"
        "Это влияет на язык, структуру предложений и глубину изложения:",
        parse_mode="HTML",
        reply_markup=kb_writing_style(),
    )
    await state.set_state(WorkState.writing_style)
    await cb.answer()


@dp.callback_query(F.data.in_(["style_smart", "style_classic"]))
async def h_writing_style(cb: CallbackQuery, state: FSMContext) -> None:
    style = cb.data.replace("style_", "", 1)
    await state.update_data(writing_style=style)

    data = await state.get_data()
    mode = data.get("mode", "free")
    text = _mode_free_text() if mode == "free" else _mode_paid_text()
    style_label = "🎓 Умный стиль выбран" if style == "smart" else "📝 Классический стиль выбран"

    await cb.message.edit_text(
        f"✅ {style_label}\n\n{text}",
        parse_mode="HTML",
        reply_markup=kb_cancel(),
    )
    await state.set_state(WorkState.topic)
    await cb.answer()


@dp.message(WorkState.topic)
async def h_topic(message: Message, state: FSMContext) -> None:
    topic = (message.text or "").strip()
    if len(topic) < 5:
        await message.answer(
            "❌ Тема слишком короткая. Опишите тему подробнее (минимум 5 символов).",
            reply_markup=kb_cancel(),
        )
        return

    await state.update_data(topic=topic)
    await message.answer(
        f"✅ Тема принята: <b>{topic[:80]}</b>\n\n"
        "📎 <b>Есть ли у вас свои материалы?</b>\n\n"
        "Вы можете добавить план, конспект или тезисы — ИИ использует их как основу.\n"
        "Или выбрать генерацию только по теме.",
        reply_markup=kb_source_choice(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.source_choice)


@dp.callback_query(F.data == "source_yes")
async def h_source_yes(cb: CallbackQuery, state: FSMContext) -> None:
    """Пользователь выбрал 'Да — добавлю свои материалы'"""
    await cb.message.edit_text(
        "📎 <b>Отправьте ваши материалы</b>\n\n"
        "Вставьте план, тезисы, конспект или любой текст — ИИ использует его как основу.\n"
        "<i>Максимум 12 000 символов.</i>",
        parse_mode="HTML",
        reply_markup=kb_cancel(),  # <-- ИСПРАВЛЕНО
    )
    await state.set_state(WorkState.source_content)
    await cb.answer()

@dp.callback_query(F.data == "source_no")
async def h_source_no(cb: CallbackQuery, state: FSMContext) -> None:
    """Пользователь выбрал 'Нет — только по теме'"""
    await state.update_data(source_content="")
    await cb.message.edit_text(
        "🏛 <b>Тип учебного заведения</b>\n\nВыберите из списка:",
        reply_markup=kb_institution(),  # <-- ИСПРАВЛЕНО: используем правильную клавиатуру
        parse_mode="HTML",
    )
    await state.set_state(WorkState.institution_type)
    await cb.answer()

@dp.message(WorkState.source_content)
async def h_source_content(message: Message, state: FSMContext) -> None:
    content = (message.text or "").strip()
    await state.update_data(source_content=content)
    await message.answer(
        f"✅ Материалы приняты ({len(content)} символов).\n\n"
        "🏛 <b>Тип учебного заведения:</b>",
        reply_markup=kb_institution(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.institution_type)


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — УЧЕБНОЕ ЗАВЕДЕНИЕ
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data.startswith("inst_"))
async def h_inst(cb: CallbackQuery, state: FSMContext) -> None:
    inst = cb.data.replace("inst_", "", 1)
    await state.update_data(institution_type=inst)

    if inst == "custom":
        await state.update_data(org_type="")
        await cb.message.edit_text(
            "✏️ <b>Введите полное название учебного заведения</b>\n\n"
            "<i>Например: Новосибирский государственный технический университет</i>",
            parse_mode="HTML",
            reply_markup=kb_cancel(),
        )
        await state.set_state(WorkState.institution)
    else:
        info = INSTITUTION_TYPES.get(inst, INSTITUTION_TYPES["school"])
        await cb.message.edit_text(
            "🏛 <b>Введите тип организации</b>\n\n"
            f"Пример:\n<i>{info['org_example']}</i>\n\n"
            "Или скопируйте пример целиком.",
            parse_mode="HTML",
            reply_markup=kb_cancel(),
        )
        await state.set_state(WorkState.org_type)

    await cb.answer()


@dp.message(WorkState.org_type)
async def h_org_type(message: Message, state: FSMContext) -> None:
    await state.update_data(org_type=(message.text or "").strip())
    data = await state.get_data()
    info = INSTITUTION_TYPES.get(data.get("institution_type", "school"), INSTITUTION_TYPES["school"])
    await message.answer(
        "🏫 <b>Введите название учебного заведения</b>\n\n"
        f"Пример:\n<i>{info['name_example']}</i>",
        parse_mode="HTML",
        reply_markup=kb_cancel(),
    )
    await state.set_state(WorkState.institution)


@dp.message(WorkState.institution)
async def h_institution(message: Message, state: FSMContext) -> None:
    await state.update_data(institution=(message.text or "").strip())
    await message.answer(
        "👥 <b>Введите класс или группу</b>\n\n<i>Например: 10А, ИТ-21, гр. 315</i>",
        parse_mode="HTML",
        reply_markup=kb_cancel(),
    )
    await state.set_state(WorkState.group)


@dp.message(WorkState.group)
async def h_group(message: Message, state: FSMContext) -> None:
    await state.update_data(group=(message.text or "").strip())
    await message.answer(
        "👤 <b>Введите ФИО автора</b>\n\n<i>Например: Иванов Иван Иванович</i>",
        parse_mode="HTML",
        reply_markup=kb_cancel(),
    )
    await state.set_state(WorkState.author)


@dp.message(WorkState.author)
async def h_author(message: Message, state: FSMContext) -> None:
    author = (message.text or "").strip()
    if len(author.split()) < 2:
        await message.answer(
            "❌ Введите ФИО полностью (минимум имя и фамилия).",
            reply_markup=kb_cancel(),
        )
        return
    await state.update_data(author=author)
    await message.answer(
        "👨‍🏫 <b>Введите ФИО преподавателя</b>\n\n<i>Например: Петров Пётр Петрович</i>",
        parse_mode="HTML",
        reply_markup=kb_cancel(),
    )
    await state.set_state(WorkState.teacher)


@dp.message(WorkState.teacher)
async def h_teacher(message: Message, state: FSMContext) -> None:
    await state.update_data(teacher=(message.text or "").strip())
    await message.answer(
        "📚 <b>Выберите дисциплину (предмет)</b>",
        reply_markup=kb_subject(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.subject)


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — ПРЕДМЕТ, ГОРОД, СТРАНИЦЫ
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data.startswith("subj_"))
async def h_subject_cb(cb: CallbackQuery, state: FSMContext) -> None:
    subj = cb.data.replace("subj_", "", 1)
    if subj == "other":
        await cb.message.edit_text(
            "✏️ <b>Введите название предмета</b>",
            parse_mode="HTML",
            reply_markup=kb_cancel(),
        )
        await state.set_state(WorkState.subject)
    else:
        await state.update_data(subject=subj)
        await cb.message.edit_text(
            f"✅ Предмет: <b>{subj}</b>\n\n"
            "🌆 <b>Выберите город</b>:",
            reply_markup=kb_city(),
            parse_mode="HTML",
        )
        await state.set_state(WorkState.city)
    await cb.answer()


@dp.message(WorkState.subject)
async def h_subject_text(message: Message, state: FSMContext) -> None:
    await state.update_data(subject=(message.text or "").strip())
    await message.answer(
        "🌆 <b>Выберите город:</b>",
        reply_markup=kb_city(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.city)


@dp.callback_query(F.data.startswith("city_"))
async def h_city_cb(cb: CallbackQuery, state: FSMContext) -> None:
    city = cb.data.replace("city_", "", 1)
    if city == "other":
        await cb.message.edit_text(
            "✏️ <b>Введите название города</b>",
            parse_mode="HTML",
            reply_markup=kb_cancel(),
        )
        await state.set_state(WorkState.city)
    else:
        await state.update_data(city=city)
        await _ask_pages(cb.message, state, cb.from_user.id)
        await state.set_state(WorkState.pages)
    await cb.answer()


@dp.message(WorkState.city)
async def h_city_text(message: Message, state: FSMContext) -> None:
    await state.update_data(city=(message.text or "").strip())
    await _ask_pages(message, state, message.from_user.id)
    await state.set_state(WorkState.pages)


async def _ask_pages(event: Message, state: FSMContext, user_id: int) -> None:
    data     = await state.get_data()
    doc_type = data.get("doc_type", "referat")
    dt       = DOC_TYPES.get(doc_type, DOC_TYPES["referat"])
    mode     = data.get("mode", "free")

    if mode == "free" and not is_vip(user_id):
        max_p = min(dt["max_pages"], FREE_MAX_PAGES)
        note  = f"\n⚠️ <i>В бесплатном режиме максимум {FREE_MAX_PAGES} стр.</i>"
    else:
        max_p = dt["max_pages"]
        note  = ""

    await event.answer(
        f"📄 <b>Количество страниц</b>\n\n"
        f"Допустимо: <b>{dt['min_pages']}–{max_p}</b> страниц{note}\n\n"
        f"Введите число:",
        parse_mode="HTML",
        reply_markup=kb_cancel(),
    )


@dp.message(WorkState.pages)
async def h_pages(message: Message, state: FSMContext) -> None:
    data     = await state.get_data()
    doc_type = data.get("doc_type", "referat")
    dt       = DOC_TYPES.get(doc_type, DOC_TYPES["referat"])

    try:
        pages = int((message.text or "").strip())
    except Exception:
        await message.answer("❌ Введите целое число.", reply_markup=kb_cancel())
        return

    mode = data.get("mode", "free")

    # Клипируем до допустимого диапазона
    min_p = dt["min_pages"]
    if mode == "free" and not is_vip(message.from_user.id):
        max_p = min(dt["max_pages"], FREE_MAX_PAGES)
    else:
        max_p = dt["max_pages"]

    if pages < min_p or pages > max_p:
        await message.answer(
            f"❌ Допустимое количество страниц: <b>{min_p}–{max_p}</b>.\n"
            f"Вы ввели: {pages}. Попробуйте ещё раз.",
            parse_mode="HTML",
            reply_markup=kb_cancel(),
        )
        return

    await state.update_data(pages=pages)
    await message.answer(
        f"✅ Страниц: <b>{pages}</b>\n\n"
        "🔢 <b>Нумерация страниц — где расположить?</b>",
        reply_markup=kb_page_number(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.page_number_position)


# ═══════════════════════════════════════════════════════════════
#  ХЭНДЛЕРЫ — НУМЕРАЦИЯ, МОДЕЛЬ, ОПЛАТА
# ═══════════════════════════════════════════════════════════════

@dp.callback_query(F.data.in_(["pagepos_bottom", "pagepos_top"]))
async def h_pagepos(cb: CallbackQuery, state: FSMContext) -> None:
    pos = "top_center" if cb.data == "pagepos_top" else "bottom_center"
    await state.update_data(page_number_position=pos)

    data  = await state.get_data()
    pages = int(data.get("pages", 10))
    mode  = data.get("mode", "free")

    # ── Бесплатный режим ──
    if mode == "free":
        if pages > FREE_MAX_PAGES and not is_vip(cb.from_user.id):
            await cb.message.edit_text(
                f"🚫 <b>В бесплатном режиме максимум {FREE_MAX_PAGES} страниц.</b>\n\n"
                "Используйте платный режим для большего объёма.",
                parse_mode="HTML",
            )
            await state.clear()
            await cb.answer()
            return

        ok, reason = check_user_limit(cb.from_user.id, "free")
        if not ok:
            await cb.message.edit_text(reason, parse_mode="HTML")
            await state.clear()
            await cb.answer()
            return

        await cb.message.edit_text(
            "🚀 <b>Запускаю генерацию...</b>\n\nПодождите, это займёт несколько минут.",
            parse_mode="HTML",
        )
        await cb.answer()
        await generate_and_send(cb.message, state, model_key=FREE_MODEL_KEY, pay_mode="free")
        return

    # ── Платный режим ──
    ok, reason = check_user_limit(cb.from_user.id, "paid")
    if not ok and not is_vip(cb.from_user.id):
        await cb.message.edit_text(reason, parse_mode="HTML")
        await state.clear()
        await cb.answer()
        return

    await cb.message.edit_text(
        "🤖 <b>Выберите ИИ-модель</b>\n\n"
        "Цена указана в звёздах Telegram за страницу.\n"
        "Все модели генерируют полноценный академический текст.",
        reply_markup=kb_models(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.model)
    await cb.answer()


@dp.callback_query(F.data.startswith("model_"))
async def h_model(cb: CallbackQuery, state: FSMContext) -> None:
    model_key = cb.data.replace("model_", "", 1)
    if model_key not in AI_MODELS or not AI_MODELS[model_key].get("api_key"):
        await cb.answer("⚠️ Модель временно недоступна", show_alert=True)
        return

    data  = await state.get_data()
    pages = int(data.get("pages", 10))
    model = AI_MODELS[model_key]
    total = int(model["price_per_page"]) * pages

    await state.update_data(model_key=model_key)

    if is_vip(cb.from_user.id):
        await cb.message.edit_text(
            f"👑 <b>VIP — оплата не требуется</b>\n\n"
            f"Модель: {model['name']}\n"
            f"Страниц: {pages}\n\n"
            f"🚀 Запускаю генерацию...",
            parse_mode="HTML",
        )
        await cb.answer()
        await generate_and_send(cb.message, state, model_key=model_key, pay_mode="paid")
        return

    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text=f"⭐ Оплатить {total} звёзд", callback_data=f"pay_{total}")],
            [InlineKeyboardButton(text="← Выбрать другую модель",   callback_data="back_to_models")],
        ]
    )
    await cb.message.edit_text(
        f"💳 <b>Подтверждение оплаты</b>\n\n"
        f"┌─────────────────────────\n"
        f"│ 🤖 Модель:  {model['name']}\n"
        f"│ 📄 Страниц: {pages}\n"
        f"│ 💰 Цена:    {model['price_per_page']}⭐ × {pages} = <b>{total}⭐</b>\n"
        f"└─────────────────────────\n\n"
        f"Нажмите кнопку для оплаты через Telegram Stars:",
        reply_markup=kb,
        parse_mode="HTML",
    )
    await state.set_state(WorkState.payment)
    await cb.answer()


@dp.callback_query(F.data == "back_to_models")
async def h_back_to_models(cb: CallbackQuery, state: FSMContext) -> None:
    await cb.message.edit_text(
        "🤖 <b>Выберите ИИ-модель</b>",
        reply_markup=kb_models(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.model)
    await cb.answer()


@dp.callback_query(F.data.startswith("pay_"))
async def h_pay(cb: CallbackQuery, state: FSMContext) -> None:
    stars     = int(cb.data.replace("pay_", "", 1))
    data      = await state.get_data()
    model_key = data.get("model_key", FREE_MODEL_KEY)
    payload   = json.dumps(
        {"model_key": model_key, "pages": data.get("pages", 10)},
        ensure_ascii=False,
    )

    await cb.bot.send_invoice(
        chat_id=cb.from_user.id,
        title="Академическая работа по ГОСТ",
        description=f"{AI_MODELS[model_key]['name']} · {data.get('pages', 10)} стр. · {data.get('topic','')[:40]}",
        payload=payload,
        provider_token="",
        currency="XTR",
        prices=[LabeledPrice(label="Работа по ГОСТ", amount=stars)],
        reply_markup=InlineKeyboardMarkup(
            inline_keyboard=[[InlineKeyboardButton(text=f"⭐ Оплатить {stars} звёзд", pay=True)]]
        ),
    )
    await cb.answer()


@dp.pre_checkout_query()
async def h_pre_checkout(q: PreCheckoutQuery) -> None:
    await q.answer(ok=True)


@dp.message(F.successful_payment)
async def h_payment_ok(message: Message, state: FSMContext) -> None:
    try:
        payload   = json.loads(message.successful_payment.invoice_payload)
        model_key = payload.get("model_key", FREE_MODEL_KEY)
    except Exception:
        model_key = FREE_MODEL_KEY

    await message.answer(
        "✅ <b>Оплата прошла успешно!</b>\n\n"
        "🚀 Начинаю генерацию работы...",
        parse_mode="HTML",
    )
    await generate_and_send(message, state, model_key=model_key, pay_mode="paid")


@dp.callback_query(F.data == "final_new")
async def h_final_new(cb: CallbackQuery, state: FSMContext) -> None:
    await state.clear()
    first = cb.from_user.first_name or "пользователь"
    await cb.message.edit_text(
        _welcome_text(cb.from_user.id, first),
        reply_markup=kb_doc_type(),
        parse_mode="HTML",
    )
    await state.set_state(WorkState.doc_type)
    await cb.answer()


# ═══════════════════════════════════════════════════════════════
#  ГЛАВНАЯ ГЕНЕРАЦИЯ
# ═══════════════════════════════════════════════════════════════

async def generate_and_send(
    event: Message,
    state: FSMContext,
    model_key: str,
    pay_mode: str,
) -> None:
    """
    Основная функция генерации: запрашивает названия глав,
    генерирует текст блоками, собирает DOCX, отправляет.
    """
    async with GEN_SEMAPHORE:
        data = await state.get_data()

        doc_type = data.get("doc_type", "referat")
        dt       = DOC_TYPES.get(doc_type, DOC_TYPES["referat"])
        pages    = int(data.get("pages", 10))
        topic    = data.get("topic", "")
        subject  = data.get("subject", "")

        # Определяем количество глав по типу документа
        if doc_type in ("esse", "doklad"):
            num_chapters = 0  # у них своя структура
        elif doc_type in ("kursovaya", "final_project"):
            num_chapters = 3
        else:
            num_chapters = 2  # реферат, контрольная, свой

        # Считаем шаги для прогресс-бара
        # 1 = названия глав, num_chapters+3 = блоки текста, 1 = DOCX, 1 = отправка
        extra_blocks = 2 if doc_type in ("esse",) else (num_chapters + 3 if num_chapters else 4)
        total_steps  = 1 + extra_blocks + 2

        progress_msg = await event.answer(
            "⏳ <b>Инициализация...</b>",
            parse_mode="HTML",
        )
        prog = Progress(
            msg=progress_msg,
            title=f"Создаю {dt['word'].lower()}",
            total_steps=total_steps,
            model_name=AI_MODELS.get(model_key, {}).get("name", ""),
        )

        # Запускаем фоновый цикл анимации
        stop_anim = asyncio.Event()
        anim_task = asyncio.create_task(prog.animate_loop(stop_anim))

        try:
            record_user_generation(event.chat.id, pay_mode)

            gost = get_gost_config(doc_type, event.chat.id)
            gost["page_number_position"] = data.get(
                "page_number_position",
                gost.get("page_number_position", "bottom_center"),
            )

            # ── Шаг 1: Генерируем названия глав ──
            await prog.update(label="🧠 Придумываю названия глав...", force=True)
            if num_chapters > 0:
                chapter_titles = await generate_chapter_titles(
                    model_key, doc_type, topic, subject, num_chapters
                )
            else:
                chapter_titles = []
            await prog.update(step_done=True)

            writing_style = data.get("writing_style", "classic")

            # ── Шаги 2–N: Генерируем текст ──
            await prog.update(label="✍️ Генерирую текст разделов...")
            parts = await generate_text_blocks(
                topic=topic,
                pages=pages,
                doc_type=doc_type,
                subject=subject,
                model_key=model_key,
                source=data.get("source_content", ""),
                chapter_titles=chapter_titles,
                writing_style=writing_style,
                prog=prog,
            )

            # ── Сборка структуры ──
            blocks = generate_structure(doc_type, parts, chapter_titles)

            # ── DOCX ──
            await prog.update(label="📄 Собираю DOCX-документ...", step_done=True)
            docx_raw = build_docx_bytes(data, blocks, gost)

            # ── LibreOffice ──
            await prog.update(label="🔄 Обновляю содержание (LibreOffice)...", step_done=True)
            work_dir = os.path.join(os.getcwd(), "_out")
            os.makedirs(work_dir, exist_ok=True)
            ts      = datetime.now().strftime("%Y%m%d_%H%M%S")
            tmp_in  = os.path.join(work_dir, f"tmp_{event.chat.id}_{ts}.docx")
            tmp_out = os.path.join(work_dir, f"final_{event.chat.id}_{ts}.docx")

            with open(tmp_in, "wb") as f:
                f.write(docx_raw)

            updated    = libreoffice_update_docx(tmp_in, tmp_out)
            final_path = tmp_out if updated else tmp_in

            # ── Имя файла ──
            safe_topic = re.sub(r'[<>"/:\\|?*]', "", topic[:35]).replace(" ", "_")
            fname      = f"{dt['word'].replace(' ', '_')}_{safe_topic}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

            with open(final_path, "rb") as f:
                final_bytes = f.read()

            # ── Отправляем ──
            await prog.update(label="📤 Отправляю файл...", step_done=True)

            toc_status = "✅ обновлено автоматически" if updated else "⚠️ обновите вручную в Word"
            style_label = "🎓 Умный" if data.get("writing_style") == "smart" else "📝 Классический"
            caption = (
                f"🎉 <b>{dt['word']} ГОТОВ!</b>\n\n"
                f"┌─────────────────────────\n"
                f"│ 📖 Тема: {topic[:60]}\n"
                f"│ 📄 Страниц: ~{pages}\n"
                f"│ 🤖 ИИ: {AI_MODELS.get(model_key, {}).get('name', model_key)}\n"
                f"│ ✍️ Стиль: {style_label}\n"
                f"│ 📐 Шрифт: {gost.get('font_name')} {gost.get('font_size')}pt\n"
                f"│ ↕️ Интервал: {gost.get('line_spacing')}\n"
                f"│ 📑 Содержание: {toc_status}\n"
                f"└─────────────────────────"
            )

            await event.answer_document(
                BufferedInputFile(final_bytes, filename=fname),
                caption=caption,
                parse_mode="HTML",
            )

        except Exception as e:
            print(f"[GEN ERROR] {e}")
            await prog.finish(
                "❌ <b>Ошибка генерации</b>\n\n"
                f"Причина: {str(e)[:200]}\n\n"
                "Попробуйте ещё раз или выберите другую модель (/start)."
            )
            await state.clear()
            return

        finally:
            stop_anim.set()
            anim_task.cancel()
            try:
                await anim_task
            except asyncio.CancelledError:
                pass

        await prog.delete()

        await event.answer(
            "✅ <b>Работа готова!</b>\n\n"
            "Если содержание не обновилось — откройте файл в Word и нажмите:\n"
            "<code>Ctrl+A → F9 → Обновить всё поле</code>",
            reply_markup=kb_final(),
            parse_mode="HTML",
        )
        await state.clear()

        # Чистим временные файлы
        for tmp in (tmp_in, tmp_out):
            try:
                if os.path.exists(tmp):
                    os.remove(tmp)
            except Exception:
                pass


# ═══════════════════════════════════════════════════════════════
#  ЗАПУСК
# ═══════════════════════════════════════════════════════════════

async def main() -> None:
    print("═" * 62)
    print("  🤖  ГОСТ-АССИСТЕНТ v2.0")
    print("═" * 62)
    print(f"  LibreOffice : {shutil.which('soffice') or '❌ не найден'}")
    print(f"  DeepSeek    : {'✅' if DEEPSEEK_KEY else '❌ нет ключа'}")
    print(f"  OpenRouter  : {'✅' if OPENROUTER_KEY else '❌ нет ключа'}")
    print(f"  Groq        : {'✅' if GROQ_KEY else '❌ нет ключа'}")
    print(f"  VIP users   : {VIP_USERS or 'нет'}")
    print(f"  Free limit  : {FREE_DAILY_LIMIT}/день, max {FREE_MAX_PAGES} стр.")
    print(f"  Paid limit  : {'безлимит' if PAID_DAILY_LIMIT == 0 else str(PAID_DAILY_LIMIT)+'/день'}")
    print("═" * 62)

    await _start_http_server()

    while True:
        try:
            print("[BOT] Запускаю polling...")
            await dp.start_polling(bot)
        except Exception as e:
            print(f"[BOT] Ошибка поллинга: {e}")
            await asyncio.sleep(5)


if __name__ == "__main__":
    asyncio.run(main())
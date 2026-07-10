# -*- coding: utf-8 -*-
"""Смоук-тест gost_images: вставка по ГОСТ, валидация, авто-фикс брака."""
import io

from docx import Document
from docx.shared import Pt

from gost_images import (
    FigureNumberer, add_gost_figure, autofix_figures, clean_caption_title,
    format_caption, make_reference, prepare_image, validate_figures,
    calc_display_size,
)


def _png_bytes(w=800, h=500, color=(120, 60, 200)):
    from PIL import Image
    buf = io.BytesIO()
    Image.new("RGB", (w, h), color).save(buf, format="PNG")
    return buf.getvalue()


def test_caption_formatting():
    assert format_caption("1", "схема процесса.") == "Рисунок 1 — Схема процесса"
    assert format_caption("2.1", "**Рис. 3 - динамика ВВП**") == "Рисунок 2.1 — Динамика ВВП"
    assert clean_caption_title("рис. 5: график") == "График"
    assert make_reference("3") == "(рисунок 3)"
    assert make_reference("3", "full") == "в соответствии с рисунком 3"
    print("OK: подписи")


def test_numberer():
    n = FigureNumberer()
    assert [n.next(), n.next(), n.next()] == ["1", "2", "3"]
    nc = FigureNumberer(per_chapter=True)
    assert [nc.next(1), nc.next(1), nc.next(2)] == ["1.1", "1.2", "2.1"]
    assert nc.next_appendix("А") == "А.1"
    print("OK: нумерация")


def test_prepare_and_size():
    ok, data, w, h = prepare_image(_png_bytes())
    assert ok and w == 800 and h == 500
    ok2, *_ = prepare_image(b"not an image at all, definitely broken bytes xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx")
    assert not ok2
    ok3, *_ = prepare_image(_png_bytes(50, 50))  # слишком мелкая
    assert not ok3
    w_cm, h_cm = calc_display_size(4000, 3000, 20.0)  # шире полей → ужмётся
    assert w_cm <= 16.5 and h_cm <= 18.0
    print("OK: подготовка изображений")


def test_insert_and_validate():
    doc = Document()
    p = doc.add_paragraph(
        "Структура процесса представлена ниже. Динамика показателей "
        "рассматривается в первом разделе работы.")
    num = FigureNumberer()
    n1 = num.next()
    r = add_gost_figure(doc, _png_bytes(), n1, "схема процесса управления.",
                        ensure_reference=p)
    assert r is not None
    img_par, cap_par = r
    assert cap_par.text == "Рисунок 1 — Схема процесса управления"
    assert "(рисунок 1)" in p.text  # ссылка дописана ДО рисунка
    p2 = doc.add_paragraph("Далее приводится сравнение подходов.")
    n2 = num.next()
    add_gost_figure(doc, _png_bytes(600, 400, (10, 150, 90)), n2,
                    "Сравнение подходов", ensure_reference=p2)
    issues = validate_figures(doc)
    errors = [i for i in issues if i["severity"] == "error"]
    assert not errors, f"неожиданные ошибки: {errors}"
    doc.save("/tmp/gost_ok.docx")
    print(f"OK: вставка и валидация (замечаний: {len(issues)})")


def test_validate_catches_bad():
    doc = Document()
    doc.add_paragraph("Текст без ссылки на рисунок.")
    par = doc.add_paragraph()
    par.add_run().add_picture(io.BytesIO(_png_bytes()), width=Pt(200))
    doc.add_paragraph("Рис. 1. схема.")  # брак: «Рис.», точки, не по центру
    issues = validate_figures(doc)
    titles = [i["title"] for i in issues]
    assert any("без подписи" in t for t in titles), titles
    print(f"OK: брак ловится ({titles})")

    fixes = autofix_figures(doc)
    assert fixes, "авто-фикс не сработал"
    issues2 = validate_figures(doc)
    errors2 = [i for i in issues2 if i["severity"] == "error"
               and "подписи" in i["title"]]
    assert not errors2, f"после фикса осталось: {errors2}"
    cap = [p.text for p in doc.paragraphs if p.text.startswith("Рисунок")]
    assert cap == ["Рисунок 1 — Схема"], cap
    print(f"OK: авто-фикс ({fixes})")


def test_full_checker():
    from checker import validate_docx
    res = validate_docx("/tmp/gost_ok.docx")
    fig_errors = [str(e) for e in res.errors + res.warnings if "исунок" in str(e)]
    # ссылки на рисунки есть → ошибок по рисункам быть не должно
    assert not any("без подписи" in s or "Нет ссылки" in s for s in fig_errors), fig_errors
    print(f"OK: checker интегрирован (score={res.score}, fig_issues={fig_errors})")


if __name__ == "__main__":
    test_caption_formatting()
    test_numberer()
    test_prepare_and_size()
    test_insert_and_validate()
    test_validate_catches_bad()
    test_full_checker()
    print("\nВСЕ ТЕСТЫ ПРОЙДЕНЫ ✅")
